#!/usr/bin/env python3
#yashdate
"""
GPA's Library Management System
Version v3.7 - College-wide Library Management System for Government Polytechnic Awasari
"""

# Load .env FIRST, before any other imports that might need DATABASE_URL
try:
    from dotenv import load_dotenv
    # Look for .env in project root (one level up from LibraryApp/)
    import os as _os
    _env_path = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))), '.env')
    if _os.path.exists(_env_path):
        load_dotenv(_env_path)
        print(f"[OK] Loaded environment from {_env_path}")
    else:
        # Also try current directory
        load_dotenv()
except ImportError:
    pass

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime, timedelta
import sqlite3
import os
import sys
from tkinter import font
import webbrowser
import subprocess
import platform
from io import BytesIO
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import json
import threading
import time
import socket
import qrcode
from PIL import ImageTk, Image

# Pandas for data operations
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except Exception:
    PANDAS_AVAILABLE = False
    pd = None

# ReportLab for PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    REPORTLAB_AVAILABLE = True
except Exception as e:
    REPORTLAB_AVAILABLE = False
    print(f"[WARNING] ReportLab not available: {e}")

# Add Web-Extension directory to path to allow import
sys.path.append(os.path.join(os.path.dirname(__file__), 'Web-Extension'))

# Optional: Web Portal support
try:
    from student_portal import app as flask_app  # type: ignore
    from waitress import serve  # type: ignore
    WEB_PORTAL_AVAILABLE = True
except Exception as e:
    import traceback
    traceback.print_exc()
    print(f"Portal Import Error: {e}")
    flask_app = None
    serve = None
    WEB_PORTAL_AVAILABLE = False
    print("Web portal not available - student portal features will be disabled")

# Optional: Word export support
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

# Calendar date picker support
try:
    from tkcalendar import DateEntry
except Exception:
    DateEntry = None

# Matplotlib for charts and analysis
try:
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    from matplotlib.figure import Figure
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except Exception:
    MATPLOTLIB_AVAILABLE = False
    print("matplotlib not available - Analysis tab will show installation prompt")

# Advanced Excel export support
try:
    import xlsxwriter
    XLSXWRITER_AVAILABLE = True
except Exception:
    XLSXWRITER_AVAILABLE = False

# Add the current directory to path for imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from database import Database
from login_loader import LoginLoader
from autocomplete_widget import AutocompleteEntry

# Performance Optimization Modules
try:
    from database_pool import get_pool, ConnectionPool
    from email_batch_service import EmailBatchService, send_overdue_emails_async
    from sync_manager import create_sync_manager, SyncManager
    from config_manager import get_config, ConfigManager
    PERFORMANCE_MODULES_AVAILABLE = True
except Exception as e:
    print(f"Performance modules not available: {e}")
    PERFORMANCE_MODULES_AVAILABLE = False
    ConnectionPool = None
    EmailBatchService = None
    SyncManager = None
    ConfigManager = None

# ---------------------------------------------------------------------------
# Application Version (update this each time you create a new packaged build)
# ---------------------------------------------------------------------------
APP_VERSION = "v5.0_FINAL"
FINE_PER_DAY = 5  # monetary units per day late
import hashlib
# Fixed loan period (teacher requirement): exactly 7 days between borrow and due date
LOAN_PERIOD_DAYS = 7
# Admin login credentials (required at startup)
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "gpa")
# Default hashed admin password (sha256 of "gpa123")
ADMIN_PASSWORD_HASH = os.getenv("ADMIN_PASSWORD_HASH", "d84a5a7b330fefd799295540c51d1a01774d7d3c7a57ba9f31b3601ee259f3fc")

class LibraryApp:
    def _sync_runtime_settings_now(self):
        """Best-effort fast push of system settings to cloud (non-blocking)."""
        try:
            if not getattr(self, 'sync_manager', None):
                return

            def _do_sync_settings():
                try:
                    self.sync_manager.sync_now(direction='local_to_remote', tables_override=['system_settings'])
                except Exception as e:
                    print(f"[Settings Sync] Immediate settings sync failed: {e}")

            threading.Thread(target=_do_sync_settings, daemon=True).start()
        except Exception as e:
            print(f"[Settings Sync] Unable to start immediate settings sync: {e}")

    def run_in_background_thread(self, target, callback, **kwargs):
        """Helper to run a function in a background thread and callback on main thread."""
        def wrapper():
            try:
                result = target(**kwargs)
                # Schedule callback on main thread
                self.root.after(0, lambda: callback(result))
            except Exception as e:
                print(f"Background thread error: {e}")
                self.root.after(0, lambda err=e: callback(err)) # Pass error to callback
        
        threading.Thread(target=wrapper, daemon=True).start()

    def _fetch_analysis_data(self, days=30, enrollment_no=None, book_id=None, branch=None):
        """Fetch all analysis data in a background thread. Supports branch filter."""
        data = {}
        conn = None
        try:
            conn = self.db.get_connection()
            cursor = conn.cursor()
            start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
            today = datetime.now().strftime('%Y-%m-%d')

            # Build optional branch JOIN/WHERE clause fragments
            branch_join = ""
            branch_where = ""
            branch_params = []
            if branch and branch != "All":
                branch_join = " JOIN students s_br ON br.enrollment_no = s_br.enrollment_no"
                branch_where = " AND s_br.department = ?"
                branch_params = [branch]
            
            # 1. Borrow Status (branch-filtered)
            if branch and branch != "All":
                cursor.execute(
                    "SELECT CASE WHEN br.status = 'borrowed' THEN 'Currently Issued' ELSE 'Available' END as status, "
                    "COUNT(DISTINCT b.book_id) as count FROM books b "
                    "LEFT JOIN borrow_records br ON b.book_id = br.book_id AND br.status = 'borrowed' "
                    "LEFT JOIN students s_f ON br.enrollment_no = s_f.enrollment_no "
                    "WHERE (br.id IS NULL OR s_f.department = ?) "
                    "GROUP BY status", (branch,))
            else:
                cursor.execute("SELECT CASE WHEN br.status = 'borrowed' THEN 'Currently Issued' ELSE 'Available' END as status, COUNT(DISTINCT b.book_id) as count FROM books b LEFT JOIN borrow_records br ON b.book_id = br.book_id AND br.status = 'borrowed' GROUP BY status")
            data['borrow_status'] = cursor.fetchall()

            # 2. Student Activity by Year (branch-filtered)
            if branch and branch != "All":
                cursor.execute(
                    "SELECT s.year, COUNT(br.id) as borrow_count FROM students s "
                    "LEFT JOIN borrow_records br ON s.enrollment_no = br.enrollment_no AND br.borrow_date >= ? "
                    "WHERE s.department = ? "
                    "GROUP BY s.year HAVING COUNT(br.id) > 0 ORDER BY borrow_count DESC",
                    (start_date, branch))
            else:
                cursor.execute("SELECT s.year, COUNT(br.id) as borrow_count FROM students s LEFT JOIN borrow_records br ON s.enrollment_no = br.enrollment_no AND br.borrow_date >= ? GROUP BY s.year HAVING COUNT(br.id) > 0 ORDER BY borrow_count DESC", (start_date,))
            data['student_activity'] = cursor.fetchall()

            # 3. Inventory & Overdue (branch-filtered overdue)
            cursor.execute("SELECT COALESCE(SUM(total_copies),0) as total, COALESCE(SUM(available_copies),0) as available FROM books")
            total_copies, total_available = cursor.fetchone()
            if branch and branch != "All":
                cursor.execute(
                    "SELECT COUNT(*) FROM borrow_records br JOIN students s_br ON br.enrollment_no = s_br.enrollment_no "
                    "WHERE br.status='borrowed' AND br.due_date < ? AND s_br.department = ?", (today, branch))
            else:
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE status='borrowed' AND due_date < ?", (today,))
            overdue = cursor.fetchone()[0] or 0
            data['inventory'] = {'total_copies': total_copies, 'total_available': total_available, 'overdue': overdue}

            # 4. Daily Trend (branch-filtered)
            if branch and branch != "All":
                cursor.execute(
                    "SELECT br.borrow_date, COUNT(*) FROM borrow_records br "
                    "JOIN students s_br ON br.enrollment_no = s_br.enrollment_no "
                    "WHERE br.borrow_date >= ? AND s_br.department = ? "
                    "GROUP BY br.borrow_date ORDER BY br.borrow_date",
                    (start_date, branch))
            else:
                cursor.execute("SELECT borrow_date, COUNT(*) FROM borrow_records WHERE borrow_date >= ? GROUP BY borrow_date ORDER BY borrow_date", (start_date,))
            data['daily_trend'] = cursor.fetchall()

            # 5. Popular Books (branch-filtered)
            if branch and branch != "All":
                cursor.execute(
                    "SELECT b.title, COUNT(br.id) as borrow_count FROM books b "
                    "INNER JOIN borrow_records br ON b.book_id = br.book_id "
                    "JOIN students s_br ON br.enrollment_no = s_br.enrollment_no "
                    "WHERE br.borrow_date >= ? AND s_br.department = ? "
                    "GROUP BY b.book_id, b.title ORDER BY borrow_count DESC, b.title ASC LIMIT 10",
                    (start_date, branch))
            else:
                cursor.execute("SELECT b.title, COUNT(br.id) as borrow_count FROM books b INNER JOIN borrow_records br ON b.book_id = br.book_id WHERE br.borrow_date >= ? GROUP BY b.book_id, b.title ORDER BY borrow_count DESC, b.title ASC LIMIT 10", (start_date,))
            data['popular_books'] = cursor.fetchall()

            # 6. Least Popular (branch-filtered)
            if branch and branch != "All":
                cursor.execute(
                    "SELECT b.title, COALESCE(COUNT(br.id), 0) as borrow_count FROM books b "
                    "LEFT JOIN borrow_records br ON b.book_id = br.book_id AND br.borrow_date >= ? "
                    "LEFT JOIN students s_br ON br.enrollment_no = s_br.enrollment_no "
                    "WHERE (br.id IS NULL OR s_br.department = ?) "
                    "GROUP BY b.book_id, b.title "
                    "HAVING COALESCE(COUNT(br.id), 0) = ("
                    "SELECT MIN(cnt) FROM ("
                    "SELECT COALESCE(COUNT(br2.id), 0) as cnt FROM books b2 "
                    "LEFT JOIN borrow_records br2 ON b2.book_id = br2.book_id AND br2.borrow_date >= ? "
                    "LEFT JOIN students s_br2 ON br2.enrollment_no = s_br2.enrollment_no "
                    "WHERE (br2.id IS NULL OR s_br2.department = ?) "
                    "GROUP BY b2.book_id) t) ORDER BY b.title ASC LIMIT 10",
                    (start_date, branch, start_date, branch))
            else:
                cursor.execute("SELECT b.title, COALESCE(COUNT(br.id), 0) as borrow_count FROM books b LEFT JOIN borrow_records br ON b.book_id = br.book_id AND br.borrow_date >= ? GROUP BY b.book_id, b.title HAVING COALESCE(COUNT(br.id), 0) = (SELECT MIN(cnt) FROM (SELECT COALESCE(COUNT(br2.id), 0) as cnt FROM books b2 LEFT JOIN borrow_records br2 ON b2.book_id = br2.book_id AND br2.borrow_date >= ? GROUP BY b2.book_id) t) ORDER BY b.title ASC LIMIT 10", (start_date, start_date))
            data['least_popular'] = cursor.fetchall()
            
            # 7. Summary Stats (branch-filtered)
            if branch and branch != "All":
                cursor.execute(
                    "SELECT COUNT(*) FROM borrow_records br JOIN students s_br ON br.enrollment_no = s_br.enrollment_no "
                    "WHERE br.borrow_date >= ? AND s_br.department = ?", (start_date, branch))
                total_borrowings = cursor.fetchone()[0]
                cursor.execute(
                    "SELECT COUNT(*) FROM borrow_records br JOIN students s_br ON br.enrollment_no = s_br.enrollment_no "
                    "WHERE br.return_date >= ? AND br.return_date IS NOT NULL AND s_br.department = ?", (start_date, branch))
                total_returns = cursor.fetchone()[0]
                cursor.execute(
                    "SELECT COUNT(DISTINCT br.enrollment_no) FROM borrow_records br JOIN students s_br ON br.enrollment_no = s_br.enrollment_no "
                    "WHERE br.borrow_date >= ? AND s_br.department = ?", (start_date, branch))
                active_students = cursor.fetchone()[0]
                cursor.execute(
                    "SELECT br.return_date, br.due_date FROM borrow_records br "
                    "JOIN students s_br ON br.enrollment_no = s_br.enrollment_no "
                    "WHERE br.return_date > br.due_date AND br.return_date IS NOT NULL AND br.return_date >= ? AND s_br.department = ?",
                    (start_date, branch))
                fines_data = cursor.fetchall()
            else:
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE borrow_date >= ?", (start_date,))
                total_borrowings = cursor.fetchone()[0]
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE return_date >= ? AND return_date IS NOT NULL", (start_date,))
                total_returns = cursor.fetchone()[0]
                cursor.execute("SELECT COUNT(DISTINCT enrollment_no) FROM borrow_records WHERE borrow_date >= ?", (start_date,))
                active_students = cursor.fetchone()[0]
                cursor.execute("SELECT return_date, due_date FROM borrow_records WHERE return_date > due_date AND return_date IS NOT NULL AND return_date >= ?", (start_date,))
                fines_data = cursor.fetchall()
            data['summary'] = {
                'total_borrowings': total_borrowings,
                'total_returns': total_returns,
                'overdue_count': overdue,
                'active_students': active_students,
                'fines_data': fines_data
            }

            # 8. Branch-wise Borrowing Distribution (always fetch, regardless of branch filter)
            cursor.execute(
                "SELECT s.department, COUNT(br.id) as borrow_count FROM students s "
                "JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
                "WHERE br.borrow_date >= ? AND s.department IS NOT NULL AND s.department != '' "
                "GROUP BY s.department ORDER BY borrow_count DESC", (start_date,))
            data['branch_activity'] = cursor.fetchall()

            # 9. Branch-wise Overdue Distribution (always fetch)
            cursor.execute(
                "SELECT s.department, COUNT(br.id) as overdue_count FROM students s "
                "JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
                "WHERE br.status = 'borrowed' AND br.due_date < ? AND s.department IS NOT NULL AND s.department != '' "
                "GROUP BY s.department ORDER BY overdue_count DESC", (today,))
            data['branch_overdue'] = cursor.fetchall()

            # 10. Focused Insights
            if enrollment_no:
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE enrollment_no=? AND borrow_date>=?", (enrollment_no, start_date))
                s_total = cursor.fetchone()[0]
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE enrollment_no=? AND borrow_date>=? AND status='borrowed'", (enrollment_no, start_date))
                s_active = cursor.fetchone()[0]
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE enrollment_no=? AND return_date>=? AND return_date IS NOT NULL", (enrollment_no, start_date))
                s_returned = cursor.fetchone()[0]
                data['student_specific'] = {'total': s_total, 'active': s_active, 'returned': s_returned}
            
            if book_id:
                cursor.execute("SELECT title, total_copies, available_copies FROM books WHERE book_id=?", (book_id,))
                b_row = cursor.fetchone()
                data['book_specific'] = {'row': b_row} if b_row else None

            # Fetch names for filters to avoid blocking main thread later
            names = {}
            if enrollment_no:
                cursor.execute("SELECT name FROM students WHERE enrollment_no=?", (enrollment_no,))
                res = cursor.fetchone()
                names['student_name'] = res[0] if res else None
            
            if book_id:
                cursor.execute("SELECT title FROM books WHERE book_id=?", (book_id,))
                res = cursor.fetchone()
                names['book_title'] = res[0] if res else None
            data['filter_names'] = names

            return data
        except Exception as e:
            print(f"Data fetch error: {e}")
            return e
        finally:
            if conn:
                try:
                    conn.close()
                except Exception:
                    pass

    def import_students_excel_new(self):
        """Import students from Excel with year selection (robust new workflow)."""
        try:
            # Safety: destroy any lingering progress dialog that might be grabbing input
            if hasattr(self, 'import_progress_dialog'):
                try:
                    self.import_progress_dialog.destroy()
                except Exception:
                    pass
                self.import_progress_dialog = None

            # 1. Year + Branch selection dialog
            year_dialog = tk.Toplevel(self.root)
            year_dialog.title("Select Student Year")
            year_dialog.geometry("400x450")
            year_dialog.resizable(False, False)
            year_dialog.configure(bg='white')
            year_dialog.transient(self.root)
            year_dialog.grab_set()
            year_dialog.lift()
            year_dialog.focus_force()
            year_dialog.update_idletasks()
            x = (year_dialog.winfo_screenwidth() // 2) - (200)
            y = (year_dialog.winfo_screenheight() // 2) - (225)
            year_dialog.geometry(f"+{x}+{y}")

            selected_year = tk.StringVar(value="1st Year")
            confirmed = [False]

            header_frame = tk.Frame(year_dialog, bg=self.colors['secondary'], height=60)
            header_frame.pack(fill=tk.X)
            header_frame.pack_propagate(False)
            tk.Label(
                header_frame,
                text="📚 Import Students",
                font=('Segoe UI', 14, 'bold'),
                bg=self.colors['secondary'],
                fg='white'
            ).pack(pady=15)

            content_frame = tk.Frame(year_dialog, bg='white', padx=20, pady=20)
            content_frame.pack(fill=tk.BOTH, expand=True)
            tk.Label(
                content_frame,
                text="Which year are these students from?",
                font=('Segoe UI', 11),
                bg='white',
                fg=self.colors['accent']
            ).pack(pady=(0, 10))
            for year in ["1st Year", "2nd Year", "3rd Year", "Pass Out"]:
                tk.Radiobutton(
                    content_frame, text=year, variable=selected_year, value=year,
                    font=('Segoe UI', 10), bg='white', fg=self.colors['text'],
                    selectcolor='white', activebackground='white'
                ).pack(anchor='w', pady=3)

            tk.Label(
                content_frame, text="Default Branch (if not in Excel):",
                font=('Segoe UI', 11), bg='white', fg=self.colors['accent']
            ).pack(pady=(10, 5), anchor='w')

            selected_branch = tk.StringVar(value="Computer Engineering")
            branch_combo = ttk.Combobox(
                content_frame, textvariable=selected_branch,
                values=["Computer Engineering", "Information Technology", "Civil Engineering",
                        "Electronics & Telecommunication", "Mechanical Engineering",
                        "Automobile Engineering", "Electrical Engineering"],
                state="readonly", width=30
            )
            branch_combo.pack(anchor='w', pady=(0, 5))

            button_frame = tk.Frame(year_dialog, bg='white', padx=20, pady=20)
            button_frame.pack(side=tk.BOTTOM, fill=tk.X)

            def on_confirm():
                confirmed[0] = True
                year_dialog.destroy()

            def on_cancel():
                year_dialog.destroy()

            tk.Button(
                button_frame, text="Continue", command=on_confirm,
                bg=self.colors['secondary'], fg='white', font=('Segoe UI', 12, 'bold'),
                padx=30, pady=12, relief='flat', cursor='hand2'
            ).pack(side=tk.RIGHT, padx=(5, 0))
            tk.Button(
                button_frame, text="Cancel", command=on_cancel,
                bg='#cccccc', fg='#333333', font=('Segoe UI', 11),
                padx=20, pady=10, relief='flat', cursor='hand2'
            ).pack(side=tk.RIGHT)

            self.root.wait_window(year_dialog)
            if not confirmed[0]:
                return

            default_year   = selected_year.get()
            default_branch = selected_branch.get()

            # 2. File picker
            self.root.lift()
            self.root.focus_force()
            self.root.update()
            file_path = filedialog.askopenfilename(
                parent=self.root,
                title="Select Students Excel file",
                filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
            )
            self.root.update()
            if not file_path:
                return

            # 3. Progress dialog
            self.import_progress_dialog = tk.Toplevel(self.root)
            self.import_progress_dialog.title("Importing Students")
            self.import_progress_dialog.geometry("320x160")
            self.import_progress_dialog.resizable(False, False)
            self.import_progress_dialog.transient(self.root)
            self.import_progress_dialog.grab_set()
            px = (self.root.winfo_screenwidth() // 2) - 160
            py = (self.root.winfo_screenheight() // 2) - 80
            self.import_progress_dialog.geometry(f"+{px}+{py}")
            tk.Label(
                self.import_progress_dialog,
                text=f"Importing students from:\n{os.path.basename(file_path)}\n\nPlease wait...",
                font=('Segoe UI', 10), pady=15
            ).pack()
            pb = ttk.Progressbar(self.import_progress_dialog, mode='indeterminate')
            pb.pack(fill=tk.X, padx=20, pady=10)
            pb.start(10)
            self.import_progress_dialog.update()

            # 4. Run in background thread
            self.run_in_background_thread(
                self._import_students_worker,
                self._on_import_complete,
                file_path=file_path,
                default_year=default_year,
                default_branch=default_branch
            )

        except Exception as e:
            import traceback
            messagebox.showerror("Import Error", f"Could not open import dialog:\n{e}\n\n{traceback.format_exc()[-300:]}")

    def import_transactions_excel_report(self):
        """Import Transactions/Borrow Records from a report-style Excel (e.g., exported Records report)."""
        if 'pd' not in globals() and not PANDAS_AVAILABLE:
            messagebox.showerror("Missing Dependency", "Pandas is required for Excel import (pandas + openpyxl).")
            return

        try:
            self.root.lift()
            self.root.focus_force()
            self.root.update()
            file_path = filedialog.askopenfilename(
                parent=self.root,
                title="Select Transactions/Records Excel file",
                filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
            )
            self.root.update()
            if not file_path:
                return

            # Progress dialog
            self.import_progress_dialog = tk.Toplevel(self.root)
            self.import_progress_dialog.title("Importing Transactions")
            self.import_progress_dialog.geometry("360x170")
            self.import_progress_dialog.resizable(False, False)
            self.import_progress_dialog.transient(self.root)
            self.import_progress_dialog.grab_set()
            px = (self.root.winfo_screenwidth() // 2) - 180
            py = (self.root.winfo_screenheight() // 2) - 85
            self.import_progress_dialog.geometry(f"+{px}+{py}")
            tk.Label(
                self.import_progress_dialog,
                text=f"Importing transactions from:\n{os.path.basename(file_path)}\n\nPlease wait...",
                font=('Segoe UI', 10), pady=15
            ).pack()
            pb = ttk.Progressbar(self.import_progress_dialog, mode='indeterminate')
            pb.pack(fill=tk.X, padx=20, pady=10)
            pb.start(10)
            self.import_progress_dialog.update()

            self.run_in_background_thread(
                self._import_transactions_worker,
                self._on_transactions_import_complete,
                file_path=file_path
            )
        except Exception as e:
            import traceback
            messagebox.showerror("Import Error", f"Could not start transactions import:\n{e}\n\n{traceback.format_exc()[-300:]}")

    def _import_transactions_worker(self, file_path):
        """Worker: import borrow_records from a report-style Excel export."""
        import pandas as pd
        import re
        from datetime import datetime

        def _s(v):
            s = str(v).strip()
            return '' if s.lower() == 'nan' else s

        def _norm_id(v):
            """Normalize numeric-looking IDs like 24210270282.0 -> 24210270282."""
            s = _s(v)
            if not s:
                return ''
            s = s.replace('\u00A0', ' ').strip()
            if re.fullmatch(r'\d+\.0+', s):
                return s.split('.')[0]
            if 'e' in s.lower():
                try:
                    from decimal import Decimal
                    d = Decimal(s)
                    if d == d.to_integral_value():
                        return format(d.quantize(1), 'f').split('.')[0]
                except Exception:
                    pass
            return s

        def _to_date(v):
            s = _s(v)
            if not s:
                return None
            if s.lower().startswith('not'):
                return None
            try:
                dt = pd.to_datetime(s, errors='coerce')
                if pd.isna(dt):
                    return None
                return dt.strftime('%Y-%m-%d')
            except Exception:
                return None

        def _parse_fine(v):
            s = _s(v)
            if not s:
                return 0
            m = re.search(r'(\d+)', s.replace(',', ''))
            return int(m.group(1)) if m else 0

        def _find_header_row(path, keywords=('enrollment', 'book id', 'issue date')):
            raw = pd.read_excel(path, header=None, dtype=str)
            kw = [k.lower() for k in keywords]
            for i in range(min(len(raw), 250)):
                row = [str(x).strip().lower() for x in raw.iloc[i].tolist()]
                score = sum(any(k in cell for k in kw) for cell in row)
                if score >= 2:
                    return i
            return 0

        summary = {'added': 0, 'skipped': 0, 'duplicates': 0, 'errors': 0, 'error_list': []}
        try:
            hdr = _find_header_row(file_path)
            df = pd.read_excel(file_path, header=hdr, dtype=str)
            if df is None or df.empty:
                return Exception("The Excel file has no transactions rows.")

            # Normalize columns
            df.columns = (
                df.columns.astype(str)
                .str.strip().str.lower()
                .str.replace(r'[^a-z0-9]+', '_', regex=True)
                .str.strip('_')
            )

            col_aliases = {
                'enrollment': 'enrollment_no',
                'enrollment_no': 'enrollment_no',
                'enrollment_number': 'enrollment_no',
                'student_name': 'student_name',
                'name': 'student_name',
                'branch': 'branch',
                'department': 'branch',
                'book_id': 'book_id',
                'bookid': 'book_id',
                'book_title': 'book_title',
                'title': 'book_title',
                'issue_date': 'borrow_date',
                'borrow_date': 'borrow_date',
                'due_date': 'due_date',
                'return_date': 'return_date',
                'status': 'status',
                'fine': 'fine',
            }
            df.rename(columns={k: v for k, v in col_aliases.items() if k in df.columns}, inplace=True)

            required = ['enrollment_no', 'book_id', 'borrow_date', 'due_date', 'status']
            missing = [c for c in required if c not in df.columns]
            if missing:
                return Exception(f"Missing required columns: {', '.join(missing)}\nFound: {', '.join(df.columns.tolist())}")

            conn = self.db.get_connection()
            cur = conn.cursor()
            try:
                cur.execute('BEGIN')
            except Exception:
                pass

            # Cache existing keys to detect duplicates
            cur.execute("SELECT enrollment_no, accession_no, borrow_date FROM borrow_records")
            existing_keys = set()
            for r in cur.fetchall() or []:
                existing_keys.add((str(r[0] or ''), str(r[1] or ''), str(r[2] or '')))

            # Cache existing students/books
            cur.execute("SELECT enrollment_no FROM students")
            existing_students = {str(r[0]) for r in cur.fetchall() or []}
            cur.execute("SELECT book_id, barcode FROM books")
            existing_books = {}
            barcode_index = {}
            for bid, bc in cur.fetchall() or []:
                bid_s = str(bid)
                existing_books[bid_s] = True
                for t in [x.strip() for x in str(bc or '').split(',') if x.strip()]:
                    barcode_index[t] = bid_s

            # Academic year (optional)
            try:
                academic_year = self.db.get_active_academic_year()
            except Exception:
                academic_year = None

            for idx, row in df.iterrows():
                row_no = idx + 2
                try:
                    enr = _norm_id(row.get('enrollment_no', ''))
                    if not enr:
                        summary['skipped'] += 1
                        continue
                    student_name = _s(row.get('student_name', ''))
                    branch = _s(row.get('branch', ''))

                    acc = _norm_id(row.get('book_id', ''))
                    if not acc:
                        summary['skipped'] += 1
                        continue
                    book_title = _s(row.get('book_title', ''))

                    borrow_date = _to_date(row.get('borrow_date', ''))
                    due_date = _to_date(row.get('due_date', ''))
                    return_date = _to_date(row.get('return_date', ''))
                    status_raw = _s(row.get('status', '')).lower()
                    fine = _parse_fine(row.get('fine', ''))

                    status = 'borrowed' if 'borrow' in status_raw or 'issue' in status_raw or 'active' in status_raw else 'returned'
                    if status == 'borrowed':
                        return_date = None

                    key = (enr, acc, borrow_date or '')
                    if key in existing_keys:
                        summary['duplicates'] += 1
                        continue

                    # Ensure student exists (create minimal if missing)
                    if enr not in existing_students:
                        # Year is unknown from this report; store empty/placeholder
                        cur.execute(
                            "INSERT OR IGNORE INTO students (enrollment_no, name, email, phone, department, year, updated_at) VALUES (?, ?, '', '', ?, '', CURRENT_TIMESTAMP)",
                            (enr, student_name or enr, branch or 'Computer Engineering')
                        )
                        existing_students.add(enr)

                    # Resolve or create book
                    real_book_id = None
                    if acc in existing_books:
                        real_book_id = acc
                    elif acc in barcode_index:
                        real_book_id = barcode_index[acc]

                    if not real_book_id:
                        # Create placeholder book with this accession id
                        real_book_id = acc
                        total = 1
                        avail = 0 if status == 'borrowed' else 1
                        cur.execute(
                            "INSERT OR IGNORE INTO books (book_id, title, author, isbn, category, total_copies, available_copies, barcode, price, updated_at) "
                            "VALUES (?, ?, '', '', 'Technology', ?, ?, ?, 0, CURRENT_TIMESTAMP)",
                            (real_book_id, book_title or f"Book {acc}", total, avail, acc)
                        )
                        existing_books[real_book_id] = True
                        barcode_index[acc] = real_book_id

                    # Insert borrow record
                    cur.execute(
                        "INSERT INTO borrow_records (enrollment_no, book_id, accession_no, borrow_date, due_date, return_date, status, fine, academic_year, updated_at) "
                        "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)",
                        (enr, real_book_id, acc, borrow_date, due_date, return_date, status, fine, academic_year)
                    )
                    existing_keys.add(key)
                    summary['added'] += 1
                except Exception as e:
                    summary['errors'] += 1
                    summary['error_list'].append(f"Row {row_no}: {e}")

            conn.commit()

            # Recalculate available_copies to match current borrowed records
            try:
                self.db.verify_data_integrity()
            except Exception:
                pass

            return summary
        except Exception as e:
            return e
        finally:
            try:
                conn.close()
            except Exception:
                pass

    def _on_transactions_import_complete(self, result):
        """Callback for transactions import completion."""
        try:
            if hasattr(self, 'import_progress_dialog') and self.import_progress_dialog:
                self.import_progress_dialog.destroy()
        except Exception:
            pass
        self.import_progress_dialog = None

        if isinstance(result, Exception):
            messagebox.showerror("Import Error", f"Failed to import transactions:\n{result}")
            return

        if not isinstance(result, dict):
            messagebox.showerror("Import Error", f"Unexpected result: {result}")
            return

        msg = (
            f"✅ Transactions import completed!\n\n"
            f"Added      : {result.get('added', 0)}\n"
            f"Duplicates : {result.get('duplicates', 0)}\n"
            f"Skipped    : {result.get('skipped', 0)}\n"
            f"Errors     : {result.get('errors', 0)}"
        )
        if result.get('error_list'):
            msg += "\n\nFirst errors:\n" + "\n".join(result['error_list'][:5])

        messagebox.showinfo("Import Results", msg)
        try:
            self.refresh_all_data()
        except Exception:
            pass

    def _import_students_worker(self, file_path, default_year, default_branch="Computer Engineering"):
        """Worker function for importing students from Excel"""
        import pandas as pd
        import math
        import re

        def _s(v):
            """Safe string: strip, return '' for nan/None."""
            s = str(v).strip()
            return '' if s.lower() == 'nan' else s

        def _norm_numeric_text(v):
            """Normalize numeric-looking Excel values like 24210270282.0 -> 24210270282.
            Keeps text as-is otherwise.
            """
            s = _s(v)
            if not s:
                return ''
            s = s.replace('\u00A0', ' ').strip()
            # Common pandas float-to-string artifact
            if re.fullmatch(r'\d+\.0+', s):
                return s.split('.')[0]
            # Scientific notation: 2.4210270282E+10
            if 'e' in s.lower():
                try:
                    from decimal import Decimal
                    d = Decimal(s)
                    if d == d.to_integral_value():
                        # Format as plain integer string
                        return format(d.quantize(1), 'f').split('.')[0]
                except Exception:
                    pass
            return s

        summary = {'added': 0, 'skipped': 0, 'duplicate': 0, 'errors': 0, 'error_list': []}
        try:
            # Read once normally; if it's a report-export (headers not in row 0), we'll re-read with detected header row.
            # IMPORTANT: dtype=str prevents Enrollment No becoming float and showing '.0'
            df = pd.read_excel(file_path, dtype=str)
            if df.empty:
                return Exception("The Excel file is empty.")

            # Normalize column names: lowercase + collapse whitespace + replace space with _
            df.columns = [c.strip().lower().replace(' ', '_') for c in df.columns.astype(str)]

            # Flexible column aliases — map ANY synonym → canonical name
            aliases = {
                # enrollment
                'enrollment':          'enrollment_no',
                'enrollmentno':        'enrollment_no',
                'enrollment_number':   'enrollment_no',
                'enroll_no':           'enrollment_no',
                'enroll':              'enrollment_no',
                'roll_no':             'enrollment_no',
                'rollno':              'enrollment_no',
                'roll_number':         'enrollment_no',
                'student_id':          'enrollment_no',
                'sr_no':               'enrollment_no',
                'sr.no':               'enrollment_no',
                'sr.no.':              'enrollment_no',
                # name
                'student_name':        'name',
                'full_name':           'name',
                'student':             'name',
                # email
                'email_id':            'email',
                'mail':                'email',
                'e_mail':              'email',
                # phone
                'mobile':              'phone',
                'mobile_no':           'phone',
                'contact':             'phone',
                'contact_no':          'phone',
                'phone_no':            'phone',
                'phone_number':        'phone',
                # department / branch
                'branch':              'department',
                'dept':                'department',
                'dept.':               'department',
                'branch_name':         'department',
                # year
                'year_of_study':       'year',
                'class':               'year',
                'sem':                 'year',
            }
            df.rename(columns={k: v for k, v in aliases.items() if k in df.columns}, inplace=True)

            # If required columns are missing, try to auto-detect header row (common in exported reports)
            def _find_header_row_in_report(path):
                try:
                    raw = pd.read_excel(path, header=None, dtype=str)
                    if raw is None or len(raw) == 0:
                        return None
                    for i in range(min(len(raw), 200)):
                        row = [str(x).strip().lower() for x in raw.iloc[i].tolist()]
                        has_enroll = any(('enrollment' in c and ('no' in c or 'number' in c)) or c in ('enrollment', 'enrollment_no') for c in row)
                        has_name = any(('name' == c) or ('student_name' == c) or ('student name' in c) or (c == 'student') for c in row)
                        if has_enroll and has_name:
                            return i
                except Exception:
                    return None
                return None

            missing_now = [c for c in ('enrollment_no', 'name') if c not in df.columns]
            if missing_now:
                hdr = _find_header_row_in_report(file_path)
                if hdr is not None:
                    df = pd.read_excel(file_path, header=hdr, dtype=str)
                    if df.empty:
                        return Exception("The Excel file has no data rows.")
                    df.columns = [c.strip().lower().replace(' ', '_') for c in df.columns.astype(str)]
                    df.rename(columns={k: v for k, v in aliases.items() if k in df.columns}, inplace=True)

            # Auto-detect enrollment column: first column whose values look like enrollment numbers
            if 'enrollment_no' not in df.columns:
                for col in df.columns:
                    sample = df[col].dropna().astype(str).head(5).tolist()
                    if any(len(v.strip()) >= 4 for v in sample):
                        df.rename(columns={col: 'enrollment_no'}, inplace=True)
                        break

            # Auto-detect name column: first string column that isn't enrollment
            if 'name' not in df.columns:
                for col in df.columns:
                    if col == 'enrollment_no':
                        continue
                    sample = df[col].dropna().astype(str).head(5).tolist()
                    if any(any(c.isalpha() for c in v) for v in sample):
                        df.rename(columns={col: 'name'}, inplace=True)
                        break

            # Final validation
            missing = [c for c in ('enrollment_no', 'name') if c not in df.columns]
            if missing:
                detected = ', '.join(df.columns.tolist())
                return Exception(
                    f"Could not find columns: {', '.join(missing)}\n\n"
                    f"Columns found in your Excel:\n{detected}\n\n"
                    f"Please rename your columns to 'Enrollment No' and 'Name'."
                )

            # Normalize Year text from Excel/default
            def _norm_year(y):
                s = _s(y)
                if not s:
                    return ''
                s2 = ' '.join(s.split()).strip().lower()
                mapping = {
                    '1st year': '1st', 'first year': '1st', 'fy': '1st', '1st': '1st',
                    '2nd year': '2nd', 'second year': '2nd', 'sy': '2nd', '2nd': '2nd',
                    '3rd year': '3rd', 'third year': '3rd', 'ty': '3rd', '3rd': '3rd',
                    'pass out': 'Pass Out', 'passout': 'Pass Out', 'po': 'Pass Out'
                }
                return mapping.get(s2, s)

            for idx, row in df.iterrows():
                row_no = idx + 2
                try:
                    enrollment = _norm_numeric_text(row.get('enrollment_no', ''))
                    name       = _s(row.get('name', ''))
                    email      = _s(row.get('email', ''))
                    phone      = _norm_numeric_text(row.get('phone', ''))
                    dept_raw   = _s(row.get('department', ''))
                    department = dept_raw if dept_raw else default_branch
                    year_raw   = _s(row.get('year', ''))
                    year       = _norm_year(year_raw) if year_raw else _norm_year(default_year)

                    if not enrollment or not name:
                        summary['skipped'] += 1
                        continue

                    success, message = self.db.add_student(enrollment, name, email, phone, department, year)
                    if success:
                        summary['added'] += 1
                    else:
                        if 'exists' in message.lower() or 'duplicate' in message.lower():
                            summary['duplicate'] += 1
                        else:
                            summary['errors'] += 1
                            summary['error_list'].append(f"Row {row_no}: {message}")
                except Exception as e:
                    summary['errors'] += 1
                    summary['error_list'].append(f"Row {row_no}: {e}")

            return summary
        except Exception as e:
            return e

    def _on_import_complete(self, result):
        """Callback for import completion — always runs on main thread via root.after()"""
        # Always close progress dialog first
        try:
            if hasattr(self, 'import_progress_dialog') and self.import_progress_dialog:
                self.import_progress_dialog.destroy()
        except Exception:
            pass
        self.import_progress_dialog = None

        if isinstance(result, Exception):
            messagebox.showerror("Import Error", f"Failed to import:\n{result}")
            return

        if not isinstance(result, dict):
            messagebox.showerror("Import Error", f"Unexpected result: {result}")
            return

        summary = result
        msg = (
            f"✅ Import completed!\n\n"
            f"Added     : {summary['added']}\n"
            f"Duplicates: {summary['duplicate']}\n"
            f"Skipped   : {summary['skipped']}  (empty rows)\n"
            f"Errors    : {summary['errors']}"
        )
        if summary.get('error_list'):
            msg += "\n\nFirst errors:\n" + "\n".join(summary['error_list'][:5])

        messagebox.showinfo("Import Results", msg)

        if summary['added'] > 0 or summary['duplicate'] > 0:
            self.refresh_students()
            if hasattr(self, 'refresh_dashboard'):
                self.refresh_dashboard()

    def __init__(self, root):
        self.root = root
        # Set window title (remove version as requested)
        self.root.title("📚 GPA's Library Management System")
        self.root.geometry("1400x900")
        self.root.state('zoomed')  # Maximize window
        
        # Professional light color scheme (3 colors) - IMPROVED
        self.colors = {
            'primary': '#ffffff',      # Pure white (backgrounds)
            'secondary': '#2E86AB',    # Professional blue (buttons, accents)
            'accent': '#0F3460',       # Dark blue (text, headers)
            'text': '#333333'          # Dark gray text
        }
        
        self.root.configure(bg=self.colors['primary'])
        
        # Initialize database
        self.db = Database()
        
        # Initialize performance optimization systems
        if PERFORMANCE_MODULES_AVAILABLE:
            # Config should remain available even if optional modules fail.
            self.config_manager = get_config()
            try:
                self.connection_pool = get_pool(self.db)
                self.email_batch_service = EmailBatchService(
                    max_workers=self.config_manager.get_email_config()['max_workers'],
                    batch_size=self.config_manager.get_email_config()['batch_size']
                )
                
                # Setup Sync Manager for Local-First + Cloud Sync Architecture
                self.sync_manager = create_sync_manager(self.db)

                if self.sync_manager:
                    # Sync interval: 5 minutes keeps cloud reasonably fresh
                    sync_interval = 5 
                    if hasattr(self.config_manager, 'get_sync_interval'):
                        sync_interval = self.config_manager.get_sync_interval()
                    # Enforce fast-sync bounds for better responsiveness
                    try:
                        sync_interval = int(sync_interval)
                    except Exception:
                        sync_interval = 2
                    sync_interval = max(1, min(sync_interval, 5))
                    
                    # Start the background daemon (does initial pull + periodic bidirectional sync)
                    self.sync_manager.auto_sync_daemon(sync_interval)
                    print(f"[OK] Local-first mode: SQLite primary, Supabase sync every {sync_interval} min")

                print(f"[OK] Performance optimization modules loaded successfully")
            except Exception as e:
                print(f"[WARNING] Performance modules initialization failed: {e}")
                self.connection_pool = None
                self.email_batch_service = None
                self.sync_manager = None
        else:
            self.config_manager = None
            self.connection_pool = None
            self.email_batch_service = None
            self.sync_manager = None
        
        # Run data integrity check on startup (Background Thread to prevent freezing)
        def _check_integrity_thread():
            print("Running database integrity check...")
            try:
                integrity_result = self.db.verify_data_integrity()
                if integrity_result['status'] == 'issues_found':
                    print(f"[WARNING]  Found {integrity_result['total_issues']} integrity issues")
                    print(f"[OK] Auto-fixed {integrity_result['total_fixes']} issues")
                    if integrity_result['total_issues'] > integrity_result['total_fixes']:
                        print(f"[WARNING]  Some issues require manual attention - check Admin panel")
                elif integrity_result['status'] == 'ok':
                    print(f"[OK] Database integrity verified - all checks passed")
                else:
                    print(f"[ERROR] Integrity check error: {integrity_result.get('error', 'Unknown')}")
            except Exception as e:
                print(f"Error during integrity check: {e}")

        threading.Thread(target=_check_integrity_thread, daemon=True).start()
        
        # Auto-resume sync if overdue (after app restart)
        if PERFORMANCE_MODULES_AVAILABLE and self.sync_manager:
            def _check_and_resume_sync():
                try:
                    time.sleep(3)  # Wait for app to fully initialize
                    sync_log_path = os.path.join(os.path.dirname(self.db.db_path), 'sync_log.json')
                    if os.path.exists(sync_log_path):
                        with open(sync_log_path, 'r') as f:
                            data = json.load(f)
                            last_sync = data.get('last_sync', '2000-01-01 00:00:00')
                            last_sync_dt = datetime.strptime(last_sync, '%Y-%m-%d %H:%M:%S')
                            minutes_since = (datetime.now() - last_sync_dt).total_seconds() / 60
                            sync_interval = self.config_manager.get_sync_interval()
                            
                            if minutes_since >= sync_interval:
                                print(f"[WARNING]  Sync overdue ({int(minutes_since)} min since last sync)")
                                print("🔄 Triggering catch-up sync...")
                                result = self.sync_manager.sync_now(direction='both')
                                if 'error' not in result:
                                    print(f"[OK] Catch-up sync completed: {result.get('records_synced', 0)} records")
                                else:
                                    print(f"[ERROR] Catch-up sync failed: {result.get('error')}")
                except Exception as e:
                    print(f"Error checking sync status: {e}")
            
            threading.Thread(target=_check_and_resume_sync, daemon=True).start()

        # Notify user if calendar support missing
        if DateEntry is None:
            print("tkcalendar not installed - falling back to manual date entry dialog.")
        
        # Configure styles
        self.setup_styles()
        
        # Search/filter related variables
        self.student_search_var = tk.StringVar()
        self.student_year_filter = tk.StringVar(value="All")
        self.student_branch_filter = tk.StringVar(value="All")
        self.book_search_var = tk.StringVar()
        self.book_category_filter = tk.StringVar(value="All")
        self.record_search_var = tk.StringVar()
        self.record_type_filter = tk.StringVar(value="All")
        self.record_branch_filter = tk.StringVar(value="All")
        
        # Email settings
        self.email_settings = self.load_email_settings()
        
        # Library settings (configurable fine, loan period, etc.)
        self.library_settings = self.load_library_settings()
        # Mirror fine_per_day into synced DB settings so web portal can read same value.
        try:
            self.db.set_system_setting('fine_per_day', self.library_settings.get('fine_per_day', 5))
            self._sync_runtime_settings_now()
        except Exception as e:
            print(f"[Settings Sync] Failed to mirror fine_per_day on startup: {e}")
        
        # Start reminder email scheduler if enabled
        if self.email_settings.get('reminder_enabled', False):
            self.schedule_reminder_emails()


        # Student Portal Thread
        self.portal_thread = None
        self.portal_port = 5001  # Changed to 5001 to avoid conflicts with other apps/previous instances

        # Launch login interface
        self.create_login_interface()

    def setup_styles(self):
        """Configure ttk styles (restored after refactor)."""
        try:
            style = ttk.Style()
            # Try a modern theme if available
            for theme in ("clam", "vista", "default"):
                try:
                    style.theme_use(theme)
                    break
                except Exception:
                    continue

            # Treeview base
            style.configure(
                'Treeview',
                background='white',
                foreground='#222222',
                fieldbackground='white',
                rowheight=26,
                font=('Segoe UI', 10)
            )
            style.configure(
                'Treeview.Heading',
                font=('Segoe UI', 10, 'bold'),
                background=self.colors['secondary'],
                foreground='white'
            )
            style.map('Treeview', background=[('selected', '#cfe9ff')])

            # Buttons (ttk) - ensure consistent focus/hover colors if later used
            style.configure('TButton', font=('Segoe UI', 10, 'bold'))
        except Exception as e:
            # Fail silently – styling is not critical for functionality
            print(f"Style setup warning: {e}")
    
    def load_email_settings(self):
        """Load email settings from JSON file"""
        settings_file = os.path.join(os.path.dirname(__file__), 'email_settings.json')
        if hasattr(sys, '_MEIPASS'):
            # Running as EXE, store settings next to executable
            settings_file = os.path.join(os.path.dirname(sys.executable), 'email_settings.json')
        
        default_settings = {
            'smtp_server': 'smtp.gmail.com',
            'smtp_port': 587,
            'sender_email': '',
            'sender_password': '',
            'enabled': False,
            'reminder_enabled': False,
            'reminder_days_before': 2
        }
        
        if os.path.exists(settings_file):
            try:
                with open(settings_file, 'r') as f:
                    loaded_settings = json.load(f)
                    # Merge with defaults to ensure all keys exist
                    default_settings.update(loaded_settings)
                    return default_settings
            except:
                return default_settings
        return default_settings
    
    def save_email_settings(self, settings):
        """Save email settings to JSON file"""
        settings_file = os.path.join(os.path.dirname(__file__), 'email_settings.json')
        if hasattr(sys, '_MEIPASS'):
            settings_file = os.path.join(os.path.dirname(sys.executable), 'email_settings.json')
        
        try:
            with open(settings_file, 'w') as f:
                json.dump(settings, f, indent=4)
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save email settings.\n\n{e}")
            return False
    
    def load_library_settings(self):
        """Load library settings from JSON file (configurable fine, loan period, etc.)"""
        settings_file = os.path.join(os.path.dirname(__file__), 'library_settings.json')
        if hasattr(sys, '_MEIPASS'):
            settings_file = os.path.join(os.path.dirname(sys.executable), 'library_settings.json')
        
        import os
        default_settings = {
            'fine_per_day': 5,
            'loan_period_days': 7,
            'max_books_per_student': 5,
            'admin_password_hash': ADMIN_PASSWORD_HASH
        }
        
        if os.path.exists(settings_file):
            try:
                with open(settings_file, 'r') as f:
                    loaded_settings = json.load(f)
                    default_settings.update(loaded_settings)
                    return default_settings
            except:
                return default_settings
        else:
            # Create settings file with defaults if it doesn't exist
            try:
                with open(settings_file, 'w') as f:
                    json.dump(default_settings, f, indent=4)
            except:
                pass
        return default_settings
    
    def save_library_settings(self, settings):
        """Save library settings to JSON file"""
        settings_file = os.path.join(os.path.dirname(__file__), 'library_settings.json')
        if hasattr(sys, '_MEIPASS'):
            settings_file = os.path.join(os.path.dirname(sys.executable), 'library_settings.json')
        
        try:
            with open(settings_file, 'w') as f:
                json.dump(settings, f, indent=4)
            # Update the instance variable
            self.library_settings = settings
            # Mirror to shared DB settings for portal parity
            try:
                self.db.set_system_setting('fine_per_day', settings.get('fine_per_day', 5))
                self._sync_runtime_settings_now()
            except Exception as e:
                print(f"[Settings Sync] Failed to mirror fine_per_day: {e}")
            return True
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save library settings.\n\n{e}")
            return False
    
    def get_fine_per_day(self):
        """Get configurable fine per day from library settings"""
        return self.library_settings.get('fine_per_day', 5)
    
    def get_loan_period_days(self):
        """Get configurable loan period from library settings"""
        return self.library_settings.get('loan_period_days', 7)
    
    def get_max_books_per_student(self):
        """Get configurable max books per student from library settings"""
        return self.library_settings.get('max_books_per_student', 5)

    def start_auto_refresh_service(self):
        """Start the background auto-refresh loop"""
        # Run immediately once, then schedule loop
        self.root.after(1000, self._refresh_all_data)
    
    def _refresh_all_data(self):
        """Silent Background Refresh - No Blocking Dialog"""
        if hasattr(self, '_is_refreshing') and self._is_refreshing:
            return
        self._is_refreshing = True

        def get_sync_state_from_log():
            """Read sync state from sync_log.json ('in_progress', 'completed', etc.)."""
            try:
                sync_log_path = os.path.join(os.path.dirname(self.db.db_path), 'sync_log.json')
                if os.path.exists(sync_log_path):
                    with open(sync_log_path, 'r') as f:
                        data = json.load(f)
                    return str(data.get('status', '')).strip().lower()
            except Exception:
                pass
            return ''
        
        # Update status indicator
        def set_status(text, color='#38BDF8'): # Light blue for active
            if hasattr(self, 'sync_status_label'):
                self.sync_status_label.config(text=text, fg=color)
        
        set_status("⚡ Syncing...", '#38BDF8')
        
        def run_refresh_sequence():
            try:
                # STEP 1: UI refresh only (auto-sync daemon handles periodic sync)
                self.root.after(100, step_students)

            except Exception as e:
                finish_refresh()

        def step_students():
            try:
                if hasattr(self, 'refresh_students'): self.refresh_students()
                self.root.after(50, step_books)
            except: self.root.after(50, step_books)

        def step_books():
            try:
                if hasattr(self, 'refresh_books'): self.refresh_books()
                self.root.after(50, step_transactions)
            except: self.root.after(50, step_transactions)

        def step_transactions():
            try:
                if hasattr(self, 'refresh_borrowed'): self.refresh_borrowed()
                if hasattr(self, 'load_pending_requests'): self.load_pending_requests()
                if hasattr(self, 'load_deletion_requests'): self.load_deletion_requests()
                self.root.after(50, step_dashboard)
            except: self.root.after(50, step_dashboard)

        def step_dashboard():
            try:
                if hasattr(self, 'refresh_dashboard'): self.refresh_dashboard()
            except: pass
            finish_refresh()

        def finish_refresh():
            self._is_refreshing = False
            # Show syncing indicator when sync is active; otherwise show last updated time.
            sync_state = get_sync_state_from_log()
            if sync_state == 'in_progress':
                set_status("🔄 Syncing...", '#38BDF8')
            else:
                now = datetime.now().strftime("%I:%M %p")
                set_status(f"✔️ Updated {now}", '#94a3b8')
            
            # Schedule next refresh in 2 minutes (120 seconds)
            self.root.after(120000, self._refresh_all_data)

        # Start the sequence
        run_refresh_sequence()

    def open_email_settings(self):
        """Open email settings dialog with tabs for configuration and history"""
        settings_win = tk.Toplevel(self.root)
        settings_win.title("📧 Email Settings - Library Management System")
        settings_win.geometry("700x800")
        settings_win.configure(bg='#f5f7fa')
        settings_win.transient(self.root)
        settings_win.grab_set()
        settings_win.resizable(False, False)
        
        # Center the window
        settings_win.update_idletasks()
        x = (settings_win.winfo_screenwidth() // 2) - (700 // 2)
        y = (settings_win.winfo_screenheight() // 2) - (800 // 2)
        settings_win.geometry(f"+{x}+{y}")
        
        # Main container with shadow effect
        container = tk.Frame(settings_win, bg='#f5f7fa')
        container.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # White card
        card = tk.Frame(container, bg='white', relief='flat', bd=0)
        card.pack(fill=tk.BOTH, expand=True)
        
        # Add subtle shadow
        shadow_frame = tk.Frame(container, bg='#d0d7e2', relief='flat')
        shadow_frame.place(in_=card, relx=0.005, rely=0.005, relwidth=1, relheight=1)
        card.lift()
        
        # Create notebook for tabs
        notebook = ttk.Notebook(card)
        notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Tab 1: Configuration
        config_tab = tk.Frame(notebook, bg='white')
        notebook.add(config_tab, text="  ⚙️ Configuration  ")
        
        # Tab 2: Email History
        history_tab = tk.Frame(notebook, bg='white')
        notebook.add(history_tab, text="  📋 Email History  ")
        
        # Build configuration tab
        self._build_config_tab(config_tab, settings_win)
        
        # Build history tab
        self._build_history_tab(history_tab)
    
    def _build_config_tab(self, parent, settings_win):
        """Build the configuration tab content"""
        # Create canvas and scrollbar for scrolling
        canvas = tk.Canvas(parent, bg='white', highlightthickness=0)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='white')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Enable mouse wheel scrolling
        def _on_mousewheel(event):
            try:
                if canvas.winfo_exists():
                    canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            except Exception:
                pass
        
        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")
        
        canvas.bind("<Enter>", bind_mousewheel)
        canvas.bind("<Leave>", unbind_mousewheel)
        
        main_frame = tk.Frame(scrollable_frame, bg='white', padx=30, pady=25)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header with icon
        header_frame = tk.Frame(main_frame, bg='white')
        header_frame.pack(fill=tk.X, pady=(0, 8))
        
        title = tk.Label(header_frame, text="📧 Email Configuration", 
                        font=('Segoe UI', 18, 'bold'), bg='white', fg=self.colors['accent'])
        title.pack(anchor='w')
        
        subtitle = tk.Label(header_frame, text="Setup automatic email delivery for overdue notices", 
                           font=('Segoe UI', 10), bg='white', fg='#666')
        subtitle.pack(anchor='w', pady=(2, 0))
        
        # Separator line
        separator = tk.Frame(main_frame, height=2, bg='#e1e8ed')
        separator.pack(fill=tk.X, pady=(10, 20))
        
        # Info box
        info_box = tk.Frame(main_frame, bg='#e7f3ff', relief='flat', bd=1)
        info_box.pack(fill=tk.X, pady=(0, 20), padx=5)
        
        info_inner = tk.Frame(info_box, bg='#e7f3ff')
        info_inner.pack(fill=tk.X, padx=12, pady=10)
        
        tk.Label(info_inner, text="ℹ️", font=('Segoe UI', 14), bg='#e7f3ff', fg='#0066cc').pack(side=tk.LEFT, padx=(0, 8))
        info_text = tk.Label(info_inner, 
                            text="Use Gmail to send overdue letters automatically.\nYou'll need a Gmail App Password (not your regular password).",
                            font=('Segoe UI', 9), bg='#e7f3ff', fg='#0066cc', justify=tk.LEFT)
        info_text.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Enable checkbox with better styling
        enable_frame = tk.Frame(main_frame, bg='#f8f9fa', relief='flat', bd=1)
        enable_frame.pack(fill=tk.X, pady=(0, 20), padx=5)
        
        enable_inner = tk.Frame(enable_frame, bg='#f8f9fa')
        enable_inner.pack(fill=tk.X, padx=12, pady=12)
        
        enable_var = tk.BooleanVar(value=self.email_settings.get('enabled', False))
        enable_check = tk.Checkbutton(enable_inner, text="✓ Enable automatic email sending",
                                     variable=enable_var, font=('Segoe UI', 11, 'bold'),
                                     bg='#f8f9fa', activebackground='#f8f9fa', fg='#28a745',
                                     selectcolor='white')
        enable_check.pack(anchor='w')
        
        # Form fields
        form_frame = tk.Frame(main_frame, bg='white')
        form_frame.pack(fill=tk.BOTH, expand=True)
        
        def create_field(parent, label_text, help_text, is_password=False, default_value=''):
            field_container = tk.Frame(parent, bg='white')
            field_container.pack(fill=tk.X, pady=(0, 16))
            
            # Label with help text
            label_frame = tk.Frame(field_container, bg='white')
            label_frame.pack(fill=tk.X, pady=(0, 4))
            
            label = tk.Label(label_frame, text=label_text, 
                           font=('Segoe UI', 10, 'bold'), bg='white', fg='#333')
            label.pack(side=tk.LEFT)
            
            if help_text:
                help_label = tk.Label(label_frame, text=f"({help_text})", 
                                    font=('Segoe UI', 9), bg='white', fg='#888')
                help_label.pack(side=tk.LEFT, padx=(5, 0))
            
            # Entry field
            entry = tk.Entry(field_container, font=('Segoe UI', 11), 
                           relief='solid', bd=1, 
                           bg='#f8f9fa' if not is_password else '#fff3cd',
                           fg='#333')
            if is_password:
                entry.config(show='●')
            entry.insert(0, default_value)
            entry.pack(fill=tk.X, ipady=8)
            
            # Focus effects
            def on_focus_in(e):
                entry.config(bg='white', relief='solid', bd=2)
            def on_focus_out(e):
                entry.config(bg='#f8f9fa' if not is_password else '#fff3cd', relief='solid', bd=1)
            
            entry.bind('<FocusIn>', on_focus_in)
            entry.bind('<FocusOut>', on_focus_out)
            
            return entry
        
        # SMTP Server (readonly-ish, pre-filled)
        smtp_server = create_field(form_frame, "SMTP Server", "Gmail's mail server", 
                                   default_value=self.email_settings.get('smtp_server', 'smtp.gmail.com'))
        smtp_server.config(fg='#666')
        
        # SMTP Port (readonly-ish, pre-filled)
        smtp_port = create_field(form_frame, "SMTP Port", "Gmail uses port 587", 
                                default_value=str(self.email_settings.get('smtp_port', 587)))
        smtp_port.config(fg='#666')
        
        # Gmail Address (main input)
        sender_email = create_field(form_frame, "Your Gmail Address", "e.g., yourname@gmail.com", 
                                    default_value=self.email_settings.get('sender_email', ''))
        
        # App Password (secure input)
        sender_password = create_field(form_frame, "Gmail App Password", "16-character password from Google", 
                                      is_password=True,
                                      default_value=self.email_settings.get('sender_password', ''))
        
        # Separator for reminder section
        reminder_separator = tk.Frame(main_frame, height=2, bg='#e1e8ed')
        reminder_separator.pack(fill=tk.X, pady=(20, 20))
        
        # Reminder settings section
        reminder_header = tk.Label(main_frame, text="📅 Automatic Reminders", 
                                  font=('Segoe UI', 14, 'bold'), bg='white', fg=self.colors['accent'])
        reminder_header.pack(anchor='w', pady=(0, 5))
        
        reminder_subtitle = tk.Label(main_frame, text="Send automatic reminder emails before books become overdue", 
                                    font=('Segoe UI', 9), bg='white', fg='#666')
        reminder_subtitle.pack(anchor='w', pady=(0, 15))
        
        # Reminder enable checkbox
        reminder_enable_frame = tk.Frame(main_frame, bg='#f0f8ff', relief='flat', bd=1)
        reminder_enable_frame.pack(fill=tk.X, pady=(0, 15), padx=5)
        
        reminder_enable_inner = tk.Frame(reminder_enable_frame, bg='#f0f8ff')
        reminder_enable_inner.pack(fill=tk.X, padx=12, pady=12)
        
        reminder_enabled_var = tk.BooleanVar(value=self.email_settings.get('reminder_enabled', False))
        reminder_check = tk.Checkbutton(reminder_enable_inner, text="✉️ Send automatic reminder emails before due date",
                                       variable=reminder_enabled_var, font=('Segoe UI', 10, 'bold'),
                                       bg='#f0f8ff', activebackground='#f0f8ff', fg='#0066cc',
                                       selectcolor='white')
        reminder_check.pack(anchor='w')
        
        # Days before field
        days_frame = tk.Frame(main_frame, bg='white')
        days_frame.pack(fill=tk.X, pady=(0, 10))
        
        days_label = tk.Label(days_frame, text="Send reminder", 
                             font=('Segoe UI', 10), bg='white', fg='#333')
        days_label.pack(side=tk.LEFT)
        
        reminder_days_var = tk.StringVar(value=str(self.email_settings.get('reminder_days_before', 2)))
        days_entry = tk.Entry(days_frame, textvariable=reminder_days_var, 
                             font=('Segoe UI', 11), width=5,
                             relief='solid', bd=1, bg='#f8f9fa', justify='center')
        days_entry.pack(side=tk.LEFT, padx=8)
        
        days_suffix = tk.Label(days_frame, text="days before due date", 
                              font=('Segoe UI', 10), bg='white', fg='#333')
        days_suffix.pack(side=tk.LEFT)
        
        # Schedule info
        schedule_label = tk.Label(main_frame, text="⏰ Reminders are sent daily at 9:00 AM automatically", 
                                 font=('Segoe UI', 9, 'italic'), bg='white', fg='#666')
        schedule_label.pack(anchor='w', pady=(5, 10))
        
        # Test button
        def test_reminders():
            settings_win.withdraw()  # Hide settings window temporarily
            self.send_reminder_emails_now()
            settings_win.deiconify()  # Show it again
        
        test_btn = tk.Button(main_frame, text="🔔 Send Reminders Now (Test)", 
                           font=('Segoe UI', 10),
                           bg='#28a745', fg='white', 
                           padx=20, pady=8,
                           cursor='hand2', relief='flat',
                           command=test_reminders)
        test_btn.pack(anchor='w', pady=(5, 15))
        
        # Help link with better styling
        help_frame = tk.Frame(main_frame, bg='white')
        help_frame.pack(fill=tk.X, pady=(5, 20))
        
        help_link = tk.Label(help_frame, text="🔗 How to generate Gmail App Password? Click here", 
                            font=('Segoe UI', 9, 'underline'), bg='white', fg='#0066cc', cursor='hand2')
        help_link.pack(anchor='w')
        help_link.bind('<Button-1>', lambda e: webbrowser.open('https://support.google.com/accounts/answer/185833'))
        help_link.bind('<Enter>', lambda e: help_link.config(fg='#0052a3'))
        help_link.bind('<Leave>', lambda e: help_link.config(fg='#0066cc'))
        
        # Bottom separator
        separator2 = tk.Frame(main_frame, height=1, bg='#e1e8ed')
        separator2.pack(fill=tk.X, pady=(0, 20))
        
        # Buttons with better styling
        button_frame = tk.Frame(main_frame, bg='white')
        button_frame.pack(fill=tk.X)
        
        def save_settings():
            settings = {
                'smtp_server': smtp_server.get().strip() or 'smtp.gmail.com',
                'smtp_port': int(smtp_port.get().strip() or 587),
                'sender_email': sender_email.get().strip(),
                'sender_password': sender_password.get().strip(),
                'enabled': enable_var.get(),
                'reminder_enabled': reminder_enabled_var.get(),
                'reminder_days_before': int(reminder_days_var.get() or 2)
            }
            
            if settings['enabled'] and not settings['sender_email']:
                messagebox.showwarning("⚠️ Missing Information", 
                                     "Please enter your Gmail address to enable email sending!")
                sender_email.focus()
                return
            
            if settings['enabled'] and not settings['sender_password']:
                messagebox.showwarning("⚠️ Missing Information", 
                                     "Please enter your Gmail App Password to enable email sending!")
                sender_password.focus()
                return
            
            if self.save_email_settings(settings):
                old_reminder_enabled = self.email_settings.get('reminder_enabled', False)
                new_reminder_enabled = settings['reminder_enabled']
                
                self.email_settings = settings
                
                # Start or stop reminder scheduler based on changes
                if new_reminder_enabled and not old_reminder_enabled:
                    self.schedule_reminder_emails()
                
                messagebox.showinfo("✅ Success", 
                                  "Email settings saved successfully!\n\n"
                                  "You can now send overdue letters automatically." + 
                                  ("\n\nAutomatic reminders are now active!" if new_reminder_enabled else ""))
                settings_win.destroy()
        
        def on_save_hover_in(e):
            save_btn.config(bg='#236b9e')
        def on_save_hover_out(e):
            save_btn.config(bg=self.colors['secondary'])
        
        save_btn = tk.Button(button_frame, text="💾 Save Settings", 
                           font=('Segoe UI', 11, 'bold'),
                           bg=self.colors['secondary'], fg='white', 
                           padx=30, pady=10,
                           cursor='hand2', relief='flat', 
                           activebackground='#236b9e',
                           command=save_settings)
        save_btn.pack(side=tk.LEFT, padx=(0, 10))
        save_btn.bind('<Enter>', on_save_hover_in)
        save_btn.bind('<Leave>', on_save_hover_out)
        
        def on_cancel_hover_in(e):
            cancel_btn.config(bg='#999999')
        def on_cancel_hover_out(e):
            cancel_btn.config(bg='#cccccc')
        
        cancel_btn = tk.Button(button_frame, text="✕ Cancel", 
                             font=('Segoe UI', 11),
                             bg='#cccccc', fg='#333', 
                             padx=30, pady=10,
                             cursor='hand2', relief='flat',
                             activebackground='#999999',
                             command=settings_win.destroy)
        cancel_btn.pack(side=tk.LEFT)
        cancel_btn.bind('<Enter>', on_cancel_hover_in)
        cancel_btn.bind('<Leave>', on_cancel_hover_out)
        
        # Focus on email field if empty
        if not sender_email.get():
            sender_email.focus()
    
    def _build_history_tab(self, parent):
        """Build the email history tab content"""
        main_frame = tk.Frame(parent, bg='white', padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title = tk.Label(main_frame, text="📋 Sent Email History", 
                        font=('Segoe UI', 16, 'bold'), bg='white', fg=self.colors['accent'])
        title.pack(pady=(0, 10))
        
        subtitle = tk.Label(main_frame, text="Track all overdue letters sent via email", 
                           font=('Segoe UI', 10), bg='white', fg='#666')
        subtitle.pack(pady=(0, 15))
        
        # Treeview for history
        tree_frame = tk.Frame(main_frame, bg='white')
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        # Scrollbars
        v_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        h_scroll = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)
        
        # Treeview
        columns = ('Date/Time', 'Student', 'Enrollment', 'Email', 'Book', 'Status')
        history_tree = ttk.Treeview(tree_frame, columns=columns, show='headings',
                                    yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set,
                                    height=15)
        
        v_scroll.config(command=history_tree.yview)
        h_scroll.config(command=history_tree.xview)
        
        # Define columns
        history_tree.heading('Date/Time', text='Date/Time')
        history_tree.heading('Student', text='Student Name')
        history_tree.heading('Enrollment', text='Enrollment No')
        history_tree.heading('Email', text='Email Address')
        history_tree.heading('Book', text='Book Title')
        history_tree.heading('Status', text='Status')
        
        history_tree.column('Date/Time', width=140, anchor='center')
        history_tree.column('Student', width=150)
        history_tree.column('Enrollment', width=120, anchor='center')
        history_tree.column('Email', width=180)
        history_tree.column('Book', width=200)
        history_tree.column('Status', width=100, anchor='center')
        
        # Pack
        history_tree.grid(row=0, column=0, sticky='nsew')
        v_scroll.grid(row=0, column=1, sticky='ns')
        h_scroll.grid(row=1, column=0, sticky='ew')
        
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Load history
        self._load_email_history(history_tree)
        
        # Buttons
        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(pady=(15, 0))
        
        refresh_btn = tk.Button(btn_frame, text="🔄 Refresh", font=('Segoe UI', 10, 'bold'),
                               bg=self.colors['secondary'], fg='white', padx=20, pady=8,
                               cursor='hand2', relief='flat',
                               command=lambda: self._load_email_history(history_tree))
        refresh_btn.pack(side=tk.LEFT, padx=5)
        
        clear_btn = tk.Button(btn_frame, text="🗑️ Clear History", font=('Segoe UI', 10),
                             bg='#dc3545', fg='white', padx=20, pady=8,
                             cursor='hand2', relief='flat',
                             command=lambda: self._clear_email_history(history_tree))
        clear_btn.pack(side=tk.LEFT, padx=5)
    
    def _load_email_history(self, tree):
        """Load email history from JSON file"""
        # Clear existing items
        for item in tree.get_children():
            tree.delete(item)
        
        history_file = os.path.join(os.path.dirname(__file__), 'email_history.json')
        if hasattr(sys, '_MEIPASS'):
            history_file = os.path.join(os.path.dirname(sys.executable), 'email_history.json')
        
        if os.path.exists(history_file):
            try:
                with open(history_file, 'r') as f:
                    history = json.load(f)
                
                # Sort by date (newest first)
                history = sorted(history, key=lambda x: x.get('timestamp', ''), reverse=True)
                
                for entry in history:
                    status_icon = '✅' if entry.get('success') else '❌'
                    tree.insert('', 'end', values=(
                        entry.get('timestamp', 'N/A'),
                        entry.get('student_name', 'N/A'),
                        entry.get('enrollment_no', 'N/A'),
                        entry.get('student_email', 'N/A'),
                        entry.get('book_title', 'N/A'),
                        status_icon
                    ))
            except:
                pass
    
    def _clear_email_history(self, tree):
        """Clear all email history"""
        if not messagebox.askyesno("Confirm", "Are you sure you want to clear all email history?\n\nThis cannot be undone."):
            return
        
        history_file = os.path.join(os.path.dirname(__file__), 'email_history.json')
        if hasattr(sys, '_MEIPASS'):
            history_file = os.path.join(os.path.dirname(sys.executable), 'email_history.json')
        
        try:
            if os.path.exists(history_file):
                os.remove(history_file)
            self._load_email_history(tree)
            messagebox.showinfo("Success", "Email history cleared successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to clear history.\n\n{e}")
    
    def _log_email_sent(self, enrollment_no, student_name, student_email, book_title, success, error_message=''):
        """Log sent email to history file"""
        history_file = os.path.join(os.path.dirname(__file__), 'email_history.json')
        if hasattr(sys, '_MEIPASS'):
            history_file = os.path.join(os.path.dirname(sys.executable), 'email_history.json')
        
        # Load existing history
        history = []
        if os.path.exists(history_file):
            try:
                with open(history_file, 'r') as f:
                    history = json.load(f)
            except:
                history = []
        
        # Add new entry
        entry = {
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'enrollment_no': enrollment_no,
            'student_name': student_name,
            'student_email': student_email,
            'book_title': book_title,
            'success': success,
            'error_message': error_message
        }
        history.append(entry)
        
        # Save history
        try:
            with open(history_file, 'w') as f:
                json.dump(history, f, indent=4)
        except:
            pass
    
    def send_email_with_attachment(self, recipient_email, subject, body, attachment_path):
        """Send email with Word document attachment"""
        if not self.email_settings.get('enabled', False):
            return False, "Email sending is not enabled. Configure in Settings."
        
        if not recipient_email:
            return False, "Student email address is not available."
        
        try:
            # Create message
            msg = MIMEMultipart()
            msg['From'] = self.email_settings['sender_email']
            msg['To'] = recipient_email
            msg['Subject'] = subject
            
            # Add body
            msg.attach(MIMEText(body, 'plain'))
            
            # Attach file if exists
            if attachment_path and os.path.exists(attachment_path):
                with open(attachment_path, 'rb') as f:
                    attachment = MIMEApplication(f.read(), _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document')
                    attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment_path))
                    msg.attach(attachment)
            
            # Connect and send
            server = smtplib.SMTP(self.email_settings['smtp_server'], self.email_settings['smtp_port'])
            server.starttls()
            server.login(self.email_settings['sender_email'], self.email_settings['sender_password'])
            server.send_message(msg)
            server.quit()
            
            return True, "Email sent successfully!"
            
        except smtplib.SMTPAuthenticationError:
            return False, "Authentication failed. Please check your email and app password."
        except smtplib.SMTPException as e:
            return False, f"SMTP error: {str(e)}"
        except Exception as e:
            return False, f"Failed to send email: {str(e)}"
    
    def schedule_reminder_emails(self):
        """Start background thread for daily reminder email checks"""
        def run_scheduler():
            while True:
                try:
                    # Calculate seconds until next 9 AM
                    now = datetime.now()
                    target = now.replace(hour=9, minute=0, second=0, microsecond=0)
                    if now >= target:
                        # If already past 9 AM today, schedule for 9 AM tomorrow
                        target += timedelta(days=1)
                    
                    wait_seconds = (target - now).total_seconds()
                    print(f"[Auto-Reminder] Next reminder check scheduled at {target.strftime('%Y-%m-%d %H:%M:%S')}")
                    
                    # Wait until 9 AM
                    time.sleep(wait_seconds)
                    
                    # Check and send reminders
                    if self.email_settings.get('reminder_enabled', False):
                        self.check_and_send_reminders()
                    
                except Exception as e:
                    print(f"[Auto-Reminder] Scheduler error: {e}")
                    time.sleep(3600)  # Wait 1 hour and retry on error
        
        # Start daemon thread
        thread = threading.Thread(target=run_scheduler, daemon=True)
        thread.start()
        print("[Auto-Reminder] Background scheduler started.")
    
    def check_and_send_reminders(self):
        """Check for books due soon and send reminder emails"""
        if not self.email_settings.get('reminder_enabled', False):
            return
        
        days_before = self.email_settings.get('reminder_days_before', 2)
        
        try:
            # Use the same database connection as the rest of the app
            conn = self.db.get_connection()
            cursor = conn.cursor()
            
            # Calculate target due date (N days from now)
            target_date = (datetime.now() + timedelta(days=days_before)).strftime('%Y-%m-%d')
            
            # Find books due on target date that are not yet returned
            query = """
                SELECT b.enrollment_no, s.name, s.email, bk.title, b.due_date
                FROM borrow_records b
                JOIN students s ON b.enrollment_no = s.enrollment_no
                JOIN books bk ON b.book_id = bk.book_id
                WHERE b.due_date = ? AND b.return_date IS NULL AND s.email IS NOT NULL AND s.email != ''
            """
            cursor.execute(query, (target_date,))
            records = cursor.fetchall()
            conn.close()
            
            if not records:
                print(f"[Auto-Reminder] No books due on {target_date}. No reminders sent.")
                return
            
            # Send reminder emails
            success_count = 0
            fail_count = 0
            
            for enrollment, name, email, book_title, due_date in records:
                # Compose friendly reminder email
                subject = f"Reminder: Book Due Soon - {book_title}"
                body = f"""Dear {name},

This is a friendly reminder that you have a book due in {days_before} day(s).

📚 Book Details:
   Title: {book_title}
   Due Date: {due_date}

Please return the book on or before the due date to avoid overdue fines.

If you have already returned the book, please disregard this message.

Thank you,
GPA's Library Management System
Government Polytechnic Awasari (Kh)"""
                
                # Send email (no attachment for reminders)
                success, message = self.send_email_with_attachment(email, subject, body, None)
                
                # Log attempt via admin activity (log_email_history is not defined; use admin log instead)
                status_str = "OK" if success else f"FAILED: {message}"
                self._log_admin_activity(
                    "Reminder Email",
                    f"To {name} ({email}) for '{book_title}' — {status_str}"
                )
                
                if success:
                    success_count += 1
                    print(f"[Auto-Reminder] [OK] Sent to {name} ({email})")
                else:
                    fail_count += 1
                    print(f"[Auto-Reminder] [ERROR] Failed for {name}: {message}")
            
            print(f"[Auto-Reminder] Summary: {success_count} sent, {fail_count} failed")
            
        except Exception as e:
            print(f"[Auto-Reminder] Error checking reminders: {e}")
    
    def send_reminder_emails_now(self):
        """Manually trigger reminder email check (for testing)"""
        if not self.email_settings.get('reminder_enabled', False):
            messagebox.showwarning("Reminders Disabled", "Automatic reminders are not enabled.\n\nPlease enable them in Email Settings first.")
            return
        
        if not self.email_settings.get('enabled', False):
            messagebox.showwarning("Email Disabled", "Email sending is not enabled.\n\nPlease enable it in Email Settings first.")
            return
        
        # Show progress message
        progress = tk.Toplevel(self.root)
        progress.title("Sending Reminders")
        progress.geometry("400x120")
        progress.configure(bg='white')
        progress.transient(self.root)
        progress.resizable(False, False)
        
        tk.Label(progress, text="🔔 Checking for books due soon...", font=('Segoe UI', 12), bg='white').pack(pady=20)
        status_label = tk.Label(progress, text="Please wait...", font=('Segoe UI', 10), bg='white', fg='#666')
        status_label.pack(pady=10)
        
        progress.update()
        
        # Run check in background (MUST NOT call messagebox from background thread — use root.after)
        def run_check():
            self.check_and_send_reminders()
            def _on_done():
                try:
                    progress.destroy()
                except Exception:
                    pass
                messagebox.showinfo("Reminders Sent", "Reminder emails have been sent to students with books due soon.\n\nCheck the Email History tab for details.")
            self.root.after(0, _on_done)
        
        thread = threading.Thread(target=run_check, daemon=True)
        thread.start()

    def create_login_interface(self):
        """Render the login screen with Refined Split Screen design"""
        for w in self.root.winfo_children():
            w.destroy()

        self.root.configure(bg='white')
        
        # Main container (Split Screen)
        main_container = tk.Frame(self.root, bg='white')
        main_container.pack(fill=tk.BOTH, expand=True)
        main_container.grid_columnconfigure(0, weight=4) # Left side (40%)
        main_container.grid_columnconfigure(1, weight=6) # Right side (60%)
        main_container.grid_rowconfigure(0, weight=1)

        # --- LEFT SIDE: Branding (Premium Gradient + Circular Badge) ---
        
        left_canvas = tk.Canvas(main_container, bd=0, highlightthickness=0)
        left_canvas.grid(row=0, column=0, sticky='nsew')
        
        # Function to draw a vertical gradient
        def draw_gradient(canvas, color1, color2):
            # Parse colors
            r1, g1, b1 = canvas.winfo_rgb(color1)
            r2, g2, b2 = canvas.winfo_rgb(color2)
            r_ratio = (r2 - r1) / 1000 # arbitrary height, will resize dynamically
            g_ratio = (g2 - g1) / 1000
            b_ratio = (b2 - b1) / 1000
            
            w = canvas.winfo_width()
            h = canvas.winfo_height()
            
            # We can't draw thousands of lines efficiently every resize.
            # Instead, we'll draw ~100 bands
            steps = 100
            for i in range(steps):
                nr = int(r1 + (r_ratio * (i/steps) * 65535)) # interpolated
                ng = int(g1 + (g_ratio * (i/steps) * 65535))
                nb = int(b1 + (b_ratio * (i/steps) * 65535))
                
                # tkinter rgb is 16-bit, convert to hex
                color = f'#{nr>>8:02x}{ng>>8:02x}{nb>>8:02x}'
                
                y0 = (i / steps) * h
                y1 = ((i + 1) / steps) * h
                canvas.create_rectangle(0, y0, w, y1, fill=color, outline='', tags='gradient')

        # Colors for gradient: Deep Blue -> Darker Blue
        grad_top = '#1E40AF' # Blue-800
        grad_bot = '#0F172A' # Slate-900

        # --- Content Placement ---
        def update_canvas_layout(event):
            w = event.width
            h = event.height
            cx = w // 2
            cy = h // 2
            
            left_canvas.delete('all')
            
            # 1. Background Gradient
            # Limit steps to height for performance, simplified linear interpolation
            if h < 10: return
            
            # Simplified Gradient Drawing (Top to Bottom)
            # Drawing 50 rectangles for smoothness without lag
            r1, g1, b1 = int(grad_top[1:3], 16), int(grad_top[3:5], 16), int(grad_top[5:7], 16)
            r2, g2, b2 = int(grad_bot[1:3], 16), int(grad_bot[3:5], 16), int(grad_bot[5:7], 16)
            
            steps = 50
            for i in range(steps):
                p = i / steps
                r = int(r1 + (r2 - r1) * p)
                g = int(g1 + (g2 - g1) * p)
                b = int(b1 + (b2 - b1) * p)
                color = f'#{r:02x}{g:02x}{b:02x}'
                y0 = i * (h / steps)
                y1 = (i + 1) * (h / steps)
                left_canvas.create_rectangle(0, y0, w, y1, fill=color, outline='')

            # 2. Logo Badge (Circular White Background)
            # This looks intentional and premium
            badge_r = 90
            left_canvas.create_oval(cx - badge_r, cy - 140 - badge_r, 
                                  cx + badge_r, cy - 140 + badge_r, 
                                  fill='white', outline='white')
            
            # 3. Logo
            if hasattr(self, 'login_logo_photo') and self.login_logo_photo:
                left_canvas.create_image(cx, cy - 140, image=self.login_logo_photo)
            else:
                left_canvas.create_text(cx, cy - 140, text="🏫", font=('Segoe UI', 80), fill='#0F172A')
            
            # 4. Text Content
            left_canvas.create_text(cx, cy + 10, text="Government Polytechnic", 
                                  font=('Segoe UI', 24, 'bold'), fill='white', justify='center')
            
            left_canvas.create_text(cx, cy + 55, text="Awasari (Kh)", 
                                  font=('Segoe UI', 24, 'bold'), fill='white', justify='center')
            
            # Divider
            left_canvas.create_line(cx - 50, cy + 95, cx + 50, cy + 95, fill='#38BDF8', width=3)
            
            # System Name
            left_canvas.create_text(cx, cy + 140, text="GPA's Library", 
                                  font=('Segoe UI', 16, 'bold'), fill='#CBD5E1', justify='center')
            left_canvas.create_text(cx, cy + 170, text="Management System", 
                                  font=('Segoe UI', 16, 'bold'), fill='white', justify='center')

        # Load Logo
        try:
            from PIL import Image, ImageTk
            if hasattr(sys, '_MEIPASS'):
                base_dir = sys._MEIPASS
            else:
                base_dir = os.path.dirname(__file__)
            
            logo_path = None
            for candidate in ("logo.png", "college_logo.png", "college_logo.jpg"):
                p = os.path.join(base_dir, candidate)
                if os.path.exists(p):
                    logo_path = p
                    break
            
            if logo_path:
                img = Image.open(logo_path)
                # Slightly smaller than badge radius
                img.thumbnail((140, 140), Image.Resampling.LANCZOS)
                self.login_logo_photo = ImageTk.PhotoImage(img)
        except Exception:
            self.login_logo_photo = None

        left_canvas.bind('<Configure>', update_canvas_layout)

        # --- RIGHT SIDE: Login Form ---
        right_frame = tk.Frame(main_container, bg='#F8FAFC') # Very light gray for contrast
        right_frame.grid(row=0, column=1, sticky='nsew')

        # Login "Card" Panel
        login_panel = tk.Frame(right_frame, bg='white', padx=60, pady=60)
        login_panel.place(relx=0.5, rely=0.5, anchor='center', relwidth=0.65)
        
        # Shadow effect helper (simple border for now)
        login_panel.configure(highlightbackground='#E2E8F0', highlightthickness=1)

        # Header
        tk.Label(
            login_panel, 
            text="Welcome Back", 
            font=('Segoe UI', 28, 'bold'), 
            bg='white', 
            fg='#1E293B'
        ).pack(anchor='w', pady=(0, 10))
        
        tk.Label(
            login_panel, 
            text="Sign in to your dashboard", 
            font=('Segoe UI', 11), 
            bg='white', 
            fg='#64748B'
        ).pack(anchor='w', pady=(0, 40))

        # --- Form Fields ---
        
        def create_entry(label_text, is_password=False):
            tk.Label(
                login_panel, 
                text=label_text, 
                font=('Segoe UI', 10, 'bold'), 
                bg='white', 
                fg='#475569'
            ).pack(anchor='w', pady=(0, 8))
            
            entry = tk.Entry(
                login_panel, 
                font=('Segoe UI', 11), 
                bg='#F1F5F9', 
                fg='#0F172A', 
                relief='flat',
                highlightthickness=1,
                highlightbackground='#E2E8F0',
                highlightcolor='#38BDF8' # Sky blue focus
            )
            if is_password:
                entry.config(show='●')
            
            entry.pack(fill=tk.X, ipady=10, pady=(0, 20))
            return entry

        user_entry = create_entry("Username")
        pass_entry = create_entry("Password", is_password=True)

        # Forgot Password
        fp_link = tk.Label(
            login_panel, 
            text="Forgot Password?", 
            font=('Segoe UI', 9, 'bold'), 
            bg='white', 
            fg='#2563EB', 
            cursor='hand2'
        )
        fp_link.pack(anchor='e', pady=(0, 30))
        fp_link.bind('<Button-1>', lambda e: messagebox.showinfo("Forgot Password", "Contact administrator for password reset.", parent=self.root))

        # Login Logic
        def do_login():
            username = user_entry.get().strip()
            password = pass_entry.get().strip()
            
            # Check password via hash
            hashed_input = hashlib.sha256(password.encode()).hexdigest()
            # Stored is expected to be hash, if not we fallback to ADMIN_PASSWORD_HASH
            stored_hash = self.library_settings.get('admin_password_hash', ADMIN_PASSWORD_HASH)

            # Legacy migration check: if library_settings still has plaintext 'admin_password',
            # we should check against that and potentially migrate it (or just compare).
            # For security, we'll hash the input to check against stored hash.
            # If the user hasn't changed it via settings, we use the default hash.
            is_valid_pwd = False
            if 'admin_password' in self.library_settings:
                 # It's an old plaintext password in settings json, check directly
                 if password == self.library_settings['admin_password']:
                     is_valid_pwd = True
            elif hashed_input == stored_hash:
                 is_valid_pwd = True

            if username == ADMIN_USERNAME and is_valid_pwd:
                # Cleanup and launch main
                for w in self.root.winfo_children(): w.destroy()
                self.create_main_interface()
            else:
                messagebox.showerror('Access Denied', 'Invalid username or password!', parent=self.root)

        # Button
        login_btn = tk.Button(
            login_panel,
            text='Sign In',
            font=('Segoe UI', 11, 'bold'),
            bg='#0284c7', 
            fg='white',
            bd=0, 
            relief='flat',
            cursor='hand2',
            command=do_login,
            activebackground='#0369a1',
            activeforeground='white'
        )
        login_btn.pack(fill=tk.X, ipady=12)

        # Version
        tk.Label(
            login_panel, 
            text=f"Library Management System v{APP_VERSION}", 
            font=('Segoe UI', 8), 
            bg='white', 
            fg='#94A3B8'
        ).pack(pady=(40, 0))

        # Enter key binidng
        def handle_enter(event):
            try:
                if user_entry.winfo_exists() and pass_entry.winfo_exists():
                    do_login()
            except: pass
            
        self.root.bind('<Return>', handle_enter)
        user_entry.focus()
    
    def create_main_interface(self):
        """Create the main application interface"""
        # Remove any leftover global key bindings from the login screen
        try:
            self.root.unbind('<Return>')
        except Exception:
            pass
        # Clear root
        for widget in self.root.winfo_children():
            widget.destroy()
        
        # Main container
        main_container = tk.Frame(self.root, bg=self.colors['primary'])
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # Header
        self.create_header(main_container)
        
        # Create notebook for tabs with larger tab sizes
        self.notebook = ttk.Notebook(main_container)
        
        # Configure notebook style for larger tabs
        style = ttk.Style()
        style.configure('TNotebook.Tab', 
                       padding=[20, 12],  # Increased padding for larger tabs
                       font=('Segoe UI', 11, 'bold'))  # Larger font
        
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=20, pady=(10, 20))
        
        # Create tabs
        self.create_dashboard_tab()
        self.create_students_tab()
        self.create_books_tab()
        self.create_transactions_tab()
        self.create_records_tab()  # New records tab
        self.create_analysis_tab()  # New analysis tab with charts
        self.create_reports_tab()  # Reports tab with Excel and PDF export
        self.create_admin_tab()  # Admin Settings tab
        self.create_student_portal_tab()  # Student Portal tab
        
        # Set focus to dashboard
        self.notebook.select(0)
        
        # Initial data load
        self.refresh_all_data()
    
    def create_header(self, parent):
        """Create application header"""
        # Slightly taller header to avoid any text clipping
        header_frame = tk.Frame(parent, bg=self.colors['secondary'], height=120)
        header_frame.pack(fill=tk.X, padx=0, pady=(0, 0))
        header_frame.pack_propagate(False)

        # Add subtle shadow effect
        shadow_frame = tk.Frame(parent, bg='#d1d1d1', height=2)
        shadow_frame.pack(fill=tk.X, padx=15)

        # Logo and title container
        logo_title_frame = tk.Frame(header_frame, bg=self.colors['secondary'])
        # Reduce vertical padding so content fits comfortably
        logo_title_frame.pack(expand=True, fill=tk.BOTH, padx=25, pady=10)
        # Use grid inside this frame for precise alignment (logo | title | actions)
        logo_title_frame.grid_columnconfigure(0, weight=0)
        logo_title_frame.grid_columnconfigure(1, weight=1)
        logo_title_frame.grid_columnconfigure(2, weight=0)
        logo_title_frame.grid_rowconfigure(0, weight=1)

        # Logo - Proper size
        logo_frame = tk.Frame(logo_title_frame, bg='white', width=80, height=80)
        logo_frame.grid(row=0, column=0, sticky='w', padx=(0, 20))
        logo_frame.pack_propagate(False)

        # Try to load college logo image, fallback to emoji
        try:
            from PIL import Image, ImageTk
            import os
            # First check for logo.png, then college_logo.png, then college_logo.jpg
            logo_path = os.path.join(os.path.dirname(__file__), 'logo.png')
            if not os.path.exists(logo_path):
                logo_path = os.path.join(os.path.dirname(__file__), 'college_logo.png')
            if not os.path.exists(logo_path):
                logo_path = os.path.join(os.path.dirname(__file__), 'college_logo.jpg')

            if os.path.exists(logo_path):
                # Load original image and preserve its rectangular aspect ratio
                img = Image.open(logo_path)
                # Fit within 78x78 box while maintaining aspect ratio
                max_size = (78, 78)
                img.thumbnail(max_size, Image.Resampling.LANCZOS)
                self.logo_photo = ImageTk.PhotoImage(img)

                logo_label = tk.Label(
                    logo_frame,
                    image=self.logo_photo,
                    bg='white'
                )
            else:
                raise FileNotFoundError("Logo file not found")
        except Exception:
            # Fallback to emoji if image loading fails
            logo_label = tk.Label(
                logo_frame,
                text="📚",
                font=('Segoe UI', 32, 'bold'),
                bg='white',
                fg=self.colors['secondary'],
                justify='center'
            )
        logo_label.pack(expand=True)

        # Title - Single line header with proper spacing
        title_frame = tk.Frame(logo_title_frame, bg=self.colors['secondary'])
        title_frame.grid(row=0, column=1, sticky='nsew')
        # Center the title vertically in its cell
        title_frame.grid_rowconfigure(0, weight=1)
        title_frame.grid_columnconfigure(0, weight=1)

        # Get current active academic year
        active_year = self.db.get_active_academic_year()
        if not active_year:
            active_year = "2025-2026"  # Default fallback
        
        # Convert format from "2025-2026" to "25-26"
        if "-" in active_year:
            years = active_year.split("-")
            if len(years) == 2:
                # Extract last 2 digits of each year
                year1 = years[0][-2:]  # "2025" -> "25"
                year2 = years[1][-2:]  # "2026" -> "26"
                display_year = f"{year1}-{year2}"
            else:
                display_year = active_year
        else:
            display_year = active_year
        
        # Title with academic year - GPA's Library Management System
        main_title_label = tk.Label(
            title_frame,
            text=f"GPA's Library Management System",
            font=('Segoe UI', 26, 'bold'),  # Increased from 22 to 26
            bg=self.colors['secondary'],
            fg='white',
            anchor='w'
        )
        # Use grid so it vertically centers without extra padding
        main_title_label.grid(row=0, column=0, sticky='w', padx=(0, 0), pady=0)
        
        # Academic year label (smaller size) - store as instance variable for updates
        self.academic_year_label = tk.Label(
            title_frame,
            text=f"Academic Year: {display_year}",
            font=('Segoe UI', 16, 'bold'),  # Decreased from 22 to 16
            bg=self.colors['secondary'],
            fg='#FFD700',  # Gold color for emphasis
            cursor='hand2'
        )
        self.academic_year_label.grid(row=1, column=0, sticky='w', padx=(0, 0), pady=(5, 0))

        # User info
        user_frame = tk.Frame(logo_title_frame, bg=self.colors['secondary'])
        user_frame.grid(row=0, column=2, sticky='e', padx=(25, 0))
        

        
        # Top-right row: Developer label + inline Promote link
        user_top_row = tk.Frame(user_frame, bg=self.colors['secondary'])
        user_top_row.pack(side=tk.RIGHT, anchor='e', pady=(18, 2))
        
        user_label = tk.Label(
            user_top_row,
            text="👨‍💻 Developer",
            font=('Segoe UI', 13, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            cursor='hand2'
        )
        user_label.pack(side=tk.LEFT)
        user_label.bind('<Button-1>', lambda e: self.show_developer_info())
        promote_btn_hdr = tk.Button(
            user_top_row,
            text="⬆️ Promote Student Years…",
            font=('Segoe UI', 10, 'bold'),
            bg='#0d6efd',
            fg='white',
            relief='flat',
            padx=10,
            pady=2,
            cursor='hand2',
            activebackground='#0b5ed7',
            activeforeground='white',
            command=self._prompt_and_promote
        )
        promote_btn_hdr.pack(side=tk.LEFT, padx=(10, 0))
        
        # Email Settings button
        email_settings_btn = tk.Button(
            user_top_row,
            text='📧 Email',
            font=('Segoe UI', 10, 'bold'),
            bg='#28a745',
            fg='white',
            relief='flat',
            padx=10,
            pady=2,
            cursor='hand2',
            activebackground='#218838',
            activeforeground='white',
            command=self.open_email_settings
        )
        email_settings_btn.pack(side=tk.LEFT, padx=(10, 0))
        
        # Auto-Sync Status Indicator (Replacing Refresh Button)
        self.sync_status_label = tk.Label(
            user_top_row,
            text='⚡ Auto-sync',
            font=('Segoe UI', 9),
            bg=self.colors['secondary'],
            fg='#94a3b8' # Grayish
        )
        self.sync_status_label.pack(side=tk.LEFT, padx=(15, 0))
        
        # Start Auto-Refresh Service
        self.start_auto_refresh_service()

    # Removed "Clear All Data" button from header as requested
    def show_developer_info(self):
        """Enhanced developer info dialog with social links"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Developer Info")
        dialog.configure(bg='white')
        dialog.resizable(False, False)
        dialog.geometry("480x480")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.update_idletasks()
        dialog.geometry(f"+{self.root.winfo_rootx()+360}+{self.root.winfo_rooty()+240}")
        
        header = tk.Label(dialog, text="👨‍💻 Developer", font=('Segoe UI',14,'bold'), bg='white', fg=self.colors['accent'])
        header.pack(pady=(14,6))
        
        body = tk.Frame(dialog, bg='white')
        body.pack(fill=tk.BOTH, expand=True, padx=24, pady=4)
        
        items = [
            ("Name","Yash Vijay Date"),
            ("Enrollment","24210270230"),
            ("Branch","Computer Engineering"),
            ("Year","2nd Year")
        ]
        kf=('Segoe UI',10,'bold'); vf=('Segoe UI',10)
        for r,(k,v) in enumerate(items):
            tk.Label(body,text=f"{k}:",font=kf,bg='white',fg=self.colors['accent']).grid(row=r,column=0,sticky='w',padx=(0,10),pady=3)
            tk.Label(body,text=v,font=vf,bg='white',fg='#222').grid(row=r,column=1,sticky='w',pady=3)
        body.grid_columnconfigure(0,weight=0); body.grid_columnconfigure(1,weight=1)
        
        # Social Links Section
        social_frame = tk.Frame(dialog, bg='#f8f9fa', relief='flat', bd=1)
        social_frame.pack(fill=tk.X, padx=24, pady=(12,8))
        
        tk.Label(social_frame, text="🔗 Connect With Me", font=('Segoe UI',10,'bold'), bg='#f8f9fa', fg=self.colors['accent']).pack(pady=(8,4))
        
        links = [
            ("GitHub", "github.com/YashDate31"),
            ("LinkedIn", "linkedin.com/in/yash-date-a361a8329")
        ]
        
        for platform, url in links:
            link_btn = tk.Button(social_frame, text=f"  {platform}  ", font=('Segoe UI',9), 
                                bg='white', fg=self.colors['secondary'], relief='flat', 
                                cursor='hand2', padx=8, pady=2,
                                command=lambda u=url: self._open_social_link(u))
            link_btn.pack(pady=2)
        
        social_frame.pack_configure(pady=(8,4))
        
        # Special Thanks Section
        thanks_frame = tk.Frame(dialog, bg='#fffbf0', relief='flat', bd=1)
        thanks_frame.pack(fill=tk.X, padx=24, pady=(8,8))
        
        tk.Label(thanks_frame, text="🙏 Special Thanks", font=('Segoe UI',10,'bold'), bg='#fffbf0', fg='#d97706').pack(pady=(8,4))
        
        tk.Label(thanks_frame, text="Yash Magar", font=('Segoe UI',10,'bold'), bg='#fffbf0', fg='#222').pack(pady=2)
        tk.Label(thanks_frame, text="For valuable contributions and support", font=('Segoe UI',8), bg='#fffbf0', fg='#666').pack(pady=(0,4))
        
        thanks_links = [
            ("LinkedIn", "linkedin.com/in/yash-magar-55a184395"),
            ("Email", "yashajaymagar01@gmail.com")
        ]
        
        for platform, url in thanks_links:
            if platform == "Email":
                link_btn = tk.Button(thanks_frame, text=f"  {platform}  ", font=('Segoe UI',9), 
                                    bg='white', fg='#d97706', relief='flat', 
                                    cursor='hand2', padx=8, pady=2,
                                    command=lambda u=url: self._open_email(u))
            else:
                link_btn = tk.Button(thanks_frame, text=f"  {platform}  ", font=('Segoe UI',9), 
                                    bg='white', fg='#d97706', relief='flat', 
                                    cursor='hand2', padx=8, pady=2,
                                    command=lambda u=url: self._open_social_link(u))
            link_btn.pack(pady=2)
        
        thanks_frame.pack_configure(pady=(4,4))
        
        tk.Button(dialog,text='Close',font=('Segoe UI',10,'bold'),bg=self.colors['secondary'],fg='white',relief='flat',padx=16,pady=6,cursor='hand2',command=dialog.destroy,activebackground=self.colors['accent'],activeforeground='white').pack(pady=(4,14))
    
    def _open_social_link(self, url):
        """Open social media link in browser"""
        import webbrowser
        if not url.startswith('http'):
            url = f"https://{url}"
        webbrowser.open(url)
    
    def _open_email(self, email):
        """Open email client"""
        import webbrowser
        webbrowser.open(f"mailto:{email}")

    def clear_all_data_ui(self):
        """Clear all demo/user data with a confirmation prompt."""
        if not messagebox.askyesno(
            "Confirm",
            "Remove ALL data from this software?\n\n"
            "This will wipe:\n"
            "• Local database (students, books, records, etc.)\n"
            "• Supabase database (cloud)\n\n"
            "This cannot be undone.",
            icon='warning'
        ):
            return
        ok, msg = self.db.clear_all_data_everywhere(clear_supabase=True, clear_portal_db=True)
        if ok:
            messagebox.showinfo("Done", msg)
            self.refresh_all_data()
        else:
            messagebox.showerror("Error", msg)

    def _prompt_and_promote(self):
        """Prompt for password and, if correct, promote student years.
        Required password: Admin password from library settings
        """
        from tkinter import simpledialog
        pwd = simpledialog.askstring("Authentication Required", "Enter admin password to promote students:", show='*')
        if pwd is None:
            return

        hashed_input = hashlib.sha256(pwd.strip().encode()).hexdigest()
        is_valid_pwd = False
        if 'admin_password' in self.library_settings:
            if pwd.strip() == self.library_settings['admin_password']:
                is_valid_pwd = True
        elif hashed_input == self.library_settings.get('admin_password_hash', ADMIN_PASSWORD_HASH):
            is_valid_pwd = True

        if not is_valid_pwd:
            messagebox.showerror("Access Denied", "Incorrect password. Promotion aborted.")
            return
        self.promote_student_years()

    def clear_all_data_ui(self):
        """UI handler to clear all data from database after confirmations"""
        from tkinter import messagebox, simpledialog
        # First confirmation
        if not messagebox.askyesno(
            "Confirm Data Wipe",
            "This will permanently remove ALL data from:\n"
            "• Local database\n"
            "• Supabase database (cloud)\n\n"
            "Are you absolutely sure?",
            icon='warning'
        ):
            return
        # Password prompt (must match admin password)
        pwd = simpledialog.askstring(
            "Enter Admin Password",
            "Enter admin password to continue (cancel to abort):",
            show='*'
        )
        if pwd is None:
            messagebox.showinfo("Cancelled", "Data wipe cancelled.")
            return

        hashed_input = hashlib.sha256(pwd.strip().encode()).hexdigest()

        is_valid_pwd = False
        if 'admin_password' in self.library_settings:
            if pwd.strip() == self.library_settings['admin_password']:
                 is_valid_pwd = True
        elif hashed_input == self.library_settings.get('admin_password_hash', ADMIN_PASSWORD_HASH):
            is_valid_pwd = True

        if not is_valid_pwd:
            messagebox.showerror("Incorrect", "Wrong password. Data wipe aborted.")
            return
        # Second, stronger confirmation
        if not messagebox.askyesno(
            "Final Confirmation",
            "Last chance! This action cannot be undone.\n\nProceed with complete LOCAL + SUPABASE wipe?",
            icon='warning'
        ):
            return
        try:
            success, msg = self.db.clear_all_data_everywhere(clear_supabase=True, clear_portal_db=True)
            if success:
                # Refresh UI tables
                try:
                    self.refresh_all_data()
                except Exception:
                    pass
                messagebox.showinfo("Data Cleared", msg or "All data removed successfully.")
            else:
                messagebox.showerror("Error", msg or "Failed to clear data.")
        except Exception as e:
            messagebox.showerror("Error", f"Unexpected error: {e}")
    
    def create_dashboard_tab(self):
        """Create dashboard tab with statistics"""
        dashboard_frame = tk.Frame(self.notebook, bg=self.colors['primary'])
        self.notebook.add(dashboard_frame, text="📊 Dashboard")
        
        # Statistics cards container
        self.stats_container = tk.Frame(dashboard_frame, bg=self.colors['primary'])
        self.stats_container.pack(fill=tk.X, padx=20, pady=20)
        
        # Show loading placeholder; refresh_all_data() will populate stats shortly
        loading_label = tk.Label(
            self.stats_container,
            text="Loading statistics...",
            font=('Segoe UI', 12),
            bg=self.colors['primary'],
            fg='#888888'
        )
        loading_label.pack(pady=20)
        
        # Current Issued Books (Dashboard Table)
        borrowed_frame = tk.LabelFrame(
            dashboard_frame,
            text="📋 Currently Issued Books",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent'],
            padx=10,
            pady=10
        )
        borrowed_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))

        columns = ('Enrollment No', 'Student Name', 'Branch', 'Book ID', 'Book Name', 'Issue Date', 'Due Date', 'Days Left')
        self.dashboard_borrowed_tree = ttk.Treeview(borrowed_frame, columns=columns, show='headings', height=12)
        col_widths = {
            'Enrollment No': 120, 'Student Name': 180, 'Branch': 130,
            'Book ID': 100, 'Book Name': 200, 'Issue Date': 110, 'Due Date': 110, 'Days Left': 110
        }
        for col in columns:
            self.dashboard_borrowed_tree.heading(col, text=col)
            self.dashboard_borrowed_tree.column(col, width=col_widths.get(col, 120))

        v_scrollbar = ttk.Scrollbar(borrowed_frame, orient=tk.VERTICAL, command=self.dashboard_borrowed_tree.yview)
        h_scrollbar = ttk.Scrollbar(borrowed_frame, orient=tk.HORIZONTAL, command=self.dashboard_borrowed_tree.xview)
        self.dashboard_borrowed_tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)
        self.dashboard_borrowed_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

        # Populate dashboard issued books
        self.refresh_dashboard_borrowed()

    def create_reports_tab(self):
        """Create enhanced reports tab with calendar date pickers and improved UI"""
        reports_frame = tk.Frame(self.notebook, bg='#f0f2f5')
        self.notebook.add(reports_frame, text="📄 Reports")
        
        # Main scrollable container
        canvas = tk.Canvas(reports_frame, bg='#f0f2f5', highlightthickness=0)
        scrollbar = ttk.Scrollbar(reports_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='#f0f2f5')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        def resize_frame(event):
            canvas.itemconfig(canvas_window, width=event.width)
        canvas.bind('<Configure>', resize_frame)
        
        # Mouse wheel scrolling
        def _on_mousewheel(event):
            try:
                if canvas.winfo_exists():
                    canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            except Exception:
                pass
        
        def bind_mousewheel(widget):
            widget.bind("<MouseWheel>", _on_mousewheel)
            for child in widget.winfo_children():
                bind_mousewheel(child)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Enhanced Header with gradient effect simulation
        header_frame = tk.Frame(scrollable_frame, bg='white', relief='flat', bd=0)
        header_frame.pack(fill=tk.X, padx=30, pady=(20, 15))
        
        # Add subtle shadow
        shadow = tk.Frame(scrollable_frame, bg='#d0d0d0', height=2)
        shadow.pack(fill=tk.X, padx=30)
        
        header_content = tk.Frame(header_frame, bg='white')
        header_content.pack(fill=tk.X, padx=30, pady=20)
        
        title_container = tk.Frame(header_content, bg='white')
        title_container.pack(anchor='w')
        
        tk.Label(
            title_container,
            text="📄",
            font=('Segoe UI', 36),
            bg='white',
            fg=self.colors['secondary']
        ).pack(side=tk.LEFT)
        
        title_text_frame = tk.Frame(title_container, bg='white')
        title_text_frame.pack(side=tk.LEFT, padx=(15, 0))
        
        tk.Label(
            title_text_frame,
            text="Reports & Export Center",
            font=('Segoe UI', 26, 'bold'),
            bg='white',
            fg='#1a1a2e'
        ).pack(anchor='w')
        
        tk.Label(
            title_text_frame,
            text="📊 Export comprehensive reports • 📅 Use calendar to select dates • 🎯 No mandatory filters",
            font=('Segoe UI', 11),
            bg='white',
            fg='#666'
        ).pack(anchor='w', pady=(3, 0))
        
        # Content area
        content_frame = tk.Frame(scrollable_frame, bg='#f0f2f5')
        content_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=10)
        
        # Helper function to create enhanced report cards with calendar
        def create_report_card(parent, title, icon, description, color, report_type):
            # Card with shadow
            card_container = tk.Frame(parent, bg='#f0f2f5')
            
            shadow_frame = tk.Frame(card_container, bg='#c8c8c8')
            shadow_frame.pack(padx=3, pady=3, fill=tk.BOTH, expand=True)
            
            card = tk.Frame(shadow_frame, bg='white', padx=30, pady=25)
            card.pack(fill=tk.BOTH, expand=True)
            
            # Card header with icon
            header = tk.Frame(card, bg='white')
            header.pack(fill=tk.X, pady=(0, 12))
            
            # Larger, rounded icon
            icon_frame = tk.Frame(header, bg=color, width=50, height=50)
            icon_frame.pack(side=tk.LEFT)
            icon_frame.pack_propagate(False)
            
            icon_label = tk.Label(
                icon_frame,
                text=icon,
                font=('Segoe UI', 22),
                bg=color,
                fg='white'
            )
            icon_label.pack(expand=True)
            
            title_frame = tk.Frame(header, bg='white')
            title_frame.pack(side=tk.LEFT, padx=(15, 0), fill=tk.X, expand=True)
            
            tk.Label(
                title_frame,
                text=title,
                font=('Segoe UI', 16, 'bold'),
                bg='white',
                fg='#1a1a2e'
            ).pack(anchor='w')
            
            tk.Label(
                title_frame,
                text=description,
                font=('Segoe UI', 10),
                bg='white',
                fg='#777',
                justify='left',
                wraplength=350
            ).pack(anchor='w', pady=(2, 0))
            
            # Separator
            sep = tk.Frame(card, bg='#e0e0e0', height=1)
            sep.pack(fill=tk.X, pady=(0, 15))
            
            # Filters section - OPTIONAL
            filters_section = tk.Frame(card, bg='#f8f9fa', relief='flat')
            filters_section.pack(fill=tk.X, pady=(0, 15))
            
            filters_header = tk.Frame(filters_section, bg='#f8f9fa')
            filters_header.pack(fill=tk.X, padx=15, pady=(12, 8))
            
            tk.Label(
                filters_header,
                text="🔍 Optional Filters",
                font=('Segoe UI', 11, 'bold'),
                bg='#f8f9fa',
                fg='#333'
            ).pack(side=tk.LEFT)
            
            tk.Label(
                filters_header,
                text="(Leave empty to export all data)",
                font=('Segoe UI', 9, 'italic'),
                bg='#f8f9fa',
                fg='#888'
            ).pack(side=tk.LEFT, padx=(10, 0))
            
            # Date range with calendar pickers
            date_frame = tk.Frame(filters_section, bg='#f8f9fa')
            date_frame.pack(fill=tk.X, padx=15, pady=(0, 10))
            
            # From Date
            from_container = tk.Frame(date_frame, bg='#f8f9fa')
            from_container.pack(side=tk.LEFT, padx=(0, 20))
            
            tk.Label(
                from_container,
                text="📅 From Date:",
                font=('Segoe UI', 9, 'bold'),
                bg='#f8f9fa',
                fg='#555'
            ).pack(anchor='w')
            
            try:
                from_cal = DateEntry(
                    from_container,
                    width=15,
                    background=color,
                    foreground='white',
                    borderwidth=2,
                    font=('Segoe UI', 10),
                    date_pattern='yyyy-mm-dd'
                )
                from_cal.pack(pady=(3, 0))
                # Clear initial date
                from_cal.delete(0, tk.END)
            except Exception:
                from_cal = tk.Entry(from_container, font=('Segoe UI', 10), width=15)
                from_cal.pack(pady=(3, 0))
            
            # To Date
            to_container = tk.Frame(date_frame, bg='#f8f9fa')
            to_container.pack(side=tk.LEFT)
            
            tk.Label(
                to_container,
                text="📅 To Date:",
                font=('Segoe UI', 9, 'bold'),
                bg='#f8f9fa',
                fg='#555'
            ).pack(anchor='w')
            
            try:
                to_cal = DateEntry(
                    to_container,
                    width=15,
                    background=color,
                    foreground='white',
                    borderwidth=2,
                    font=('Segoe UI', 10),
                    date_pattern='yyyy-mm-dd'
                )
                to_cal.pack(pady=(3, 0))
                # Clear initial date
                to_cal.delete(0, tk.END)
            except Exception:
                to_cal = tk.Entry(to_container, font=('Segoe UI', 10), width=15)
                to_cal.pack(pady=(3, 0))
            
            # Quick date presets
            preset_frame = tk.Frame(date_frame, bg='#f8f9fa')
            preset_frame.pack(side=tk.LEFT, padx=(20, 0))
            
            tk.Label(
                preset_frame,
                text="⚡ Quick:",
                font=('Segoe UI', 9, 'bold'),
                bg='#f8f9fa',
                fg='#555'
            ).pack(anchor='w')
            
            preset_btns = tk.Frame(preset_frame, bg='#f8f9fa')
            preset_btns.pack(pady=(3, 0))
            
            def set_last_7_days():
                from_cal.delete(0, tk.END)
                to_cal.delete(0, tk.END)
                from_cal.insert(0, (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d'))
                to_cal.insert(0, datetime.now().strftime('%Y-%m-%d'))
            
            def set_last_30_days():
                from_cal.delete(0, tk.END)
                to_cal.delete(0, tk.END)
                from_cal.insert(0, (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
                to_cal.insert(0, datetime.now().strftime('%Y-%m-%d'))
            
            def set_this_year():
                from_cal.delete(0, tk.END)
                to_cal.delete(0, tk.END)
                from_cal.insert(0, f"{datetime.now().year}-01-01")
                to_cal.insert(0, datetime.now().strftime('%Y-%m-%d'))
            
            def clear_dates():
                from_cal.delete(0, tk.END)
                to_cal.delete(0, tk.END)
            
            for text, cmd, bg_col in [
                ("7 Days", set_last_7_days, '#007bff'),
                ("30 Days", set_last_30_days, '#007bff'),
                ("This Year", set_this_year, '#007bff'),
                ("Clear", clear_dates, '#6c757d')
            ]:
                btn = tk.Button(
                    preset_btns,
                    text=text,
                    font=('Segoe UI', 8),
                    bg=bg_col,
                    fg='white',
                    relief='flat',
                    padx=8,
                    pady=3,
                    cursor='hand2',
                    command=cmd
                )
                btn.pack(side=tk.LEFT, padx=2)
                
                # Hover effect
                def on_enter(e, b=btn, col=bg_col):
                    if col == '#6c757d':
                        b.config(bg='#5a6268')
                    else:
                        b.config(bg='#0056b3')
                
                def on_leave(e, b=btn, col=bg_col):
                    b.config(bg=col)
                
                btn.bind('<Enter>', on_enter)
                btn.bind('<Leave>', on_leave)
            
            # Additional type-specific filters
            filter_var = tk.StringVar(value="All")
            branch_var = tk.StringVar(value="All")
            
            if report_type in ["students", "books", "transactions"]:
                type_filter_frame = tk.Frame(filters_section, bg='#f8f9fa')
                type_filter_frame.pack(fill=tk.X, padx=15, pady=(0, 12))
                
                if report_type == "students":
                    label_text = "👥 Year Filter:"
                    values = ["All", "1st Year", "2nd Year", "3rd Year", "Pass Out"]
                elif report_type == "books":
                    label_text = "📚 Category Filter:"
                    values = ["All", "Technology", "Textbook", "Research"]
                else:  # transactions
                    label_text = "📖 Status Filter:"
                    values = ["All", "Active", "Returned", "Overdue"]
                
                tk.Label(
                    type_filter_frame,
                    text=label_text,
                    font=('Segoe UI', 9, 'bold'),
                    bg='#f8f9fa',
                    fg='#555'
                ).pack(side=tk.LEFT, padx=(0, 10))
                
                filter_combo = ttk.Combobox(
                    type_filter_frame,
                    textvariable=filter_var,
                    values=values,
                    state="readonly",
                    width=20,
                    font=('Segoe UI', 10)
                )
                filter_combo.pack(side=tk.LEFT)
            
            # Branch filter for students, transactions, overdue reports
            if report_type in ["students", "transactions", "overdue"]:
                branch_filter_frame = tk.Frame(filters_section, bg='#f8f9fa')
                branch_filter_frame.pack(fill=tk.X, padx=15, pady=(0, 12))
                
                tk.Label(
                    branch_filter_frame,
                    text="🏢 Branch Filter:",
                    font=('Segoe UI', 9, 'bold'),
                    bg='#f8f9fa',
                    fg='#555'
                ).pack(side=tk.LEFT, padx=(0, 10))
                
                branch_combo = ttk.Combobox(
                    branch_filter_frame,
                    textvariable=branch_var,
                    values=["All", "Computer Engineering", "Information Technology", "Civil Engineering",
                            "Electronics & Telecommunication", "Mechanical Engineering",
                            "Automobile Engineering", "Electrical Engineering"],
                    state="readonly",
                    width=28,
                    font=('Segoe UI', 10)
                )
                branch_combo.pack(side=tk.LEFT)
            
            # Export buttons section
            export_section = tk.Frame(card, bg='white')
            export_section.pack(fill=tk.X, pady=(5, 0))
            
            def get_filter_values():
                date_from = from_cal.get().strip() if hasattr(from_cal, 'get') else ""
                date_to = to_cal.get().strip() if hasattr(to_cal, 'get') else ""
                return date_from, date_to, filter_var.get(), branch_var.get()
            
            # Preview button
            preview_btn = tk.Button(
                export_section,
                text="👁️ Preview Data",
                font=('Segoe UI', 11, 'bold'),
                bg='#17a2b8',
                fg='white',
                relief='flat',
                padx=25,
                pady=12,
                cursor='hand2',
                command=lambda: self._preview_report(report_type, *get_filter_values())
            )
            preview_btn.pack(side=tk.LEFT, padx=(0, 10))
            
            # Excel button
            excel_btn = tk.Button(
                export_section,
                text="📊 Export Excel",
                font=('Segoe UI', 11, 'bold'),
                bg='#28a745',
                fg='white',
                relief='flat',
                padx=25,
                pady=12,
                cursor='hand2',
                command=lambda: self._export_report(report_type, 'excel', *get_filter_values())
            )
            excel_btn.pack(side=tk.LEFT, padx=(0, 10))
            
            # PDF button
            pdf_btn = tk.Button(
                export_section,
                text="📑 Export PDF",
                font=('Segoe UI', 11, 'bold'),
                bg='#dc3545',
                fg='white',
                relief='flat',
                padx=25,
                pady=12,
                cursor='hand2',
                command=lambda: self._export_report(report_type, 'pdf', *get_filter_values())
            )
            pdf_btn.pack(side=tk.LEFT)
            
            # Hover effects with smooth transitions
            def create_hover(btn, normal_bg, hover_bg):
                def on_enter(e):
                    btn.config(bg=hover_bg)
                def on_leave(e):
                    btn.config(bg=normal_bg)
                btn.bind('<Enter>', on_enter)
                btn.bind('<Leave>', on_leave)
            
            create_hover(preview_btn, '#17a2b8', '#138496')
            create_hover(excel_btn, '#28a745', '#218838')
            create_hover(pdf_btn, '#dc3545', '#c82333')
            
            return card_container
        
        # Create report cards in 2-column grid
        # Row 1: Students and Books
        row1 = tk.Frame(content_frame, bg='#f0f2f5')
        row1.pack(fill=tk.BOTH, expand=True, pady=(0, 20))
        row1.grid_columnconfigure(0, weight=1)
        row1.grid_columnconfigure(1, weight=1)
        
        students_card = create_report_card(
            row1,
            "Students Report",
            "👥",
            "Complete roster of all registered students with enrollment details and academic information",
            '#2E86AB',
            'students'
        )
        students_card.grid(row=0, column=0, sticky='nsew', padx=(0, 10))
        
        books_card = create_report_card(
            row1,
            "Books Catalog",
            "📚",
            "Comprehensive library inventory with book status, categories, and availability",
            '#28a745',
            'books'
        )
        books_card.grid(row=0, column=1, sticky='nsew', padx=(10, 0))
        
        # Row 2: Transactions and Overdue
        row2 = tk.Frame(content_frame, bg='#f0f2f5')
        row2.pack(fill=tk.BOTH, expand=True, pady=(0, 20))
        row2.grid_columnconfigure(0, weight=1)
        row2.grid_columnconfigure(1, weight=1)
        
        transactions_card = create_report_card(
            row2,
            "Transactions Log",
            "📖",
            "Detailed history of all book loans, returns, and current borrowing status",
            '#6f42c1',
            'transactions'
        )
        transactions_card.grid(row=0, column=0, sticky='nsew', padx=(0, 10))
        
        overdue_card = create_report_card(
            row2,
            "Overdue Analysis",
            "⚠️",
            "Active overdue books with student contacts, days late, and calculated fines",
            '#dc3545',
            'overdue'
        )
        overdue_card.grid(row=0, column=1, sticky='nsew', padx=(10, 0))
        
        # Row 3: Promotion and Activity
        row3 = tk.Frame(content_frame, bg='#f0f2f5')
        row3.pack(fill=tk.BOTH, expand=True, pady=(0, 20))
        row3.grid_columnconfigure(0, weight=1)
        row3.grid_columnconfigure(1, weight=1)
        
        promotion_card = create_report_card(
            row3,
            "Promotion History",
            "⬆️",
            "Complete record of student year progressions with academic year tracking",
            '#17a2b8',
            'promotions'
        )
        promotion_card.grid(row=0, column=0, sticky='nsew', padx=(0, 10))
        
        activity_card = create_report_card(
            row3,
            "Admin Activity Audit",
            "📋",
            "Comprehensive audit trail of all system operations and administrative actions",
            '#ffc107',
            'admin_activity'
        )
        activity_card.grid(row=0, column=1, sticky='nsew', padx=(10, 0))
        
        # Bind mousewheel after all widgets created
        self.root.after(100, lambda: bind_mousewheel(scrollable_frame))
    
    def _preview_report(self, report_type, date_from, date_to, filter_value, branch_filter="All"):
        """Preview report data in a dialog before exporting"""
        try:
            # Get report data
            if report_type == 'students':
                data = self._get_students_report_data(filter_value, date_from, date_to, branch_filter)
                title = "Students Report Preview"
                columns = ['Enrollment No', 'Name', 'Email', 'Phone', 'Branch', 'Year', 'Registration Date']
            elif report_type == 'books':
                data = self._get_books_report_data(filter_value, date_from, date_to)
                title = "Books Report Preview"
                columns = ['Book ID', 'Title', 'Author', 'Category', 'Status', 'Condition', 'Added Date']
            elif report_type == 'transactions':
                data = self._get_transactions_report_data(filter_value, date_from, date_to, branch_filter)
                title = "Transactions Report Preview"
                columns = ['Enrollment', 'Student', 'Branch', 'Book ID', 'Book', 'Issue', 'Due', 'Return', 'Status', 'Fine']
            elif report_type == 'overdue':
                data = self._get_overdue_report_data(date_from, date_to, branch_filter)
                title = "Overdue Books Preview"
                columns = ['Enrollment', 'Student', 'Branch', 'Phone', 'Book ID', 'Title', 'Issue', 'Due', 'Days', 'Fine']
            elif report_type == 'promotions':
                data = self._get_promotions_report_data(date_from, date_to)
                title = "Promotion History Preview"
                columns = ['Date', 'Action', 'Students', 'From', 'To', 'Details']
            elif report_type == 'admin_activity':
                data = self._get_admin_activity_report_data(date_from, date_to)
                title = "Admin Activity Preview"
                columns = ['Timestamp', 'Action', 'Details', 'User']
            else:
                return
            
            if not data:
                messagebox.showinfo("Preview", "No data available with the selected filters.")
                return
            
            # Create preview dialog
            preview_dialog = tk.Toplevel(self.root)
            preview_dialog.title(title)
            preview_dialog.geometry("1000x600")
            preview_dialog.configure(bg='white')
            preview_dialog.transient(self.root)
            
            # Header
            header = tk.Frame(preview_dialog, bg=self.colors['secondary'], height=60)
            header.pack(fill=tk.X)
            header.pack_propagate(False)
            
            tk.Label(
                header,
                text=f"👁️ {title}",
                font=('Segoe UI', 16, 'bold'),
                bg=self.colors['secondary'],
                fg='white'
            ).pack(side=tk.LEFT, padx=20, pady=15)
            
            tk.Label(
                header,
                text=f"Total Records: {len(data)}",
                font=('Segoe UI', 12),
                bg=self.colors['secondary'],
                fg='#FFD700'
            ).pack(side=tk.RIGHT, padx=20)
            
            # Tree frame
            tree_frame = tk.Frame(preview_dialog, bg='white')
            tree_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
            
            # Create treeview
            tree = ttk.Treeview(tree_frame, columns=columns, show='headings', height=20)
            
            # Configure columns
            for col in columns:
                tree.heading(col, text=col)
                tree.column(col, width=100, anchor='center')
            
            # Add scrollbars
            vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
            hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
            tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
            
            tree.grid(row=0, column=0, sticky='nsew')
            vsb.grid(row=0, column=1, sticky='ns')
            hsb.grid(row=1, column=0, sticky='ew')
            
            tree_frame.grid_rowconfigure(0, weight=1)
            tree_frame.grid_columnconfigure(0, weight=1)
            
            # Insert data
            for row_data in data:
                tree.insert('', tk.END, values=row_data)
            
            # Close button
            tk.Button(
                preview_dialog,
                text="✓ Close Preview",
                font=('Segoe UI', 11, 'bold'),
                bg=self.colors['secondary'],
                fg='white',
                relief='flat',
                padx=30,
                pady=10,
                cursor='hand2',
                command=preview_dialog.destroy
            ).pack(pady=(0, 20))
            
        except Exception as e:
            messagebox.showerror("Preview Error", f"Failed to preview report:\n{str(e)}")
            print(f"Preview error: {e}")
    
    def _log_admin_activity(self, action, details):
        """Log admin activity to database asynchronously (fire-and-forget).

        Running this in a background thread prevents it from opening a competing
        SQLite connection inside a UI callback that fires immediately after a
        borrow/return commit -- the secondary trigger path for 'database is locked'.
        """
        def _log_worker():
            try:
                conn = self.db.get_connection()
                c = conn.cursor()
                timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                c.execute(
                    'INSERT INTO admin_activity (timestamp, action, details, admin_user) VALUES (?, ?, ?, ?)',
                    (timestamp, action, details, 'Admin')
                )
                conn.commit()
                conn.close()
            except Exception as e:
                print(f"Error logging admin activity: {e}")
        import threading as _threading
        _threading.Thread(target=_log_worker, daemon=True).start()
    
    def _export_report(self, report_type, format_type, date_from, date_to, filter_value, branch_filter="All"):
        """Export report to Excel or PDF format"""
        try:
            # Clean date inputs
            if date_from == "YYYY-MM-DD" or not date_from:
                date_from = None
            if date_to == "YYYY-MM-DD" or not date_to:
                date_to = None
            
            # Get report data based on type
            if report_type == 'students':
                data = self._get_students_report_data(filter_value, date_from, date_to, branch_filter)
                title = "Students Report"
                columns = ['Enrollment No', 'Name', 'Email', 'Phone', 'Branch', 'Year', 'Registration Date']
            
            elif report_type == 'books':
                data = self._get_books_report_data(filter_value, date_from, date_to)
                title = "Books Report"
                columns = ['Book ID', 'Title', 'Author', 'Category', 'Status', 'Condition', 'Added Date']
            
            elif report_type == 'transactions':
                data = self._get_transactions_report_data(filter_value, date_from, date_to, branch_filter)
                title = "Transactions Report"
                columns = ['Enrollment No', 'Student Name', 'Branch', 'Book ID', 'Book Title', 'Issue Date', 'Due Date', 'Return Date', 'Status', 'Fine']
            
            elif report_type == 'overdue':
                data = self._get_overdue_report_data(date_from, date_to, branch_filter)
                title = "Overdue Books Report"
                columns = ['Enrollment No', 'Student Name', 'Branch', 'Phone', 'Book ID', 'Book Title', 'Issue Date', 'Due Date', 'Days Overdue', 'Fine Amount']
            
            elif report_type == 'promotions':
                data = self._get_promotions_report_data(date_from, date_to)
                title = "Promotion History Report"
                columns = ['Date', 'Action', 'Students Affected', 'From Year', 'To Year', 'Details']
            
            elif report_type == 'admin_activity':
                data = self._get_admin_activity_report_data(date_from, date_to)
                title = "Admin Activity Log"
                columns = ['Timestamp', 'Action', 'Details', 'Admin User']
            
            else:
                messagebox.showerror("Error", "Invalid report type")
                return
            
            if not data:
                messagebox.showwarning("No Data", "No data available for the selected filters.")
                return
            
            # Export based on format
            if format_type == 'excel':
                self._export_to_excel(data, columns, title, report_type, filter_value, date_from, date_to, branch_filter)
            elif format_type == 'pdf':
                self._export_to_pdf(data, columns, title, report_type, filter_value, date_from, date_to, branch_filter)
            
            # Log the export activity
            filter_info = f"Filter: {filter_value}" if filter_value and filter_value != "All" else ""
            branch_info = f"Branch: {branch_filter}" if branch_filter and branch_filter != "All" else ""
            date_info = f"From: {date_from} To: {date_to}" if date_from or date_to else ""
            self._log_admin_activity(f"Report Export - {title}", f"Format: {format_type.upper()}, {filter_info} {branch_info} {date_info}")
            
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export report:\n{str(e)}")
            print(f"Export error: {e}")
    
    def _get_students_report_data(self, year_filter, date_from, date_to, branch_filter="All"):
        """Get students data for report with optimized query"""
        try:
            conn = self.db.get_connection()
            c = conn.cursor()
            
            # Optimized query with proper column selection - separate branch and year columns
            query = """SELECT
                        enrollment_no,
                        name,
                        email,
                        phone,
                        COALESCE(department, '') as branch,
                        COALESCE(year, '') as year_val,
                        COALESCE(date_registered, '') as registration_date
                       FROM students
                       WHERE 1=1"""
            params = []
            
            if year_filter and year_filter != "All":
                query += " AND year = ?"
                params.append(year_filter)
            
            if branch_filter and branch_filter != "All":
                query += " AND department = ?"
                params.append(branch_filter)

            # Optional date range filter for registration date
            if date_from:
                query += " AND date_registered >= ?"
                params.append(date_from)
            if date_to:
                query += " AND date_registered <= ?"
                params.append(date_to)
            
            query += " ORDER BY department, year, name"
            
            c.execute(query, params)
            rows = c.fetchall()
            conn.close()
            
            # Ensure data exists
            if not rows:
                print("No students found in database")
            return [list(row) for row in rows]
        except Exception as e:
            print(f"Error getting students report data: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def _get_books_report_data(self, category_filter, date_from, date_to):
        """Get books data for report with optimized query"""
        try:
            conn = self.db.get_connection()
            c = conn.cursor()
            
            # Optimized query: check if book is borrowed using status column (no book_condition column)
            query = """SELECT DISTINCT b.book_id, b.title, b.author, b.category, 
                      CASE WHEN EXISTS(
                          SELECT 1 FROM borrow_records br 
                          WHERE br.book_id = b.book_id AND br.return_date IS NULL
                      ) THEN 'Borrowed' ELSE 'Available' END as status,
                      'Good' as condition, COALESCE(b.date_added, '') as added_date
                      FROM books b
                      WHERE 1=1"""
            params = []
            
            if category_filter and category_filter != "All":
                query += " AND b.category = ?"
                params.append(category_filter)

            # Optional date range filter for added date
            if date_from:
                query += " AND b.date_added >= ?"
                params.append(date_from)
            if date_to:
                query += " AND b.date_added <= ?"
                params.append(date_to)
            
            query += " ORDER BY b.category, b.title"
            
            c.execute(query, params)
            rows = c.fetchall()
            conn.close()
            
            if not rows:
                print("No books found in database")
            return [list(row) for row in rows]
        except Exception as e:
            print(f"Error getting books report data: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def _get_transactions_report_data(self, status_filter, date_from, date_to, branch_filter="All"):
        """Get transactions data for report"""
        try:
            conn = self.db.get_connection()
            c = conn.cursor()
            
            # Use borrow_records table with branch column
            query = """SELECT 
                br.enrollment_no, s.name, COALESCE(s.department, '') as branch,
                br.book_id, b.title,
                br.borrow_date, br.due_date, br.return_date,
                br.status
                FROM borrow_records br
                LEFT JOIN students s ON br.enrollment_no = s.enrollment_no
                LEFT JOIN books b ON br.book_id = b.book_id
                WHERE 1=1"""
            params = []
            
            if status_filter and status_filter != "All":
                if status_filter == "Active":
                    query += " AND br.return_date IS NULL AND CURRENT_DATE <= br.due_date"
                elif status_filter == "Returned":
                    query += " AND br.return_date IS NOT NULL"
                elif status_filter == "Overdue":
                    query += " AND br.return_date IS NULL AND CURRENT_DATE > br.due_date"
            
            if branch_filter and branch_filter != "All":
                query += " AND s.department = ?"
                params.append(branch_filter)
            
            if date_from:
                query += " AND br.borrow_date >= ?"
                params.append(date_from)
            
            if date_to:
                query += " AND br.borrow_date <= ?"
                params.append(date_to)
            
            query += " ORDER BY br.borrow_date DESC"
            
            c.execute(query, params)
            rows = c.fetchall()
            conn.close()
            
            # Calculate fines dynamically (same logic as get_all_records)
            from datetime import datetime as _dt
            today = _dt.now().date()
            fine_per_day = self.get_fine_per_day()
            result = []
            for row in rows:
                row = list(row)
                enroll, name, branch, book_id, title, borrow_date, due_date, return_date_raw, status = row[:9]
                
                # Determine display status
                if return_date_raw is not None:
                    display_status = 'Returned'
                else:
                    try:
                        due_d = _dt.strptime(str(due_date), '%Y-%m-%d').date()
                        display_status = 'Overdue' if today > due_d else 'Active'
                    except Exception:
                        display_status = 'Active'
                
                # Calculate fine
                fine = 0
                try:
                    due_d = _dt.strptime(str(due_date), '%Y-%m-%d').date()
                    if status == 'borrowed' and return_date_raw is None:
                        if today > due_d:
                            fine = (today - due_d).days * fine_per_day
                    elif return_date_raw is not None:
                        if hasattr(return_date_raw, 'year'):
                            ret_d = return_date_raw
                            if isinstance(ret_d, datetime):
                                ret_d = ret_d.date()
                        else:
                            ret_d = _dt.strptime(str(return_date_raw), '%Y-%m-%d').date()
                        if ret_d > due_d:
                            fine = (ret_d - due_d).days * fine_per_day
                except Exception:
                    pass
                
                # Format return date for display
                return_display = str(return_date_raw) if return_date_raw is not None else 'None'
                
                result.append([enroll, name, branch, book_id, title, borrow_date, due_date, return_display, display_status, fine])
            
            return result
        except Exception as e:
            print(f"Error getting transactions report data: {e}")
            return []
    
    def _get_overdue_report_data(self, date_from, date_to, branch_filter="All"):
        """Get overdue books data for report"""
        try:
            conn = self.db.get_connection()
            c = conn.cursor()
            
            query = """SELECT 
                br.enrollment_no, s.name, COALESCE(s.department, '') as branch,
                s.phone, br.book_id, b.title,
                br.borrow_date, br.due_date
                FROM borrow_records br
                LEFT JOIN students s ON br.enrollment_no = s.enrollment_no
                LEFT JOIN books b ON br.book_id = b.book_id
                WHERE br.return_date IS NULL"""
            
            params = []
            
            if branch_filter and branch_filter != "All":
                query += " AND s.department = ?"
                params.append(branch_filter)
            
            if date_from:
                query += " AND br.borrow_date >= ?"
                params.append(date_from)
            
            if date_to:
                query += " AND br.borrow_date <= ?"
                params.append(date_to)
            
            query += " ORDER BY br.due_date ASC"
            
            c.execute(query, params)
            rows = c.fetchall()
            conn.close()
            
            # Calculate days overdue and fine in Python (works for both SQLite and PostgreSQL)
            from datetime import datetime as _dt
            today = _dt.now().date()
            fine_per_day = self.get_fine_per_day()
            result = []
            for row in rows:
                row = list(row)
                enroll, name, branch, phone, book_id, title, borrow_date, due_date = row[:8]
                try:
                    due_d = _dt.strptime(str(due_date), '%Y-%m-%d').date()
                    days_overdue = (today - due_d).days
                except Exception:
                    days_overdue = 0
                if days_overdue > 0:
                    fine = days_overdue * fine_per_day
                    result.append([enroll, name, branch, phone, book_id, title, borrow_date, due_date, days_overdue, fine])
            
            # Sort by most overdue first
            result.sort(key=lambda x: x[8], reverse=True)
            return result
        except Exception as e:
            print(f"Error getting overdue report data: {e}")
            return []
    
    def _get_promotions_report_data(self, date_from, date_to):
        """Get promotion history data for report"""
        try:
            conn = self.db.get_connection()
            c = conn.cursor()
            
            # Try to create table if it doesn't exist
            try:
                c.execute('''CREATE TABLE IF NOT EXISTS promotion_history (
                    id SERIAL PRIMARY KEY,
                    promotion_date TEXT NOT NULL,
                    action TEXT NOT NULL,
                    students_affected INTEGER,
                    from_year TEXT,
                    to_year TEXT,
                    details TEXT
                )''')
                conn.commit()
            except:
                pass
            
            query = """SELECT promotion_date, action, students_affected, from_year, to_year, details
                      FROM promotion_history WHERE 1=1"""
            params = []
            
            if date_from:
                query += " AND promotion_date >= ?"
                params.append(date_from)
            
            if date_to:
                query += " AND promotion_date <= ?"
                params.append(date_to)
            
            query += " ORDER BY promotion_date DESC"
            
            try:
                c.execute(query, params)
                rows = c.fetchall()
                conn.close()
                return [list(row) for row in rows]
            except:
                conn.close()
                return []
        except Exception as e:
            print(f"Error getting promotions report data: {e}")
            return []
    
    def _get_admin_activity_report_data(self, date_from, date_to):
        """Get admin activity log data for report"""
        try:
            conn = self.db.get_connection()
            c = conn.cursor()
            
            # Try to create table if it doesn't exist (works for both SQLite and PostgreSQL)
            try:
                c.execute('''CREATE TABLE IF NOT EXISTS admin_activity (
                    id SERIAL PRIMARY KEY,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    action TEXT NOT NULL,
                    details TEXT,
                    admin_user TEXT,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )''')
                conn.commit()
            except:
                # Table might already exist or creation failed, continue anyway
                pass
            
            query = """SELECT timestamp, action, details, admin_user
                      FROM admin_activity WHERE 1=1"""
            params = []
            
            if date_from:
                query += " AND timestamp >= ?"
                params.append(date_from)
            
            if date_to:
                query += " AND timestamp <= ?"
                params.append(date_to)
            
            query += " ORDER BY timestamp DESC"
            
            try:
                c.execute(query, params)
                rows = c.fetchall()
                conn.close()
                return [list(row) for row in rows]
            except:
                # Table doesn't exist or query failed, return empty
                conn.close()
                return []
        except Exception as e:
            print(f"Error getting admin activity report data: {e}")
            return []
    
    def _export_to_excel(self, data, columns, title, report_type, filter_value, date_from, date_to, branch_filter="All"):
        """Export report data to Excel format with GPAK branding"""
        try:
            #
            if not data:
                messagebox.showwarning("No Data", "No data available to export.")
                return
                
            # Ask user for save location
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            default_filename = f"GPAK_{title.replace(' ', '_')}_{timestamp}.xlsx"
            
            filepath = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                initialfile=default_filename,
                title=f"Save {title}"
            ) 
            #this is to export data from the excel sheet into the proper data.........
            
            if not filepath:
                return
            
            # Create DataFrame
            df = pd.DataFrame(data, columns=columns)
            
            # Create Excel writer with xlsxwriter engine
            with pd.ExcelWriter(filepath, engine='xlsxwriter') as writer:
                # Write data to sheet starting from row 8 (leave space for header)
                df.to_excel(writer, sheet_name='Report', index=False, startrow=8)
                
                # Get workbook and worksheet objects
                workbook = writer.book
                worksheet = writer.sheets['Report']
                
                # Define formats
                college_header_format = workbook.add_format({
                    'bold': True,
                    'font_size': 18,
                    'align': 'center',
                    'valign': 'vcenter',
                    'fg_color': '#003366',
                    'font_color': 'white',
                    'border': 1
                })
                
                subtitle_format = workbook.add_format({
                    'bold': True,
                    'font_size': 11,
                    'align': 'center',
                    'valign': 'vcenter',
                    'fg_color': '#0066CC',
                    'font_color': 'white'
                })
                
                title_format = workbook.add_format({
                    'bold': True,
                    'font_size': 14,
                    'align': 'center',
                    'valign': 'vcenter',
                    'fg_color': '#2E86AB',
                    'font_color': 'white',
                    'border': 1
                })
                
                header_format = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'align': 'center',
                    'fg_color': '#4DA6FF',
                    'font_color': 'white',
                    'border': 1,
                    'font_size': 10
                })
                
                data_format = workbook.add_format({
                    'border': 1,
                    'valign': 'vcenter'
                })
                
                info_format = workbook.add_format({
                    'italic': True,
                    'font_size': 9,
                    'align': 'center',
                    'font_color': '#666666'
                })
                
                # Add college logo if exists
                logo_path = os.path.join(os.path.dirname(__file__), 'logo.png')
                try:
                    if os.path.exists(logo_path):
                        worksheet.insert_image('A1', logo_path, {'x_scale': 0.15, 'y_scale': 0.15})
                except Exception as e:
                    print(f"Logo insertion failed: {e}")
                
                # Write college header
                end_col = chr(65 + len(columns) - 1)
                worksheet.merge_range(f'B1:{end_col}1', 'GOVERNMENT POLYTECHNIC AWASARI KHURD', college_header_format)
                worksheet.merge_range(f'B2:{end_col}2', 'Tal. Ambegaon, Dist. Pune - 410503', subtitle_format)
                worksheet.set_row(0, 25)
                worksheet.set_row(1, 20)
                
                # Library title
                worksheet.merge_range(f'A3:{end_col}3', 'LIBRARY MANAGEMENT SYSTEM', title_format)
                worksheet.set_row(2, 20)
                
                # Report title
                worksheet.merge_range(f'A5:{end_col}5', title, title_format)
                worksheet.set_row(4, 18)
                
                # Write metadata
                metadata = f"Generated: {datetime.now().strftime('%d-%b-%Y %I:%M %p')}"
                if filter_value and filter_value != "All":
                    metadata += f" | Filter: {filter_value}"
                if branch_filter and branch_filter != "All":
                    metadata += f" | Branch: {branch_filter}"
                if date_from or date_to:
                    metadata += f" | Date Range: {date_from or 'Start'} to {date_to or 'End'}"
                
                worksheet.merge_range(f'A6:{end_col}6', metadata, info_format)
                worksheet.merge_range(f'A7:{end_col}7', f"Total Records: {len(data)}", info_format)
                
                # Apply header format to column headers
                for col_num, value in enumerate(columns):
                    worksheet.write(8, col_num, value, header_format)
                
                # Apply data format to all data cells
                for row_num in range(len(data)):
                    for col_num in range(len(columns)):
                        worksheet.write(9 + row_num, col_num, data[row_num][col_num], data_format)
                
                # Set column widths dynamically
                for i, col in enumerate(columns):
                    max_len = len(str(col))
                    for row in data:
                        if i < len(row):
                            max_len = max(max_len, len(str(row[i])))
                    worksheet.set_column(i, i, min(max_len + 2, 30))
                
                # Set row heights
                worksheet.set_row(8, 25)  # Header row height
            
            messagebox.showinfo("Success", f"Report exported successfully!\n\nFile: {os.path.basename(filepath)}\nLocation: {os.path.dirname(filepath)}")
            
            # Ask if user wants to open the file
            if messagebox.askyesno("Open File", "Do you want to open the exported file now?"):
                os.startfile(filepath)
        
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export to Excel:\n{str(e)}")
            print(f"Excel export error: {e}")
            import traceback
            traceback.print_exc()
    
    def _export_to_pdf(self, data, columns, title, report_type, filter_value, date_from, date_to, branch_filter="All"):
        """Export report data to PDF format with GPAK branding"""
        try:
            if not data:
                messagebox.showwarning("No Data", "No data available to export.")
                return
                
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.units import inch
            from reportlab.platypus import Image as RLImage
            
            # Ask user for save location
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            default_filename = f"GPAK_{title.replace(' ', '_')}_{timestamp}.pdf"
            
            filepath = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                initialfile=default_filename,
                title=f"Save {title}"
            )
            
            if not filepath:
                return
            
            # Use landscape orientation for better table fit
            page_size = landscape(A4)
            
            # Create PDF with custom margins
            doc = SimpleDocTemplate(
                filepath, 
                pagesize=page_size,
                topMargin=0.5*inch,
                bottomMargin=0.4*inch,
                leftMargin=0.4*inch,
                rightMargin=0.4*inch
            )
            elements = []
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles for IARE branding
            college_name_style = ParagraphStyle(
                'CollegeName',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=rl_colors.HexColor('#003366'),
                spaceAfter=4,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            college_location_style = ParagraphStyle(
                'CollegeLocation',
                parent=styles['Normal'],
                fontSize=10,
                textColor=rl_colors.HexColor('#0066CC'),
                spaceAfter=8,
                alignment=TA_CENTER,
                fontName='Helvetica'
            )
            
            library_title_style = ParagraphStyle(
                'LibraryTitle',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=rl_colors.HexColor('#2E86AB'),
                spaceAfter=12,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            report_title_style = ParagraphStyle(
                'ReportTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=rl_colors.HexColor('#2E86AB'),
                spaceAfter=12,
                spaceBefore=8,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            info_style = ParagraphStyle(
                'Info',
                parent=styles['Normal'],
                fontSize=9,
                textColor=rl_colors.HexColor('#666666'),
                spaceAfter=4,
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            
            # Add college logo
            logo_path = os.path.join(os.path.dirname(__file__), 'logo.png')
            try:
                if os.path.exists(logo_path):
                    logo = RLImage(logo_path, width=1*inch, height=1*inch)
                    elements.append(logo)
                    elements.append(Spacer(1, 0.1*inch))
            except Exception as e:
                print(f"Logo insertion failed: {e}")
            
            # Add college header
            elements.append(Paragraph('GOVERNMENT POLYTECHNIC AWASARI KHURD', college_name_style))
            elements.append(Paragraph('Tal. Ambegaon, Dist. Pune - 410503', college_location_style))
            elements.append(Paragraph('LIBRARY MANAGEMENT SYSTEM', library_title_style))
            
            # Add horizontal line
            from reportlab.platypus import HRFlowable
            elements.append(HRFlowable(width="100%", thickness=2, color=rl_colors.HexColor('#2E86AB')))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add report title
            elements.append(Paragraph(title, report_title_style))
            elements.append(Spacer(1, 0.15*inch))
            
            # Add metadata
            metadata = f"<b>Generated:</b> {datetime.now().strftime('%d-%b-%Y %I:%M %p')}"
            if filter_value and filter_value != "All":
                metadata += f" | <b>Filter:</b> {filter_value}"
            if branch_filter and branch_filter != "All":
                metadata += f" | <b>Branch:</b> {branch_filter}"
            if date_from or date_to:
                metadata += f"<br/><b>Date Range:</b> {date_from or 'Start'} to {date_to or 'End'}"
            
            elements.append(Paragraph(metadata, info_style))
            elements.append(Paragraph(f"<b>Total Records:</b> {len(data)}", info_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Prepare table data with proper formatting and Paragraph wrapping
            table_data = [columns]
            for row in data:
                # Convert all values to strings and handle None
                formatted_row = [str(val) if val is not None else 'N/A' for val in row]
                table_data.append(formatted_row)
            
            # Calculate column widths dynamically based on content - use landscape width
            page_width = page_size[0] - 0.8*inch  # Available width for landscape A4
            
            # Fixed column widths for common report types to prevent merging
            num_cols = len(columns)
            if num_cols <= 4:
                col_widths = [page_width / num_cols] * num_cols
            elif num_cols <= 6:
                col_widths = [page_width / num_cols] * num_cols
            else:
                # For many columns, calculate proportionally with minimums
                col_widths = []
                min_width = 0.7*inch
                
                # Calculate weights based on content
                col_weights = []
                for col_idx in range(len(columns)):
                    max_len = len(str(columns[col_idx]))
                    for row in data[:50]:  # Sample first 50 rows for speed
                        if col_idx < len(row):
                            max_len = max(max_len, len(str(row[col_idx])))
                    col_weights.append(min(max_len, 40))
                
                total_weight = sum(col_weights)
                
                # Distribute width proportionally
                for weight in col_weights:
                    col_width = max((weight / total_weight) * page_width, min_width)
                    col_widths.append(col_width)
                
                # Ensure total doesn't exceed page width
                total_col_width = sum(col_widths)
                if total_col_width > page_width:
                    scale = page_width / total_col_width
                    col_widths = [w * scale for w in col_widths]
            
            # Wrap text in cells using Paragraph for better formatting
            from reportlab.platypus import Paragraph as RLParagraph
            cell_style = ParagraphStyle(
                'CellStyle',
                parent=styles['Normal'],
                fontSize=7,
                leading=9,
                fontName='Helvetica',
                alignment=TA_LEFT,
                wordWrap='CJK'
            )
            
            wrapped_table_data = []
            for row_idx, row in enumerate(table_data):
                wrapped_row = []
                for cell in row:
                    # Wrap long text in Paragraph for automatic line breaks
                    if row_idx == 0:  # Header row
                        wrapped_row.append(str(cell))
                    else:
                        # For data cells, use Paragraph to enable wrapping
                        cell_text = str(cell).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        wrapped_row.append(RLParagraph(cell_text, cell_style))
                wrapped_table_data.append(wrapped_row)
            
            # Create table with dynamic widths
            table = Table(wrapped_table_data, colWidths=col_widths, repeatRows=1)
            
            # Professional table styling
            table_style = TableStyle([
                # Header row styling
                ('BACKGROUND', (0, 0), (-1, 0), rl_colors.HexColor('#003366')),
                ('TEXTCOLOR', (0, 0), (-1, 0), rl_colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                
                # Data rows styling
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 5),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                
                # Grid and borders
                ('GRID', (0, 0), (-1, -1), 0.5, rl_colors.grey),
                ('BOX', (0, 0), (-1, -1), 1.5, rl_colors.HexColor('#003366')),
                ('LINEBELOW', (0, 0), (-1, 0), 2, rl_colors.HexColor('#0066CC')),
                
                # Alternating row colors
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor('#F0F0F0')]),
                
                # Enable word wrap
                ('WORDWRAP', (0, 0), (-1, -1), True)
            ])
            
            table.setStyle(table_style)
            elements.append(table)
            
            # Add footer
            elements.append(Spacer(1, 0.3*inch))
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=8,
                textColor=rl_colors.HexColor('#999999'),
                alignment=TA_CENTER,
                fontName='Helvetica-Oblique'
            )
            elements.append(Paragraph(f'Report generated by GPAK Library Management System v{APP_VERSION}', footer_style))
            
            # Build PDF
            doc.build(elements)
            
            messagebox.showinfo("Success", f"Report exported successfully!\n\nFile: {os.path.basename(filepath)}\nLocation: {os.path.dirname(filepath)}")
            
            # Ask if user wants to open the file
            if messagebox.askyesno("Open File", "Do you want to open the exported file now?"):
                os.startfile(filepath)
        
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export to PDF:\n{str(e)}")
            print(f"PDF export error: {e}")
            import traceback
            traceback.print_exc()

    def create_admin_tab(self):
        """Create Admin Settings tab with premium UI design for librarians"""
        admin_frame = tk.Frame(self.notebook, bg='#f0f2f5')
        self.notebook.add(admin_frame, text="⚙️ Admin")
        
        # Create scrollable container
        canvas = tk.Canvas(admin_frame, bg='#f0f2f5', highlightthickness=0)
        scrollbar = ttk.Scrollbar(admin_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='#f0f2f5')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Resize canvas window with canvas
        def resize_frame(event):
            canvas.itemconfig(canvas_window, width=event.width)
        canvas.bind('<Configure>', resize_frame)
        
        # Mouse wheel scrolling
        def _on_mousewheel(event):
            try:
                if canvas.winfo_exists():
                    canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            except Exception:
                pass
        
        def bind_mousewheel(widget):
            widget.bind("<MouseWheel>", _on_mousewheel)
            for child in widget.winfo_children():
                bind_mousewheel(child)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # ============== HEADER ==============
        header_frame = tk.Frame(scrollable_frame, bg='#f0f2f5')
        header_frame.pack(fill=tk.X, padx=40, pady=(30, 20))
        
        # Title with icon
        title_container = tk.Frame(header_frame, bg='#f0f2f5')
        title_container.pack(anchor='w')
        
        tk.Label(
            title_container,
            text="⚙️",
            font=('Segoe UI', 32),
            bg='#f0f2f5',
            fg=self.colors['secondary']
        ).pack(side=tk.LEFT)
        
        title_text_frame = tk.Frame(title_container, bg='#f0f2f5')
        title_text_frame.pack(side=tk.LEFT, padx=(15, 0))
        
        tk.Label(
            title_text_frame,
            text="Admin Settings",
            font=('Segoe UI', 24, 'bold'),
            bg='#f0f2f5',
            fg='#1a1a2e'
        ).pack(anchor='w')
        
        tk.Label(
            title_text_frame,
            text="Configure library rules and system settings",
            font=('Segoe UI', 12),
            bg='#f0f2f5',
            fg='#666'
        ).pack(anchor='w')
        
        # ============== MAIN CONTENT GRID ==============
        content_frame = tk.Frame(scrollable_frame, bg='#f0f2f5')
        content_frame.pack(fill=tk.BOTH, expand=True, padx=40, pady=10)
        
        # Helper function to create premium cards
        def create_card(parent, title, icon, accent_color):
            # Outer frame for shadow effect
            outer = tk.Frame(parent, bg='#d0d0d0', padx=2, pady=2)
            
            card = tk.Frame(outer, bg='white', padx=25, pady=20)
            card.pack(fill=tk.BOTH, expand=True)
            
            # Card header
            header = tk.Frame(card, bg='white')
            header.pack(fill=tk.X, pady=(0, 15))
            
            # Icon circle
            icon_label = tk.Label(
                header,
                text=icon,
                font=('Segoe UI', 18),
                bg=accent_color,
                fg='white',
                width=2,
                height=1
            )
            icon_label.pack(side=tk.LEFT)
            
            tk.Label(
                header,
                text=title,
                font=('Segoe UI', 14, 'bold'),
                bg='white',
                fg='#1a1a2e'
            ).pack(side=tk.LEFT, padx=(12, 0))
            
            return outer, card
        
        # Helper for styled entry with label
        def create_styled_setting(parent, label, setting_key, unit, min_val=1, max_val=100):
            row = tk.Frame(parent, bg='white')
            row.pack(fill=tk.X, pady=10)
            
            tk.Label(
                row,
                text=label,
                font=('Segoe UI', 11),
                bg='white',
                fg='#444',
                anchor='w'
            ).pack(side=tk.LEFT)
            
            # Right side container
            right_container = tk.Frame(row, bg='white')
            right_container.pack(side=tk.RIGHT)
            
            current_val = self.library_settings.get(setting_key, 5)
            value_var = tk.StringVar(value=str(current_val))
            
            entry = tk.Entry(
                right_container,
                textvariable=value_var,
                font=('Segoe UI', 11, 'bold'),
                width=6,
                justify='center',
                relief='solid',
                bd=1,
                bg='#f8f9fa'
            )
            entry.pack(side=tk.LEFT, padx=5)
            
            if unit:
                tk.Label(
                    right_container,
                    text=unit,
                    font=('Segoe UI', 10),
                    bg='white',
                    fg='#888'
                ).pack(side=tk.LEFT, padx=(0, 10))
            
            def save_this_setting():
                try:
                    new_val = int(value_var.get())
                    if new_val < min_val or new_val > max_val:
                        messagebox.showerror("Invalid", f"Value must be {min_val}-{max_val}")
                        return
                    self.library_settings[setting_key] = new_val
                    if self.save_library_settings(self.library_settings):
                        messagebox.showinfo("Saved", f"{label} updated to {new_val}")
                        entry.config(bg='#d4edda')  # Green flash
                        self.root.after(1000, lambda: entry.config(bg='#f8f9fa'))
                except ValueError:
                    messagebox.showerror("Error", "Enter a valid number")
            
            save_btn = tk.Button(
                right_container,
                text="Save",
                font=('Segoe UI', 9, 'bold'),
                bg='#28a745',
                fg='white',
                relief='flat',
                padx=12,
                pady=2,
                cursor='hand2',
                command=save_this_setting
            )
            save_btn.pack(side=tk.LEFT)
            
            # Hover effect
            def on_enter(e): save_btn.config(bg='#218838')
            def on_leave(e): save_btn.config(bg='#28a745')
            save_btn.bind('<Enter>', on_enter)
            save_btn.bind('<Leave>', on_leave)
            
            return value_var
        
        # ============== ROW 1: Library Rules + Security ==============
        row1 = tk.Frame(content_frame, bg='#f0f2f5')
        row1.pack(fill=tk.X, pady=10)
        
        # --- LIBRARY RULES CARD ---
        rules_outer, rules_card = create_card(row1, "Library Rules", "📚", '#2E86AB')
        rules_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        create_styled_setting(rules_card, "Fine Per Day", "fine_per_day", "₹/day", 0, 500)
        create_styled_setting(rules_card, "Loan Period", "loan_period_days", "days", 1, 30)
        create_styled_setting(rules_card, "Max Books/Student", "max_books_per_student", "books", 1, 20)
        
        # --- SECURITY CARD ---
        sec_outer, sec_card = create_card(row1, "Security", "🔐", '#dc3545')
        sec_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        def open_password_dialog():
            dialog = tk.Toplevel(self.root)
            dialog.title("🔑 Change Admin Password")
            dialog.geometry("420x420")
            dialog.configure(bg='white')
            dialog.transient(self.root)
            dialog.grab_set()
            dialog.resizable(False, False)
            
            # Center
            dialog.update_idletasks()
            x = (dialog.winfo_screenwidth() - 420) // 2
            y = (dialog.winfo_screenheight() - 420) // 2
            dialog.geometry(f"+{x}+{y}")
            
            # Header
            header_bg = tk.Frame(dialog, bg='#2E86AB', height=60)
            header_bg.pack(fill=tk.X)
            header_bg.pack_propagate(False)
            
            tk.Label(
                header_bg,
                text="🔑 Change Admin Password",
                font=('Segoe UI', 16, 'bold'),
                bg='#2E86AB',
                fg='white'
            ).pack(expand=True)
            
            form_frame = tk.Frame(dialog, bg='white', padx=30, pady=25)
            form_frame.pack(fill=tk.BOTH, expand=True)
            
            # Form fields
            def create_field(parent, label_text):
                tk.Label(parent, text=label_text, font=('Segoe UI', 11, 'bold'), 
                        bg='white', fg='#333').pack(anchor='w', pady=(10, 3))
                entry = tk.Entry(parent, font=('Segoe UI', 12), show='●', 
                               relief='solid', bd=1, bg='#f8f9fa')
                entry.pack(fill=tk.X, ipady=8)
                return entry
            
            current_entry = create_field(form_frame, "Current Password")
            new_entry = create_field(form_frame, "New Password")
            confirm_entry = create_field(form_frame, "Confirm New Password")
            
            def do_save():
                # Verify current password
                pwd = current_entry.get()
                hashed_input = hashlib.sha256(pwd.encode()).hexdigest()
                is_valid_pwd = False

                if 'admin_password' in self.library_settings:
                    if pwd == self.library_settings['admin_password']:
                        is_valid_pwd = True
                elif hashed_input == self.library_settings.get('admin_password_hash', ADMIN_PASSWORD_HASH):
                    is_valid_pwd = True

                if not is_valid_pwd:
                    messagebox.showerror("Error", "Current password is incorrect", parent=dialog)
                    return
                if new_entry.get() != confirm_entry.get():
                    messagebox.showerror("Error", "New passwords do not match", parent=dialog)
                    return
                if len(new_entry.get()) < 4:
                    messagebox.showerror("Error", "Password must be at least 4 characters", parent=dialog)
                    return
                
                # Delete old plaintext password if it existed
                if 'admin_password' in self.library_settings:
                    del self.library_settings['admin_password']

                # Save new password
                self.library_settings['admin_password_hash'] = hashlib.sha256(new_entry.get().encode()).hexdigest()
                if self.save_library_settings(self.library_settings):
                    messagebox.showinfo("Success", "Password updated successfully!", parent=dialog)
                    dialog.destroy()
                else:
                    messagebox.showerror("Error", "Failed to save password settings", parent=dialog)
            
            # Buttons
            btn_frame = tk.Frame(form_frame, bg='white')
            btn_frame.pack(fill=tk.X, pady=(25, 0))
            
            save_btn = tk.Button(
                btn_frame,
                text="💾 Save Password",
                font=('Segoe UI', 11, 'bold'),
                bg='#28a745',
                fg='white',
                relief='flat',
                padx=25,
                pady=10,
                cursor='hand2',
                command=do_save
            )
            save_btn.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 5))
            
            cancel_btn = tk.Button(
                btn_frame,
                text="Cancel",
                font=('Segoe UI', 11),
                bg='#6c757d',
                fg='white',
                relief='flat',
                padx=25,
                pady=10,
                cursor='hand2',
                command=dialog.destroy
            )
            cancel_btn.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(5, 0))
        
        # Password button
        pwd_btn = tk.Button(
            sec_card,
            text="🔑  Change Admin Password",
            font=('Segoe UI', 11, 'bold'),
            bg='#2E86AB',
            fg='white',
            relief='flat',
            padx=20,
            pady=12,
            cursor='hand2',
            command=open_password_dialog
        )
        pwd_btn.pack(fill=tk.X, pady=(5, 10))
        
        # Hover effect
        def pwd_enter(e): pwd_btn.config(bg='#1d6a8a')
        def pwd_leave(e): pwd_btn.config(bg='#2E86AB')
        pwd_btn.bind('<Enter>', pwd_enter)
        pwd_btn.bind('<Leave>', pwd_leave)
        
        # Info text
        tk.Label(
            sec_card,
            text="Admin credentials are used to login\nto the desktop application",
            font=('Segoe UI', 10),
            bg='white',
            fg='#888',
            justify='left'
        ).pack(anchor='w', pady=(5, 0))
        
        # ============== ROW 2: Backup + Email ==============
        row2 = tk.Frame(content_frame, bg='#f0f2f5')
        row2.pack(fill=tk.X, pady=10)
        
        # --- BACKUP CARD ---
        backup_outer, backup_card = create_card(row2, "Backup & Data", "💾", '#28a745')
        backup_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        def do_backup():
            try:
                import shutil
                from tkinter import filedialog
                src = os.path.join(os.path.dirname(__file__), 'library.db')
                if not os.path.exists(src):
                    messagebox.showerror("Error", "Database file not found")
                    return
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                dest = filedialog.asksaveasfilename(
                    defaultextension=".db",
                    filetypes=[("Database", "*.db")],
                    initialfile=f"library_backup_{timestamp}.db",
                    title="Save Backup"
                )
                if dest:
                    shutil.copy2(src, dest)
                    messagebox.showinfo("Success", f"Backup saved to:\n{dest}")
            except Exception as e:
                messagebox.showerror("Error", f"Backup failed: {e}")
        
        backup_btn = tk.Button(
            backup_card,
            text="💾  Backup Database Now",
            font=('Segoe UI', 11, 'bold'),
            bg='#28a745',
            fg='white',
            relief='flat',
            padx=20,
            pady=12,
            cursor='hand2',
            command=do_backup
        )
        backup_btn.pack(fill=tk.X, pady=(5, 10))
        
        tk.Label(
            backup_card,
            text="Create a backup copy of library.db\nRecommended: weekly backups to USB",
            font=('Segoe UI', 10),
            bg='white',
            fg='#888',
            justify='left'
        ).pack(anchor='w')
        
        # --- EMAIL CARD ---
        email_outer, email_card = create_card(row2, "Email Settings", "📧", '#6f42c1')
        email_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        email_btn = tk.Button(
            email_card,
            text="📧  Configure Email & Reminders",
            font=('Segoe UI', 11, 'bold'),
            bg='#6f42c1',
            fg='white',
            relief='flat',
            padx=20,
            pady=12,
            cursor='hand2',
            command=self.open_email_settings
        )
        email_btn.pack(fill=tk.X, pady=(5, 10))
        
        tk.Label(
            email_card,
            text="SMTP settings for Gmail\nAuto-reminders run daily at 9 AM",
            font=('Segoe UI', 10),
            bg='white',
            fg='#888',
            justify='left'
        ).pack(anchor='w')
        
        # ============== ROW 3: System Info + Danger Zone ==============
        row3 = tk.Frame(content_frame, bg='#f0f2f5')
        row3.pack(fill=tk.X, pady=10)
        
        # --- SYSTEM INFO CARD ---
        info_outer, info_card = create_card(row3, "System Information", "ℹ️", '#17a2b8')
        info_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        info_text = f"""
Application: Library Management System 2.0
Version: {APP_VERSION}
Built by: Yash Date
Release Date: December 23, 2025
Database: library.db
Portal: Port {self.portal_port}

Current Settings:
  • Fine: ₹{self.library_settings.get('fine_per_day', 5)}/day
  • Loan: {self.library_settings.get('loan_period_days', 7)} days
  • Max Books: {self.library_settings.get('max_books_per_student', 5)}
        """.strip()
        
        self.admin_info_label = tk.Label(
            info_card,
            text=info_text,
            font=('Consolas', 10),
            bg='#f8f9fa',
            fg='#333',
            justify='left',
            padx=15,
            pady=12,
            relief='flat',
            anchor='nw'
        )
        self.admin_info_label.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # --- DANGER ZONE CARD ---
        danger_outer, danger_card = create_card(row3, "Danger Zone", "⚠️", '#dc3545')
        danger_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        danger_btn = tk.Button(
            danger_card,
            text="🗑️  Clear All Data",
            font=('Segoe UI', 11, 'bold'),
            bg='#dc3545',
            fg='white',
            relief='flat',
            padx=20,
            pady=12,
            cursor='hand2',
            command=self.clear_all_data_ui
        )
        danger_btn.pack(fill=tk.X, pady=(5, 10))
        
        tk.Label(
            danger_card,
            text="⚠️ Permanently deletes all data\nRequires password confirmation",
            font=('Segoe UI', 10),
            bg='white',
            fg='#dc3545',
            justify='left'
        ).pack(anchor='w')
        
        # ============== ROW 4: PERFORMANCE & SYNC ==============
        if PERFORMANCE_MODULES_AVAILABLE:
            row4 = tk.Frame(content_frame, bg='#f0f2f5')
            row4.pack(fill=tk.X, pady=10)
            
            # --- SYNC CARD ---
            sync_outer, sync_card = create_card(row4, "Database Sync", "🔄", '#6c5ce7')
            sync_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
            
            # Sync status label
            sync_status_label = tk.Label(
                sync_card,
                text="Last sync: Never",
                font=('Segoe UI', 10),
                bg='white',
                fg='#666'
            )
            sync_status_label.pack(anchor='w', pady=(0, 10))
            
            # Update sync status
            def update_sync_status():
                try:
                    log_file = os.path.join(os.path.dirname(__file__), 'sync_log.json')
                    if os.path.exists(log_file):
                        with open(log_file, 'r') as f:
                            log_data = json.load(f)
                            last_sync = log_data.get('last_sync', 'Never')
                            status = log_data.get('status', 'unknown')
                            sync_status_label['text'] = f"Last sync: {last_sync} ({status})"
                    else:
                        sync_status_label['text'] = "Last sync: Never"
                except:
                    sync_status_label['text'] = "Last sync: Unknown"
            
            update_sync_status()
            
            # Manual sync button
            def do_manual_sync():
                if not self.sync_manager:
                    messagebox.showinfo(
                        "Sync Not Configured",
                        "Remote database sync is not configured yet.\n\n"
                        "Your data is safely stored in the local database (library.db).\n\n"
                        "To enable remote sync (optional):\n"
                        "1) Create a .env file with DATABASE_URL=postgresql://user:pass@host:port/db\n"
                        "2) Install psycopg2: pip install psycopg2-binary\n"
                        "3) Restart the application\n\n"
                        "This feature is optional — the application works perfectly offline."
                    )
                    return
                
                # Show progress dialog
                progress_win = tk.Toplevel(self.root)
                progress_win.title("Database Sync")
                progress_win.geometry("400x200")
                progress_win.transient(self.root)
                progress_win.grab_set()
                
                # Center window
                progress_win.update_idletasks()
                x = (progress_win.winfo_screenwidth() // 2) - 200
                y = (progress_win.winfo_screenheight() // 2) - 100
                progress_win.geometry(f"+{x}+{y}")
                
                tk.Label(
                    progress_win,
                    text="🔄 Syncing database...",
                    font=('Segoe UI', 12, 'bold'),
                    pady=20
                ).pack()
                
                status_label = tk.Label(
                    progress_win,
                    text="Syncing local to remote...",
                    font=('Segoe UI', 10)
                )
                status_label.pack(pady=10)
                
                progress_win.update()
                
                # Run sync in background
                def run_sync():
                    try:
                        result = self.sync_manager.sync_now(direction='both')
                        return result
                    except Exception as e:
                        return {'error': str(e)}
                
                def sync_callback(result):
                    progress_win.destroy()
                    if 'error' in result:
                        messagebox.showerror("Sync Error", f"Failed to sync database:\n{result['error']}")
                    else:
                        update_sync_status()
                        messagebox.showinfo(
                            "Sync Complete",
                            f"✅ Database synced successfully!\n\n"
                            f"Records synced: {result.get('records_synced', 0)}\n"
                            f"Direction: {result.get('direction', 'both')}"
                        )
                
                self.run_in_background_thread(run_sync, sync_callback)
            
            sync_btn_bg = '#2E86AB' if self.sync_manager else '#e74c3c'
            sync_btn_text = "🔄  Sync Now (Local ↔ Remote)" if self.sync_manager else "⚠️  Sync Now (Not Configured)"
            
            manual_sync_btn = tk.Button(
                sync_card,
                text=sync_btn_text,
                font=('Segoe UI', 11, 'bold'),
                bg=sync_btn_bg,
                fg='white',
                relief='flat',
                padx=20,
                pady=12,
                cursor='hand2',
                command=do_manual_sync,
                activebackground='#1a6d8a' if self.sync_manager else '#c0392b',
                activeforeground='white'
            )
            manual_sync_btn.pack(fill=tk.X, pady=(5, 10))
            
            # Hover effect
            sync_hover_bg = '#1a6d8a' if self.sync_manager else '#c0392b'
            def sync_enter(e): manual_sync_btn.config(bg=sync_hover_bg)
            def sync_leave(e): manual_sync_btn.config(bg=sync_btn_bg)
            manual_sync_btn.bind('<Enter>', sync_enter)
            manual_sync_btn.bind('<Leave>', sync_leave)
            
            # Sync interval configuration
            interval_frame = tk.Frame(sync_card, bg='white')
            interval_frame.pack(fill=tk.X, pady=(10, 5))
            
            tk.Label(
                interval_frame,
                text="⏱️ Auto-Sync Interval:",
                font=('Segoe UI', 10, 'bold'),
                bg='white',
                fg='#333'
            ).pack(side=tk.LEFT, padx=(0, 10))
            
            # Ensure a config manager exists to persist interval even if sync isn't configured yet.
            if not getattr(self, 'config_manager', None):
                self.config_manager = get_config()

            interval_var = tk.StringVar(value=str(self.config_manager.get_sync_interval()))
            interval_combo = ttk.Combobox(
                interval_frame,
                textvariable=interval_var,
                values=['2', '5', '10', '15', '20', '30', '45', '60', '90', '120'],
                state='readonly',
                width=8,
                font=('Segoe UI', 10)
            )
            interval_combo.pack(side=tk.LEFT, padx=5)
            
            tk.Label(
                interval_frame,
                text="minutes",
                font=('Segoe UI', 10),
                bg='white',
                fg='#666'
            ).pack(side=tk.LEFT, padx=(0, 10))
            
            def save_interval():
                try:
                    minutes = int(interval_var.get())
                    if self.config_manager.set_sync_interval(minutes):
                        messagebox.showinfo(
                            "Settings Saved",
                            f"✅ Auto-sync interval set to {minutes} minutes\n\n"
                            f"⚠️ Restart the application for changes to take effect."
                        )
                    else:
                        messagebox.showerror("Error", "Invalid interval value")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save interval: {e}")
            
            save_interval_btn = tk.Button(
                interval_frame,
                text="💾 Save",
                font=('Segoe UI', 9, 'bold'),
                bg='#27ae60',
                fg='white',
                relief='flat',
                padx=12,
                pady=4,
                cursor='hand2',
                command=save_interval
            )
            save_interval_btn.pack(side=tk.LEFT, padx=5)
            
            # Info text
            auto_sync_status = "enabled" if self.config_manager.is_sync_enabled() else "disabled"
            sync_availability = "configured" if self.sync_manager else "not configured"
            tk.Label(
                sync_card,
                text=(
                    f"Auto-sync: {auto_sync_status}\n"
                    f"Remote sync: {sync_availability}\n"
                    "Local database ↔ Remote PostgreSQL\n"
                    "💡 Application must be running for sync to work"
                ),
                font=('Segoe UI', 10),
                bg='white',
                fg='#888',
                justify='left'
            ).pack(anchor='w', pady=(5, 0))
            
            # --- CONNECTION POOL STATS CARD ---
            pool_outer, pool_card = create_card(row4, "Performance Stats", "⚡", '#00b894')
            pool_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            
            # Connection pool stats
            pool_stats_label = tk.Label(
                pool_card,
                text="Loading stats...",
                font=('Segoe UI', 10),
                bg='white',
                fg='#666',
                justify='left'
            )
            pool_stats_label.pack(anchor='w', pady=(0, 10))
            
            def update_pool_stats():
                try:
                    if self.connection_pool:
                        stats = self.connection_pool.get_stats()
                        stats_text = (
                            f"Active connections: {stats['active_connections']}\n"
                            f"Available: {stats['available_connections']}\n"
                            f"Total created: {stats['total_connections_created']}\n"
                            f"Total requests: {stats['total_requests']}"
                        )
                        pool_stats_label['text'] = stats_text
                    else:
                        pool_stats_label['text'] = "Connection pool not initialized"
                except Exception as e:
                    pool_stats_label['text'] = f"Error: {e}"
            
            update_pool_stats()
            
            # Refresh stats button
            refresh_stats_btn = tk.Button(
                pool_card,
                text="🔄  Refresh Stats",
                font=('Segoe UI', 10),
                bg='#00b894',
                fg='white',
                relief='flat',
                padx=15,
                pady=8,
                cursor='hand2',
                command=update_pool_stats
            )
            refresh_stats_btn.pack(fill=tk.X, pady=(5, 10))
            
            # Info
            tk.Label(
                pool_card,
                text="✅ Performance optimizations active\nConnection pooling • Batch emails",
                font=('Segoe UI', 10),
                bg='white',
                fg='#888',
                justify='left'
            ).pack(anchor='w', pady=(5, 0))
        
        # Bind mousewheel after all widgets created
        self.root.after(100, lambda: bind_mousewheel(scrollable_frame))

    def refresh_dashboard_borrowed(self):
        """Refresh dashboard borrowed books table (Async)"""
        # Clear current items immediately or wait? 
        # Better to wait until data is ready to avoid flicker, or show "Loading..."
        self.run_in_background_thread(
            lambda: self.db.get_borrowed_books(),
            self._dashboard_borrowed_callback
        )

    def _dashboard_borrowed_callback(self, result):
        if isinstance(result, Exception):
            print(f"Error refreshing dashboard borrowed: {result}")
            return
            
        # Update Treeview
        if hasattr(self, 'dashboard_borrowed_tree'):
            # Configure tags for color coding
            # Text color only (no background highlight)
            self.dashboard_borrowed_tree.tag_configure('overdue', foreground='#dc3545')
            self.dashboard_borrowed_tree.tag_configure('due_soon', foreground='#856404')
            self.dashboard_borrowed_tree.tag_configure('ok', foreground='#155724')
            
            for item in self.dashboard_borrowed_tree.get_children():
                try:
                    self.dashboard_borrowed_tree.delete(item)
                except:
                    pass
            
            for record in result:
                enrollment_no = record[0]
                student_name = record[1]
                department = record[2]
                book_id = record[4]
                book_name = record[5]
                borrow_date = record[7]
                due_date = record[8]
                # Calculate days left and format nicely
                try:
                    from datetime import datetime
                    days_diff = (datetime.strptime(due_date, '%Y-%m-%d') - datetime.now()).days
                    
                    if days_diff < 0:
                        # Overdue - show in red with "X days late"
                        days_display = f"{abs(days_diff)} days late"
                        tag = 'overdue'
                    elif days_diff == 0:
                        days_display = "Due Today!"
                        tag = 'due_soon'
                    elif days_diff <= 3:
                        days_display = f"{days_diff} days left"
                        tag = 'due_soon'
                    else:
                        days_display = f"{days_diff} days left"
                        tag = 'ok'
                except:
                    days_display = 'N/A'
                    tag = ''
                
                self.dashboard_borrowed_tree.insert('', 'end', values=(enrollment_no, student_name, department, book_id, book_name, borrow_date, due_date, days_display), tags=(tag,))

    def refresh_stats_async(self):
        """Fetch stats in background and update UI"""
        self.run_in_background_thread(
            lambda: self.get_library_statistics(),
            self._update_stats_callback
        )

    def _update_stats_callback(self, stats):
        if isinstance(stats, Exception):
            print(f"Stats fetch failed: {stats}")
            return
        
        if hasattr(self, 'stats_container'):
             # Clear existing stats
            for widget in self.stats_container.winfo_children():
                widget.destroy()
            # Recreate with new data
            self.create_stats_cards(self.stats_container, stats)

    def create_stats_cards(self, parent, stats=None):
        """Create statistics cards. If stats is None, trigger async fetch."""
        if stats is None:
            # Trigger async fetch and return (display loading or empty initially)
            self.refresh_stats_async()
            return

        cards_data = [
            ("📚 Total Books", stats['total_titles'], self.colors['secondary']),
            ("📦 Total Copies", stats['total_books'], '#8e44ad'),  # Purple for copies
            ("✅ Available Copies", stats['available_books'], '#28a745'),
            ("📖 Issued Books", stats['borrowed_books'], '#ffc107'),
            ("👥 Total Students", stats['total_students'], '#17a2b8')
        ]
        
        for i, (title, value, color) in enumerate(cards_data):
            card = tk.Frame(parent, bg='white', relief='raised', bd=2)
            # Adjust padding based on position
            pad_x = (0, 15 if i < len(cards_data) - 1 else 0)
            card.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=pad_x)
            
            # Icon/Title
            title_label = tk.Label(
                card,
                text=title,
                font=('Segoe UI', 11, 'bold'), # slightly smaller font to fit 5 cards
                bg='white',
                fg=self.colors['accent']
            )
            title_label.pack(pady=(15, 5))
            
            # Value
            value_label = tk.Label(
                card,
                text=str(value),
                font=('Segoe UI', 24, 'bold'), # slightly smaller font to fit 5 cards
                bg='white',
                fg=color
            )
            value_label.pack(pady=(0, 15))
    
    def get_library_statistics(self):
        """Get library statistics (helper for worker thread)."""
        conn = None
        try:
            conn = self.db.get_connection()
            cursor = conn.cursor()

            # Unique titles
            cursor.execute("SELECT COUNT(*) FROM books")
            total_titles = cursor.fetchone()[0]

            # Total *physical* copies across all titles
            cursor.execute("SELECT COALESCE(SUM(total_copies), 0) FROM books")
            total_books = cursor.fetchone()[0]

            # Available physical copies (books not currently issued)
            cursor.execute("SELECT COALESCE(SUM(available_copies), 0) FROM books")
            available_books = cursor.fetchone()[0]

            # Currently issued (active borrows)
            cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE status = 'borrowed'")
            borrowed_books = cursor.fetchone()[0]

            # Overdue books (still out past due date)
            today = datetime.now().strftime('%Y-%m-%d')
            cursor.execute(
                "SELECT COUNT(*) FROM borrow_records WHERE status = 'borrowed' AND due_date < ?",
                (today,)
            )
            overdue_books = cursor.fetchone()[0]

            # Total students
            cursor.execute("SELECT COUNT(*) FROM students")
            total_students = cursor.fetchone()[0]

            return {
                'total_titles': total_titles,
                'total_books': total_books,
                'available_books': available_books,
                'borrowed_books': borrowed_books,
                'overdue_books': overdue_books,
                'total_students': total_students
            }
        except Exception as e:
            print(f"Error getting statistics: {e}")
            return {
                'total_titles': 0,
                'total_books': 0,
                'available_books': 0,
                'borrowed_books': 0,
                'overdue_books': 0,
                'total_students': 0
            }
        finally:
            if conn:
                try:
                    conn.close()
                except Exception:
                    pass
    
    def create_students_tab(self):
        """Create students management tab"""
        students_frame = tk.Frame(self.notebook, bg=self.colors['primary'])
        self.notebook.add(students_frame, text="👥 Students")
        
        # Top frame for search and actions
        top_frame = tk.Frame(students_frame, bg=self.colors['primary'])
        top_frame.pack(fill=tk.X, padx=20, pady=20)
        
        # Search frame
        search_frame = tk.LabelFrame(
            top_frame,
            text="🔍 Search & Filter Students",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        search_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        # Search controls
        search_controls = tk.Frame(search_frame, bg=self.colors['primary'])
        search_controls.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(search_controls, text="Search:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        search_entry = tk.Entry(search_controls, textvariable=self.student_search_var, font=('Segoe UI', 10), width=25)
        search_entry.pack(side=tk.LEFT, padx=(5, 15))
        # Live search while typing
        search_entry.bind('<KeyRelease>', lambda e: self.search_students())
        
        tk.Label(search_controls, text="Year:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        # Include 'Pass Out' instead of '4th' for final year
        year_combo = ttk.Combobox(search_controls, textvariable=self.student_year_filter, values=["All", "1st", "2nd", "3rd", "Pass Out"], state="readonly", width=12)
        year_combo.pack(side=tk.LEFT, padx=5)
        year_combo.bind('<<ComboboxSelected>>', lambda e: self.search_students())

        tk.Label(search_controls, text="Branch:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT, padx=(10, 0))
        branch_combo = ttk.Combobox(search_controls, textvariable=self.student_branch_filter,
            values=["All", "Computer Engineering", "Information Technology", "Civil Engineering", "Electronics & Telecommunication", "Mechanical Engineering", "Automobile Engineering", "Electrical Engineering"],
            state="readonly", width=25)
        branch_combo.pack(side=tk.LEFT, padx=5)
        branch_combo.bind('<<ComboboxSelected>>', lambda e: self.search_students())
        
        # Actions frame
        actions_frame = tk.LabelFrame(
            top_frame,
            text="⚡ Actions",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        actions_frame.pack(side=tk.RIGHT)
        
        # Action buttons
        buttons_frame = tk.Frame(actions_frame, bg=self.colors['primary'])
        buttons_frame.pack(padx=10, pady=10)
        
        add_student_btn = tk.Button(
            buttons_frame,
            text="➕ Add Student",
            font=('Segoe UI', 10, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.show_add_student_dialog,
            cursor='hand2'
        )
        add_student_btn.pack(side=tk.LEFT, padx=(0, 5))

        import_students_btn = tk.Button(
            buttons_frame,
            text="📥 Import Students (Excel)",
            font=('Segoe UI', 10, 'bold'),
            bg='#6f42c1',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.import_students_excel_new,
            cursor='hand2'
        )
        import_students_btn.pack(side=tk.LEFT, padx=5)
        
        
        # Students list
        students_list_frame = tk.LabelFrame(
            students_frame,
            text="📋 Students List",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent'],
            padx=10,
            pady=10
        )
        students_list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Students treeview
        columns = ('Enrollment No', 'Name', 'Email', 'Phone', 'Branch', 'Year')
        self.students_tree = ttk.Treeview(students_list_frame, columns=columns, show='headings', height=15)
        col_widths = {'Enrollment No': 120, 'Name': 150, 'Email': 180, 'Phone': 120, 'Branch': 130, 'Year': 80}
        for col in columns:
            self.students_tree.heading(col, text=col)
            self.students_tree.column(col, width=col_widths[col])
        students_v_scrollbar = ttk.Scrollbar(students_list_frame, orient=tk.VERTICAL, command=self.students_tree.yview)
        students_h_scrollbar = ttk.Scrollbar(students_list_frame, orient=tk.HORIZONTAL, command=self.students_tree.xview)
        self.students_tree.configure(yscrollcommand=students_v_scrollbar.set, xscrollcommand=students_h_scrollbar.set)
        self.students_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        students_v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        students_h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Add double-click binding for delete option
        self.students_tree.bind('<Double-1>', self.on_student_double_click)
    
    def create_books_tab(self):
        """Create books management tab"""
        books_frame = tk.Frame(self.notebook, bg=self.colors['primary'])
        self.notebook.add(books_frame, text="📚 Books")
        
        # Top frame for search and actions
        top_frame = tk.Frame(books_frame, bg=self.colors['primary'])
        top_frame.pack(fill=tk.X, padx=20, pady=20)
        
        # Search frame
        search_frame = tk.LabelFrame(
            top_frame,
            text="🔍 Search & Filter Books",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        search_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        # Search controls
        search_controls = tk.Frame(search_frame, bg=self.colors['primary'])
        search_controls.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(search_controls, text="Search:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        book_search_entry = tk.Entry(search_controls, textvariable=self.book_search_var, font=('Segoe UI', 10), width=25)
        book_search_entry.pack(side=tk.LEFT, padx=(5, 15))
        # Live book search while typing
        book_search_entry.bind('<KeyRelease>', lambda e: self.search_books())
        
        tk.Label(search_controls, text="Category:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        category_combo = ttk.Combobox(search_controls, textvariable=self.book_category_filter, 
                                    values=["All", "Technology", "Textbook", "Research"], state="readonly", width=12)
        category_combo.pack(side=tk.LEFT, padx=5)
        category_combo.bind('<<ComboboxSelected>>', lambda e: self.search_books())
        
        # Actions frame
        actions_frame = tk.LabelFrame(
            top_frame,
            text="⚡ Actions",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        actions_frame.pack(side=tk.RIGHT)
        
        # Action buttons
        buttons_frame = tk.Frame(actions_frame, bg=self.colors['primary'])
        buttons_frame.pack(padx=10, pady=10)
        
        add_book_btn = tk.Button(
            buttons_frame,
            text="➕ Add Book",
            font=('Segoe UI', 10, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.show_add_book_dialog,
            cursor='hand2'
        )
        add_book_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        import_books_btn = tk.Button(
            buttons_frame,
            text="📥 Import Excel",
            font=('Segoe UI', 10, 'bold'),
            bg='#6f42c1',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.import_books_from_excel,
            cursor='hand2'
        )
        import_books_btn.pack(side=tk.LEFT, padx=5)
        
        export_books_btn = tk.Button(
            buttons_frame,
            text="📊 Export to Excel",
            font=('Segoe UI', 10, 'bold'),
            bg='#28a745',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.export_books_to_excel,
            cursor='hand2'
        )
        export_books_btn.pack(side=tk.LEFT, padx=5)
        # Books list
        books_list_frame = tk.LabelFrame(
            books_frame,
            text="📚 Books Collection",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent'],
            padx=10,
            pady=10
        )
        books_list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Books treeview
        columns = ('Book ID', 'Title', 'Author', 'ISBN', 'Category', 'Total Copies', 'Available', 'Price (₹)')
        self.books_tree = ttk.Treeview(books_list_frame, columns=columns, show='headings', height=15)
        
        # Define headings and widths
        column_widths = {'Book ID': 100, 'Title': 250, 'Author': 180, 'ISBN': 120, 
                        'Category': 100, 'Total Copies': 100, 'Available': 100, 'Price (₹)': 90}
        
        for col in columns:
            self.books_tree.heading(col, text=col)
            self.books_tree.column(col, width=column_widths[col])
        
        # Scrollbars
        books_v_scrollbar = ttk.Scrollbar(books_list_frame, orient=tk.VERTICAL, command=self.books_tree.yview)
        books_h_scrollbar = ttk.Scrollbar(books_list_frame, orient=tk.HORIZONTAL, command=self.books_tree.xview)
        self.books_tree.configure(yscrollcommand=books_v_scrollbar.set, xscrollcommand=books_h_scrollbar.set)
        
        # Pack treeview and scrollbars
        self.books_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        books_v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        books_h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Add double-click binding for delete option
        self.books_tree.bind('<Double-1>', self.on_book_double_click)
    
    def create_transactions_tab(self):
        """Create transactions tab with full-page scrolling (borrow + return + list)"""
        transactions_frame = tk.Frame(self.notebook, bg=self.colors['primary'])
        self.notebook.add(transactions_frame, text="📋 Transactions")

        # Outer canvas + single scrollbar for whole page
        outer_container = tk.Frame(transactions_frame, bg=self.colors['primary'])
        outer_container.pack(fill=tk.BOTH, expand=True)

        canvas = tk.Canvas(outer_container, bg=self.colors['primary'], highlightthickness=0)
        v_scroll = ttk.Scrollbar(outer_container, orient=tk.VERTICAL, command=canvas.yview)
        h_scroll = ttk.Scrollbar(outer_container, orient=tk.HORIZONTAL, command=canvas.xview)
        canvas.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        # Horizontal scrollbar will be packed only if needed (dynamic)
        def _toggle_hbar():
            bbox = canvas.bbox('all')
            if not bbox:
                return
            content_width = bbox[2] - bbox[0]
            visible_width = canvas.winfo_width()
            if content_width > visible_width and not getattr(h_scroll, '_visible', False):
                h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
                h_scroll._visible = True
            elif content_width <= visible_width and getattr(h_scroll, '_visible', False):
                h_scroll.pack_forget()
                h_scroll._visible = False
        h_scroll._visible = False

        # Internal frame that will hold all sections
        main_container = tk.Frame(canvas, bg=self.colors['primary'])
        _mc_window = canvas.create_window((0, 0), window=main_container, anchor='nw')

        def _update_scroll_region(event):
            canvas.configure(scrollregion=canvas.bbox('all'))
            _toggle_hbar()
        main_container.bind('<Configure>', _update_scroll_region)

        # Also re-evaluate horizontal bar on canvas resize
        def _resize_content(event):
            # Match internal frame width to canvas width to remove empty right side
            canvas.itemconfig(_mc_window, width=canvas.winfo_width())
            _toggle_hbar()
        canvas.bind('<Configure>', _resize_content)

        # Mouse wheel scrolling (Windows/Linux) with Shift for horizontal
        def _tx_units(delta):
            return -1 if delta > 0 else (1 if delta < 0 else 0)
        # Handlers when bound directly to the scrollable canvas/main container
        def _tx_vwheel(event):
            u = _tx_units(event.delta)
            if u:
                canvas.yview_scroll(u, 'units')
            return 'break'
        def _tx_hwheel(event):
            u = _tx_units(event.delta)
            if u:
                canvas.xview_scroll(u, 'units')
            return 'break'
        # Child-binding variants: scroll page but don't cancel widget defaults
        def _tx_vwheel_child(event):
            u = _tx_units(getattr(event, 'delta', 0))
            if u:
                canvas.yview_scroll(u, 'units')
            # no return => allow widget default if any
        def _tx_hwheel_child(event):
            u = _tx_units(getattr(event, 'delta', 0))
            if u:
                canvas.xview_scroll(u, 'units')
            # no return => allow widget default if any
        canvas.bind('<MouseWheel>', _tx_vwheel)
        canvas.bind('<Shift-MouseWheel>', _tx_hwheel)
        main_container.bind('<MouseWheel>', _tx_vwheel)
        main_container.bind('<Shift-MouseWheel>', _tx_hwheel)
        # Pointer-scoped global binds while cursor is over the Transactions tab
        def _enter_tx(_e=None):
            try:
                transactions_frame.bind_all('<MouseWheel>', _tx_vwheel, add='+')
                transactions_frame.bind_all('<Shift-MouseWheel>', _tx_hwheel, add='+')
            except Exception:
                pass
        def _leave_tx(_e=None):
            try:
                transactions_frame.unbind_all('<MouseWheel>')
                transactions_frame.unbind_all('<Shift-MouseWheel>')
            except Exception:
                pass
        transactions_frame.bind('<Enter>', _enter_tx)
        transactions_frame.bind('<Leave>', _leave_tx)
        # Focus canvas when pointer enters so wheel events target it
        canvas.bind('<Enter>', lambda e: canvas.focus_set())
        main_container.bind('<Enter>', lambda e: canvas.focus_set())
        # Linux-style wheel events fallback
        def _tx_btn4(_e=None):
            canvas.yview_scroll(-1, 'units'); return 'break'
        def _tx_btn5(_e=None):
            canvas.yview_scroll(1, 'units'); return 'break'
        canvas.bind('<Button-4>', _tx_btn4)
        canvas.bind('<Button-5>', _tx_btn5)

        # Recursively bind wheel events to children so hovering any element scrolls the page
        def _bind_wheel_recursive(widget):
            try:
                # Skip Treeview to avoid double-scrolling its own content
                if isinstance(widget, ttk.Treeview):
                    return
            except Exception:
                pass
            try:
                widget.bind('<MouseWheel>', _tx_vwheel_child, add='+')
                widget.bind('<Shift-MouseWheel>', _tx_hwheel_child, add='+')
            except Exception:
                pass
            for child in getattr(widget, 'winfo_children', lambda: [])():
                _bind_wheel_recursive(child)

        # Defer binding until content is populated; schedule after idle
        try:
            widget_ref = main_container
            widget_ref.after(200, lambda: _bind_wheel_recursive(widget_ref))
        except Exception:
            pass
        
        # =================================================================
        # ISSUE BOOK SECTION - Fixed UI Layout
        # =================================================================
        borrow_frame = tk.LabelFrame(
            main_container,
            text="📚 Issue Book",
            font=('Segoe UI', 16, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent'],
            padx=20,
            pady=20
        )
        borrow_frame.pack(fill=tk.X, expand=True, pady=(0, 25))

        # Quick actions (Import Transactions)
        tx_actions = tk.Frame(borrow_frame, bg=self.colors['primary'])
        tx_actions.pack(fill=tk.X, pady=(0, 10))
        tk.Button(
            tx_actions,
            text="📥 Import Transactions (Excel)",
            font=('Segoe UI', 10, 'bold'),
            bg='#6f42c1',
            fg='white',
            relief='flat',
            padx=12,
            pady=6,
            cursor='hand2',
            command=self.import_transactions_excel_report
        ).pack(side=tk.RIGHT)
        
        # Borrow form container
        borrow_form = tk.Frame(borrow_frame, bg=self.colors['primary'])
        borrow_form.pack(fill=tk.X, padx=20, pady=10)
        
        # Row 1: Student and Book Input - Properly Spaced
        input_row = tk.Frame(borrow_form, bg=self.colors['primary'])
        input_row.pack(fill=tk.X, pady=(0, 20))
        
        # Student Enrollment Column
        student_col = tk.Frame(input_row, bg=self.colors['primary'])
        student_col.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 20))
        
        tk.Label(student_col, text="Student Enrollment No:", 
                font=('Segoe UI', 12, 'bold'), 
                bg=self.colors['primary'], 
                fg=self.colors['accent']).pack(anchor='w', pady=(0, 8))
        
        # Autocomplete cache to avoid hitting DB on every keystroke.
        # AUDIT FIX (Issue 10): Cache stores plain dicts, NOT sqlite3.Row objects.
        # sqlite3.Row objects reference the underlying C cursor state and can become
        # undefined once the DB connection is closed. Converting to dict immediately
        # after fetch makes the cached data independent of the connection lifetime.
        import time as _time
        _ac_cache = {'students': None, 'students_ts': 0, 'books': None, 'books_ts': 0}
        _AC_CACHE_TTL = 30  # seconds

        def _rows_to_dicts(rows):
            """Convert sqlite3.Row list → plain dict list for safe long-lived caching."""
            try:
                return [dict(r) for r in rows]
            except Exception:
                return list(rows)

        # Autocomplete for student enrollment
        def get_student_suggestions(query):
            try:
                now = _time.time()
                if _ac_cache['students'] is None or (now - _ac_cache['students_ts']) > _AC_CACHE_TTL:
                    _ac_cache['students'] = _rows_to_dicts(self.db.get_students())
                    _ac_cache['students_ts'] = now
                students = _ac_cache['students']
                matches = []
                query_lower = query.lower()
                for s in students:
                    enrollment = str(s.get('enrollment_no', s.get('enrollment', ''))).lower()
                    name = str(s.get('name', '')).lower()
                    if query_lower in enrollment or query_lower in name:
                        matches.append(s)
                return matches[:10]  # Limit to 10 suggestions
            except Exception as e:
                print(f"[Autocomplete] Student suggestion error: {e}")
                return []
        
        def format_student_display(student):
            return f"{student['enrollment']} - {student['name']} ({student['year']}, {student['section']})"
        
        self.borrow_enrollment_entry = AutocompleteEntry(
            student_col,
            data_callback=get_student_suggestions,
            display_callback=format_student_display,
            width=25
        )
        self.borrow_enrollment_entry.pack(fill=tk.X, pady=(0, 8))
        self.borrow_enrollment_entry.bind('<KeyRelease>', lambda e: self.show_student_details('borrow'))

        # Student details display
        self.borrow_student_details = tk.Label(student_col, 
                                             text="", 
                                             font=('Segoe UI', 10), 
                                             bg=self.colors['primary'], 
                                             fg='#666666',
                                             wraplength=300)
        self.borrow_student_details.pack(anchor='w')
        
        # Book ID Column
        book_col = tk.Frame(input_row, bg=self.colors['primary'])
        book_col.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(20, 0))
        
        tk.Label(book_col, text="Book ID / Barcode:", 
                font=('Segoe UI', 12, 'bold'), 
                bg=self.colors['primary'], 
                fg=self.colors['accent']).pack(anchor='w', pady=(0, 8))
        
        # Autocomplete for book ID - also searches by barcode
        def get_book_suggestions(query):
            try:
                now = _time.time()
                if _ac_cache['books'] is None or (now - _ac_cache['books_ts']) > _AC_CACHE_TTL:
                    _ac_cache['books'] = _rows_to_dicts(self.db.get_books())
                    _ac_cache['books_ts'] = now
                books = _ac_cache['books']
                matches = []
                query_lower = query.lower()
                for b in books:
                    book_id = str(b.get('book_id', '')).lower()
                    title = str(b.get('title', '')).lower()
                    author = str(b.get('author', '')).lower()
                    barcode = str(b.get('barcode') or '').lower()
                    if query_lower in book_id or query_lower in title or query_lower in author or query_lower in barcode:
                        matches.append(b)
                return matches[:10]
            except Exception as e:
                print(f"[Autocomplete] Book suggestion error: {e}")
                return []
        
        def format_book_display(book):
            barcode_str = (f" [Barcode: {book['barcode']}]" if book['barcode'] else "")
            return f"{book['book_id']} - {book['title']} by {book['author']}{barcode_str}"
        
        # Helper to find book by barcode and auto-fill book_id
        def on_borrow_barcode_scan(event):
            if event.keysym == 'Return':
                query = self.borrow_book_id_entry.get().strip()
                if query:
                    # Use cached books (plain dicts) if available
                    now = _time.time()
                    if _ac_cache['books'] is None or (now - _ac_cache['books_ts']) > _AC_CACHE_TTL:
                        _ac_cache['books'] = _rows_to_dicts(self.db.get_books())
                        _ac_cache['books_ts'] = now
                    books = _ac_cache['books']
                    for b in books:
                        barcode = str(b.get('barcode') or '')
                        if barcode and barcode.lower() == query.lower():
                            # Found barcode match - auto-fill with book_id
                            self.borrow_book_id_entry.delete(0, tk.END)
                            self.borrow_book_id_entry.insert(0, b['book_id'])
                            self.show_book_details('borrow')
                            print(f"[Barcode Scan] Found book: {b['book_id']} - {b['title']}")
                            return
        
        self.borrow_book_id_entry = AutocompleteEntry(
            book_col,
            data_callback=get_book_suggestions,
            display_callback=format_book_display,
            width=25
        )
        self.borrow_book_id_entry.pack(fill=tk.X, pady=(0, 8))
        self.borrow_book_id_entry.bind('<KeyRelease>', lambda e: self.show_book_details('borrow'))
        self.borrow_book_id_entry.bind('<Return>', on_borrow_barcode_scan)
        
        # Book details display
        self.borrow_book_details = tk.Label(book_col, 
                                          text="", 
                                          font=('Segoe UI', 10), 
                                          bg=self.colors['primary'], 
                                          fg='#666666',
                                          wraplength=300)
        self.borrow_book_details.pack(anchor='w')
        
    # Row 2: Issue & Due Dates and Action Button - Better Layout
        action_row = tk.Frame(borrow_form, bg=self.colors['primary'])
        action_row.pack(fill=tk.X, pady=(15, 0))
        
        # Issue Date Section (uses DateEntry if available)
        borrow_date_section = tk.Frame(action_row, bg=self.colors['primary'])
        borrow_date_section.pack(side=tk.LEFT, fill=tk.X, expand=True)
        tk.Label(borrow_date_section, text="Issue Date:", font=('Segoe UI', 12, 'bold'), bg=self.colors['primary'], fg=self.colors['accent']).pack(anchor='w', pady=(0,8))
        borrow_date_input_frame = tk.Frame(borrow_date_section, bg=self.colors['primary'])
        borrow_date_input_frame.pack(anchor='w')
        if DateEntry:
            self.borrow_borrow_date_entry = DateEntry(borrow_date_input_frame, width=13, date_pattern='yyyy-mm-dd', state='readonly')
            self.borrow_borrow_date_entry.pack(side=tk.LEFT)
            # Auto-update due date when borrow date picked
            self.borrow_borrow_date_entry.bind('<<DateEntrySelected>>', self.on_borrow_date_changed)
        else:
            self.borrow_borrow_date_entry = tk.Entry(borrow_date_input_frame, font=('Segoe UI',12), width=15, relief='solid', bd=2)
            self.borrow_borrow_date_entry.pack(side=tk.LEFT)
            tk.Button(
                borrow_date_input_frame,
                text="📅",
                font=('Segoe UI', 14),
                bg=self.colors['secondary'],
                fg='white',
                relief='flat',
                padx=10,
                pady=6,
                command=lambda: self.show_date_picker(self.borrow_borrow_date_entry),
                cursor='hand2'
            ).pack(side=tk.LEFT, padx=(8,0))
            # Prevent manual typing; force picker
            self.borrow_borrow_date_entry.bind('<Key>', lambda e: 'break')

        # Due Date Section (DateEntry if available)
        date_section = tk.Frame(action_row, bg=self.colors['primary'])
        date_section.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(20,0))
        tk.Label(date_section, text="Due Date:", font=('Segoe UI', 12, 'bold'), bg=self.colors['primary'], fg=self.colors['accent']).pack(anchor='w', pady=(0,8))
        date_input_frame = tk.Frame(date_section, bg=self.colors['primary'])
        date_input_frame.pack(anchor='w')
        if DateEntry:
            self.borrow_due_date_entry = DateEntry(date_input_frame, width=13, date_pattern='yyyy-mm-dd', state='readonly')
            self.borrow_due_date_entry.pack(side=tk.LEFT)
            # Intercept any manual selection attempts
            self.borrow_due_date_entry.bind('<<DateEntrySelected>>', lambda e: self.on_due_date_attempt_change())
        else:
            self.borrow_due_date_entry = tk.Entry(date_input_frame, font=('Segoe UI', 12), width=15, relief='solid', bd=2)
            self.borrow_due_date_entry.pack(side=tk.LEFT)
            tk.Button(
                date_input_frame,
                text="📅",
                font=('Segoe UI', 14),
                bg=self.colors['secondary'],
                fg='white',
                relief='flat',
                padx=10,
                pady=6,
                command=lambda: self.show_date_picker(self.borrow_due_date_entry),
                cursor='hand2'
            ).pack(side=tk.LEFT, padx=(8,0))
            # Block manual typing; will be auto-set
            self.borrow_due_date_entry.bind('<Key>', lambda e: 'break')
        
        # Action Button Section - Right Side
        button_section = tk.Frame(action_row, bg=self.colors['primary'])
        button_section.pack(side=tk.RIGHT, padx=(20, 0))
        
        # Issue button - Prominent and well-positioned
        borrow_btn = tk.Button(
            button_section,
            text="📚 Issue Book",
            font=('Segoe UI', 14, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=30,
            pady=12,
            command=self.borrow_book,
            cursor='hand2'
        )
        borrow_btn.pack(pady=(25, 0))
        
        # Set default due date (exact 7-day loan period per requirement)
        from datetime import datetime, timedelta
        today_str = datetime.now().strftime('%Y-%m-%d')
        default_due = (datetime.now() + timedelta(days=7)).strftime('%Y-%m-%d')
        try:
            if DateEntry:
                self.borrow_borrow_date_entry.set_date(datetime.now())
                self.borrow_due_date_entry.set_date(datetime.now() + timedelta(days=7))
                # Ensure readonly state remains
                try:
                    self.borrow_borrow_date_entry.config(state='readonly')
                    self.borrow_due_date_entry.config(state='readonly')
                except Exception:
                    pass
            else:
                self.borrow_borrow_date_entry.delete(0, tk.END)
                self.borrow_borrow_date_entry.insert(0, today_str)
                self.borrow_due_date_entry.delete(0, tk.END)
                self.borrow_due_date_entry.insert(0, default_due)
        except Exception:
            pass
        
        # =================================================================
        # VISUAL SEPARATOR
        # =================================================================
        separator = tk.Frame(main_container, height=3, bg=self.colors['accent'])
        separator.pack(fill=tk.X, pady=30)
        
        # =================================================================
        # RETURN BOOK SECTION
        # =================================================================
        return_outer = tk.Frame(main_container, bg=self.colors['primary'], highlightthickness=0)
        return_outer.pack(fill=tk.X, pady=(0, 25))
        heading_bar = tk.Frame(return_outer, bg=self.colors['primary'])
        heading_bar.pack(fill=tk.X)
        heading_label = tk.Label(
            heading_bar,
            text="🔄 Return Book",
            font=('Segoe UI', 16, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        heading_label.pack(anchor='w')
        # Return input row - directly in outer frame now
        return_input_row = tk.Frame(return_outer, bg=self.colors['primary'])
        return_input_row.pack(fill=tk.X, pady=(0, 20))
        
        # Student section
        return_student_col = tk.Frame(return_input_row, bg=self.colors['primary'])
        return_student_col.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 20))
        
        tk.Label(return_student_col, text="Student Enrollment No:", 
                font=('Segoe UI', 12, 'bold'), 
                bg=self.colors['primary'], 
                fg=self.colors['accent']).pack(anchor='w', pady=(0, 8))
        
        # Autocomplete for return student enrollment
        def get_return_student_suggestions(query):
            try:
                now = _time.time()
                if _ac_cache['students'] is None or (now - _ac_cache['students_ts']) > _AC_CACHE_TTL:
                    _ac_cache['students'] = _rows_to_dicts(self.db.get_students())
                    _ac_cache['students_ts'] = now
                students = _ac_cache['students']
                matches = []
                query_lower = query.lower()
                for s in students:
                    enrollment = str(s.get('enrollment_no', s.get('enrollment', ''))).lower()
                    name = str(s.get('name', '')).lower()
                    if query_lower in enrollment or query_lower in name:
                        matches.append(s)
                return matches[:10]
            except Exception as e:
                print(f"[Autocomplete] Return student suggestion error: {e}")
                return []
        
        def format_return_student_display(student):
            return f"{student['enrollment']} - {student['name']} ({student['year']}, {student['section']})"
        
        self.return_enrollment_entry = AutocompleteEntry(
            return_student_col,
            data_callback=get_return_student_suggestions,
            display_callback=format_return_student_display,
            width=25
        )
        self.return_enrollment_entry.pack(fill=tk.X, pady=(0, 8))
        self.return_enrollment_entry.bind('<KeyRelease>', lambda e: self.show_student_details('return'))
        
        self.return_student_details = tk.Label(return_student_col, 
                                             text="", 
                                             font=('Segoe UI', 10), 
                                             bg=self.colors['primary'], 
                                             fg='#666666',
                                             wraplength=300)
        self.return_student_details.pack(anchor='w')
        
        # Book section
        return_book_col = tk.Frame(return_input_row, bg=self.colors['primary'])
        return_book_col.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(20, 0))
        
        tk.Label(return_book_col, text="Book ID / Barcode:", 
                font=('Segoe UI', 12, 'bold'), 
                bg=self.colors['primary'], 
                fg=self.colors['accent']).pack(anchor='w', pady=(0, 8))
        
        # Autocomplete for return book ID - also searches by barcode
        def get_return_book_suggestions(query):
            try:
                now = _time.time()
                if _ac_cache['books'] is None or (now - _ac_cache['books_ts']) > _AC_CACHE_TTL:
                    _ac_cache['books'] = _rows_to_dicts(self.db.get_books())
                    _ac_cache['books_ts'] = now
                books = _ac_cache['books']
                matches = []
                query_lower = query.lower()
                for b in books:
                    book_id = str(b.get('book_id', '')).lower()
                    title = str(b.get('title', '')).lower()
                    author = str(b.get('author', '')).lower()
                    barcode = str(b.get('barcode') or '').lower()
                    if query_lower in book_id or query_lower in title or query_lower in author or query_lower in barcode:
                        matches.append(b)
                return matches[:10]
            except Exception as e:
                print(f"[Autocomplete] Return book suggestion error: {e}")
                return []
        
        def format_return_book_display(book):
            barcode_str = (f" [Barcode: {book['barcode']}]" if book['barcode'] else "")
            return f"{book['book_id']} - {book['title']} by {book['author']}{barcode_str}"
        
        # Helper to find book by barcode and auto-fill book_id for return
        def on_return_barcode_scan(event):
            if event.keysym == 'Return':
                query = self.return_book_id_entry.get().strip()
                if query:
                    # Use cached books if available
                    now = _time.time()
                    if _ac_cache['books'] is None or (now - _ac_cache['books_ts']) > _AC_CACHE_TTL:
                        _ac_cache['books'] = self.db.get_books()
                        _ac_cache['books_ts'] = now
                    books = _ac_cache['books']
                    for b in books:
                        barcode = str(b['barcode'] if b['barcode'] else '')
                        if barcode and barcode.lower() == query.lower():
                            # Found barcode match - auto-fill with book_id
                            self.return_book_id_entry.delete(0, tk.END)
                            self.return_book_id_entry.insert(0, b['book_id'])
                            self.show_book_details('return')
                            print(f"[Barcode Scan] Found book: {b['book_id']} - {b['title']}")
                            return
        
        self.return_book_id_entry = AutocompleteEntry(
            return_book_col,
            data_callback=get_return_book_suggestions,
            display_callback=format_return_book_display,
            width=25
        )
        self.return_book_id_entry.pack(fill=tk.X, pady=(0, 8))
        self.return_book_id_entry.bind('<KeyRelease>', lambda e: self.show_book_details('return'))
        self.return_book_id_entry.bind('<Return>', on_return_barcode_scan)
        
        self.return_book_details = tk.Label(return_book_col, 
                                          text="", 
                                          font=('Segoe UI', 10), 
                                          bg=self.colors['primary'], 
                                          fg='#666666',
                                          wraplength=300)
        self.return_book_details.pack(anchor='w')
        
        # Return button - Centered and prominent
        return_button_frame = tk.Frame(return_outer, bg=self.colors['primary'])
        # Return date row
        return_date_row = tk.Frame(return_outer, bg=self.colors['primary'])
        return_date_row.pack(fill=tk.X, pady=(10, 5))
        tk.Label(return_date_row, text="Return Date:", font=('Segoe UI', 12, 'bold'), bg=self.colors['primary'], fg=self.colors['accent']).pack(side=tk.LEFT, padx=(0,8))
        if DateEntry:
            self.return_date_entry = DateEntry(return_date_row, width=13, date_pattern='yyyy-mm-dd', state='readonly')
            self.return_date_entry.pack(side=tk.LEFT)
            try:
                self.return_date_entry.set_date(datetime.now())
            except Exception:
                pass
        else:
            self.return_date_entry = tk.Entry(return_date_row, font=('Segoe UI', 12), width=15, relief='solid', bd=2)
            self.return_date_entry.pack(side=tk.LEFT)
            tk.Button(return_date_row, text="📅", font=('Segoe UI', 14), bg=self.colors['secondary'], fg='white', relief='flat', padx=10, pady=6, cursor='hand2', command=lambda: self.show_date_picker(self.return_date_entry)).pack(side=tk.LEFT, padx=(8,0))
            # Set default return date to today
            try:
                from datetime import datetime as _dt
                self.return_date_entry.delete(0, tk.END)
                self.return_date_entry.insert(0, _dt.now().strftime('%Y-%m-%d'))
            except Exception:
                pass
            self.return_date_entry.bind('<Key>', lambda e: 'break')
        return_button_frame.pack(pady=(10, 0))
        
        return_btn = tk.Button(
            return_button_frame,
            text="🔄 Return Book",
            font=('Segoe UI', 14, 'bold'),
            bg='#28a745',
            fg='white',
            relief='flat',
            padx=30,
            pady=12,
            command=self.return_book,
            cursor='hand2'
        )
        return_btn.pack()
        
        # =================================================================
        # CURRENTLY ISSUED BOOKS SECTION - Improved Layout
        # =================================================================
        borrowed_frame = tk.LabelFrame(
            main_container,
            text="📋 Currently Issued Books",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent'],
            padx=10,
            pady=10
        )
        borrowed_frame.pack(fill=tk.BOTH, expand=True, pady=(0,40))
        
        borrowed_columns = ('Student', 'Branch', 'Book ID', 'Book Title', 'Issue Date', 'Due Date', 'Days Left')
        self.borrowed_tree = ttk.Treeview(borrowed_frame, columns=borrowed_columns, show='headings', height=10)
        borrowed_widths = {'Student': 150, 'Branch': 130, 'Book ID': 100, 'Book Title': 220, 'Issue Date': 110, 'Due Date': 110, 'Days Left': 90}
        
        for col in borrowed_columns:
            self.borrowed_tree.heading(col, text=col)
            self.borrowed_tree.column(col, width=borrowed_widths[col])
        
        borrowed_v_scrollbar = ttk.Scrollbar(borrowed_frame, orient=tk.VERTICAL, command=self.borrowed_tree.yview)
        borrowed_h_scrollbar = ttk.Scrollbar(borrowed_frame, orient=tk.HORIZONTAL, command=self.borrowed_tree.xview)
        self.borrowed_tree.configure(yscrollcommand=borrowed_v_scrollbar.set, xscrollcommand=borrowed_h_scrollbar.set)
        
        self.borrowed_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        borrowed_v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        borrowed_h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
    
    # Commented out - no longer needed since we removed combobox autocomplete
    # def filter_student_suggestions(self, event, mode):
    #     """Filter student suggestions as user types"""
    #     try:
    #         current_text = event.widget.get().lower().strip()
    #         if len(current_text) < 1:  # Start filtering after 1 character
    #             return
    #         
    #         # Get all students
    #         students = self.db.get_students()
    #         # Filter students that match the current text (enrollment number or name)
    #         matching_students = []
    #         for s in students:
    #             enrollment = str(s[0]).lower()
    #             name = str(s[1]).lower()
    #             if current_text in enrollment or current_text in name:
    #                 matching_students.append(f"{s[0]} - {s[1]}")  # Format as "enrollment - name"
    #         
    #         # Update combobox values
    #         event.widget['values'] = matching_students[:10]  # Limit to 10 suggestions
    #         
    #         # Auto-open dropdown if there are matches
    #         if matching_students and len(current_text) >= 1:
    #             event.widget.event_generate('<Down>')
    #             
    #     except Exception as e:
    #         print(f"Error filtering student suggestions: {e}")
    # 
    # def filter_book_suggestions(self, event, mode):
    #     """Filter book suggestions as user types"""
    #     try:
    #         current_text = event.widget.get().lower().strip()
    #         if len(current_text) < 1:  # Start filtering after 1 character
    #             return
    #         
    #         # Get all books
    #         books = self.db.get_books()
    #         # Filter books that match the current text (book ID or title)
    #         matching_books = []
    #         for b in books:
    #             book_id = str(b[0]).lower()
    #             title = str(b[1]).lower()
    #             if current_text in book_id or current_text in title:
    #                 matching_books.append(f"{b[0]} - {b[1]}")  # Format as "ID - title"
    #         
    #         # Update combobox values
    #         event.widget['values'] = matching_books[:10]  # Limit to 10 suggestions
    #         
    #         # Auto-open dropdown if there are matches
    #         if matching_books and len(current_text) >= 1:
    #             event.widget.event_generate('<Down>')
    #             
    #     except Exception as e:
    #         print(f"Error filtering book suggestions: {e}")
    
    def show_student_details(self, mode):
        """Show student details with debouncing and async fetch"""
        if mode == 'borrow':
            enrollment_no = self.borrow_enrollment_entry.get().strip()
        else:
            enrollment_no = self.return_enrollment_entry.get().strip()
        
        # Debouncing
        if hasattr(self, '_student_details_timer') and self._student_details_timer:
            self.root.after_cancel(self._student_details_timer)
        
        if not enrollment_no:
            if mode == 'borrow':
                self.borrow_student_details.config(text="")
            else:
                self.return_student_details.config(text="")
            return

        # Delay before fetching (500ms)
        self._student_details_timer = self.root.after(500, 
            lambda: self.run_in_background_thread(
                lambda: self.db.get_student_by_enrollment(enrollment_no),
                lambda result: self._show_student_details_callback(result, mode)
            )
        )

    def _show_student_details_callback(self, student, mode):
        """Callback to update student details label"""
        if isinstance(student, Exception):
            details = "Error loading student data."
        elif student:
            # student tuple: (id, enrollment_no, name, email, phone, department, year, ...)
            details = f"Name: {student[2]} | Branch: {student[5]} | Email: {student[3]} | Phone: {student[4]} | Year: {student[6]}"
        else:
            details = "Student not found."

        if mode == 'borrow':
            self.borrow_student_details.config(text=details)
        else:
            self.return_student_details.config(text=details)
    
    def show_book_details(self, mode):
        """Show book details with debouncing and async fetch"""
        if mode == 'borrow':
            book_id = self.borrow_book_id_entry.get().strip()
        else:
            book_id = self.return_book_id_entry.get().strip()
        
        # Debouncing
        if hasattr(self, '_book_details_timer') and self._book_details_timer:
            self.root.after_cancel(self._book_details_timer)

        if not book_id:
            if mode == 'borrow':
                self.borrow_book_details.config(text="")
            else:
                self.return_book_details.config(text="")
            return
            
        # Delay before fetching (500ms)
        self._book_details_timer = self.root.after(500, 
            lambda: self.run_in_background_thread(
                lambda: self.db.get_book_by_id(book_id),
                lambda result: self._show_book_details_callback(result, mode)
            )
        )

    def _show_book_details_callback(self, book, mode):
        """Callback to update book details label"""
        if isinstance(book, Exception):
            details = "Error loading book data."
        elif book:
            # book tuple: (id, book_id, title, author, isbn, category, total_copies, available, ...)
            details = f"Title: {book[2]} | Author: {book[3]} | Available: {book[7]}"
        else:
            details = "Book not found."

        # Update the appropriate book details label
        if mode == 'borrow':
            self.borrow_book_details.config(text=details)
        else:
            self.return_book_details.config(text=details)
    
    def on_student_double_click(self, event):
        """Handle double-click on student to show edit dialog"""
        item = self.students_tree.selection()[0] if self.students_tree.selection() else None
        if not item:
            return

        student_data = self.students_tree.item(item, 'values')
        if not student_data:
            return

        enrollment_no = student_data[0]
        
        # Get full student details from database
        try:
            students = self.db.get_students(enrollment_no)
            student = None
            for s in students:
                if str(s[1]) == str(enrollment_no):  # s[1] is enrollment_no
                    student = s
                    break
            
            if not student:
                messagebox.showerror("Error", "Student not found")
                return
            
            # Show edit dialog
            self.show_edit_student_dialog(student)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load student details: {e}")
    
    def show_edit_student_dialog(self, student):
        """Show dialog to edit student information"""
        # student tuple: (id, enrollment_no, name, email, phone, department, year, date_registered)
        dialog = tk.Toplevel(self.root)
        dialog.title("Edit Student")
        dialog.geometry("500x550")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 100, self.root.winfo_rooty() + 100))
        
        # Title
        title_label = tk.Label(
            dialog,
            text="✏️ Edit Student",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack(pady=(20, 30))
        
        # Form frame
        form_frame = tk.Frame(dialog, bg='white')
        form_frame.pack(expand=True, padx=40)
        
        # Form fields
        fields = [
            ("Enrollment No:", "enrollment", student[1], False),  # Read-only
            ("Name:", "name", student[2], True),
            ("Email:", "email", student[3] or '', True),
            ("Phone:", "phone", student[4] or '', True),
            ("Department:", "department", student[5] or 'Computer', False),  # Read-only
            ("Year:", "year", student[6] or '1st', True)
        ]
        
        entries = {}
        
        for i, (label_text, field_name, default_value, editable) in enumerate(fields):
            # Label
            label = tk.Label(
                form_frame,
                text=label_text,
                font=('Segoe UI', 11, 'bold'),
                bg='white',
                fg=self.colors['accent']
            )
            label.grid(row=i, column=0, sticky='w', pady=(0, 15), padx=(0, 20))
            
            # Entry or Combobox
            if field_name == "year":
                entry = ttk.Combobox(
                    form_frame,
                    values=["1st", "2nd", "3rd", "Pass Out"],
                    font=('Segoe UI', 11),
                    width=25,
                    state="readonly" if editable else "disabled"
                )
                entry.set(default_value)
                entry.grid(row=i, column=1, pady=(0, 15))
            else:
                entry = tk.Entry(
                    form_frame,
                    font=('Segoe UI', 11),
                    width=28,
                    relief='solid',
                    bd=2
                )
                entry.insert(0, default_value)
                if not editable:
                    entry.config(state='readonly', bg='#f0f0f0')
                entry.grid(row=i, column=1, pady=(0, 15))
            
            entries[field_name] = entry
        
        # Buttons frame
        btn_frame = tk.Frame(dialog, bg='white')
        btn_frame.pack(pady=20)
        
        def save_changes():
            # Validate required fields
            if not entries['name'].get().strip():
                messagebox.showerror("Error", "Name is required!")
                return
            
            # Update student
            success, message = self.db.update_student(
                entries['enrollment'].get(),
                entries['name'].get().strip(),
                entries['email'].get().strip(),
                entries['phone'].get().strip(),
                entries['department'].get().strip(),
                entries['year'].get()
            )
            
            if success:
                messagebox.showinfo("Success", message)
                dialog.destroy()
                self.refresh_students()
                self.refresh_dashboard()
            else:
                messagebox.showerror("Error", message)
        
        def delete_student():
            if messagebox.askyesno(
                "Confirm Delete",
                f"Delete student: {entries['name'].get()}?\n\nThis action cannot be undone!",
                icon='warning'
            ):
                student_name = entries['name'].get()
                enrollment = entries['enrollment'].get()
                success, message = self.db.delete_student(enrollment)
                if success:
                    # Log the activity
                    self._log_admin_activity(
                        "Student Deleted",
                        f"Deleted student: {student_name} (Enrollment: {enrollment})"
                    )
                    messagebox.showinfo("Success", "Student deleted successfully")
                    dialog.destroy()
                    self.refresh_students()
                    self.refresh_dashboard()
                else:
                    messagebox.showerror("Error", message)
        
        # Save button
        save_btn = tk.Button(
            btn_frame,
            text="💾 Save Changes",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=save_changes,
            cursor='hand2'
        )
        save_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Delete button
        delete_btn = tk.Button(
            btn_frame,
            text="🗑️ Delete",
            font=('Segoe UI', 12, 'bold'),
            bg='#dc3545',
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=delete_student,
            cursor='hand2'
        )
        delete_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Cancel button
        cancel_btn = tk.Button(
            btn_frame,
            text="❌ Cancel",
            font=('Segoe UI', 12, 'bold'),
            bg='#6c757d',
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=dialog.destroy,
            cursor='hand2'
        )
        cancel_btn.pack(side=tk.LEFT)

    def on_book_double_click(self, event):
        """Handle double-click on book — show Book Details & Copies dialog."""
        item = self.books_tree.selection()[0] if self.books_tree.selection() else None
        if not item:
            return
        book_data = self.books_tree.item(item, 'values')
        if not book_data:
            return
        book_id = book_data[0]
        try:
            books = self.db.get_books(book_id)
            book = None
            for b in books:
                if str(b[1]) == str(book_id):
                    book = b
                    break
            if not book:
                messagebox.showerror("Error", "Book not found")
                return
            self.show_book_copies_dialog(book)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load book details: {e}")

    def show_book_copies_dialog(self, book):
        """Show a dialog with all physical copy IDs (accession numbers) of a book,
        plus the option to edit the book.
        book tuple: (id, book_id, title, author, isbn, category, total_copies, available_copies, date_added, barcode, price)
        """
        # DB columns: 0=id 1=book_id 2=title 3=author 4=isbn 5=category
        #             6=total_copies 7=available_copies 8=date_added 9=barcode 10=price
        book_id_str    = str(book[1])
        title_str      = str(book[2])
        author_str     = str(book[3])
        total_copies   = int(book[6]) if book[6] else 0
        avail_copies   = int(book[7]) if book[7] else 0
        barcode_field  = str(book[9]) if len(book) > 9 and book[9] else ''
        price_val      = book[10] if len(book) > 10 and book[10] else 0

        # Parse individual accession numbers from the barcode CSV field
        # Filter out non-numeric tokens like "TO" that may be present from old imports
        def _is_acc_no(s):
            try: int(float(s)); return True
            except: return False
        if barcode_field and barcode_field.lower() not in ('none', 'nan', ''):
            acc_list = [x.strip() for x in barcode_field.split(',') if x.strip() and _is_acc_no(x.strip())]
        else:
            # Fallback: expand range in book_id if it looks like "8264-8266"
            acc_list = []
            if '-' in book_id_str:
                parts = book_id_str.split('-')
                if len(parts) == 2:
                    try:
                        start, end = int(parts[0]), int(parts[1])
                        acc_list = [str(n) for n in range(start, end + 1)]
                    except Exception:
                        acc_list = [book_id_str]
            if not acc_list:
                acc_list = [book_id_str] if book_id_str else []

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Book Details — {title_str[:60]}")
        dialog.geometry("560x520")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.resizable(False, False)
        cx = (self.root.winfo_screenwidth() // 2) - 280
        cy = (self.root.winfo_screenheight() // 2) - 260
        dialog.geometry(f"+{cx}+{cy}")

        # ── Header ──────────────────────────────────────────────────────
        header = tk.Frame(dialog, bg=self.colors['secondary'], height=65)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        tk.Label(header, text="📚 Book Details & Copy IDs",
                 font=('Segoe UI', 14, 'bold'), bg=self.colors['secondary'], fg='white').pack(pady=15)

        # ── Info cards ───────────────────────────────────────────────────
        info_frame = tk.Frame(dialog, bg='white', padx=20, pady=10)
        info_frame.pack(fill=tk.X)

        def info_row(parent, label, value, row):
            tk.Label(parent, text=label, font=('Segoe UI', 10, 'bold'),
                     bg='white', fg=self.colors['accent'], width=14, anchor='w').grid(
                     row=row, column=0, sticky='w', pady=3)
            tk.Label(parent, text=value, font=('Segoe UI', 10),
                     bg='white', fg=self.colors['text'], anchor='w', wraplength=360).grid(
                     row=row, column=1, sticky='w', pady=3, padx=(5, 0))

        info_row(info_frame, "Book ID / Range:", book_id_str, 0)
        info_row(info_frame, "Title:", title_str, 1)
        info_row(info_frame, "Author:", author_str, 2)
        info_row(info_frame, "Total Copies:", str(total_copies), 3)
        info_row(info_frame, "Available:", str(avail_copies), 4)
        issued = total_copies - avail_copies
        info_row(info_frame, "Issued:", str(issued), 5)
        price_str = f"₹{float(price_val):.2f}" if price_val else "—"
        info_row(info_frame, "Price:", price_str, 6)

        # ── Copy IDs list ────────────────────────────────────────────────
        copies_frame = tk.LabelFrame(dialog, text=f"📋 All Copy/Accession Numbers ({len(acc_list)} copies)",
                                     font=('Segoe UI', 10, 'bold'), bg='white',
                                     fg=self.colors['accent'], padx=10, pady=8)
        copies_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 10))

        # Grid of accession numbers (8 per row)
        grid_frame = tk.Frame(copies_frame, bg='white')
        grid_frame.pack(fill=tk.BOTH, expand=True)

        cols_per_row = 8
        for idx, acc_no in enumerate(acc_list):
            r, c = divmod(idx, cols_per_row)
            chip = tk.Label(
                grid_frame, text=acc_no,
                font=('Segoe UI', 9, 'bold'),
                bg=self.colors['secondary'], fg='white',
                relief='flat', padx=6, pady=3,
                cursor='arrow'
            )
            chip.grid(row=r, column=c, padx=3, pady=3, sticky='w')

        if not acc_list:
            tk.Label(grid_frame, text="No individual copy IDs stored.",
                     font=('Segoe UI', 10, 'italic'), bg='white',
                     fg='#888888').pack(pady=10)

        # ── Buttons ──────────────────────────────────────────────────────
        btn_frame = tk.Frame(dialog, bg='white', padx=20, pady=10)
        btn_frame.pack(fill=tk.X, side=tk.BOTTOM)

        def open_edit():
            dialog.destroy()
            self.show_edit_book_dialog(book)

        tk.Button(btn_frame, text="✏️ Edit Book", command=open_edit,
                  bg=self.colors['accent'], fg='white', font=('Segoe UI', 10, 'bold'),
                  relief='flat', padx=15, pady=8, cursor='hand2').pack(side=tk.LEFT)
        tk.Button(btn_frame, text="Close", command=dialog.destroy,
                  bg='#cccccc', fg='#333333', font=('Segoe UI', 10),
                  relief='flat', padx=15, pady=8, cursor='hand2').pack(side=tk.RIGHT)
    
    def show_edit_book_dialog(self, book):
        """Show dialog to edit book information"""
        # book tuple: (id, book_id, title, author, isbn, category, total_copies, available_copies, date_added)
        dialog = tk.Toplevel(self.root)
        dialog.title("Edit Book")
        dialog.geometry("500x600")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 100, self.root.winfo_rooty() + 100))
        
        # Title
        title_label = tk.Label(
            dialog,
            text="✏️ Edit Book",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack(pady=(20, 30))
        
        # Form frame
        form_frame = tk.Frame(dialog, bg='white')
        form_frame.pack(expand=True, padx=40)
        
        # Get book categories from your predefined list
        book_categories = [
            "Core CS", "Web Development", "Programming", "Database",
            "AI/ML", "Networking", "Operating Systems", "Algorithms",
            "Software Engineering", "Mobile Development", "Cloud Computing",
            "Cybersecurity", "IoT", "Competitive Programming", "Project Guides", "Others"
        ]
        
        # Form fields — use named key access (not integer index) to guard against schema column order changes
        def _bkey(key, default=''):
            try:
                return book[key] if book[key] is not None else default
            except (IndexError, KeyError):
                return default
        fields = [
            ("Book ID:", "book_id", _bkey('book_id'), False),  # Read-only
            ("Title:", "title", _bkey('title'), True),
            ("Author:", "author", _bkey('author'), True),
            ("ISBN:", "isbn", _bkey('isbn'), True),
            ("Category:", "category", _bkey('category') or 'Others', True),
            ("Total Copies:", "total_copies", str(_bkey('total_copies', 1)), True),
            ("Available Copies:", "available_copies", str(_bkey('available_copies', 0)), False),  # Read-only
            ("Barcode:", "barcode", _bkey('barcode'), True)
        ]
        
        entries = {}
        
        for i, (label_text, field_name, default_value, editable) in enumerate(fields):
            # Label
            label = tk.Label(
                form_frame,
                text=label_text,
                font=('Segoe UI', 11, 'bold'),
                bg='white',
                fg=self.colors['accent']
            )
            label.grid(row=i, column=0, sticky='w', pady=(0, 15), padx=(0, 20))
            
            # Entry or Combobox
            if field_name == "category" and editable:
                entry = ttk.Combobox(
                    form_frame,
                    values=book_categories,
                    font=('Segoe UI', 11),
                    width=25,
                    state="readonly"
                )
                entry.set(default_value)
                entry.grid(row=i, column=1, pady=(0, 15))
            else:
                entry = tk.Entry(
                    form_frame,
                    font=('Segoe UI', 11),
                    width=28,
                    relief='solid',
                    bd=2
                )
                entry.insert(0, default_value)
                if not editable:
                    entry.config(state='readonly', bg='#f0f0f0')
                entry.grid(row=i, column=1, pady=(0, 15))
            
            entries[field_name] = entry
        
        # Info label about available copies
        info_label = tk.Label(
            form_frame,
            text="Note: Available copies will be adjusted automatically based on borrowed books.",
            font=('Segoe UI', 9, 'italic'),
            bg='white',
            fg='#666',
            wraplength=350
        )
        info_label.grid(row=len(fields), column=0, columnspan=2, pady=(5, 0))
        
        # Buttons frame
        btn_frame = tk.Frame(dialog, bg='white')
        btn_frame.pack(pady=20)
        
        def save_changes():
            # Validate required fields
            if not all([entries['title'].get().strip(), entries['author'].get().strip()]):
                messagebox.showerror("Error", "Title and Author are required!")
                return
            
            # Validate total copies is a positive integer
            try:
                total_copies = int(entries['total_copies'].get())
                if total_copies < 1:
                    messagebox.showerror("Error", "Total copies must be at least 1!")
                    return
            except ValueError:
                messagebox.showerror("Error", "Total copies must be a valid number!")
                return
            
            # Update book
            success, message = self.db.update_book(
                entries['book_id'].get(),
                entries['title'].get().strip(),
                entries['author'].get().strip(),
                entries['isbn'].get().strip(),
                entries['category'].get(),
                total_copies,
                entries['barcode'].get().strip() if entries.get('barcode') else ''
            )
            
            if success:
                messagebox.showinfo("Success", message)
                dialog.destroy()
                self.refresh_books()
                self.refresh_dashboard()
            else:
                messagebox.showerror("Error", message)
        
        def delete_book():
            if messagebox.askyesno(
                "Confirm Delete",
                f"Delete book: {entries['title'].get()}?\n\nThis action cannot be undone!",
                icon='warning'
            ):
                book_title = entries['title'].get()
                book_id = entries['book_id'].get()
                success, message = self.db.delete_book(book_id)
                if success:
                    # Log the activity
                    self._log_admin_activity(
                        "Book Deleted",
                        f"Deleted book: {book_title} (ID: {book_id})"
                    )
                    messagebox.showinfo("Success", "Book deleted successfully")
                    dialog.destroy()
                    self.refresh_books()
                    self.refresh_dashboard()
                else:
                    messagebox.showerror("Error", message)
        
        # Save button
        save_btn = tk.Button(
            btn_frame,
            text="💾 Save Changes",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=save_changes,
            cursor='hand2'
        )
        save_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Delete button
        delete_btn = tk.Button(
            btn_frame,
            text="🗑️ Delete",
            font=('Segoe UI', 12, 'bold'),
            bg='#dc3545',
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=delete_book,
            cursor='hand2'
        )
        delete_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Cancel button
        cancel_btn = tk.Button(
            btn_frame,
            text="❌ Cancel",
            font=('Segoe UI', 12, 'bold'),
            bg='#6c757d',
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=dialog.destroy,
            cursor='hand2'
        )
        cancel_btn.pack(side=tk.LEFT)
        
    # ...existing code...
        # ...existing code...
        # ...existing code...
    
    def create_records_tab(self):
        """Create comprehensive records tab"""
        records_frame = tk.Frame(self.notebook, bg=self.colors['primary'])
        self.notebook.add(records_frame, text="📊 Records")
        
        # Top frame for search, filters and actions
        top_frame = tk.Frame(records_frame, bg=self.colors['primary'])
        top_frame.pack(fill=tk.X, padx=20, pady=20)
        
        # Search and filter frame
        search_filter_frame = tk.LabelFrame(
            top_frame,
            text="🔍 Search & Filter Records",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        search_filter_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        # Search controls
        search_controls = tk.Frame(search_filter_frame, bg=self.colors['primary'])
        search_controls.pack(fill=tk.X, padx=10, pady=10)
        
        # Row 1: Search and type filter
        row1 = tk.Frame(search_controls, bg=self.colors['primary'])
        row1.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(row1, text="Search:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        record_search_entry = tk.Entry(row1, textvariable=self.record_search_var, font=('Segoe UI', 10), width=25)
        record_search_entry.pack(side=tk.LEFT, padx=(5, 15))
        
        tk.Label(row1, text="Type:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        type_combo = ttk.Combobox(row1, textvariable=self.record_type_filter, 
                                  values=["All", "Issued", "Returned", "Overdue"], state="readonly", width=12)
        type_combo.pack(side=tk.LEFT, padx=5)
        type_combo.bind('<<ComboboxSelected>>', lambda e: self.search_records())
        
        # Branch filter
        tk.Label(row1, text="Branch:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT, padx=(15, 0))
        branch_values = ["All", "Computer Engineering", "Information Technology", "Civil Engineering",
                         "Electronics & Telecommunication", "Mechanical Engineering",
                         "Automobile Engineering", "Electrical Engineering"]
        record_branch_combo = ttk.Combobox(row1, textvariable=self.record_branch_filter,
                                           values=branch_values, state="readonly", width=20)
        record_branch_combo.pack(side=tk.LEFT, padx=5)
        record_branch_combo.bind('<<ComboboxSelected>>', lambda e: self.search_records())
        
        # Row 2: Date filters and Academic Year
        row2 = tk.Frame(search_controls, bg=self.colors['primary'])
        row2.pack(fill=tk.X)
        
        tk.Label(row2, text="From Date:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        if DateEntry:
            self.record_from_date = DateEntry(row2, width=11, date_pattern='yyyy-mm-dd')
            self.record_from_date.pack(side=tk.LEFT, padx=(5, 15))
        else:
            self.record_from_date = tk.Entry(row2, font=('Segoe UI', 10), width=12)
            self.record_from_date.pack(side=tk.LEFT, padx=(5, 15))
        
        tk.Label(row2, text="To Date:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        if DateEntry:
            self.record_to_date = DateEntry(row2, width=11, date_pattern='yyyy-mm-dd')
            self.record_to_date.pack(side=tk.LEFT, padx=(5, 15))
        else:
            self.record_to_date = tk.Entry(row2, font=('Segoe UI', 10), width=12)
            self.record_to_date.pack(side=tk.LEFT, padx=(5, 15))
        
        # Academic Year filter
        tk.Label(row2, text="Academic Year:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10)).pack(side=tk.LEFT)
        self.record_academic_year_var = tk.StringVar(value="All")
        
        # Get academic years from database and convert format
        academic_years_raw = self.db.get_all_academic_years()
        academic_years = ["All"]
        for year in academic_years_raw:
            # Convert format from "2025-2026" to "25-26"
            if "-" in year:
                years = year.split("-")
                if len(years) == 2:
                    year1 = years[0][-2:]  # "2025" -> "25"
                    year2 = years[1][-2:]  # "2026" -> "26"
                    academic_years.append(f"{year1}-{year2}")
                else:
                    academic_years.append(year)
            else:
                academic_years.append(year)
        
        self.academic_year_combo = ttk.Combobox(row2, textvariable=self.record_academic_year_var, 
                                          values=academic_years, state="readonly", width=15)
        self.academic_year_combo.pack(side=tk.LEFT, padx=(5, 15))
        self.academic_year_combo.bind('<<ComboboxSelected>>', lambda e: self.search_records())
        
        filter_btn = tk.Button(
            row2,
            text="🔍 Filter",
            font=('Segoe UI', 9, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=10,
            pady=5,
            command=self.search_records,
            cursor='hand2'
        )
        filter_btn.pack(side=tk.LEFT, padx=5)
        
        clear_filter_btn = tk.Button(
            row2,
            text="🗑️ Clear",
            font=('Segoe UI', 9, 'bold'),
            bg='#6c757d',
            fg='white',
            relief='flat',
            padx=10,
            pady=5,
            command=self.clear_record_filters,
            cursor='hand2'
        )
        clear_filter_btn.pack(side=tk.LEFT, padx=5)
        
        # Row 3: Quick date filters
        row3 = tk.Frame(search_controls, bg=self.colors['primary'])
        row3.pack(fill=tk.X, pady=(10, 0))
        
        tk.Label(row3, text="Quick Filter:", bg=self.colors['primary'], fg=self.colors['accent'], font=('Segoe UI', 10, 'bold')).pack(side=tk.LEFT, padx=(0, 10))
        
        def filter_last_days(days):
            """Filter records for last N days"""
            from datetime import datetime, timedelta
            today = datetime.now()
            from_date = (today - timedelta(days=days)).strftime('%Y-%m-%d')
            to_date = today.strftime('%Y-%m-%d')
            
            # Set date fields
            if DateEntry:
                self.record_from_date.set_date(today - timedelta(days=days))
                self.record_to_date.set_date(today)
            else:
                self.record_from_date.delete(0, tk.END)
                self.record_from_date.insert(0, from_date)
                self.record_to_date.delete(0, tk.END)
                self.record_to_date.insert(0, to_date)
            
            # Apply filter
            self.search_records()
        
        last_7_btn = tk.Button(
            row3,
            text="📅 Last 7 Days",
            font=('Segoe UI', 9),
            bg='#17a2b8',
            fg='white',
            relief='flat',
            padx=10,
            pady=5,
            command=lambda: filter_last_days(7),
            cursor='hand2'
        )
        last_7_btn.pack(side=tk.LEFT, padx=2)
        
        last_15_btn = tk.Button(
            row3,
            text="📅 Last 15 Days",
            font=('Segoe UI', 9),
            bg='#17a2b8',
            fg='white',
            relief='flat',
            padx=10,
            pady=5,
            command=lambda: filter_last_days(15),
            cursor='hand2'
        )
        last_15_btn.pack(side=tk.LEFT, padx=2)
        
        last_30_btn = tk.Button(
            row3,
            text="📅 Last 30 Days",
            font=('Segoe UI', 9),
            bg='#17a2b8',
            fg='white',
            relief='flat',
            padx=10,
            pady=5,
            command=lambda: filter_last_days(30),
            cursor='hand2'
        )
        last_30_btn.pack(side=tk.LEFT, padx=2)
        
        # Actions frame
        actions_frame = tk.LabelFrame(
            top_frame,
            text="⚡ Actions",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        actions_frame.pack(side=tk.RIGHT)
        
        # Action buttons
        buttons_frame = tk.Frame(actions_frame, bg=self.colors['primary'])
        buttons_frame.pack(padx=10, pady=10)
        
        export_records_btn = tk.Button(
            buttons_frame,
            text="📊 Export Records",
            font=('Segoe UI', 10, 'bold'),
            bg='#28a745',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.export_records_to_excel,
            cursor='hand2'
        )
        export_records_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        share_records_btn = tk.Button(
            buttons_frame,
            text="📤 Share Data",
            font=('Segoe UI', 10, 'bold'),
            bg='#6f42c1',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.share_data_dialog,
            cursor='hand2'
        )
        share_records_btn.pack(side=tk.LEFT, padx=5)

        # New: Overdue notice letter generation button
        overdue_letter_btn = tk.Button(
            buttons_frame,
            text="📄 Overdue Letter (Word)",
            font=('Segoe UI', 10, 'bold'),
            bg='#d35400',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.export_overdue_notice_letter_word,
            cursor='hand2'
        )
        overdue_letter_btn.pack(side=tk.LEFT, padx=5)
        # Optional Excel version button retained for users wanting spreadsheet format
        # Excel overdue letter button removed per user request (Word only)
        
        # Tip label for double-click feature
        tip_frame = tk.Frame(actions_frame, bg=self.colors['primary'])
        tip_frame.pack(padx=10, pady=(0, 10))
        
        tip_label = tk.Label(
            tip_frame,
            text="💡 Tip: Double-click on overdue student to send letter",
            font=('Segoe UI', 9, 'italic'),
            bg='#fff3cd',
            fg='#856404',
            relief='solid',
            bd=1,
            padx=10,
            pady=5
        )
        tip_label.pack()
        
        # Records list
        records_list_frame = tk.LabelFrame(
            records_frame,
            text="📋 Complete Transaction Records",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent'],
            padx=10,
            pady=10
        )
        records_list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Records treeview
        record_columns = ('Enrollment No', 'Student Name', 'Branch', 'Book ID', 'Book Title', 'Issue Date', 'Due Date', 'Return Date', 'Status', 'Fine')
        self.records_tree = ttk.Treeview(records_list_frame, columns=record_columns, show='headings', height=15)
        record_widths = {'Enrollment No': 120, 'Student Name': 150, 'Branch': 120, 'Book ID': 100, 'Book Title': 200, 'Issue Date': 100, 'Due Date': 100, 'Return Date': 100, 'Status': 80, 'Fine': 80}
        for col in record_columns:
            self.records_tree.heading(col, text=col)
            self.records_tree.column(col, width=record_widths[col])
        # Scrollbars for records
        records_v_scrollbar = ttk.Scrollbar(records_list_frame, orient=tk.VERTICAL, command=self.records_tree.yview)
        records_h_scrollbar = ttk.Scrollbar(records_list_frame, orient=tk.HORIZONTAL, command=self.records_tree.xview)
        self.records_tree.configure(yscrollcommand=records_v_scrollbar.set, xscrollcommand=records_h_scrollbar.set)
        # Pack records treeview and scrollbars
        self.records_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        records_v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        records_h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

        # Free scrolling in Records: vertical and Shift+Wheel horizontal, pointer-scoped
        def _rec_units(delta):
            return -1 if delta > 0 else (1 if delta < 0 else 0)
        def _records_vwheel(event):
            u = _rec_units(event.delta)
            if u:
                self.records_tree.yview_scroll(u, 'units')
            return 'break'
        def _records_hwheel(event):
            u = _rec_units(event.delta)
            if u:
                self.records_tree.xview_scroll(u, 'units')
            return 'break'
        self.records_tree.bind('<MouseWheel>', _records_vwheel)
        self.records_tree.bind('<Shift-MouseWheel>', _records_hwheel)
        
        # Bind double-click event for sending overdue letter
        self.records_tree.bind('<Double-1>', self.on_record_double_click)

        # Initial population of records (avoid empty first view)
        try:
            self.clear_record_filters()  # resets inputs and calls search_records
        except Exception:
            # Fallback: directly populate with all records
            try:
                self.populate_records_tree(self.get_all_records())
            except Exception:
                pass
    
    def show_date_picker(self, entry_widget):
        """Show a simple date picker dialog"""
        from tkinter import simpledialog
        
        # Create a simple date input dialog
        date_str = simpledialog.askstring(
            "Select Date",
            "Enter date (YYYY-MM-DD):",
            initialvalue=entry_widget.get() or datetime.now().strftime('%Y-%m-%d')
        )
        
        if date_str:
            try:
                # Validate date format
                datetime.strptime(date_str, '%Y-%m-%d')
                entry_widget.delete(0, tk.END)
                entry_widget.insert(0, date_str)
            except ValueError:
                messagebox.showerror("Error", "Invalid date format! Please use YYYY-MM-DD")
    
    def show_add_student_dialog(self):
        """Show add student dialog with branch selection"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Add New Student")
        dialog.geometry("500x450")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 100, self.root.winfo_rooty() + 100))
        
        # Title
        title_label = tk.Label(
            dialog,
            text="👥 Add New Student",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack(pady=(20, 30))
        
        # Form frame
        form_frame = tk.Frame(dialog, bg='white')
        form_frame.pack(expand=True, padx=40)
        
        # Form fields
        fields = [
            ("Enrollment No:", "enrollment"),
            ("Full Name:", "name"),
            ("Email:", "email"),
            ("Phone:", "phone"),
            ("Year:", "year")
        ]
        
        entries = {}
        
        for i, (label_text, field_name) in enumerate(fields):
            # Label
            label = tk.Label(
                form_frame,
                text=label_text,
                font=('Segoe UI', 11, 'bold'),
                bg='white',
                fg=self.colors['accent']
            )
            label.grid(row=i, column=0, sticky='w', pady=(0, 15), padx=(0, 20))
            
            # Entry or Combobox
            if field_name == "year":
                entry = ttk.Combobox(
                    form_frame,
                    values=["1st", "2nd", "3rd"],
                    font=('Segoe UI', 11),
                    width=25,
                    state="readonly"
                )
                entry.set("1st")  # Default to 1st year
            else:
                entry = tk.Entry(
                    form_frame,
                    font=('Segoe UI', 11),
                    width=30,
                    relief='solid',
                    bd=2
                )
            
            entry.grid(row=i, column=1, pady=(0, 15))
            entries[field_name] = entry
        
        # Branch selection
        branch_label = tk.Label(
            form_frame,
            text="Branch:",
            font=('Segoe UI', 11, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        branch_label.grid(row=len(fields), column=0, sticky='w', pady=(0, 15), padx=(0, 20))

        branch_entry = ttk.Combobox(
            form_frame,
            values=["Computer Engineering", "Information Technology", "Civil Engineering", "Electronics & Telecommunication", "Mechanical Engineering", "Automobile Engineering", "Electrical Engineering"],
            font=('Segoe UI', 11),
            width=30,
            state="readonly"
        )
        branch_entry.set("Computer Engineering")  # Default
        branch_entry.grid(row=len(fields), column=1, pady=(0, 15))
        entries['branch'] = branch_entry
        
        # Buttons
        btn_frame = tk.Frame(dialog, bg='white')
        btn_frame.pack(pady=20)
        
        def save_student():
            # Validate fields
            if not all([entries['enrollment'].get(), entries['name'].get()]):
                messagebox.showerror("Error", "Enrollment No and Name are required!")
                return
            
            # Add student with selected branch
            success, message = self.db.add_student(
                entries['enrollment'].get(),
                entries['name'].get(),
                entries['email'].get(),
                entries['phone'].get(),
                entries['branch'].get(),  # Selected branch
                entries['year'].get()
            )
            
            if success:
                # Log the activity
                self._log_admin_activity(
                    "Student Added",
                    f"Added student: {entries['name'].get()} (Enrollment: {entries['enrollment'].get()}, Year: {entries['year'].get()})"
                )
                messagebox.showinfo("Success", message)
                dialog.destroy()
                self.refresh_students()
                # Refresh dashboard so total student count updates immediately
                try:
                    self.refresh_dashboard()
                except Exception:
                    pass
            else:
                messagebox.showerror("Error", message)
        
        save_btn = tk.Button(
            btn_frame,
            text="💾 Save Student",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=save_student,
            cursor='hand2'
        )
        save_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        cancel_btn = tk.Button(
            btn_frame,
            text="❌ Cancel",
            font=('Segoe UI', 12, 'bold'),
            bg='#6c757d',
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=dialog.destroy,
            cursor='hand2'
        )
        cancel_btn.pack(side=tk.LEFT)
        
        # Focus on first field
        entries['enrollment'].focus()
    
    def show_add_book_dialog(self):
        """Show add book dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Add New Book")
        dialog.geometry("500x620")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 100, self.root.winfo_rooty() + 100))
        
        # Title
        title_label = tk.Label(
            dialog,
            text="📚 Add New Book",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack(pady=(20, 30))
        
        # Form frame
        form_frame = tk.Frame(dialog, bg='white')
        form_frame.pack(expand=True, padx=40)
        
        # Form fields
        fields = [
            ("Book ID (Range e.g. 1001-1010):", "book_id"),
            ("Title:", "title"),
            ("Author:", "author"),
            ("ISBN:", "isbn"),
            ("Category:", "category"),
            ("Total Copies:", "copies"),
            ("Price (₹):", "price"),
            ("Barcode / Copy IDs (auto from Book ID range):", "barcode")
        ]
        
        entries = {}
        
        # Auto-generate checkbox variable
        auto_gen_var = tk.BooleanVar(value=False)
        
        for i, (label_text, field_name) in enumerate(fields):
            # Label
            label = tk.Label(
                form_frame,
                text=label_text,
                font=('Segoe UI', 11, 'bold'),
                bg='white',
                fg=self.colors['accent']
            )
            label.grid(row=i, column=0, sticky='w', pady=(0, 15), padx=(0, 20))
            
            # Entry or Combobox
            if field_name == "category":
                entry = ttk.Combobox(
                    form_frame,
                    values=[
                        "Core CS", "Web Development", "Programming", "Database",
                        "AI/ML", "Networking", "Operating Systems", "Algorithms",
                        "Software Engineering", "Mobile Development", "Cloud Computing",
                        "Cybersecurity", "IoT", "Competitive Programming", "Project Guides", "Others"
                    ],
                    font=('Segoe UI', 11),
                    width=25,
                    state="readonly"
                )
                entry.set("Core CS")  # Default category
                entry.grid(row=i, column=1, pady=(0, 15))
            elif field_name == "book_id":
                # Book ID entry with auto-generate option
                id_frame = tk.Frame(form_frame, bg='white')
                book_id_entry = tk.Entry(
                    id_frame,
                    font=('Segoe UI', 11),
                    width=30,
                    relief='solid',
                    bd=2
                )
                book_id_entry.pack(side=tk.TOP, fill=tk.X)
                
                def toggle_auto_gen():
                    if auto_gen_var.get():
                        next_id = self.db.get_next_book_id()
                        book_id_entry.delete(0, tk.END)
                        book_id_entry.insert(0, next_id)
                        book_id_entry.config(state='readonly')
                    else:
                        book_id_entry.config(state='normal')
                        book_id_entry.delete(0, tk.END)
                
                auto_check = tk.Checkbutton(
                    id_frame,
                    text="Auto-generate Book ID",
                    variable=auto_gen_var,
                    command=toggle_auto_gen,
                    font=('Segoe UI', 9),
                    bg='white',
                    fg=self.colors['text']
                )
                auto_check.pack(side=tk.TOP, anchor='w', pady=(3, 0))
                
                id_frame.grid(row=i, column=1, pady=(0, 15))
                entry = book_id_entry  # Store the actual entry widget
            else:
                entry = tk.Entry(
                    form_frame,
                    font=('Segoe UI', 11),
                    width=30,
                    relief='solid',
                    bd=2
                )
                if field_name == "copies":
                    entry.insert(0, "1")  # Default to 1 copy
                    entry.config(validate='key', validatecommand=(dialog.register(lambda P: P.isdigit() or P == ""), '%P'))
                elif field_name == "price":
                    entry.insert(0, "0")  # Default price 0
                entry.grid(row=i, column=1, pady=(0, 15))
            
            entries[field_name] = entry
        
        # Buttons
        btn_frame = tk.Frame(dialog, bg='white')
        btn_frame.pack(pady=20)
        
        def save_book():
            # Validate required fields (only book_id and title are required)
            required_fields = ['book_id', 'title']
            for field in required_fields:
                if not entries[field].get().strip():
                    messagebox.showerror("Error", f"{field.replace('_', ' ').title()} is required!")
                    return
            # Validate copies
            try:
                copies = int(entries['copies'].get())
                if copies <= 0:
                    raise ValueError()
            except ValueError:
                messagebox.showerror("Error", "Total copies must be a positive number!")
                return
            # Validate price
            try:
                price = float(entries['price'].get() or 0)
                if price < 0:
                    raise ValueError()
            except ValueError:
                messagebox.showerror("Error", "Price must be a valid non-negative number!")
                return

            # Book ID == Barcode (Copy IDs) in your library:
            # Accept Book ID as a sequential range "1001-1010" and auto-calc copies.
            import re

            def _parse_range_or_csv(text):
                """Return (list_of_accessions, canonical_range_str_or_none).
                Accepts: '1001-1010' or '1001,1002,1003'."""
                raw = (text or '').strip()
                if not raw:
                    return [], None
                m = re.match(r'^\s*(\d+)\s*-\s*(\d+)\s*$', raw)
                if m:
                    start = int(m.group(1))
                    end = int(m.group(2))
                    if end < start:
                        raise ValueError("Range must be sequential (start ≤ end).")
                    return [str(n) for n in range(start, end + 1)], f"{start}-{end}"
                # CSV of numbers
                parts = [p.strip() for p in raw.split(',') if p.strip()]
                out = []
                for p in parts:
                    try:
                        out.append(str(int(float(p))))
                    except Exception:
                        # Non-numeric: don't treat as accession IDs
                        return [], None
                # Dedupe keep order
                seen = set()
                out = [x for x in out if not (x in seen or seen.add(x))]
                return out, None

            book_id_raw = entries['book_id'].get().strip()
            barcode_raw = (entries.get('barcode').get() if entries.get('barcode') else '').strip()

            accession_list = []
            barcode_csv = barcode_raw
            book_id_val = book_id_raw

            try:
                accession_list, canonical_range = _parse_range_or_csv(book_id_raw)
            except ValueError as ve:
                messagebox.showerror("Error", f"Invalid Book ID range: {ve}")
                return

            # If user entered a single numeric ID and set copies>1, auto-expand to a range.
            if not accession_list:
                try:
                    if book_id_raw.isdigit() and copies > 1:
                        start = int(book_id_raw)
                        end = start + copies - 1
                        accession_list = [str(n) for n in range(start, end + 1)]
                        canonical_range = f"{start}-{end}"
                except Exception:
                    pass

            # Backward compatibility: if Book ID isn't a range/CSV, allow Barcode field range/CSV.
            if not accession_list and barcode_raw:
                try:
                    accession_list, _ = _parse_range_or_csv(barcode_raw)
                except ValueError:
                    accession_list = []

            if accession_list:
                # Canonical stored book_id should be a range string for multi-copy
                if len(accession_list) >= 2:
                    try:
                        book_id_val = f"{int(accession_list[0])}-{int(accession_list[-1])}"
                    except Exception:
                        book_id_val = canonical_range or book_id_raw
                else:
                    book_id_val = accession_list[0]

                barcode_csv = ','.join(accession_list)
                if copies != len(accession_list):
                    copies = len(accession_list)
                    try:
                        entries['copies'].delete(0, tk.END)
                        entries['copies'].insert(0, str(copies))
                    except Exception:
                        pass

                # If user typed a single numeric ID and copies>1, update the UI Book ID to the range.
                try:
                    if entries['book_id'].cget('state') != 'readonly':
                        entries['book_id'].delete(0, tk.END)
                        entries['book_id'].insert(0, book_id_val)
                except Exception:
                    pass

                # Also auto-fill the Barcode entry with the range (short) for user visibility, but store CSV in DB.
                try:
                    if entries.get('barcode') and entries['barcode'].get().strip() == '':
                        # show short range, not huge CSV
                        entries['barcode'].insert(0, book_id_val)
                except Exception:
                    pass
            # Check Book ID uniqueness (exact)
            if book_id_val:
                conn = self.db.get_connection()
                cur = conn.cursor()
                cur.execute('SELECT COUNT(*) FROM books WHERE book_id = ?', (book_id_val,))
                exists = cur.fetchone()[0]
                conn.close()
                if exists:
                    messagebox.showerror("Error", f"Book ID '{book_id_val}' already exists! Please use a unique Book ID.")
                    return

            # Enforce: no Book ID (accession) duplicates/overlaps anywhere in the library
            if accession_list:
                conn = self.db.get_connection()
                cur = conn.cursor()
                cur.execute("SELECT book_id, barcode FROM books")
                used_ids = set()
                used_ranges = []  # list of (start,end)

                for bid, bc in cur.fetchall() or []:
                    bid_s = str(bid or '').strip()
                    # From barcode CSV
                    for t in [x.strip() for x in str(bc or '').split(',') if x.strip()]:
                        try:
                            used_ids.add(str(int(float(t))))
                        except Exception:
                            continue
                    # From numeric book_id (older rows)
                    if bid_s.isdigit():
                        used_ids.add(bid_s)
                    # From book_id ranges like "1001-1010"
                    m = re.match(r'^(\d+)\s*-\s*(\d+)$', bid_s)
                    if m:
                        try:
                            s0 = int(m.group(1)); e0 = int(m.group(2))
                            if e0 >= s0:
                                used_ranges.append((s0, e0))
                        except Exception:
                            pass

                conn.close()

                # Check overlaps with ranges efficiently
                try:
                    new_start = int(accession_list[0])
                    new_end = int(accession_list[-1])
                except Exception:
                    new_start = None
                    new_end = None

                if new_start is not None and new_end is not None:
                    for (s0, e0) in used_ranges:
                        if not (new_end < s0 or new_start > e0):
                            messagebox.showerror(
                                "Error",
                                f"Book ID range overlaps with an existing range ({s0}-{e0}).\n\n"
                                f"Please choose a non-overlapping range."
                            )
                            return

                collision = [a for a in accession_list if a in used_ids]
                if collision:
                    messagebox.showerror(
                        "Error",
                        "Some Book IDs are already used by existing books:\n" + ", ".join(collision[:25])
                    )
                    return
            # Check for duplicate title
            title_val = entries['title'].get().strip()
            conn = self.db.get_connection()
            cur = conn.cursor()
            cur.execute("SELECT book_id, total_copies FROM books WHERE LOWER(title) = LOWER(?)", (title_val,))
            existing_title = cur.fetchone()
            conn.close()
            if existing_title:
                existing_id, existing_copies = existing_title
                choice = messagebox.askyesnocancel(
                    "⚠️ Duplicate Title Found",
                    f"A book with the same title already exists:\n"
                    f"  Book ID: {existing_id}\n"
                    f"  Current Copies: {existing_copies}\n\n"
                    f"YES  → Add {copies} more copies to existing book\n"
                    f"NO   → Add as a separate/new entry\n"
                    f"CANCEL → Go back",
                    icon='warning'
                )
                if choice is None:  # CANCEL
                    return
                elif choice:  # YES - merge copies
                    conn = self.db.get_connection()
                    cur = conn.cursor()
                    # Enforce global uniqueness of accession numbers when provided
                    if accession_list:
                        # Gather used accessions/ranges from ALL books
                        cur.execute("SELECT book_id, barcode FROM books")
                        used_elsewhere = set()
                        existing_acc = set()
                        used_ranges_elsewhere = []
                        for bid, bc in cur.fetchall() or []:
                            bid_s = str(bid or '').strip()
                            tokens = [t.strip() for t in str(bc or '').split(',') if t.strip()]
                            for t in tokens:
                                try:
                                    a = str(int(float(t)))
                                except Exception:
                                    continue
                                if str(bid) == str(existing_id):
                                    existing_acc.add(a)
                                else:
                                    used_elsewhere.add(a)

                            # also parse book_id numeric/range as reserved
                            if bid_s.isdigit():
                                (existing_acc if str(bid) == str(existing_id) else used_elsewhere).add(bid_s)
                            m = re.match(r'^(\d+)\s*-\s*(\d+)$', bid_s)
                            if m and str(bid) != str(existing_id):
                                try:
                                    s0 = int(m.group(1)); e0 = int(m.group(2))
                                    if e0 >= s0:
                                        used_ranges_elsewhere.append((s0, e0))
                                except Exception:
                                    pass

                        # Range overlap check against other books that store ranges
                        try:
                            new_start = int(accession_list[0]); new_end = int(accession_list[-1])
                        except Exception:
                            new_start = None; new_end = None
                        if new_start is not None and new_end is not None:
                            for (s0, e0) in used_ranges_elsewhere:
                                if not (new_end < s0 or new_start > e0):
                                    conn.close()
                                    messagebox.showerror(
                                        "Error",
                                        f"Book ID range overlaps with another book's range ({s0}-{e0})."
                                    )
                                    return

                        new_unique = [a for a in accession_list if a not in existing_acc]
                        collision = [a for a in new_unique if a in used_elsewhere]
                        if collision:
                            conn.close()
                            messagebox.showerror(
                                "Error",
                                "Some accession numbers are already used by other books:\n" + ", ".join(collision[:25])
                            )
                            return

                        merged_acc = sorted(existing_acc.union(new_unique), key=lambda x: int(x))
                        add_n = len(new_unique)
                        cur.execute(
                            "UPDATE books SET total_copies = total_copies + ?, available_copies = available_copies + ?, barcode = ?, updated_at=CURRENT_TIMESTAMP WHERE book_id = ?",
                            (add_n, add_n, ','.join(merged_acc), existing_id)
                        )
                        copies_added = add_n
                    else:
                        cur.execute(
                            "UPDATE books SET total_copies = total_copies + ?, available_copies = available_copies + ?, updated_at=CURRENT_TIMESTAMP WHERE book_id = ?",
                            (copies, copies, existing_id)
                        )
                        copies_added = copies
                    conn.commit()
                    conn.close()
                    self._log_admin_activity(
                        "Book Copies Added",
                        f"Added {copies_added} copies to existing book '{title_val}' (ID: {existing_id}). New total: {existing_copies + copies_added}"
                    )
                    messagebox.showinfo("Success", f"Added {copies_added} copies to '{title_val}'.\nNew total: {existing_copies + copies_added} copies.")
                    dialog.destroy()
                    self.refresh_books()
                    self.refresh_dashboard()
                    return
                # else NO - continue adding as separate entry

            # (Uniqueness already validated above for new insert too)
            # Add book
            success, message = self.db.add_book(
                book_id_val,
                title_val,
                entries['author'].get().strip() if entries['author'].get().strip() else '',
                entries['isbn'].get().strip(),
                entries['category'].get(),
                copies,
                barcode_csv,
                price
            )
            if success:
                # Log the activity
                self._log_admin_activity(
                    "Book Added",
                    f"Added book: {entries['title'].get().strip()} (ID: {book_id_val}, Category: {entries['category'].get()}, Copies: {copies})"
                )
                messagebox.showinfo("Success", message)
                dialog.destroy()
                # Refresh books view
                self.refresh_books()
                # Update dashboard statistics
                self.refresh_dashboard()
            else:
                messagebox.showerror("Error", message)
        
        save_btn = tk.Button(
            btn_frame,
            text="💾 Save Book",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=save_book,
            cursor='hand2'
        )
        save_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        cancel_btn = tk.Button(
            btn_frame,
            text="❌ Cancel",
            font=('Segoe UI', 12, 'bold'),
            bg='#6c757d',
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=dialog.destroy,
            cursor='hand2'
        )
        cancel_btn.pack(side=tk.LEFT)
        
        # Focus on first field
        entries['book_id'].focus()
    
    def borrow_book(self):
        """Handle book borrowing (Async)"""
        enrollment_text = self.borrow_enrollment_entry.get().strip()
        book_text = self.borrow_book_id_entry.get().strip()
        borrow_date = self.borrow_borrow_date_entry.get().strip()
        due_date = self.borrow_due_date_entry.get().strip()
        
        if not all([enrollment_text, book_text, borrow_date, due_date]):
            messagebox.showerror("Error", "All fields are required!")
            return
        
        # Extract enrollment number (before " - " if formatted, otherwise use as is)
        enrollment_no = enrollment_text.split(' - ')[0] if ' - ' in enrollment_text else enrollment_text
        
        # Extract book ID (before " - " if formatted, otherwise use as is)
        book_id = book_text.split(' - ')[0] if ' - ' in book_text else book_text

        # Validate date format (borrow & due)
        try:
            from datetime import datetime, timedelta
            datetime.strptime(borrow_date, '%Y-%m-%d')
            datetime.strptime(due_date, '%Y-%m-%d')
        except ValueError:
            messagebox.showerror("Error", "Invalid date format! Please use YYYY-MM-DD")
            return
        # Ensure due >= borrow
        from datetime import datetime as _dt
        if _dt.strptime(due_date, '%Y-%m-%d') < _dt.strptime(borrow_date, '%Y-%m-%d'):
            messagebox.showerror("Error", "Due date cannot be before issue date")
            return

        # UI-side validation for allowed range: 1..loan_period days from borrow
        try:
            bd_obj = datetime.strptime(borrow_date, '%Y-%m-%d')
            dd_obj = datetime.strptime(due_date, '%Y-%m-%d')
            diff_days = (dd_obj - bd_obj).days
            loan_period = self.get_loan_period_days()
            if not (1 <= diff_days <= loan_period):
                messagebox.showerror(
                    "Invalid Due Date",
                    f"Due date must be between 1 and {loan_period} days after the issue date."
                )
                return
        except Exception:
            pass

        # Disable button to prevent double-submit? (Optional, skipping for minimal intrusion)
        
        self.run_in_background_thread(
            self._borrow_book_worker,
            self._borrow_book_callback,
            enrollment_no=enrollment_no,
            book_id=book_id,
            borrow_date=borrow_date,
            due_date=due_date
        )

    def _borrow_book_worker(self, enrollment_no, book_id, borrow_date, due_date):
        """Worker: Pre-borrow DB checks + Execute Borrow"""
        conn = None
        try:
            # Enforce: Pass Out students cannot borrow
            conn = self.db.get_connection()
            cur = conn.cursor()
            cur.execute("SELECT year FROM students WHERE enrollment_no = ?", (enrollment_no,))
            row = cur.fetchone()
            if not row:
                conn.close()
                return (False, "Student not found!")
            
            year_val = (row[0] or '').strip().lower()
            if year_val in ("pass out", "passout"):
                conn.close()
                return (False, "Pass Out students cannot borrow books.")
            
            # Check max books limit
            max_books = self.get_max_books_per_student()
            cur.execute("SELECT COUNT(*) FROM borrow_records WHERE enrollment_no = ? AND status = 'borrowed'", (enrollment_no,))
            current_books = cur.fetchone()[0]
            conn.close()
            conn = None
            
            if current_books >= max_books:
                return (False, f"Limit Reached: Student has {current_books}/{max_books} books.")
        except Exception as e:
            print(f"Pre-borrow validation error: {e}")
            if conn:
                try:
                    conn.close()
                except:
                    pass
            return (False, f"Validation failed: {e}")

        return self.db.borrow_book(enrollment_no, book_id, borrow_date, due_date)

    def _borrow_book_callback(self, result):
        """Callback for borrow book"""
        if isinstance(result, Exception):
             messagebox.showerror("Error", f"System Error: {str(result)}")
             return
        
        success, message = result
        
        if success:
            # Extract student and book info from entries for logging
            enrollment_text = self.borrow_enrollment_entry.get().strip()
            book_text = self.borrow_book_id_entry.get().strip()
            enrollment_no = enrollment_text.split(' - ')[0] if ' - ' in enrollment_text else enrollment_text
            book_id = book_text.split(' - ')[0] if ' - ' in book_text else book_text
            
            # Log the activity
            self._log_admin_activity(
                "Book Issued",
                f"Issued book {book_id} to student {enrollment_no}"
            )
            
            messagebox.showinfo("Success", message)
            # Clear fields
            self.borrow_enrollment_entry.delete(0, tk.END)
            self.borrow_book_id_entry.delete(0, tk.END)
            # Reset due date to default (suggest max window of 7 days)
            from datetime import datetime, timedelta
            today_str = datetime.now().strftime('%Y-%m-%d')
            self.borrow_borrow_date_entry.delete(0, tk.END)
            self.borrow_borrow_date_entry.insert(0, today_str)
            self.borrow_due_date_entry.delete(0, tk.END)
            # Suggest 7-day period by default; user may change to 1..7 days
            default_due = (datetime.now() + timedelta(days=self.get_loan_period_days())).strftime('%Y-%m-%d')
            self.borrow_due_date_entry.insert(0, default_due)

            # Clear student and book details
            self.borrow_student_details.config(text="")
            self.borrow_book_details.config(text="")

            # Refresh views (Async calls where possible)
            self.refresh_borrowed() 
            self.refresh_books()
            self.refresh_dashboard()
            self.refresh_records()
        else:
            # Handle specific error titles if possible
            title = "Not Allowed" if "Pass Out" in message else "Error"
            if "Limit Reached" in message: title = "Limit Reached"
            messagebox.showerror(title, message)
    
    def return_book(self):
        """Handle book return (Async)"""
        enrollment_text = self.return_enrollment_entry.get().strip()
        book_text = self.return_book_id_entry.get().strip()
        
        if not all([enrollment_text, book_text]):
            messagebox.showerror("Error", "Both Enrollment No and Book ID are required!")
            return
        
        # Extract enrollment number (before " - " if formatted, otherwise use as is)
        enrollment_no = enrollment_text.split(' - ')[0] if ' - ' in enrollment_text else enrollment_text
        
        # Extract book ID (before " - " if formatted, otherwise use as is)
        book_id = book_text.split(' - ')[0] if ' - ' in book_text else book_text
        
        # Get optional return date from UI
        return_date_input = getattr(self, 'return_date_entry', None)
        user_return_date = return_date_input.get().strip() if return_date_input else None
        if user_return_date == '':
            user_return_date = None

        self.run_in_background_thread(
            self._return_book_worker,
            self._return_book_callback,
            enrollment_no=enrollment_no, 
            book_id=book_id, 
            user_return_date=user_return_date
        )

    def _return_book_worker(self, enrollment_no, book_id, user_return_date):
        """Worker: Execute Return + Calculate Fine"""
        success, message = self.db.return_book(enrollment_no, book_id, user_return_date)
        
        fine_data = None
        if success:
            # Late return fine calculation (DB Read)
            conn = None
            try:
                conn = self.db.get_connection()
                cur = conn.cursor()
                cur.execute("""
                    SELECT due_date, return_date FROM borrow_records
                    WHERE enrollment_no=? AND book_id=?
                    ORDER BY id DESC LIMIT 1
                """, (enrollment_no, book_id))
                row = cur.fetchone()
                conn.close()
                conn = None
                if row and row[0] and row[1]:
                    from datetime import datetime as _dt
                    due_dt = _dt.strptime(row[0], '%Y-%m-%d')
                    ret_dt = _dt.strptime(row[1], '%Y-%m-%d')
                    days_late = (ret_dt - due_dt).days
                    if days_late > 0:
                        fine_amount = days_late * self.get_fine_per_day()
                        fine_data = (days_late, fine_amount)
            except Exception as e:
                print(f"Late return fine computation failed: {e}")
                if conn:
                    try:
                        conn.close()
                    except:
                        pass
        
        return (success, message, fine_data)

    def _return_book_callback(self, result):
        """Callback for return book"""
        if isinstance(result, Exception):
             messagebox.showerror("Error", f"System Error: {str(result)}")
             return

        success, message, fine_data = result
        
        if success:
            # Extract student and book info from entries for logging
            enrollment_text = self.return_enrollment_entry.get().strip()
            book_text = self.return_book_id_entry.get().strip()
            enrollment_no = enrollment_text.split(' - ')[0] if ' - ' in enrollment_text else enrollment_text
            book_id = book_text.split(' - ')[0] if ' - ' in book_text else book_text
            
            # Log the activity
            fine_info = f" (Fine: ₹{fine_data[1]})" if fine_data else ""
            self._log_admin_activity(
                "Book Returned",
                f"Returned book {book_id} from student {enrollment_no}{fine_info}"
            )
            
            messagebox.showinfo("Success", message)
            
            # Late return fine popup
            if fine_data:
                days_late, fine_amount = fine_data
                messagebox.showwarning(
                    "Late Return",
                    f"This book is returned {days_late} day(s) late.\nFine: ₹{fine_amount}"
                )

            # Clear fields
            self.return_enrollment_entry.delete(0, tk.END)
            self.return_book_id_entry.delete(0, tk.END)
            
            # Clear student details and return date (reset to today for convenience)
            self.return_student_details.config(text="")
            return_date_input = getattr(self, 'return_date_entry', None)
            if return_date_input:
                from datetime import datetime
                return_date_input.delete(0, tk.END)
                return_date_input.insert(0, datetime.now().strftime('%Y-%m-%d'))
            
            # Refresh views (Async calls where possible)
            self.refresh_borrowed()
            self.refresh_books()
            self.refresh_dashboard()
            self.refresh_records()
        else:
            messagebox.showerror("Error", message)
    
    def import_students_from_excel(self):
        """Import students from Excel file (append; skip invalid or duplicate)."""
        # First, ask user to select the year for these students
        year_dialog = tk.Toplevel(self.root)
        year_dialog.title("Select Student Year")
        year_dialog.geometry("350x250")
        year_dialog.resizable(False, False)
        year_dialog.configure(bg='white')
        year_dialog.transient(self.root)
        year_dialog.grab_set()
        
        # Center the dialog
        year_dialog.update_idletasks()
        x = (year_dialog.winfo_screenwidth() // 2) - (350 // 2)
        y = (year_dialog.winfo_screenheight() // 2) - (250 // 2)
        year_dialog.geometry(f"+{x}+{y}")
        
        # Variable to store selected year and confirmation
        selected_year = tk.StringVar(value="1st Year")
        confirmed = [False]  # Use list to allow modification in nested function
        
        # Header
        header_frame = tk.Frame(year_dialog, bg=self.colors['secondary'], height=60)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        tk.Label(
            header_frame,
            text="📚 Import Students",
            font=('Segoe UI', 14, 'bold'),
            bg=self.colors['secondary'],
            fg='white'
        ).pack(pady=15)
        
        # Content
        content_frame = tk.Frame(year_dialog, bg='white', padx=20, pady=20)
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(
            content_frame,
            text="Which year are these students from?",
            font=('Segoe UI', 11),
            bg='white',
            fg=self.colors['accent']
        ).pack(pady=(0, 15))
        
        # Radio buttons for year selection
        years = ["1st Year", "2nd Year", "3rd Year", "Pass Out"]
        for year in years:
            rb = tk.Radiobutton(
                content_frame,
                text=year,
                variable=selected_year,
                value=year,
                font=('Segoe UI', 10),
                bg='white',
                fg=self.colors['text'],
                selectcolor='white',
                activebackground='white'
            )
            rb.pack(anchor='w', pady=5)
        
        # Buttons
        button_frame = tk.Frame(year_dialog, bg='white', padx=20, pady=(0, 20))
        button_frame.pack(fill=tk.X)
        
        def on_confirm():
            confirmed[0] = True
            year_dialog.destroy()
        
        def on_cancel():
            confirmed[0] = False
            year_dialog.destroy()
        
        tk.Button(
            button_frame,
            text="Continue",
            command=on_confirm,
            bg=self.colors['secondary'],
            fg='white',
            font=('Segoe UI', 10, 'bold'),
            padx=20,
            pady=8,
            relief='flat',
            cursor='hand2'
        ).pack(side=tk.RIGHT, padx=(5, 0))
        
        tk.Button(
            button_frame,
            text="Cancel",
            command=on_cancel,
            bg='#cccccc',
            fg='#333333',
            font=('Segoe UI', 10),
            padx=20,
            pady=8,
            relief='flat',
            cursor='hand2'
        ).pack(side=tk.RIGHT)
        
        # Wait for dialog to close
        self.root.wait_window(year_dialog)
        
        # If user cancelled, abort import
        if not confirmed[0]:
            return
        
        default_year = selected_year.get()

        # Ensure main window has focus before opening file dialog
        self.root.focus_force()
        self.root.update_idletasks()

        # Now proceed with file selection
        # Force file dialog to front
        self.root.attributes('-topmost', True)
        self.root.update()
        file_path = filedialog.askopenfilename(
            parent=self.root,
            title="Select Students Excel file",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
        )
        self.root.attributes('-topmost', False)
        self.root.update()
        if not file_path:
            return
        try:
            df = pd.read_excel(file_path)
            # Normalize columns
            df.columns = df.columns.str.lower().str.replace(' ', '_')
            # Accept some synonyms
            column_map = {
                'enrollment': 'enrollment_no',
                'enrollmentno': 'enrollment_no',
                'enrollment_number': 'enrollment_no'
            }
            df.rename(columns={k: v for k, v in column_map.items() if k in df.columns}, inplace=True)
            required = ['enrollment_no', 'name']
            missing = [c for c in required if c not in df.columns]
            if missing:
                messagebox.showerror("Error", f"Missing required columns: {', '.join(missing)}")
                return
            added = 0
            skipped = 0
            duplicate = 0
            errors = 0
            error_list = []
            for idx, row in df.iterrows():
                row_no = idx + 2
                try:
                    enrollment = str(row.get('enrollment_no', '')).strip()
                    name = str(row.get('name', '')).strip()
                    email = str(row.get('email', '')).strip()
                    phone = str(row.get('phone', '')).strip()
                    department = str(row.get('department', 'Computer')).strip() or 'Computer'
                    year = default_year  # Always use the selected year from the dialog
                    if not enrollment or not name or enrollment.lower() == 'nan' or name.lower() == 'nan':
                        skipped += 1
                        continue
                    success, message = self.db.add_student(enrollment, name, email, phone, department, year)
                    if success:
                        added += 1
                    else:
                        # Determine if it's duplicate vs generic error
                        if 'exists' in message.lower() or 'duplicate' in message.lower():
                            duplicate += 1
                        else:
                            errors += 1
                            error_list.append(f"Row {row_no}: {message}")
                except Exception as e:
                    errors += 1
                    error_list.append(f"Row {row_no}: {e}")
            summary = (
                f"Import completed!\n\nAdded: {added}\nDuplicates: {duplicate}\n"
                f"Skipped (missing enrollment/name): {skipped}\nErrors: {errors}"
            )
            if error_list:
                summary += "\n\nFirst few errors:\n" + "\n".join(error_list[:5])
                if len(error_list) > 5:
                    summary += f"\n... and {len(error_list) - 5} more errors."
            messagebox.showinfo("Import Results", summary)
            if added or duplicate:
                self.refresh_students()
                self.refresh_dashboard()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to import Excel file: {e}")
    def search_students(self):
        """Search and filter students (Async)"""
        search_term = self.student_search_var.get().lower()
        year_filter = self.student_year_filter.get()
        branch_filter = self.student_branch_filter.get()
        
        self.run_in_background_thread(
            self._search_students_worker,
            self._search_students_callback,
            search_term=search_term,
            year_filter=year_filter,
            branch_filter=branch_filter
        )

    def _search_students_worker(self, search_term, year_filter, branch_filter="All"):
        """Worker thread for searching students"""
        if hasattr(self.db, 'get_students'):
             students = self.db.get_students()
        else:
             return []

        # DB students columns:
        # (id, enrollment_no, name, email, phone, department, year, date_registered)
        filtered_students = []
        for s in students:
            if search_term:
                text = f"{s[1]} {s[2]} {s[3]} {s[4]} {s[6]} {s[5]}".lower()
                if search_term not in text:
                    continue
            if year_filter != "All":
                def normalize_year(y):
                    y = str(y or '').lower().strip()
                    if '1' in y: return '1'
                    if '2' in y: return '2'
                    if '3' in y: return '3'
                    return y
                
                student_year = (s[6] or '').strip()
                if normalize_year(year_filter) != normalize_year(student_year):
                    continue
            if branch_filter != "All":
                student_branch = (s[5] or '').strip().lower()
                if branch_filter.lower() not in student_branch:
                    continue
            filtered_students.append(s)
        return filtered_students

    def _search_students_callback(self, result):
        """Update UI with search results"""
        if isinstance(result, Exception):
            messagebox.showerror("Error", f"Error searching students: {str(result)}")
            return
        self.populate_students_tree(result)
    
    def search_books(self):
        """Search and filter books (Async)"""
        try:
            if hasattr(self, 'books_tree'):
                term = (self.book_search_var.get() or '').strip()
                category = self.book_category_filter.get()
                
                self.run_in_background_thread(
                    self._search_books_worker,
                    self._search_books_callback,
                    term=term,
                    category=category
                )
        except Exception as e:
             messagebox.showerror("Error", f"Error initiating book search: {str(e)}")

    def _search_books_worker(self, term, category):
        """Worker thread for searching books"""
        # Fetch from DB with optional LIKE filtering when term present
        books = self.db.get_books(term)
        # DB books columns:
        # (id, book_id, title, author, isbn, category, total_copies, available_copies, date_added)
        filtered = []
        for b in books:
            if category and category != 'All':
                if (b[5] or '').strip() != category:
                    continue
            filtered.append(b)
        return filtered

    def _search_books_callback(self, result):
        """Update UI with book search results"""
        if isinstance(result, Exception):
            print(f"Error searching books: {str(result)}") # Keep original silent/print error handling? Or show message?
            # Original code printed it. I'll print it but also maybe show status if I had one.
            return
        self.populate_books_tree(result)
    
    def search_records(self):
        """Search and filter records including academic year and branch (Async)"""
        search_term = self.record_search_var.get().lower()
        type_filter = self.record_type_filter.get()
        from_date = self.record_from_date.get()
        to_date = self.record_to_date.get()
        academic_year_filter = self.record_academic_year_var.get() if hasattr(self, 'record_academic_year_var') else "All"
        branch_filter = self.record_branch_filter.get() if hasattr(self, 'record_branch_filter') else "All"
        
        self.run_in_background_thread(
            self._search_records_worker,
            self._search_records_callback,
            search_term=search_term,
            type_filter=type_filter,
            from_date=from_date,
            to_date=to_date,
            academic_year_filter=academic_year_filter,
            branch_filter=branch_filter
        )

    def _search_records_worker(self, search_term, type_filter, from_date, to_date, academic_year_filter, branch_filter="All"):
        """Worker: Fetch and Filter Records"""
        records = self.get_all_records()
        
        # Filter records
        filtered_records = []
        for record in records:
            # record tuple: (enroll[0], name[1], dept[2], book_id[3], title[4],
            #                borrow_date[5], due_date[6], return_date[7], status[8], fine[9], academic_year[10])
            dept_val = record[2] if len(record) > 2 else ''
            status_val = record[8] if len(record) > 8 else ''
            fine_val = record[9] if len(record) > 9 else 0
            academic_year_val = record[10] if len(record) > 10 else 'N/A'
            
            # Apply branch filter
            if branch_filter != "All":
                if branch_filter.lower() not in str(dept_val).lower():
                    continue
            
            # Apply type filter
            if type_filter != "All":
                if type_filter == "Overdue":
                    # Overdue means borrowed and due_date has passed (independent of fine amount)
                    try:
                        due_obj = datetime.strptime(str(record[6]), '%Y-%m-%d').date()
                        overdue_by_date = (str(status_val).lower() == 'borrowed' and due_obj < datetime.now().date())
                    except Exception:
                        overdue_by_date = False
                    if not overdue_by_date:
                        continue
                elif type_filter == "Issued":
                    if str(status_val).lower() != 'borrowed':
                        continue
                elif type_filter == "Returned":
                    if str(status_val).lower() != 'returned':
                        continue
            
            # Apply academic year filter
            if academic_year_filter != "All":
                display_academic_year = academic_year_val
                if academic_year_val and "-" in str(academic_year_val):
                    years = str(academic_year_val).split("-")
                    if len(years) == 2:
                        year1 = years[0][-2:]
                        year2 = years[1][-2:]
                        display_academic_year = f"{year1}-{year2}"
                if display_academic_year != academic_year_filter:
                    continue
            
            # Apply search filter
            if search_term:
                if not any(search_term in str(field).lower() for field in record):
                    continue
            
            # Apply date filters (borrow_date is now at index 5)
            from datetime import datetime
            if from_date:
                try:
                    record_date = datetime.strptime(str(record[5]), '%Y-%m-%d')  # borrow_date index 5
                    from_date_obj = datetime.strptime(from_date, '%Y-%m-%d')
                    if record_date < from_date_obj:
                        continue
                except ValueError:
                    pass
            
            if to_date:
                try:
                    record_date = datetime.strptime(str(record[5]), '%Y-%m-%d')
                    to_date_obj = datetime.strptime(to_date, '%Y-%m-%d')
                    if record_date > to_date_obj:
                        continue
                except ValueError:
                    pass
            
            filtered_records.append(record)
        return filtered_records

    def _search_records_callback(self, result):
         if isinstance(result, Exception):
             messagebox.showerror("Error", f"Error searching records: {str(result)}")
             return
         self.populate_records_tree(result)
    
    def refresh_academic_year_filter(self):
        """Refresh academic year dropdown with latest years from database"""
        if hasattr(self, 'academic_year_combo'):
            # Get academic years from database and convert format
            academic_years_raw = self.db.get_all_academic_years()
            academic_years = ["All"]
            for year in academic_years_raw:
                # Convert format from "2025-2026" to "25-26"
                if "-" in year:
                    years = year.split("-")
                    if len(years) == 2:
                        year1 = years[0][-2:]  # "2025" -> "25"
                        year2 = years[1][-2:]  # "2026" -> "26"
                        academic_years.append(f"{year1}-{year2}")
                    else:
                        academic_years.append(year)
                else:
                    academic_years.append(year)
            
            # Update combobox values
            self.academic_year_combo['values'] = academic_years
            # If current selection is not in list, reset to "All"
            if self.record_academic_year_var.get() not in academic_years:
                self.record_academic_year_var.set("All")
    
    def clear_record_filters(self):
        """Clear all record filters"""
        self.record_search_var.set("")
        self.record_type_filter.set("All")
        if hasattr(self, 'record_branch_filter'):
            self.record_branch_filter.set("All")
        self.record_from_date.delete(0, tk.END)
        self.record_to_date.delete(0, tk.END)
        if hasattr(self, 'record_academic_year_var'):
            self.record_academic_year_var.set("All")
        self.search_records()
    
    def refresh_all_data(self):
        """Refresh all data in the application - runs in background to avoid UI freeze"""
        def do_refresh():
            try:
                # Small delay to let UI settle
                self.root.after(100, lambda: self.refresh_dashboard())
                self.root.after(200, lambda: self.refresh_students())
                self.root.after(300, lambda: self.refresh_books())
                self.root.after(400, lambda: self.refresh_borrowed())
                self.root.after(500, lambda: self.refresh_records())
            except Exception as e:
                print(f"Refresh error: {e}")
        
        # Schedule refresh after short delay to let UI draw first
        self.root.after(50, do_refresh)
    
    def refresh_dashboard(self):
        """Refresh dashboard statistics"""
        # Refresh dashboard borrowed books
        if hasattr(self, 'dashboard_borrowed_tree'):
            self.refresh_dashboard_borrowed()
        
        # Refresh statistics cards if they exist
        if hasattr(self, 'stats_container'):
            # Clear existing stats
            for widget in self.stats_container.winfo_children():
                widget.destroy()
            # Recreate stats cards
            self.create_stats_cards(self.stats_container)
        
        # Refresh academic year display
        self.refresh_academic_year_display()
        
        # Auto-refresh Analysis tab charts when data changes
        if hasattr(self, 'notebook') and hasattr(self, 'refresh_analysis'):
            try:
                self.refresh_analysis()
            except Exception:
                pass
    
    def refresh_academic_year_display(self):
        """Refresh the academic year display in the header (async)"""
        def _fetch_year():
            return self.db.get_active_academic_year()
        
        def _update_year(active_year):
            try:
                if isinstance(active_year, Exception):
                    return
                if not active_year:
                    active_year = "2025-2026"
                if hasattr(self, 'academic_year_label'):
                    if "-" in active_year:
                        years = active_year.split("-")
                        if len(years) == 2:
                            year1 = years[0][-2:]
                            year2 = years[1][-2:]
                            display_year = f"{year1}-{year2}"
                        else:
                            display_year = active_year
                    else:
                        display_year = active_year
                    self.academic_year_label.config(text=f"Academic Year: {display_year}")
            except Exception:
                pass
        
        self.run_in_background_thread(_fetch_year, _update_year)
    
    def refresh_students(self):
        """Refresh students list"""
        self.search_students()  # This will apply current filters
    
    def refresh_books(self):
        """Refresh books list"""
        self.search_books()  # This will apply current filters
    
    def refresh_borrowed(self):
        """Refresh borrowed books list (Async)"""
        # Fetch data in background, then populate tree
        self.run_in_background_thread(
            lambda: self.db.get_borrowed_books(),
            self._refresh_borrowed_callback
        )

    def _refresh_borrowed_callback(self, result):
        if isinstance(result, Exception):
            print(f"Error refreshing borrowed books: {result}")
            return
        self.populate_borrowed_tree(result)
    
    def refresh_records(self):
        """Refresh records list"""
        self.search_records()  # This will apply current filters
    
    def populate_students_tree(self, students):
        """Populate students treeview"""
        if hasattr(self, 'students_tree'):
            for item in self.students_tree.get_children():
                self.students_tree.delete(item)
            
            for student in students:
                # Map DB tuple to UI columns
                # Enrollment No, Name, Email, Phone, Branch, Year
                # DB: (id, enrollment_no, name, email, phone, department, year, date_registered)
                display_data = (student[1], student[2], student[3], student[4], student[5], student[6])
                self.students_tree.insert('', 'end', values=display_data)
    
    def populate_books_tree(self, books):
        """Populate books treeview"""
        if hasattr(self, 'books_tree'):
            for item in self.books_tree.get_children():
                self.books_tree.delete(item)
            
            for book in books:
                # DB cols: (id, book_id, title, author, isbn, category, total_copies, available_copies, date_added, barcode, price)
                # indexes:   0    1       2      3       4     5         6             7                  8           9       10
                price_val = book[10] if len(book) > 10 and book[10] is not None else 0
                price_str = f"₹{float(price_val):.2f}" if price_val else "—"
                display_data = (book[1], book[2], book[3], book[4], book[5], book[6], book[7], price_str)
                self.books_tree.insert('', 'end', values=display_data)
    
    def populate_borrowed_tree(self, borrowed):
        """Populate borrowed books treeview with enhanced data
        Expected order from get_borrowed_books():
        (enrollment_no, student_name, department, year, book_id, title, author, borrow_date, due_date)
        """
        if hasattr(self, 'borrowed_tree'):
            # Configure tags once (text-only coloring)
            try:
                self.borrowed_tree.tag_configure('overdue', foreground='#b30000')
                self.borrowed_tree.tag_configure('due_soon', foreground='#856404')
                self.borrowed_tree.tag_configure('ok', foreground='#155724')
            except Exception:
                pass

            for item in self.borrowed_tree.get_children():
                self.borrowed_tree.delete(item)
            
            for record in borrowed:
                # record indexes mapping
                enrollment_no = record[0]
                student_name = record[1]
                department    = record[2] if len(record) > 2 else ''
                book_id = record[4]
                book_title = record[5]
                borrow_date = record[7]
                due_date_val = record[8]
                # Calculate days left
                try:
                    from datetime import datetime as _dt
                    due_date = _dt.strptime(due_date_val, '%Y-%m-%d').date()
                    today = _dt.now().date()
                    delta = (due_date - today).days
                    if delta < 0:
                        days_left_str = f"Overdue {abs(delta)}d"
                        tag = 'overdue'
                    elif delta == 0:
                        days_left_str = 'Due Today'
                        tag = 'due_soon'
                    elif delta <= 3:
                        days_left_str = f"{delta}d left"
                        tag = 'due_soon'
                    else:
                        days_left_str = f"{delta}d left"
                        tag = 'ok'
                except Exception:
                    days_left_str = 'N/A'
                    tag = ''
                display_data = (student_name, department, book_id, book_title, borrow_date, due_date_val, days_left_str)
                self.borrowed_tree.insert('', 'end', values=display_data, tags=(tag,))
    
    def populate_activities_tree(self, activities):
        """Populate recent activities treeview"""
        if hasattr(self, 'activities_tree'):
            for item in self.activities_tree.get_children():
                self.activities_tree.delete(item)
            
            for activity in activities:
                self.activities_tree.insert('', 'end', values=activity)
    
    def populate_records_tree(self, records):
        """Populate records treeview"""
        if hasattr(self, 'records_tree'):
            for item in self.records_tree.get_children():
                self.records_tree.delete(item)
            
            for record in records:
                # record tuple: (enroll[0], student_name[1], department[2], book_id[3], title[4],
                #                borrow_date[5], due_date[6], return_date_str[7], status[8], fine[9], academic_year[10])
                if len(record) >= 11:
                    enroll = record[0]; student_name = record[1]; department = record[2]
                    book_id = record[3]; title = record[4]
                    borrow_date = record[5]; due_date = record[6]; return_date_str = record[7]
                    status = record[8]; fine = record[9]
                elif len(record) >= 9:
                    enroll = record[0]; student_name = record[1]; department = ''
                    book_id = record[2]; title = record[3]
                    borrow_date = record[4]; due_date = record[5]; return_date_str = record[6]
                    status = record[7]; fine = record[8]
                else:
                    continue
                # Format fine display
                fine_is_num = isinstance(fine, (int, float))
                if not fine_is_num and isinstance(fine, str):
                    try:
                        fine = int(fine)
                        fine_is_num = True
                    except ValueError:
                        pass
                # Determine overdue by due date, not by fine amount
                try:
                    due_obj = datetime.strptime(str(due_date), '%Y-%m-%d').date()
                    overdue_by_date = (str(status).lower() == 'borrowed' and due_obj < datetime.now().date())
                except Exception:
                    overdue_by_date = False

                tag = 'late' if overdue_by_date else ''
                if fine_is_num and isinstance(fine, (int, float)) and overdue_by_date:
                    fine_display = f"Rs {int(fine)} (Late)"
                else:
                    fine_display = f"Rs {int(fine)}" if fine_is_num else str(fine)
                self.records_tree.insert('', 'end',
                    values=(enroll, student_name, department, book_id, title,
                            borrow_date, due_date, return_date_str, status, fine_display),
                    tags=(tag,))
            try:
                # Text color only (no background highlight)
                self.records_tree.tag_configure('late', foreground='#b30000')
            except Exception:
                pass
    
    def on_record_double_click(self, event):
        """Handle double-click on record to send overdue letter"""
        selection = self.records_tree.selection()
        if not selection:
            return
        
        # Get the selected record data
        item = self.records_tree.item(selection[0])
        values = item['values']
        
        if not values or len(values) < 10:
            return
        
        # Extract record details (tree columns: Enrollment No, Student Name, Branch, Book ID, Book Title, Issue Date, Due Date, Return Date, Status, Fine)
        enrollment_no = str(values[0]).strip()
        student_name = str(values[1]).strip()
        branch = str(values[2]).strip()
        book_id = str(values[3]).strip()
        book_title = str(values[4]).strip()
        issue_date = str(values[5]).strip()
        due_date = str(values[6]).strip()
        return_date = str(values[7]).strip()
        status = str(values[8]).strip()
        fine_info = str(values[9]).strip()
        
        # Extract fine amount to check if overdue
        fine_amount = 0
        try:
            # Remove "Rs", "(Late)", and other text to get the number
            fine_str = fine_info.replace('Rs', '').replace('(Late)', '').strip()
            fine_amount = int(fine_str)
        except:
            fine_amount = 0
        
        # Check if this is an overdue record by due_date (independent of fine amount)
        is_overdue = False
        try:
            due_obj = datetime.strptime(str(due_date), '%Y-%m-%d').date()
            is_overdue = status.lower() == 'borrowed' and due_obj < datetime.now().date()
        except Exception:
            is_overdue = False

        if is_overdue:
            # Show dialog to confirm sending letter
            response = messagebox.askyesno(
                "Send Overdue Letter",
                f"Send overdue letter to:\n\n"
                f"Student: {student_name}\n"
                f"Branch: {branch}\n"
                f"Enrollment: {enrollment_no}\n"
                f"Book: {book_title}\n"
                f"Fine: {fine_info}\n\n"
                f"Generate Word document?",
                icon='question'
            )
            
            if response:
                self.generate_overdue_letter(
                    enrollment_no, student_name, book_id, 
                    book_title, issue_date, due_date, fine_info, branch=branch
                )
        else:
            messagebox.showinfo(
                "Not Overdue",
                "This record is not overdue. Letters can only be sent for overdue borrowed books.",
                icon='info'
            )
    
    def generate_overdue_letter(self, enrollment_no, student_name, book_id, book_title, issue_date, due_date, fine_info, branch=''):
        """Generate overdue letter as Word document"""
        if not Document:
            messagebox.showerror("Error", "python-docx is not installed. Cannot generate Word document.")
            return
        
        try:
            # Normalize values to safe strings (treeview may provide ints/floats)
            enrollment_no = str(enrollment_no).strip()
            student_name = str(student_name).strip()
            book_id = str(book_id).strip()
            book_title = str(book_title).strip()
            issue_date = str(issue_date).strip()
            due_date = str(due_date).strip()
            fine_info = str(fine_info).strip()
            branch = str(branch).strip() if branch is not None else ''

            # Get branch from DB if not provided
            if not branch:
                try:
                    conn = self.db.get_connection()
                    cur = conn.cursor()
                    cur.execute("SELECT department FROM students WHERE enrollment_no = ?", (enrollment_no,))
                    row = cur.fetchone()
                    branch = row[0] if row and row[0] else ''
                    conn.close()
                except Exception:
                    branch = ''
            
            # Calculate days overdue
            from datetime import datetime as _dt
            due_d = _dt.strptime(due_date, '%Y-%m-%d')
            today = _dt.now()
            days_overdue = (today.date() - due_d.date()).days
            
            # Extract fine amount
            fine_amount = 0
            try:
                fine_str = fine_info.replace('Rs', '').replace('(Late)', '').strip()
                fine_amount = int(fine_str)
            except:
                fine_amount = days_overdue * self.get_fine_per_day()
            
            # Create Word document
            doc = Document()
            
            # Add logo at the top if available
            logo_path = os.path.join(os.path.dirname(__file__), 'logo.png')
            if os.path.exists(logo_path):
                try:
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(logo_path, width=Pt(80))
                except Exception as e:
                    print(f"Could not add logo: {e}")
            
            # Add institutional header with colors
            def add_colored_header(text, size, rgb_color):
                from docx.shared import RGBColor
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(size)
                run.font.color.rgb = RGBColor(*rgb_color)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return p
            
            # Three-line header with gradient colors
            add_colored_header("Government Polytechnic Awasari (Kh)", 20, (31, 71, 136))
            add_colored_header("Main Library", 16, (46, 92, 138))
            
            # Add separator line
            sep_para = doc.add_paragraph("_" * 70)
            sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Spacing
            
            # Add date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f"Date: {today.strftime('%B %d, %Y')}")
            date_run.font.size = Pt(11)
            
            doc.add_paragraph()  # Spacing
            
            # Add subject
            subject = doc.add_paragraph()
            subject_run = subject.add_run('Subject: Overdue Book Notice')
            subject_run.bold = True
            subject_run.font.size = Pt(12)
            subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Spacing
            
            # Recipient details
            to_para = doc.add_paragraph()
            branch_line = f"\nBranch: {branch}" if branch else ''
            to_run = to_para.add_run(f'To,\n{student_name}\nEnrollment No: {enrollment_no}{branch_line}')
            to_run.font.size = Pt(11)
            
            doc.add_paragraph()  # Spacing
            
            # Body of letter
            body = doc.add_paragraph()
            body_text = (
                f"Dear {student_name},\n\n"
                f"This is to inform you that the following book borrowed from GPA's Library Management System "
                f"is overdue and needs to be returned immediately.\n\n"
            )
            body_run = body.add_run(body_text)
            body_run.font.size = Pt(11)
            
            # Book details table
            doc.add_paragraph('Book Details:', style='Heading 2')
            table_rows = 7 if branch else 6
            table = doc.add_table(rows=table_rows, cols=2)
            table.style = 'Light Grid Accent 1'
            
            table.cell(0, 0).text = 'Student Name:'
            table.cell(0, 1).text = student_name
            row_idx = 1
            if branch:
                table.cell(row_idx, 0).text = 'Branch:'
                table.cell(row_idx, 1).text = branch
                row_idx += 1
            table.cell(row_idx, 0).text = 'Book ID:'
            table.cell(row_idx, 1).text = str(book_id)
            table.cell(row_idx + 1, 0).text = 'Book Title:'
            table.cell(row_idx + 1, 1).text = book_title
            table.cell(row_idx + 2, 0).text = 'Issue Date:'
            table.cell(row_idx + 2, 1).text = issue_date
            table.cell(row_idx + 3, 0).text = 'Due Date:'
            table.cell(row_idx + 3, 1).text = due_date
            table.cell(row_idx + 4, 0).text = 'Days Overdue:'
            table.cell(row_idx + 4, 1).text = str(days_overdue)
            
            doc.add_paragraph()  # Spacing
            
            # Fine details
            fine_para = doc.add_paragraph()
            fine_run = fine_para.add_run(
                f"As per library rules, a fine of ₹{self.get_fine_per_day()} per day is applicable for overdue books.\n"
                f"Your current fine amount is: ₹{fine_amount}\n\n"
            )
            fine_run.font.size = Pt(11)
            fine_run.bold = True
            
            # Request
            request_para = doc.add_paragraph()
            request_text = (
                "You are hereby requested to return the book to the library at the earliest and clear the pending fine. "
                "Failure to do so may result in restrictions on future borrowing privileges.\n\n"
                "Please contact the library desk for any queries or clarifications.\n\n"
            )
            request_run = request_para.add_run(request_text)
            request_run.font.size = Pt(11)
            
            # Closing
            doc.add_paragraph()  # Spacing
            closing = doc.add_paragraph()
            closing_text = "Thank you for your cooperation.\n\nYours sincerely,\n\n"
            closing_run = closing.add_run(closing_text)
            closing_run.font.size = Pt(11)
            
            # Signature
            doc.add_paragraph()
            doc.add_paragraph("__________________________")
            signature = doc.add_paragraph()
            signature_run = signature.add_run('Librarian')
            signature_run.font.size = Pt(11)
            signature_run.bold = True
            
            dept1 = doc.add_paragraph()
            dept1.add_run('Main Library').font.size = Pt(10)
            dept2 = doc.add_paragraph()
            dept2.add_run('Government Polytechnic Awasari (Kh)').font.size = Pt(10)
            
            # Save the document to temp location first
            import tempfile
            temp_dir = tempfile.gettempdir()
            default_filename = f"Overdue_Letter_{enrollment_no}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            temp_filepath = os.path.join(temp_dir, default_filename)
            doc.save(temp_filepath)
            
            # Get student email
            student_email = self.get_student_email(enrollment_no)
            
            # Try to send email automatically if enabled
            email_sent = False
            email_message = ""
            
            if self.email_settings.get('enabled', False) and student_email:
                email_subject = f"Overdue Book Notice - {book_title}"
                branch_line = f"Branch: {branch}\n" if branch else ""
                email_body = f"""Dear {student_name},

This is an automated notification from GPA's Library Management System, Government Polytechnic Awasari (Kh).

Student Details:
Enrollment No: {enrollment_no}
{branch_line}
The following book borrowed from our library is overdue and needs to be returned immediately:

Book ID: {book_id}
Book Title: {book_title}
Issue Date: {issue_date}
Due Date: {due_date}
Days Overdue: {days_overdue}
Fine Amount: ₹{fine_amount}

As per library rules, a fine of ₹{self.get_fine_per_day()} per day is applicable for overdue books.

Please return the book to the library at the earliest and clear the pending fine. Failure to do so may result in restrictions on future borrowing privileges.

For any queries, please contact the library desk.

Thank you for your cooperation.

Librarian
Main Library
Government Polytechnic Awasari (Kh)

---
Note: This is an automated email. Please find the attached formal overdue letter.
"""
                
                success, message = self.send_email_with_attachment(
                    student_email, 
                    email_subject, 
                    email_body, 
                    temp_filepath
                )
                email_sent = success
                email_message = message
                
                # Log email attempt to history
                self._log_email_sent(
                    enrollment_no,
                    student_name,
                    student_email,
                    book_title,
                    success,
                    message if not success else ''
                )
            
            # Now ask user where to save the document
            final_filepath = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile=default_filename,
                title="Save Overdue Letter"
            )
            
            if final_filepath:
                # Copy from temp to final location
                import shutil
                shutil.copy(temp_filepath, final_filepath)
                
                # Show success message with email status
                if email_sent:
                    success_msg = (
                        f"✅ Overdue letter generated and emailed successfully!\n\n"
                        f"📧 Email sent to: {student_email}\n"
                        f"💾 Document saved to:\n{final_filepath}"
                    )
                elif self.email_settings.get('enabled', False):
                    success_msg = (
                        f"⚠️ Document saved but email failed!\n\n"
                        f"Reason: {email_message}\n\n"
                        f"💾 Document saved to:\n{final_filepath}"
                    )
                else:
                    success_msg = (
                        f"✅ Overdue letter generated successfully!\n\n"
                        f"💾 Saved to:\n{final_filepath}\n\n"
                        f"💡 Tip: Enable email in Settings to send automatically!"
                    )
                
                messagebox.showinfo("Success", success_msg, icon='info')
                
                # Ask if user wants to open the document
                if messagebox.askyesno("Open Document", "Do you want to open the document now?"):
                    try:
                        if platform.system() == 'Windows':
                            os.startfile(final_filepath)
                        elif platform.system() == 'Darwin':  # macOS
                            subprocess.call(['open', final_filepath])
                        else:  # Linux
                            subprocess.call(['xdg-open', final_filepath])
                    except Exception as e:
                        messagebox.showwarning("Warning", f"Document saved but couldn't open automatically.\nError: {e}")
            
            # Clean up temp file
            try:
                if os.path.exists(temp_filepath):
                    os.remove(temp_filepath)
            except:
                pass
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate overdue letter.\n\nError: {e}")
    
    def get_student_email(self, enrollment_no):
        """Get student email by enrollment number"""
        try:
            conn = self.db.get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT email FROM students WHERE enrollment_no = ?", (enrollment_no,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result and result[0] else None
        except:
            return None
    
    def get_student_name(self, enrollment_no):
        """Get student name by enrollment number"""
        try:
            conn = self.db.get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM students WHERE enrollment_no = ?", (enrollment_no,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else enrollment_no
        except:
            return enrollment_no
    
    def get_book_title(self, book_id):
        """Get book title by book ID"""
        try:
            conn = self.db.get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT title FROM books WHERE book_id = ?", (book_id,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else book_id
        except:
            return book_id
    
    def get_recent_activities(self):
        """Get recent activities for dashboard"""
        conn = None
        try:
            conn = self.db.get_connection()
            cursor = conn.cursor()
            
            # Get recent borrow/return activities
            cursor.execute("""
                SELECT 
                    CASE WHEN return_date IS NULL THEN 'Borrow' ELSE 'Return' END as type,
                    s.name as student_name,
                    b.title as book_title,
                    CASE WHEN return_date IS NULL THEN borrow_date ELSE return_date END as activity_date,
                    status
                FROM borrow_records br
                JOIN students s ON br.enrollment_no = s.enrollment_no
                JOIN books b ON br.book_id = b.book_id
                ORDER BY br.id DESC
                LIMIT 20
            """)
            
            activities = cursor.fetchall()
            conn.close()
            return activities
        except Exception as e:
            print(f"Error getting recent activities: {e}")
            if conn:
                try:
                    conn.close()
                except:
                    pass
            return []
    
    def get_all_records(self):
        """Get all transaction records including academic year"""
        conn = None
        try:
            conn = self.db.get_connection()
            cursor = conn.cursor()
            # Return enrollment_no, book_id, and academic_year explicitly for correct display
            cursor.execute("""
                SELECT 
                    br.enrollment_no,
                    s.name as student_name,
                    COALESCE(s.department, '') as department,
                    COALESCE(
                        NULLIF(br.accession_no, ''),
                        CASE WHEN INSTR(br.book_id, '-') > 0
                             THEN SUBSTR(br.book_id, 1, INSTR(br.book_id, '-') - 1)
                             ELSE br.book_id END
                    ) as display_book_id,
                    b.title as book_title,
                    br.borrow_date,
                    br.due_date,
                    br.return_date,
                    br.status,
                    0 as days_overdue,
                    COALESCE(br.academic_year, 'N/A') as academic_year
                FROM borrow_records br
                JOIN students s ON br.enrollment_no = s.enrollment_no
                JOIN books b ON br.book_id = b.book_id
                ORDER BY br.id DESC
            """)
            
            records = cursor.fetchall()
            conn.close()
            
            # Format records for display
            formatted_records = []
            from datetime import datetime as _dt
            today = _dt.now().date()
            for rec in records:
                (enroll, student_name, department, book_id, title, borrow_date, due_date, return_date_raw, status, _, academic_year) = rec
                
                # Handle return_date normalization (None/Date -> String)
                if return_date_raw is None:
                    return_date_str = 'Not returned'
                else:
                    return_date_str = str(return_date_raw)

                # Determine effective overdue days & fine
                try:
                    due_d = _dt.strptime(str(due_date), '%Y-%m-%d').date()
                except Exception:
                    due_d = None
                
                fine = 0
                if status == 'borrowed':
                    # still out; overdue based on today
                    if due_d and today > due_d:
                        overdue_days = (today - due_d).days
                        fine = overdue_days * self.get_fine_per_day()
                else:
                    # returned; compute late based on return_date
                    try:
                        # Use raw if date object, or parse if string
                        if hasattr(return_date_raw, 'year'):
                             ret_d = return_date_raw
                             # Postgres returns date object, but Python datetime.date doesn't strictly have comparison with None same way
                             if isinstance(ret_d, datetime): ret_d = ret_d.date() 
                        else:
                             ret_d = _dt.strptime(return_date_str, '%Y-%m-%d').date()
                             
                        if due_d and ret_d > due_d:
                            overdue_days = (ret_d - due_d).days
                            fine = overdue_days * self.get_fine_per_day()
                    except Exception:
                        pass
                # Keep fine as numeric for downstream display logic, add academic_year
                formatted_records.append((enroll, student_name, department, book_id, title, borrow_date, due_date, return_date_str, status, fine, academic_year))
            return formatted_records
        except Exception as e:
            print(f"Error getting records: {e}")
            if conn:
                try:
                    conn.close()
                except:
                    pass
            return []

    # ------------------------------------------------------------------
    # Date auto update helpers
    # ------------------------------------------------------------------
    def on_borrow_date_changed(self, event=None):
        """When issue date changes, suggest due date = issue + LOAN_PERIOD_DAYS (max),
        but allow user to choose 1..LOAN_PERIOD_DAYS days range."""
        try:
            from datetime import datetime as _dt, timedelta
            borrow_str = self.borrow_borrow_date_entry.get()
            bd = _dt.strptime(borrow_str, '%Y-%m-%d')
            new_due = (bd + timedelta(days=LOAN_PERIOD_DAYS)).strftime('%Y-%m-%d')
            if DateEntry:
                try:
                    self.borrow_due_date_entry.config(state='normal')
                except Exception:
                    pass
                try:
                    self.borrow_due_date_entry.set_date(_dt.strptime(new_due, '%Y-%m-%d'))
                except Exception:
                    self.borrow_due_date_entry.delete(0, tk.END)
                    self.borrow_due_date_entry.insert(0, new_due)
                try:
                    self.borrow_due_date_entry.config(state='readonly')
                except Exception:
                    pass
            else:
                self.borrow_due_date_entry.delete(0, tk.END)
                self.borrow_due_date_entry.insert(0, new_due)
        except Exception as e:
            print(f"Auto due date update failed: {e}")

    def on_due_date_attempt_change(self):
        """Validate due-date change: must be within 1..loan_period days of issue date."""
        try:
            from datetime import datetime as _dt
            bd = _dt.strptime(self.borrow_borrow_date_entry.get(), '%Y-%m-%d')
            dd = _dt.strptime(self.borrow_due_date_entry.get(), '%Y-%m-%d')
            diff = (dd - bd).days
            loan_period = self.get_loan_period_days()
            if not (1 <= diff <= loan_period):
                messagebox.showerror(
                    "Invalid Due Date",
                    f"Due date must be between 1 and {loan_period} days after the issue date."
                )
                # Re-suggest maximum allowed
                from datetime import timedelta
                suggested = (bd + timedelta(days=loan_period)).strftime('%Y-%m-%d')
                try:
                    self.borrow_due_date_entry.config(state='normal')
                    self.borrow_due_date_entry.delete(0, tk.END)
                    self.borrow_due_date_entry.insert(0, suggested)
                finally:
                    try:
                        self.borrow_due_date_entry.config(state='readonly')
                    except Exception:
                        pass
        except Exception:
            pass
    
    def export_students_dialog(self):
        """Show export students dialog with year selection"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Export Students to Excel")
        dialog.geometry("400x300")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 200, self.root.winfo_rooty() + 200))
        
        # Title
        title_label = tk.Label(
            dialog,
            text="📊 Export Students",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack(pady=(20, 30))
        
        # Options frame
        options_frame = tk.Frame(dialog, bg='white')
        options_frame.pack(expand=True, padx=40)
        
        # Year selection
        year_label = tk.Label(
            options_frame,
            text="Select Year to Export:",
            font=('Segoe UI', 12, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        year_label.pack(pady=(0, 10))
        
        year_var = tk.StringVar(value="All")
        year_frame = tk.Frame(options_frame, bg='white')
        year_frame.pack(pady=(0, 20))
        years = ["All", "1st", "2nd", "3rd", "Pass Out"]
        for year in years:
            rb = tk.Radiobutton(
                year_frame,
                text=year,
                variable=year_var,
                value=year,
                font=('Segoe UI', 11),
                bg='white',
                fg=self.colors['accent']
            )
            rb.pack(anchor='w')
        
        # Auto-update years option
        auto_update_var = tk.BooleanVar(value=True)
        auto_cb = tk.Checkbutton(
            options_frame,
            text="Auto-update years to current academic year before export",
            variable=auto_update_var,
            bg='white', fg=self.colors['accent'], selectcolor='white', anchor='w', justify='left', wraplength=320
        )
        auto_cb.pack(fill=tk.X, pady=(0, 10))

        # Buttons
        btn_frame = tk.Frame(dialog, bg='white')
        btn_frame.pack(pady=20)
        
        def export_data():
            selected_year = year_var.get()
            self.export_students_to_excel(selected_year, auto_update=auto_update_var.get())
            dialog.destroy()
        
        export_btn = tk.Button(
            btn_frame,
            text="📊 Export to Excel",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=export_data,
            cursor='hand2'
        )
        export_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        cancel_btn = tk.Button(
            btn_frame,
            text="❌ Cancel",
            font=('Segoe UI', 12, 'bold'),
            bg='#6c757d',
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=dialog.destroy,
            cursor='hand2'
        )
        cancel_btn.pack(side=tk.LEFT)
    
    def export_students_to_excel(self, year_filter="All", auto_update=True):
        """Export students to Excel with year filter.
        If auto_update is True, compute current year (1st..3rd) based on date_registered and an academic rollover on July 1.
        """
        try:
            students = self.db.get_students()
            
            # Filter by Computer department and year
            filtered_students = []
            for student in students:
                if student[5] == "Computer":  # Department filter
                    # Compute effective year label if requested
                    eff_year = student[6]
                    if auto_update:
                        try:
                            eff_year = self._compute_current_year_label(student[7], student[6])
                        except Exception:
                            eff_year = student[6]
                    if year_filter == "All" or eff_year == year_filter:
                        filtered_students.append({
                            'Enrollment No': student[1],
                            'Name': student[2],
                            'Email': student[3],
                            'Phone': student[4],
                            'Department': student[5],
                            'Year': eff_year,
                            'Registered': student[7]
                        })
            
            if not filtered_students:
                messagebox.showwarning("Warning", f"No students found for year: {year_filter}")
                return
            
            # Create DataFrame
            df = pd.DataFrame(filtered_students)
            
            # Save to file
            filename = f"students_{year_filter}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                # Tkinter option is 'initialfile' (not 'initialname')
                initialfile=filename
            )
            
            if file_path:
                # Add header then table using openpyxl for styling flexibility
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    sheet = 'Students'
                    # Write institutional header first
                    header_end_row = 6  # New enhanced header takes 6 rows
                    df.to_excel(writer, sheet_name=sheet, index=False, startrow=header_end_row)
                    ws = writer.book[sheet]
                    # Write enhanced header
                    self._write_excel_header_openpyxl(ws, start_row=1)
                    # Apply enhanced auto-adjust with styling
                    self._auto_adjust_column_width(ws)
                messagebox.showinfo("Success", f"Students data exported to {file_path}")
                
                # Ask if user wants to open the file
                if messagebox.askyesno("Open File", "Do you want to open the exported file?"):
                    self.open_file(file_path)
                    
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export students: {str(e)}")

    # ---------------------- Academic Year Helpers ----------------------
    def _compute_current_year_label(self, date_registered, stored_year_label):
        """Return current year label ('1st','2nd','3rd') based on date_registered and today's date.
        Assumptions: 3-year course; academic year rolls over on July 1. Falls back to stored label if missing date.
        """
        try:
            from datetime import datetime as _dt
            if not date_registered:
                return stored_year_label or '1st'
            d = _dt.strptime(str(date_registered), '%Y-%m-%d').date()
            today = _dt.now().date()
            def acad_index(x):
                return x.year if (x.month, x.day) >= (7, 1) else x.year - 1
            promotions = max(0, acad_index(today) - acad_index(d))
            base = 1
            map_to_num = {'1st': 1, '2nd': 2, '3rd': 3, 'First':1, 'Second':2, 'Third':3}
            if stored_year_label in map_to_num:
                base = map_to_num[stored_year_label]
            progressed = base + promotions
            if progressed >= 4:
                return 'Pass Out'
            new_num = min(max(progressed, 1), 3)
            return {1:'1st',2:'2nd',3:'3rd'}[new_num]
        except Exception:
            return stored_year_label or '1st'

    def export_all_students_direct(self):
        """Direct export of students with year selection for Share Data function"""
        # Create year selection dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Export Students - Select Year")
        dialog.geometry("450x550")  # Increased height significantly
        dialog.configure(bg='#f8f9fa')  # Light gray background
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.resizable(False, False)
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 180, self.root.winfo_rooty() + 100))
        
        # Main container with border
        main_frame = tk.Frame(
            dialog, 
            bg='white', 
            relief='ridge', 
            bd=2,
            padx=25,
            pady=20
        )
        main_frame.pack(fill='both', expand=True, padx=15, pady=15)
        
        # Title with icon
        title_frame = tk.Frame(main_frame, bg='white')
        title_frame.pack(pady=(0, 20))
        
        title_label = tk.Label(
            title_frame,
            text="📊 Export Students Data",
            font=('Segoe UI', 18, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack()
        
        # Subtitle
        subtitle_label = tk.Label(
            title_frame,
            text="Select year and branch to export:",
            font=('Segoe UI', 12),
            bg='white',
            fg='#6c757d'
        )
        subtitle_label.pack(pady=(5, 0))
        
        # Year selection container with background
        selection_frame = tk.Frame(
            main_frame, 
            bg='#f8f9fa',
            relief='solid',
            bd=1,
            padx=20,
            pady=15
        )
        selection_frame.pack(fill='x', pady=(0, 10))
        
        year_var = tk.StringVar(value="All")
        
        # Year options with better styling
        year_options = [
            ("📚 All Students", "All", "#4285f4"),
            ("🥇 1st Year", "1st Year", "#34a853"), 
            ("🥈 2nd Year", "2nd Year", "#fbbc04"),
            ("🥉 3rd Year", "3rd Year", "#ea4335"),
            ("🎓 Pass Out", "Pass Out", "#9333ea")
        ]
        
        for i, (text, value, color) in enumerate(year_options):
            rb = tk.Radiobutton(
                selection_frame,
                text=text,
                variable=year_var,
                value=value,
                font=('Segoe UI', 12, 'bold'),
                bg='#f8f9fa',
                fg=color,
                selectcolor='white',
                activebackground='#f8f9fa',
                activeforeground=color,
                relief='flat',
                padx=10,
                pady=6
            )
            rb.pack(anchor='w', pady=4)
        
        # Branch selection
        branch_frame = tk.Frame(main_frame, bg='white')
        branch_frame.pack(fill='x', pady=(0, 10))
        tk.Label(branch_frame, text="Branch:", font=('Segoe UI', 11, 'bold'),
                 bg='white', fg=self.colors['accent']).pack(side=tk.LEFT)
        branch_var = tk.StringVar(value="All")
        branch_opts = ["All", "Computer Engineering", "Information Technology",
                       "Civil Engineering", "Electronics & Telecommunication",
                       "Mechanical Engineering", "Automobile Engineering", "Electrical Engineering"]
        ttk.Combobox(branch_frame, textvariable=branch_var, values=branch_opts,
                     state="readonly", width=32).pack(side=tk.LEFT, padx=(10, 0))
        
        # Buttons container - FIXED positioning at bottom
        btn_container = tk.Frame(main_frame, bg='white')
        btn_container.pack(side='bottom', pady=(20, 0))
        
        def perform_export():
            selected_year = year_var.get()
            selected_branch = branch_var.get()
            dialog.destroy()
            self._export_students_by_year(selected_year, selected_branch)
        
        # Export button with gradient-like styling
        export_btn = tk.Button(
            btn_container,
            text="📊 Export to Excel",
            font=('Segoe UI', 12, 'bold'),
            bg='#28a745',
            fg='white',
            relief='flat',
            padx=25,
            pady=12,
            command=perform_export,
            cursor='hand2',
            activebackground='#218838',
            activeforeground='white'
        )
        export_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        # Cancel button
        cancel_btn = tk.Button(
            btn_container,
            text="❌ Cancel",
            font=('Segoe UI', 12, 'bold'),
            bg='#dc3545',
            fg='white',
            relief='flat',
            padx=25,
            pady=12,
            command=dialog.destroy,
            cursor='hand2',
            activebackground='#c82333',
            activeforeground='white'
        )
        cancel_btn.pack(side=tk.LEFT)
        
        # Add hover effects
        def on_enter_export(e):
            export_btn.configure(bg='#218838')
        
        def on_leave_export(e):
            export_btn.configure(bg='#28a745')
            
        def on_enter_cancel(e):
            cancel_btn.configure(bg='#c82333')
        
        def on_leave_cancel(e):
            cancel_btn.configure(bg='#dc3545')
        
        export_btn.bind("<Enter>", on_enter_export)
        export_btn.bind("<Leave>", on_leave_export)
        cancel_btn.bind("<Enter>", on_enter_cancel)
        cancel_btn.bind("<Leave>", on_leave_cancel)

    def _export_students_by_year(self, year_filter, branch_filter="All"):
        """Helper function to export students filtered by year and branch"""
        try:
            students = self.db.get_students()
            
            # Filter students by year and branch
            filtered_students = []
            for student in students:
                # student: (id[0], enrollment_no[1], name[2], email[3], phone[4], department[5], year[6], date_registered[7])
                dept = str(student[5] or '')
                yr = str(student[6] or '')
                
                # Branch filter
                if branch_filter != "All" and branch_filter.lower() not in dept.lower():
                    continue
                # Year filter
                if year_filter != "All" and yr != year_filter:
                    continue
                
                filtered_students.append({
                    'Enrollment No': student[1],
                    'Name': student[2],
                    'Email': student[3] or '',
                    'Phone': student[4] or '',
                    'Branch': dept,
                    'Year': yr,
                    'Registered On': student[7] or ''
                })
            
            if not filtered_students:
                messagebox.showwarning("Warning",
                    f"No students found for Year: {year_filter}, Branch: {branch_filter}!")
                return
            
            # Create DataFrame
            df = pd.DataFrame(filtered_students)
            
            # Save to file
            year_suffix = year_filter.replace(' ', '_') if year_filter != "All" else "all_years"
            branch_suffix = branch_filter.replace(' ', '_').replace('&', 'and') if branch_filter != "All" else "all_branches"
            filename = f"students_{year_suffix}_{branch_suffix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                initialfile=filename
            )
            
            if file_path:
                from openpyxl.styles import Font, Alignment, PatternFill
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    sheet_name = 'Students'
                    header_end_row = 7
                    df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=header_end_row)
                    ws = writer.book[sheet_name]
                    self._write_excel_header_openpyxl(ws, start_row=1)
                    # Filter info row
                    ws.cell(row=7, column=1,
                            value=f"Filter: Year = {year_filter} | Branch = {branch_filter} | Total: {len(df)} students")
                    ws.cell(row=7, column=1).font = Font(italic=True, size=9, color='555555')
                    # Style data header row
                    header_row = header_end_row + 1
                    header_fill = PatternFill(start_color='1F4788', end_color='1F4788', fill_type='solid')
                    for col_idx in range(1, len(df.columns) + 1):
                        cell = ws.cell(row=header_row, column=col_idx)
                        cell.font = Font(bold=True, color='FFFFFF', size=10)
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                    self._auto_adjust_column_width(ws)
                    
                messagebox.showinfo("Success",
                    f"✅ Exported {len(df)} students to:\n{file_path}\n\n"
                    f"Year: {year_filter} | Branch: {branch_filter}")
                if messagebox.askyesno("Open File", "Do you want to open the exported file?"):
                    self.open_file(file_path)
                    
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export students: {str(e)}")

    def promote_student_years(self):
        """Enhanced promotion with letter number, academic year creation, and history tracking"""
        # Create promotion dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Promote Student Years")
        dialog.geometry("600x500")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 100, self.root.winfo_rooty() + 100))
        
        # Title
        title_label = tk.Label(
            dialog,
            text="🎓 Promote Student Years",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack(pady=(20, 10))
        
        # Info text
        info_label = tk.Label(
            dialog,
            text="This will promote: 1st→2nd, 2nd→3rd, 3rd→Pass Out",
            font=('Segoe UI', 10),
            bg='white',
            fg=self.colors['text']
        )
        info_label.pack(pady=(0, 20))
        
        # Form frame
        form_frame = tk.Frame(dialog, bg='white')
        form_frame.pack(expand=True, padx=40)
        
        # Letter Number field
        letter_label = tk.Label(
            form_frame,
            text="Letter Number:",
            font=('Segoe UI', 11, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        letter_label.grid(row=0, column=0, sticky='w', pady=(0, 15), padx=(0, 20))
        
        letter_entry = tk.Entry(
            form_frame,
            font=('Segoe UI', 11),
            width=30,
            relief='solid',
            bd=2
        )
        letter_entry.grid(row=0, column=1, pady=(0, 15))
        
        letter_hint = tk.Label(
            form_frame,
            text="(Can contain letters, numbers, and symbols)",
            font=('Segoe UI', 8, 'italic'),
            bg='white',
            fg='#666'
        )
        letter_hint.grid(row=1, column=1, sticky='w', pady=(0, 15))
        
        # Academic Year field
        year_label = tk.Label(
            form_frame,
            text="Academic Year:",
            font=('Segoe UI', 11, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        year_label.grid(row=2, column=0, sticky='w', pady=(0, 15), padx=(0, 20))
        
        # Auto-suggest NEXT academic year for teacher convenience
        current_year = datetime.now().year
        # If we're in July-December, suggest next year
        # If we're in January-June, suggest current year
        if datetime.now().month >= 7:  # July onwards = next academic year
            suggested_year = f"{current_year + 1}-{current_year + 2}"
        else:  # January-June = current academic year
            suggested_year = f"{current_year}-{current_year + 1}"
        
        # Get all academic years from database
        all_academic_years = self.db.get_all_academic_years()
        
        # Create a list with suggested year first, then existing years
        year_options = []
        if suggested_year not in all_academic_years:
            year_options.append(suggested_year)
        year_options.extend(all_academic_years)
        
        # Also add next few years for convenience
        for i in range(1, 4):
            future_year = f"{current_year + i}-{current_year + i + 1}"
            if future_year not in year_options:
                year_options.append(future_year)
        
        # Create combobox for academic year selection
        year_entry = ttk.Combobox(
            form_frame,
            font=('Segoe UI', 11),
            width=28,
            values=year_options,
            state='normal'  # Allow typing custom years
        )
        year_entry.set(suggested_year)
        year_entry.grid(row=2, column=1, pady=(0, 15))
        
        year_hint = tk.Label(
            form_frame,
            text=f"(Suggested: {suggested_year} - You can type custom year)",
            font=('Segoe UI', 8, 'italic'),
            bg='white',
            fg='#666'
        )
        year_hint.grid(row=3, column=1, sticky='w', pady=(0, 20))
        
        # Buttons frame
        btn_frame = tk.Frame(dialog, bg='white')
        btn_frame.pack(pady=20)
        
        def perform_promotion():
            letter_number = letter_entry.get().strip()
            academic_year = year_entry.get().strip()
            
            if not letter_number:
                messagebox.showerror("Error", "Letter Number is required!")
                return
            
            if not academic_year:
                messagebox.showerror("Error", "Academic Year is required!")
                return
            
            # Confirm promotion
            if not messagebox.askyesno(
                "Confirm Promotion",
                f"This will promote all students with:\n\nLetter Number: {letter_number}\nAcademic Year: {academic_year}\n\nProceed?"
            ):
                return
            
            def _norm(label: str) -> str:
                if not label:
                    return '1st'
                s = str(label).strip().lower()
                # Accept more variants for year labels
                if s in ('1st', 'first', 'first year', '1', 'i', 'fy', 'f.y', 'f.y.', 'fe', 'fe year', '1st year', '1styear'):
                    return '1st'
                if s in ('2nd', 'second', 'second year', '2', 'ii', 'sy', 's.y', 's.y.', '2nd year', '2ndyear'):
                    return '2nd'
                if s in ('3rd', 'third', 'third year', '3', 'iii', 'ty', 't.y', 't.y.', '3rd year', '3rdyear'):
                    return '3rd'
                if 'pass' in s:
                    return 'Pass Out'
                return label

            def _promote_once(label: str) -> str:
                l = _norm(label)
                return {'1st': '2nd', '2nd': '3rd', '3rd': 'Pass Out'}.get(l, l)

            try:
                # Create/activate academic year
                self.db.create_academic_year(academic_year)
                
                conn = self.db.get_connection()
                cur = conn.cursor()
                cur.execute('SELECT enrollment_no, name, year FROM students')
                rows = cur.fetchall()
                c_12 = c_23 = c_3p = c_skip = 0
                
                # Collect promotion records and batch updates
                promotion_records = []
                batch_updates = []
                
                for en, name, yr in rows:
                    new = _promote_once(yr)
                    if new != yr:
                        batch_updates.append((new, en))
                        promotion_records.append((en, name, yr, new, letter_number, academic_year))
                        
                        if _norm(yr) == '1st' and new == '2nd':
                            c_12 += 1
                        elif _norm(yr) == '2nd' and new == '3rd':
                            c_23 += 1
                        elif _norm(yr) == '3rd' and new == 'Pass Out':
                            c_3p += 1
                    else:
                        c_skip += 1
                
                # Batch UPDATE: single executemany instead of N individual queries
                if batch_updates:
                    cur.executemany('UPDATE students SET year=? WHERE enrollment_no=?', batch_updates)
                
                conn.commit()
                conn.close()
                
                # Now add all promotion history records (after closing the main connection)
                for record in promotion_records:
                    self.db.add_promotion_history(*record)
                
                # Calculate total BEFORE using it
                total = c_12 + c_23 + c_3p
                
                # Log the promotion activity
                self._log_admin_activity(
                    "Student Year Promotion",
                    f"Promoted {total} students. 1st→2nd: {c_12}, 2nd→3rd: {c_23}, 3rd→Pass Out: {c_3p}. Letter: {letter_number}, Academic Year: {academic_year}"
                )
                
                msg = (
                    f"Promotion Complete!\n\n"
                    f"Updated Year for {total} student(s).\n\n"
                    f"1st → 2nd: {c_12}\n"
                    f"2nd → 3rd: {c_23}\n"
                    f"3rd → Pass Out: {c_3p}\n"
                    f"Unchanged: {c_skip}\n\n"
                    f"Letter Number: {letter_number}\n"
                    f"Academic Year: {academic_year}"
                )
                messagebox.showinfo("Success", msg)
                dialog.destroy()
                
                try:
                    self.refresh_students()
                    self.refresh_dashboard()
                    self.refresh_academic_year_filter()  # Refresh academic year dropdown
                except Exception:
                    pass
                    
            except Exception as e:
                messagebox.showerror("Error", f"Promotion failed: {e}")
        
        def undo_last():
            """Undo the last promotion"""
            # First check if there is anything to undo
            history = self.db.get_promotion_history()
            if not history:
                messagebox.showwarning("No History", "No promotion history found. Nothing to undo.")
                return
            
            # Show details of what will be undone
            last_promo = history[0]
            details_msg = (
                f"Undo the last promotion?\n\n"
                f"Last Promotion Details:\n"
                f"Student: {last_promo['student_name']}\n"
                f"Changed: {last_promo['old_year']} → {last_promo['new_year']}\n"
                f"Date: {last_promo['promotion_date']}\n\n"
                f"This will revert ALL students from the most recent promotion batch."
            )
            
            if messagebox.askyesno("Confirm Undo", details_msg):
                success, message = self.db.undo_last_promotion()
                if success:
                    messagebox.showinfo("Success", message)
                    try:
                        self.refresh_students()
                        self.refresh_dashboard()
                        self.refresh_academic_year_filter()  # Refresh academic year dropdown
                    except Exception:
                        pass
                else:
                    messagebox.showerror("Error", message)
        
        def view_history():
            """View promotion history"""
            self.show_promotion_history_dialog()
        
        # Promote button
        promote_btn = tk.Button(
            btn_frame,
            text="🎓 Promote Students",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['accent'],
            fg='white',
            command=perform_promotion,
            relief='flat',
            padx=20,
            pady=10,
            cursor='hand2'
        )
        promote_btn.pack(side=tk.LEFT, padx=5)
        
        # Undo button
        undo_btn = tk.Button(
            btn_frame,
            text="↩️ Undo Last",
            font=('Segoe UI', 11),
            bg='#FF9800',
            fg='white',
            command=undo_last,
            relief='flat',
            padx=20,
            pady=10,
            cursor='hand2'
        )
        undo_btn.pack(side=tk.LEFT, padx=5)
        
        # History button
        history_btn = tk.Button(
            btn_frame,
            text="📜 History",
            font=('Segoe UI', 11),
            bg=self.colors['secondary'],
            fg='white',
            command=view_history,
            relief='flat',
            padx=20,
            pady=10,
            cursor='hand2'
        )
        history_btn.pack(side=tk.LEFT, padx=5)
        
        # Cancel button
        cancel_btn = tk.Button(
            btn_frame,
            text="❌ Cancel",
            font=('Segoe UI', 11),
            bg='#666',
            fg='white',
            command=dialog.destroy,
            relief='flat',
            padx=20,
            pady=10,
            cursor='hand2'
        )
        cancel_btn.pack(side=tk.LEFT, padx=5)
    
    def show_promotion_history_dialog(self):
        """Show promotion history summary (Letter Number, Date, Time, Student Count)"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Promotion History")
        dialog.geometry("800x600")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 50, self.root.winfo_rooty() + 50))
        
        # Title
        title_label = tk.Label(
            dialog,
            text="📜 Promotion History Summary",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack(pady=(20, 10))
        
        # Tree frame
        tree_frame = tk.Frame(dialog, bg='white')
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # Scrollbars
        v_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        h_scroll = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)
        
        # Treeview - Simplified columns
        columns = ('Letter Number', 'Date', 'Time', '1st→2nd', '2nd→3rd', '3rd→Pass Out', 'Total Students')
        tree = ttk.Treeview(
            tree_frame,
            columns=columns,
            show='headings',
            yscrollcommand=v_scroll.set,
            xscrollcommand=h_scroll.set,
            height=15
        )
        
        v_scroll.config(command=tree.yview)
        h_scroll.config(command=tree.xview)
        
        # Configure columns
        column_widths = [150, 100, 80, 80, 80, 100, 100]
        for col, width in zip(columns, column_widths):
            tree.heading(col, text=col)
            tree.column(col, width=width, anchor='center')
        
        # Grid layout
        tree.grid(row=0, column=0, sticky='nsew')
        v_scroll.grid(row=0, column=1, sticky='ns')
        h_scroll.grid(row=1, column=0, sticky='ew')
        
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Load and process data
        history = self.db.get_promotion_history()
        
        # Group by promotion date/time and letter number
        from collections import defaultdict
        grouped = defaultdict(lambda: {'1st→2nd': 0, '2nd→3rd': 0, '3rd→Pass Out': 0, 'letter': '', 'datetime': ''})
        
        for record in history:
            # record: (enrollment_no, student_name, old_year, new_year, letter_number, academic_year, promotion_date)
            old_year = record[2]
            new_year = record[3]
            letter_num = record[4]
            promo_datetime = record[6]
            
            key = (letter_num, promo_datetime)
            grouped[key]['letter'] = letter_num
            grouped[key]['datetime'] = promo_datetime
            
            # Count transitions
            if old_year == '1st' and new_year == '2nd':
                grouped[key]['1st→2nd'] += 1
            elif old_year == '2nd' and new_year == '3rd':
                grouped[key]['2nd→3rd'] += 1
            elif old_year == '3rd' and new_year == 'Pass Out':
                grouped[key]['3rd→Pass Out'] += 1
        
        # Display grouped data
        summary_data = []
        for key, data in grouped.items():
            try:
                # Parse datetime
                dt = datetime.strptime(data['datetime'], '%Y-%m-%d %H:%M:%S')
                date_str = dt.strftime('%Y-%m-%d')
                time_str = dt.strftime('%H:%M:%S')
            except:
                date_str = data['datetime']
                time_str = ''
            
            total = data['1st→2nd'] + data['2nd→3rd'] + data['3rd→Pass Out']
            
            row = (
                data['letter'],
                date_str,
                time_str,
                data['1st→2nd'],
                data['2nd→3rd'],
                data['3rd→Pass Out'],
                total
            )
            summary_data.append(row)
            tree.insert('', 'end', values=row)
        
        # Button frame
        btn_frame = tk.Frame(dialog, bg='white')
        btn_frame.pack(pady=20)
        
        def download_history():
            """Download promotion history summary as Excel file"""
            if not summary_data:
                messagebox.showinfo("Info", "No promotion history to download")
                return
            
            try:
                # Create DataFrame
                df = pd.DataFrame(summary_data, columns=columns)
                
                # Save to file
                filename = f"promotion_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                file_path = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[("Excel files", "*.xlsx")],
                    initialfile=filename
                )
                
                if file_path:
                    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                        sheet = 'Promotion Summary'
                        df.to_excel(writer, sheet_name=sheet, index=False, startrow=4)
                        ws = writer.book[sheet]
                        self._write_excel_header_openpyxl(ws, start_row=1)
                        self._auto_adjust_column_width(ws)
                    
                    messagebox.showinfo("Success", f"Promotion history exported to {file_path}")
                    
                    if messagebox.askyesno("Open File", "Do you want to open the exported file?"):
                        self.open_file(file_path)
                        
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export: {str(e)}")
        
        # Download button
        download_btn = tk.Button(
            btn_frame,
            text="📥 Download History",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['accent'],
            fg='white',
            command=download_history,
            relief='flat',
            padx=20,
            pady=10,
            cursor='hand2'
        )
        download_btn.pack(side=tk.LEFT, padx=5)
        
        # Close button
        close_btn = tk.Button(
            btn_frame,
            text="❌ Close",
            font=('Segoe UI', 11),
            bg='#666',
            fg='white',
            command=dialog.destroy,
            relief='flat',
            padx=20,
            pady=10,
            cursor='hand2'
        )
        close_btn.pack(side=tk.LEFT, padx=5)
    
    def export_books_to_excel(self):
        """Export books to Excel"""
        try:
            books = self.db.get_books()
            
            # Defensive validation: check for missing Book ID or Title
            invalid_rows = []
            for b in books:
                try:
                    book_id_val = str(b[1]).strip()
                    title_val = str(b[2]).strip()
                    if not book_id_val or not title_val:
                        invalid_rows.append(b)
                except Exception:
                    invalid_rows.append(b)
            if invalid_rows:
                if not messagebox.askyesno(
                    "Validation Warning",
                    f"Detected {len(invalid_rows)} book record(s) with missing Book ID or Title. Continue export?"
                ):
                    return
            
            # Export all books
            filtered_books = []
            for book in books:
                filtered_books.append({
                    'Book ID': book[1],
                    'Title': book[2],
                    'Author': book[3],
                    'ISBN': book[4],
                    'Category': book[5],
                    'Total Copies': book[6],
                    'Available Copies': book[7],
                    'Date Added': book[8]
                })
            
            if not filtered_books:
                messagebox.showwarning("Warning", "No books found to export!")
                return
            
            # Create DataFrame
            df = pd.DataFrame(filtered_books)
            
            # Save to file
            filename = f"books_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                # Correct parameter name
                initialfile=filename
            )
            
            if file_path:
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    sheet = 'Books'
                    header_end_row = 6  # Enhanced header takes 6 rows
                    df.to_excel(writer, sheet_name=sheet, index=False, startrow=header_end_row)
                    ws = writer.book[sheet]
                    self._write_excel_header_openpyxl(ws, start_row=1)
                    # Apply enhanced auto-adjust with styling
                    self._auto_adjust_column_width(ws)
                messagebox.showinfo("Success", f"Books data exported to {file_path}")
                
                # Ask if user wants to open the file
                if messagebox.askyesno("Open File", "Do you want to open the exported file?"):
                    self.open_file(file_path)
                    
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export books: {str(e)}")
    
    def export_records_to_excel(self):
        """Export records to Excel with branch/year/status filter dialog"""
        try:
            if not hasattr(self, 'records_tree'):
                messagebox.showerror("Error", "Records view not initialized yet.")
                return

            # ------ Filter dialog -------
            dlg = tk.Toplevel(self.root)
            dlg.title("Export Records - Filter Options")
            dlg.geometry("420x330")
            dlg.configure(bg='white')
            dlg.transient(self.root)
            dlg.grab_set()
            dlg.resizable(False, False)
            dlg.geometry("+%d+%d" % (self.root.winfo_rootx() + 200, self.root.winfo_rooty() + 150))

            tk.Label(dlg, text="📊 Export Transaction Records",
                     font=('Segoe UI', 14, 'bold'), bg='white',
                     fg=self.colors['accent']).pack(pady=(20, 5))
            tk.Label(dlg, text="Select filters for export (leave All for everything):",
                     font=('Segoe UI', 10), bg='white', fg='#555').pack()

            form = tk.Frame(dlg, bg='white', padx=30)
            form.pack(pady=15, fill=tk.X)

            # Branch filter
            tk.Label(form, text="Branch:", font=('Segoe UI', 10, 'bold'),
                     bg='white', fg=self.colors['accent']).grid(row=0, column=0, sticky='w', pady=6)
            export_branch_var = tk.StringVar(value="All")
            branch_opts = ["All", "Computer Engineering", "Information Technology",
                           "Civil Engineering", "Electronics & Telecommunication",
                           "Mechanical Engineering", "Automobile Engineering", "Electrical Engineering"]
            ttk.Combobox(form, textvariable=export_branch_var, values=branch_opts,
                         state="readonly", width=28).grid(row=0, column=1, padx=10, pady=6)

            # Year filter
            tk.Label(form, text="Student Year:", font=('Segoe UI', 10, 'bold'),
                     bg='white', fg=self.colors['accent']).grid(row=1, column=0, sticky='w', pady=6)
            export_year_var = tk.StringVar(value="All")
            ttk.Combobox(form, textvariable=export_year_var,
                         values=["All", "1st", "2nd", "3rd", "Pass Out"],
                         state="readonly", width=28).grid(row=1, column=1, padx=10, pady=6)

            # Status filter
            tk.Label(form, text="Status:", font=('Segoe UI', 10, 'bold'),
                     bg='white', fg=self.colors['accent']).grid(row=2, column=0, sticky='w', pady=6)
            export_status_var = tk.StringVar(value="All")
            ttk.Combobox(form, textvariable=export_status_var,
                         values=["All", "Issued (Borrowed)", "Returned", "Overdue"],
                         state="readonly", width=28).grid(row=2, column=1, padx=10, pady=6)

            confirmed = [False]
            def on_export_confirm():
                confirmed[0] = True
                dlg.destroy()

            btn_frame = tk.Frame(dlg, bg='white')
            btn_frame.pack(pady=10)
            tk.Button(btn_frame, text="📥 Export", command=on_export_confirm,
                      bg=self.colors['secondary'], fg='white',
                      font=('Segoe UI', 10, 'bold'), padx=20, pady=8,
                      relief='flat', cursor='hand2').pack(side=tk.LEFT, padx=5)
            tk.Button(btn_frame, text="Cancel", command=dlg.destroy,
                      bg='#6c757d', fg='white', font=('Segoe UI', 10),
                      padx=15, pady=8, relief='flat', cursor='hand2').pack(side=tk.LEFT, padx=5)

            self.root.wait_window(dlg)
            if not confirmed[0]:
                return

            sel_branch = export_branch_var.get()
            sel_year = export_year_var.get()
            sel_status = export_status_var.get()

            # ------ Collect visible records from tree (10 columns now) -------
            all_tree_records = []
            for item_id in self.records_tree.get_children():
                vals = list(self.records_tree.item(item_id, 'values'))
                while len(vals) < 10:
                    vals.append('')
                all_tree_records.append(vals[:10])

            if not all_tree_records:
                messagebox.showwarning("Warning", "No records to export (list is empty)!")
                return

            # Apply additional dialog filters
            filtered = []
            for v in all_tree_records:
                # v[2]=Branch, v[8]=Status, v[9]=Fine
                if sel_branch != "All" and sel_branch.lower() not in str(v[2]).lower():
                    continue
                if sel_status != "All":
                    st = str(v[8]).lower()
                    if sel_status == "Issued (Borrowed)" and st != 'borrowed':
                        continue
                    elif sel_status == "Returned" and st != 'returned':
                        continue
                    elif sel_status == "Overdue":
                        fine_raw = str(v[9]).replace('Rs', '').replace('(Late)', '').strip()
                        try:
                            if int(fine_raw) <= 0:
                                continue
                        except Exception:
                            continue
                filtered.append(v)

            if not filtered:
                messagebox.showwarning("Warning", "No records match the selected filters!")
                return

            df = pd.DataFrame(filtered, columns=[
                'Enrollment No', 'Student Name', 'Branch', 'Book ID', 'Book Title',
                'Issue Date', 'Due Date', 'Return Date', 'Status', 'Fine'
            ])

            # Year filter (joins with student DB)
            if sel_year != "All":
                try:
                    students = self.db.get_students()
                    yr_map = {str(s[1]): str(s[6]) for s in students}
                    mask = df['Enrollment No'].astype(str).map(yr_map) == sel_year
                    df = df[mask]
                except Exception:
                    pass

            filter_desc = f"Branch: {sel_branch} | Year: {sel_year} | Status: {sel_status}"
            filename = f"records_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                initialfile=filename
            )

            if file_path:
                from openpyxl.styles import Font, Alignment, PatternFill
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    sheet = 'Records'
                    header_end_row = 7
                    df.to_excel(writer, sheet_name=sheet, index=False, startrow=header_end_row)
                    ws = writer.book[sheet]
                    self._write_excel_header_openpyxl(ws, start_row=1)
                    ws.cell(row=7, column=1, value=f"Filter Applied: {filter_desc}")
                    ws.cell(row=7, column=1).font = Font(italic=True, size=9, color='555555')
                    header_row = header_end_row + 1
                    header_fill = PatternFill(start_color='1F4788', end_color='1F4788', fill_type='solid')
                    for col_idx in range(1, len(df.columns) + 1):
                        cell = ws.cell(row=header_row, column=col_idx)
                        cell.font = Font(bold=True, color='FFFFFF', size=10)
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                    self._auto_adjust_column_width(ws)
                messagebox.showinfo("Success",
                    f"✅ Exported {len(df)} records to:\n{file_path}\n\nFilter: {filter_desc}")
                if messagebox.askyesno("Open File", "Do you want to open the exported file?"):
                    self.open_file(file_path)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to export records: {str(e)}")

    # ------------------------------------------------------------------
    # Overdue Notice Letter Export
    # ------------------------------------------------------------------
    def get_current_overdue_records(self):
        """Return list of overdue (currently issued and past due date) records.
        Each record dict with: Enrollment No, Student Name, Book ID, Book Title, Issue Date, Due Date, Days Overdue, Accrued Fine
        """
        overdue = []
        try:
            all_records = self.get_all_records()
            today = datetime.now().date()
            from datetime import datetime as _dt
            for rec in all_records:
                # record tuple: (enroll[0], name[1], dept[2], book_id[3], title[4],
                #                borrow_date[5], due_date[6], return_date[7], status[8], fine[9], academic_year[10])
                if len(rec) >= 11:
                    enroll = rec[0]; name = rec[1]; dept = rec[2]
                    book_id = rec[3]; title = rec[4]
                    borrow_date = rec[5]; due_date = rec[6]
                    status = rec[8]
                elif len(rec) >= 9:
                    enroll = rec[0]; name = rec[1]; dept = ''
                    book_id = rec[2]; title = rec[3]
                    borrow_date = rec[4]; due_date = rec[5]
                    status = rec[7]
                else:
                    continue
                    
                if status == 'borrowed':
                    try:
                        due_d = _dt.strptime(due_date, '%Y-%m-%d').date()
                        if due_d < today:
                            days_overdue = (today - due_d).days
                            overdue.append({
                                'Enrollment No': enroll,
                                'Student Name': name,
                                'Branch': dept,
                                'Book ID': book_id,
                                'Book Title': title,
                                'Issue Date': borrow_date,
                                'Due Date': due_date,
                                'Days Overdue': days_overdue,
                                'Accrued Fine': int(days_overdue) * self.get_fine_per_day()
                            })
                    except Exception as e:
                        print(f"Error processing overdue record: {e}")
                        continue
            return overdue
        except Exception as e:
            print(f"Overdue fetch error: {e}")
            return []

    def export_overdue_notice_letter(self):
        """Generate overdue notice as Excel (.xlsx) + plain text (.txt) letter.
        Excel sheet contains letter heading, body, table of overdue books, and closing with signature.
        """
        try:
            overdue = self.get_current_overdue_records()
            if not overdue:
                messagebox.showinfo("Overdue Notice", "There are currently no overdue issued books.")
                return

            today_str = datetime.now().strftime('%Y-%m-%d')
            ref_code = f"Library/Overdue/{datetime.now().strftime('%Y%m%d')}"

            body_lines = [
                "Government Polytechnic Awasari (Kh)",
                "Main Library",
                "", f"Date: {today_str}", f"Ref: {ref_code}", "",
                "Subject: Submission of Overdue Library Books",
                "", "Dear Students,", "",
                "The following students are hereby notified to immediately submit the listed library books that are now overdue.",
                f"A fine of Rs {self.get_fine_per_day()} per day has accrued (or will continue to accrue) until the books are returned.",
                "Failure to comply today will trigger disciplinary action per departmental policy.",
                "", "Overdue Book List:", ""
            ]

            df = pd.DataFrame(overdue)
            base_filename = f"overdue_notice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                initialfile=base_filename + '.xlsx'
            )
            if not file_path:
                return

            try:
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    sheet_name = 'Overdue Notice'
                    
                    # Add institutional header with logo first
                    ws = writer.book.create_sheet(sheet_name)
                    header_rows = self._write_excel_header_openpyxl(ws, start_row=1)
                    
                    # Add date and ref below header
                    current_row = header_rows + 1
                    ws.cell(row=current_row, column=1, value=f"Date: {today_str}")
                    ws.cell(row=current_row + 1, column=1, value=f"Ref: {ref_code}")
                    current_row += 3
                    
                    # Add subject and body text
                    from openpyxl.styles import Font
                    ws.cell(row=current_row, column=1, value="Subject: Submission of Overdue Library Books").font = Font(bold=True)
                    current_row += 2
                    ws.cell(row=current_row, column=1, value="Dear Students,")
                    current_row += 2
                    ws.cell(row=current_row, column=1, value="The following students are hereby notified to immediately submit the listed library books that are now overdue.")
                    current_row += 1
                    ws.cell(row=current_row, column=1, value=f"A fine of Rs {self.get_fine_per_day()} per day has accrued (or will continue to accrue) until the books are returned.")
                    current_row += 1
                    ws.cell(row=current_row, column=1, value="Failure to comply today will trigger disciplinary action per departmental policy.")
                    current_row += 2
                    ws.cell(row=current_row, column=1, value="Overdue Book List:").font = Font(bold=True)
                    current_row += 2
                    
                    # Write table data
                    startrow = current_row
                    df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=startrow)
                    
                    # Auto-adjust columns for the table
                    self._auto_adjust_column_width(ws)
                    
                    # Add closing text
                    closing_start = startrow + len(df) + 3
                    closing_lines = [
                        "", "You are directed to return the above books without further delay.",
                        "Punishment Clause: Continued non-compliance after 3 days from this notice will result in suspension of borrowing privileges for one month and a formal report to the Academic Coordinator.",
                        "", "Regards,", "", "__________________________", "Librarian", "Main Library",
                        "Government Polytechnic Awasari (Kh)"
                    ]
                    for offset, line in enumerate(closing_lines):
                        ws.cell(row=closing_start + offset, column=1, value=line)
            except Exception as e:
                messagebox.showwarning("Excel Write Warning", f"Excel formatting fallback due to: {e}\nCreating plain text letter only.")

            # Create companion text letter
            try:
                txt_path = os.path.splitext(file_path)[0] + '.txt'
                cols = list(df.columns)
                col_widths = [max(len(c), *(len(str(v)) for v in df[c].astype(str))) for c in cols]
                def fmt_row(row_vals):
                    return "  ".join(str(v).ljust(w) for v, w in zip(row_vals, col_widths))
                table_lines = [fmt_row(cols), fmt_row(['-'*w for w in col_widths])] + [fmt_row(df.iloc[i]) for i in range(len(df))]
                closing_text = [
                    "", "You are directed to return the above books without further delay.",
                    "Punishment Clause: Continued non-compliance after 3 days from this notice will result in suspension of borrowing privileges for one month and a formal report to the Academic Coordinator.",
                    "", "Regards,", "", "__________________________", "Librarian", "Main Library", "Government Polytechnic Awasari (Kh)"
                ]
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write("\n".join(body_lines + table_lines + closing_text))
            except Exception as e:
                print(f"Text companion creation failed: {e}")

            messagebox.showinfo("Overdue Notice", f"Overdue notice letter exported:\n{file_path}")
            if messagebox.askyesno("Open File", "Open the Excel letter now?"):
                self.open_file(file_path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate overdue notice: {e}")

    def export_overdue_notice_letter_word(self):
        """Generate the overdue notice as a formatted Microsoft Word (.docx) document.
        Provides proper headings, paragraphs, table, punishment clause, and signature block.
        Now also sends emails to all overdue students automatically.
        """
        try:
            overdue = self.get_current_overdue_records()
            if not overdue:
                messagebox.showinfo("Overdue Notice", "There are currently no overdue issued books.")
                return
            if Document is None:
                messagebox.showerror(
                    "Dependency Missing",
                    "python-docx is not installed. Please install it (pip install python-docx) and try again."
                )
                return
            
            # Ask if user wants to send emails to all overdue students
            send_emails = False
            _sequential_needed = False
            if self.email_settings.get('enabled', False):
                send_emails = messagebox.askyesno(
                    "Send Emails?",
                    f"Found {len(overdue)} overdue student(s).\n\n"
                    "Do you want to send overdue notice emails to ALL of them automatically?\n\n"
                    "✉️ Email will be sent to each student with their overdue letter attached.",
                    icon='question'
                )
            
            today_str = datetime.now().strftime('%Y-%m-%d')
            ref_code = f"Library/Overdue/{datetime.now().strftime('%Y%m%d')}"
            base_filename = f"overdue_notice_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile=base_filename + '.docx'
            )
            if not file_path:
                return

            doc = Document()
            
            # Add logo at the top if available
            logo_path = os.path.join(os.path.dirname(__file__), 'logo.png')
            if hasattr(sys, '_MEIPASS'):
                logo_path = os.path.join(sys._MEIPASS, 'logo.png')
            if os.path.exists(logo_path):
                try:
                    # Add logo centered at top
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(logo_path, width=Pt(60))
                except Exception as e:
                    print(f"Could not add logo: {e}")
            
            def add_center(text, bold=True, size=16, color=None):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(size)
                if color:
                    from docx.shared import RGBColor
                    run.font.color.rgb = RGBColor(*color)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Main heading - VERY LARGE with dark blue color
            add_center("Government Polytechnic Awasari (Kh)", True, 22, (31, 71, 136))
            # Subheading - Large with medium blue color
            add_center("Main Library", True, 18, (46, 92, 138))
            
            # Add a line separator
            doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date and Ref aligned to the right
            meta_p = doc.add_paragraph()
            meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            meta_p.add_run(f"Date: {today_str}\nRef: {ref_code}")
            subj = doc.add_paragraph()
            run_subj = subj.add_run("Subject: Submission of Overdue Library Books")
            run_subj.bold = True
            body_paras = [
                "Dear Students,",
                "The following students are hereby notified to immediately submit the listed library books that are now overdue.",
                f"A fine of Rs {self.get_fine_per_day()} per day has accrued (or will continue to accrue) until the books are returned.",
                "Failure to comply today will trigger disciplinary action per departmental policy.",
                "Overdue Book List:"
            ]
            for line in body_paras:
                doc.add_paragraph(line)

            columns = ["Enrollment No", "Student Name", "Branch", "Book ID", "Book Title", "Issue Date", "Due Date", "Days Overdue", "Accrued Fine"]
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = 'Light Grid Accent 1'  # Professional table style with borders
            
            # Header row - bold and formatted
            hdr = table.rows[0].cells
            for i, col in enumerate(columns):
                hdr[i].text = col
                # Make header bold
                for paragraph in hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            
            # Data rows
            for rec in overdue:
                row = table.add_row().cells
                row[0].text = str(rec['Enrollment No'])
                row[1].text = str(rec['Student Name'])
                row[2].text = str(rec.get('Branch', ''))
                row[3].text = str(rec['Book ID'])
                row[4].text = str(rec['Book Title'])
                row[5].text = str(rec['Issue Date'])
                row[6].text = str(rec['Due Date'])
                row[7].text = str(rec['Days Overdue'])
                row[8].text = "Rs " + str(rec['Accrued Fine'])  # Add Rs prefix for clarity

            doc.add_paragraph()
            clause = doc.add_paragraph()
            clause.add_run(
                "You are directed to return the above books without further delay.\n"
                "Punishment Clause: Continued non-compliance after 3 days from this notice will result in suspension "
                "of borrowing privileges for one month and a formal report to the Academic Coordinator."
            )
            doc.add_paragraph()
            doc.add_paragraph("Regards,")
            doc.add_paragraph()
            doc.add_paragraph("__________________________")
            doc.add_paragraph("Librarian")
            doc.add_paragraph("Main Library")
            doc.add_paragraph("Government Polytechnic Awasari (Kh)")

            doc.save(file_path)
            
            # Send emails to all overdue students if requested
            email_results = []
            sent_count = 0
            failed_count = 0
            _sequential_needed = False
            if send_emails:
                import tempfile
                
                # Use batch email service if available, otherwise use old method
                if self.email_batch_service and PERFORMANCE_MODULES_AVAILABLE:
                    # Prepare batch of emails using performance-optimized service
                    progress_win = tk.Toplevel(self.root)
                    progress_win.title("Sending Emails...")
                    progress_win.geometry("500x250")
                    progress_win.transient(self.root)
                    progress_win.grab_set()
                    
                    # Center the window
                    progress_win.update_idletasks()
                    x = (progress_win.winfo_screenwidth() // 2) - (250)
                    y = (progress_win.winfo_screenheight() // 2) - (125)
                    progress_win.geometry(f"+{x}+{y}")
                    
                    label = tk.Label(progress_win, text="⚡ Sending overdue emails in parallel batches...", 
                                   font=('Segoe UI', 12, 'bold'), pady=20)
                    label.pack()
                    
                    # Progress bar
                    from tkinter import ttk
                    progress_bar = ttk.Progressbar(progress_win, length=400, mode='determinate')
                    progress_bar.pack(pady=10)
                    
                    status_label = tk.Label(progress_win, text="Preparing...", font=('Segoe UI', 10), 
                                          wraplength=450, justify=tk.LEFT)
                    status_label.pack(pady=10)
                    
                    stats_label = tk.Label(progress_win, text="", font=('Segoe UI', 9, 'italic'), 
                                          fg='#666', wraplength=450, justify=tk.LEFT)
                    stats_label.pack(pady=5)
                    
                    progress_win.update()
                    
                    # Prepare emails
                    emails_to_send = []
                    temp_files = []
                    
                    for rec in overdue:
                        enrollment_no = str(rec['Enrollment No'])
                        student_name = str(rec['Student Name'])
                        book_id = str(rec['Book ID'])
                        book_title = str(rec['Book Title'])
                        issue_date = str(rec['Issue Date'])
                        due_date = str(rec['Due Date'])
                        days_overdue = str(rec['Days Overdue'])
                        fine = str(rec['Accrued Fine'])
                        
                        student_email = self.get_student_email(enrollment_no)
                        
                        if student_email:
                            # Generate individual letter
                            temp_doc = Document()
                            
                            # Add header
                            if os.path.exists(logo_path):
                                try:
                                    logo_para = temp_doc.add_paragraph()
                                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    logo_run = logo_para.add_run()
                                    logo_run.add_picture(logo_path, width=Pt(80))
                                except:
                                    pass
                            
                            def add_center_temp(text, bold=True, size=16, color=None):
                                p = temp_doc.add_paragraph()
                                run = p.add_run(text)
                                run.bold = bold
                                run.font.size = Pt(size)
                                if color:
                                    from docx.shared import RGBColor
                                    run.font.color.rgb = RGBColor(*color)
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            add_center_temp("Government Polytechnic Awasari (Kh)", True, 20, (31, 71, 136))
                            add_center_temp("Main Library", True, 16, (46, 92, 138))
                            temp_doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
                            temp_doc.add_paragraph()
                            
                            date_para = temp_doc.add_paragraph()
                            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
                            temp_doc.add_paragraph()
                            
                            subject = temp_doc.add_paragraph()
                            subject.add_run('Subject: Overdue Book Notice').bold = True
                            subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            temp_doc.add_paragraph()
                            
                            to_para = temp_doc.add_paragraph()
                            to_para.add_run(f'To,\\n{student_name}\\nEnrollment No: {enrollment_no}')
                            temp_doc.add_paragraph()
                            
                            body = temp_doc.add_paragraph()
                            body.add_run(
                                f"Dear {student_name},\\n\\n"
                                f"This is to inform you that the following book borrowed from GPA's Library Management System "
                                f"is overdue and needs to be returned immediately.\\n\\n"
                            )
                            
                            temp_doc.add_paragraph('Book Details:', style='Heading 2')
                            details_table = temp_doc.add_table(rows=5, cols=2)
                            details_table.style = 'Light Grid Accent 1'
                            details_table.cell(0, 0).text = 'Book ID:'
                            details_table.cell(0, 1).text = book_id
                            details_table.cell(1, 0).text = 'Book Title:'
                            details_table.cell(1, 1).text = book_title
                            details_table.cell(2, 0).text = 'Issue Date:'
                            details_table.cell(2, 1).text = issue_date
                            details_table.cell(3, 0).text = 'Due Date:'
                            details_table.cell(3, 1).text = due_date
                            details_table.cell(4, 0).text = 'Days Overdue:'
                            details_table.cell(4, 1).text = days_overdue
                            
                            temp_doc.add_paragraph()
                            fine_para = temp_doc.add_paragraph()
                            fine_run = fine_para.add_run(
                                f"As per library rules, a fine of ₹{self.get_fine_per_day()} per day is applicable for overdue books.\\n"
                                f"Your current fine amount is: ₹{fine}\\n\\n"
                            )
                            fine_run.bold = True
                            
                            request_para = temp_doc.add_paragraph()
                            request_para.add_run(
                                "You are hereby requested to return the book to the library at the earliest and clear the pending fine. "
                                "Failure to do so may result in restrictions on future borrowing privileges.\\n\\n"
                                "Please contact the library desk for any queries or clarifications.\\n\\n"
                            )
                            
                            temp_doc.add_paragraph()
                            temp_doc.add_paragraph("Thank you for your cooperation.\\n\\nYours sincerely,\\n\\n")
                            temp_doc.add_paragraph("__________________________")
                            temp_doc.add_paragraph("Librarian").runs[0].bold = True
                            temp_doc.add_paragraph('Main Library')
                            temp_doc.add_paragraph('Government Polytechnic Awasari (Kh)')
                            
                            # Save to temp file
                            temp_dir = tempfile.gettempdir()
                            temp_file = os.path.join(temp_dir, f"Overdue_{enrollment_no}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
                            temp_doc.save(temp_file)
                            temp_files.append(temp_file)
                            
                            # Prepare email data
                            email_subject = f"Overdue Book Notice - {book_title}"
                            email_body = f"""Dear {student_name},

This is an automated notification from GPA's Library Management System, Government Polytechnic Awasari (Kh).

The following book borrowed from our library is overdue and needs to be returned immediately:

Book ID: {book_id}
Book Title: {book_title}
Issue Date: {issue_date}
Due Date: {due_date}
Days Overdue: {days_overdue}
Fine Amount: ₹{fine}

As per library rules, a fine of ₹{self.get_fine_per_day()} per day is applicable for overdue books.

Please return the book to the library at the earliest and clear the pending fine. Failure to do so may result in restrictions on future borrowing privileges.

For any queries, please contact the library desk.

Thank you for your cooperation.

Librarian
Main Library
Government Polytechnic Awasari (Kh)

---
Note: This is an automated email. Please find the attached formal overdue letter.
"""
                            
                            emails_to_send.append({
                                'to': student_email,
                                'subject': email_subject,
                                'body': email_body,
                                'attachment': temp_file,
                                'enrollment_no': enrollment_no,
                                'student_name': student_name,
                                'book_title': book_title
                            })
                    
                    # Define progress callback
                    def update_progress(sent, total, percentage):
                        try:
                            progress_bar['value'] = percentage
                            status_label['text'] = f"Sending... {sent}/{total} emails ({percentage:.1f}%)"
                            batch_num = (sent // self.email_batch_service.batch_size) + 1
                            stats_label['text'] = f"🔄 Processing batch {batch_num} | ⚡ {self.email_batch_service.max_workers} parallel workers"
                            progress_win.update()
                        except Exception:
                            pass  # progress window may have been closed
                    
                    # Send emails using batch service
                    # Map email_settings keys to what batch service expects
                    batch_email_config = {
                        'from_email': self.email_settings.get('sender_email', ''),
                        'password': self.email_settings.get('sender_password', ''),
                        'smtp_server': self.email_settings.get('smtp_server', 'smtp.gmail.com'),
                        'smtp_port': int(self.email_settings.get('smtp_port', 587)),
                    }
                    result = self.email_batch_service.send_batch_emails(
                        emails_to_send,
                        batch_email_config,
                        update_progress
                    )
                    
                    # Log results
                    sent_count = result['sent']
                    failed_count = result['failed']
                    
                    for email_data, detail in zip(emails_to_send, result.get('details', [])):
                        success = detail.get('success', False)
                        message = detail.get('error', '') or ''
                        self._log_email_sent(
                            email_data['enrollment_no'],
                            email_data['student_name'],
                            email_data['to'],
                            email_data['book_title'],
                            success,
                            message if not success else ''
                        )
                        
                        if success:
                            email_results.append(f"✅ {email_data['student_name']} ({email_data['enrollment_no']})")
                        else:
                            email_results.append(f"❌ {email_data['student_name']} ({email_data['enrollment_no']}) - {message}")
                    
                    # Clean up temp files
                    for temp_file in temp_files:
                        try:
                            os.remove(temp_file)
                        except:
                            pass
                    
                    progress_win.destroy()
                    _sequential_needed = False
                    
                else:
                    # Fallback to original sequential method
                    _sequential_needed = True
                    progress_win = tk.Toplevel(self.root)
                    progress_win.title("Sending Emails...")
                    progress_win.geometry("500x200")
                    progress_win.transient(self.root)
                    progress_win.grab_set()
                    
                    # Center the window
                    progress_win.update_idletasks()
                    x = (progress_win.winfo_screenwidth() // 2) - (250)
                    y = (progress_win.winfo_screenheight() // 2) - (100)
                    progress_win.geometry(f"+{x}+{y}")
                    
                    label = tk.Label(progress_win, text="Sending overdue emails...", 
                                   font=('Segoe UI', 12), pady=20)
                    label.pack()
                    
                    status_label = tk.Label(progress_win, text="", font=('Segoe UI', 10), 
                                          wraplength=450, justify=tk.LEFT)
                    status_label.pack(pady=10)
                    
                    progress_win.update()
                    
                    sent_count = 0
                    failed_count = 0
                
                for idx, rec in (enumerate(overdue, 1) if _sequential_needed else []):
                    enrollment_no = str(rec['Enrollment No'])
                    student_name = str(rec['Student Name'])
                    book_id = str(rec['Book ID'])
                    book_title = str(rec['Book Title'])
                    issue_date = str(rec['Issue Date'])
                    due_date = str(rec['Due Date'])
                    days_overdue = str(rec['Days Overdue'])
                    fine = str(rec['Accrued Fine'])
                    
                    status_label.config(text=f"Sending email {idx}/{len(overdue)}\nTo: {student_name} ({enrollment_no})")
                    progress_win.update()
                    
                    # Get student email
                    student_email = self.get_student_email(enrollment_no)
                    
                    if student_email:
                        # Generate individual letter for this student
                        temp_doc = Document()
                        
                        # Add same header
                        if os.path.exists(logo_path):
                            try:
                                logo_para = temp_doc.add_paragraph()
                                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                logo_run = logo_para.add_run()
                                logo_run.add_picture(logo_path, width=Pt(80))
                            except:
                                pass
                        
                        def add_center_temp(text, bold=True, size=16, color=None):
                            p = temp_doc.add_paragraph()
                            run = p.add_run(text)
                            run.bold = bold
                            run.font.size = Pt(size)
                            if color:
                                from docx.shared import RGBColor
                                run.font.color.rgb = RGBColor(*color)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        add_center_temp("Government Polytechnic Awasari (Kh)", True, 20, (31, 71, 136))
                        add_center_temp("Main Library", True, 16, (46, 92, 138))
                        temp_doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        temp_doc.add_paragraph()
                        
                        date_para = temp_doc.add_paragraph()
                        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        date_para.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
                        temp_doc.add_paragraph()
                        
                        subject = temp_doc.add_paragraph()
                        subject.add_run('Subject: Overdue Book Notice').bold = True
                        subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        temp_doc.add_paragraph()
                        
                        to_para = temp_doc.add_paragraph()
                        to_para.add_run(f'To,\n{student_name}\nEnrollment No: {enrollment_no}')
                        temp_doc.add_paragraph()
                        
                        body = temp_doc.add_paragraph()
                        body.add_run(
                            f"Dear {student_name},\n\n"
                            f"This is to inform you that the following book borrowed from GPA's Library Management System "
                            f"is overdue and needs to be returned immediately.\n\n"
                        )
                        
                        temp_doc.add_paragraph('Book Details:', style='Heading 2')
                        details_table = temp_doc.add_table(rows=5, cols=2)
                        details_table.style = 'Light Grid Accent 1'
                        details_table.cell(0, 0).text = 'Book ID:'
                        details_table.cell(0, 1).text = book_id
                        details_table.cell(1, 0).text = 'Book Title:'
                        details_table.cell(1, 1).text = book_title
                        details_table.cell(2, 0).text = 'Issue Date:'
                        details_table.cell(2, 1).text = issue_date
                        details_table.cell(3, 0).text = 'Due Date:'
                        details_table.cell(3, 1).text = due_date
                        details_table.cell(4, 0).text = 'Days Overdue:'
                        details_table.cell(4, 1).text = days_overdue
                        
                        temp_doc.add_paragraph()
                        fine_para = temp_doc.add_paragraph()
                        fine_run = fine_para.add_run(
                            f"As per library rules, a fine of ₹{self.get_fine_per_day()} per day is applicable for overdue books.\n"
                            f"Your current fine amount is: ₹{fine}\n\n"
                        )
                        fine_run.bold = True
                        
                        request_para = temp_doc.add_paragraph()
                        request_para.add_run(
                            "You are hereby requested to return the book to the library at the earliest and clear the pending fine. "
                            "Failure to do so may result in restrictions on future borrowing privileges.\n\n"
                            "Please contact the library desk for any queries or clarifications.\n\n"
                        )
                        
                        temp_doc.add_paragraph()
                        temp_doc.add_paragraph("Thank you for your cooperation.\n\nYours sincerely,\n\n")
                        temp_doc.add_paragraph("__________________________")
                        temp_doc.add_paragraph("Librarian").runs[0].bold = True
                        temp_doc.add_paragraph('Main Library')
                        temp_doc.add_paragraph('Government Polytechnic Awasari (Kh)')
                        
                        # Save to temp file
                        temp_dir = tempfile.gettempdir()
                        temp_file = os.path.join(temp_dir, f"Overdue_{enrollment_no}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
                        temp_doc.save(temp_file)
                        
                        # Send email
                        email_subject = f"Overdue Book Notice - {book_title}"
                        email_body = f"""Dear {student_name},

This is an automated notification from GPA's Library Management System, Government Polytechnic Awasari (Kh).

The following book borrowed from our library is overdue and needs to be returned immediately:

Book ID: {book_id}
Book Title: {book_title}
Issue Date: {issue_date}
Due Date: {due_date}
Days Overdue: {days_overdue}
Fine Amount: ₹{fine}

As per library rules, a fine of ₹{self.get_fine_per_day()} per day is applicable for overdue books.

Please return the book to the library at the earliest and clear the pending fine. Failure to do so may result in restrictions on future borrowing privileges.

For any queries, please contact the library desk.

Thank you for your cooperation.

Librarian
Main Library
Government Polytechnic Awasari (Kh)

---
Note: This is an automated email. Please find the attached formal overdue letter.
"""
                        
                        success, message = self.send_email_with_attachment(
                            student_email, 
                            email_subject, 
                            email_body, 
                            temp_file
                        )
                        
                        # Log the email
                        self._log_email_sent(
                            enrollment_no,
                            student_name,
                            student_email,
                            book_title,
                            success,
                            message if not success else ''
                        )
                        
                        if success:
                            sent_count += 1
                            email_results.append(f"✅ {student_name} ({enrollment_no})")
                        else:
                            failed_count += 1
                            email_results.append(f"❌ {student_name} ({enrollment_no}) - {message}")
                        
                        # Clean up temp file
                        try:
                            os.remove(temp_file)
                        except:
                            pass
                    else:
                        failed_count += 1
                        email_results.append(f"❌ {student_name} ({enrollment_no}) - No email address")
                
                if _sequential_needed:
                    progress_win.destroy()
                
                # Show results
                result_message = f"📧 Email Sending Complete!\n\n"
                result_message += f"✅ Successfully sent: {sent_count}\n"
                result_message += f"❌ Failed: {failed_count}\n\n"
                
                if email_results:
                    result_message += "Details:\n" + "\n".join(email_results[:10])  # Show first 10
                    if len(email_results) > 10:
                        result_message += f"\n... and {len(email_results) - 10} more"
                
                result_message += f"\n\n💾 Master document saved at:\n{file_path}"
                
                messagebox.showinfo("Email Results", result_message)
            else:
                messagebox.showinfo("Overdue Notice", f"Word overdue notice exported:\n{file_path}")
            
            if messagebox.askyesno("Open File", "Open the Word letter now?"):
                self.open_file(file_path)
            # Excel/text companion removed per user request
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate Word overdue notice: {e}")
    
    def import_books_from_excel(self):
        """Import books from Excel file.
        Supports three formats detected automatically:
        1. GPA Library Format (multi-sheet: ME/EE/IT/ETC/AE/CE/computer):
           Sr No | Actual NO | Author | Tital | Publisher | TOTAL BOOK | Accession Number | ...
           Accession numbers spread across extra columns represent the RANGE of copy IDs.
        2. Book Bank Format (bookbank1 sheet):
           Book Name | Author I | Publisher Name | Book Price | Accession No
           Each row = 1 physical copy; grouped by title+author for total count.
        3. Standard Format: book_id | title | author | isbn | category | total_copies | price | barcode

        KEY RULE: Same Title + Same Author = ONE book entry.
        Copies are MERGED across all sheets/branches — no duplicate entries created.
        """
        import math

        file_path = filedialog.askopenfilename(
            title="Select Excel file to import",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
        )
        if not file_path:
            return

        # ── Helpers ─────────────────────────────────────────────────────────────
        def normalize(s):
            """Lowercase + collapse whitespace for comparison key."""
            return ' '.join(str(s or '').lower().split())

        def is_nan_val(v):
            if v is None: return True
            try: return isinstance(v, float) and math.isnan(v)
            except: return False

        def safe_str(v):
            s = str(v).strip()
            return '' if s.lower() == 'nan' else s

        def safe_int(v, default=1):
            try:
                r = int(float(str(v)))
                return r if r > 0 else default
            except: return default

        def safe_float(v, default=0.0):
            try: return float(str(v))
            except: return default

        def get_col(col_map, *keys):
            """Return first matching original column name from col_map."""
            for k in keys:
                if k in col_map:
                    return col_map[k]
            return None

        def make_accession_info(acc_list):
            """Given a list of accession number strings, return (first_id, range_str, csv_str).
            range_str: '8264-8266' if continuous, else '8264, 8266, 8270'.
            csv_str  : '8264,8265,8266' (comma-sep for DB storage).
            """
            if not acc_list:
                return '', '', ''
            csv_str = ','.join(acc_list)
            if len(acc_list) == 1:
                return acc_list[0], acc_list[0], csv_str
            try:
                nums = [int(x) for x in acc_list]
                nums_sorted = sorted(nums)
                if nums_sorted == list(range(nums_sorted[0], nums_sorted[0] + len(nums_sorted))):
                    range_str = f"{nums_sorted[0]}-{nums_sorted[-1]}"
                else:
                    range_str = ', '.join(str(n) for n in nums_sorted)
            except Exception:
                range_str = ', '.join(acc_list)
            return acc_list[0], range_str, csv_str

        def parse_gpa_sheet(df, sheet_name):
            """Parse one GPA-format sheet. Returns list of book dicts.
            Collects ALL accession numbers (first named column + subsequent unnamed overflow columns).
            """
            df_cols = list(df.columns)
            col_map = {str(c).strip().lower(): c for c in df_cols}

            col_srno      = get_col(col_map, 'sr no', 'sr. no.', 'sr .no', 'sr.no', 'srno')
            col_catalog   = get_col(col_map, 'actual no', 'actual  no', 'actual no ', 'sr no', 'srno')
            col_author    = get_col(col_map, 'author')
            col_title     = get_col(col_map, 'tital', 'title', 'book title', 'book name')
            col_publisher = get_col(col_map, ' publisher', 'publisher', 'publisher name')
            col_copies    = get_col(col_map, 'total book', 'total books', 'total_copies', 'copies')
            col_accession = get_col(col_map, 'accession number', 'accession_number', 'accession no')
            col_price     = get_col(col_map, 'price', 'book price')

            # Columns AFTER the main Accession Number column hold overflow copies
            acc_overflow_cols = []
            if col_accession and col_accession in df_cols:
                acc_idx = df_cols.index(col_accession)
                for extra_col in df_cols[acc_idx + 1:]:
                    acc_overflow_cols.append(extra_col)

            if not col_title:
                return []

            def collect_accessions_from_row(row):
                """Collect all numeric accession values from a row."""
                row_acc = []

                # First: main accession column
                if col_accession:
                    acc = row.get(col_accession, '')
                    if not is_nan_val(acc) and safe_str(acc):
                        try:
                            row_acc.append(str(int(float(str(acc).strip()))))
                        except:
                            pass

                # Then: overflow columns (Unnamed:7, Unnamed:8, …)
                for ec in acc_overflow_cols:
                    val = row.get(ec, '')
                    if is_nan_val(val):
                        continue
                    v = safe_str(val)
                    if v:
                        try:
                            row_acc.append(str(int(float(v))))
                        except:
                            # Skip text tokens like "TO"
                            pass

                # Fallback to catalog only if accession columns are empty
                if not row_acc and col_catalog:
                    cat = row.get(col_catalog, '')
                    if not is_nan_val(cat) and safe_str(cat):
                        try:
                            row_acc.append(str(int(float(str(cat).strip()))))
                        except:
                            row_acc.append(safe_str(cat))

                return row_acc

            def dedupe_keep_order(items):
                seen = set()
                out = []
                for x in items:
                    if x not in seen:
                        seen.add(x)
                        out.append(x)
                return out

            results = []
            current_book = None

            for _, row in df.iterrows():
                title_val = safe_str(row.get(col_title, ''))
                author_val = safe_str(row.get(col_author, '') if col_author else '')
                publisher_val = safe_str(row.get(col_publisher, '') if col_publisher else '')
                price_val = safe_float(row.get(col_price, 0) if col_price else 0)

                copies_raw = row.get(col_copies, None) if col_copies else None
                has_copies = (copies_raw is not None and not is_nan_val(copies_raw) and safe_str(copies_raw) != '')
                copies_val = safe_int(copies_raw, 1) if has_copies else None

                row_acc = collect_accessions_from_row(row)

                if title_val:
                    # Start a new logical book row
                    current_book = {
                        'title': title_val,
                        'author': author_val,
                        'publisher': publisher_val,
                        'price': price_val,
                        '_copies_declared': copies_val,
                        '_acc_list': row_acc[:],
                        'sheet': sheet_name
                    }
                    results.append(current_book)
                    continue

                # Continuation row: no title, but may contain remaining accession numbers
                if current_book and row_acc:
                    current_book['_acc_list'].extend(row_acc)
                    # If continuation row unexpectedly includes total copies, keep a valid declared count
                    if copies_val is not None and (current_book.get('_copies_declared') is None):
                        current_book['_copies_declared'] = copies_val

            # Finalize computed fields
            finalized = []
            for b in results:
                all_acc = dedupe_keep_order(b.get('_acc_list', []))
                first_id, range_str, csv_str = make_accession_info(all_acc)

                copies_val = b.get('_copies_declared') if b.get('_copies_declared') else len(all_acc)
                copies_val = max(copies_val, 1)
                if copies_val <= 1 and len(all_acc) > 1:
                    copies_val = len(all_acc)

                finalized.append({
                    'title': b.get('title', ''),
                    'author': b.get('author', ''),
                    'publisher': b.get('publisher', ''),
                    'copies': copies_val,
                    'book_id': range_str or first_id,
                    'accession_csv': csv_str,
                    'price': b.get('price', 0.0),
                    'sheet': b.get('sheet', sheet_name)
                })

            return finalized

        def parse_bookbank_sheet(df):
            """Parse bookbank format (each row = 1 physical copy). Returns list grouped by title+author."""
            col_map = {str(c).strip().lower(): c for c in df.columns}
            col_title     = get_col(col_map, 'book name', 'title')
            col_author    = get_col(col_map, 'author i', 'author')
            col_publisher = get_col(col_map, 'publisher name', 'publisher')
            col_price     = get_col(col_map, 'book price', 'price')
            col_accession = get_col(col_map, 'accession no', 'accession number')

            if not col_title:
                return []

            grouped = {}
            for _, row in df.iterrows():
                title_val = safe_str(row.get(col_title, ''))
                if not title_val: continue
                author_val    = safe_str(row.get(col_author, '')    if col_author    else '')
                publisher_val = safe_str(row.get(col_publisher, '') if col_publisher else '')
                price_val     = safe_float(row.get(col_price, 0)    if col_price     else 0)
                acc_val       = safe_str(row.get(col_accession, '') if col_accession else '')

                key = (normalize(title_val), normalize(author_val))
                if key not in grouped:
                    grouped[key] = {
                        'title': title_val, 'author': author_val,
                        'publisher': publisher_val,
                        'copies': 0,
                        'book_id': acc_val,
                        'price': price_val,
                        'sheet': 'bookbank',
                        '_acc_list': []
                    }
                grouped[key]['copies'] += 1   # each row = 1 physical copy

                # Collect accession numbers when present
                if acc_val:
                    try:
                        grouped[key]['_acc_list'].append(str(int(float(acc_val))))
                    except Exception:
                        pass

            finalized = []
            for b in grouped.values():
                acc_list = b.get('_acc_list', [])
                # Dedupe but keep order
                seen = set()
                acc_list = [x for x in acc_list if not (x in seen or seen.add(x))]
                first_id, range_str, csv_str = make_accession_info(acc_list)

                copies_val = len(acc_list) if acc_list else max(int(b.get('copies', 1) or 1), 1)
                finalized.append({
                    'title': b.get('title', ''),
                    'author': b.get('author', ''),
                    'publisher': b.get('publisher', ''),
                    'copies': copies_val,
                    'book_id': range_str or first_id or safe_str(b.get('book_id', '')),
                    'accession_csv': csv_str,
                    'price': b.get('price', 0.0),
                    'sheet': b.get('sheet', 'bookbank')
                })
            return finalized

        try:
            import pandas as pd  # ensure pandas is available in local scope
            xl = pd.ExcelFile(file_path)
            all_sheets = xl.sheet_names

            # ── books_map: key=(norm_title, norm_author) → merged book info ──
            books_map = {}
            sheets_processed = []

            for sheet_name in all_sheets:
                try:
                    df_raw = pd.read_excel(file_path, sheet_name=sheet_name)
                    if df_raw.empty:
                        continue

                    # Detect if header is in row 0 of data (all Unnamed columns — e.g. IT, AE sheets)
                    all_unnamed = all('unnamed' in str(c).lower() for c in df_raw.columns)
                    if all_unnamed:
                        df = pd.read_excel(file_path, sheet_name=sheet_name, header=1)
                        if df.empty:
                            continue
                    else:
                        df = df_raw

                    sheet_cols_lower = [str(c).strip().lower() for c in df.columns]
 
                    # Detect format
                    is_bookbank = ('book name' in sheet_cols_lower or
                                   'accession no' in sheet_cols_lower or
                                   'book price' in sheet_cols_lower)
                    is_gpa      = ('tital'             in sheet_cols_lower or
                                   'total book'        in sheet_cols_lower or
                                   'total books'       in sheet_cols_lower or
                                   'accession number'  in sheet_cols_lower)

                    if is_bookbank:
                        parsed = parse_bookbank_sheet(df)
                        fmt    = 'bookbank'
                    elif is_gpa:
                        parsed = parse_gpa_sheet(df, sheet_name)
                        fmt    = 'gpa'
                    else:
                        continue   # standard / unrecognised — try later

                    sheets_processed.append(f"{sheet_name}({fmt},{len(parsed)})")

                    # Merge into books_map — same title+author → accumulate copies
                    for book in parsed:
                        key = (normalize(book['title']), normalize(book['author']))
                        if key in books_map:
                            books_map[key]['copies'] += book.get('copies', 0)
                            # Merge accession lists (copy IDs) when available
                            ex_csv = safe_str(books_map[key].get('accession_csv', ''))
                            new_csv = safe_str(book.get('accession_csv', ''))
                            if new_csv:
                                ex_list = [x.strip() for x in ex_csv.split(',') if x.strip()] if ex_csv else []
                                new_list = [x.strip() for x in new_csv.split(',') if x.strip()]
                                seen = set()
                                merged = []
                                for x in ex_list + new_list:
                                    if x not in seen:
                                        seen.add(x)
                                        merged.append(x)
                                books_map[key]['accession_csv'] = ','.join(merged)
                                # If we have accession IDs, treat them as the true copy count
                                books_map[key]['copies'] = max(1, len(merged))
                        else:
                            books_map[key] = dict(book)

                except Exception as e:
                    print(f"[ImportBooks] Sheet '{sheet_name}' skipped: {e}")
                    continue

            # ── Fallback: standard format if nothing was detected as GPA ──────
            if not books_map:
                for sheet_name in all_sheets:
                    try:
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                        if df.empty: continue
                        df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_')
                        rn = {'book_title':'title','tital':'title','total_book':'total_copies',
                              'total_books':'total_copies','sr_no':'book_id'}
                        df.rename(columns={k:v for k,v in rn.items() if k in df.columns}, inplace=True)
                        if 'title' not in df.columns: continue
                        for _, row in df.iterrows():
                            title_val  = safe_str(row.get('title',''))
                            if not title_val: continue
                            author_val = safe_str(row.get('author',''))
                            key = (normalize(title_val), normalize(author_val))
                            if key not in books_map:
                                books_map[key] = {
                                    'title': title_val, 'author': author_val,
                                    'publisher': safe_str(row.get('isbn','')),
                                    'copies': safe_int(row.get('total_copies',1)),
                                    'book_id': safe_str(row.get('book_id','')),
                                    'price': safe_float(row.get('price',0)),
                                    'sheet': sheet_name
                                }
                            else:
                                books_map[key]['copies'] += safe_int(row.get('total_copies',1))
                    except Exception:
                        continue

            if not books_map:
                messagebox.showerror("Error", "No books found in the Excel file.\n"
                                     "File may be empty or in an unrecognized format.")
                return

            # ── Insert / Update DB ───────────────────────────────────────────
            success_count = 0
            merged_count  = 0
            error_count   = 0
            errors        = []

            conn = self.db.get_connection()
            cur  = conn.cursor()

            def _parse_acc_csv(csv_text):
                items = []
                if not csv_text:
                    return items
                for x in str(csv_text).split(','):
                    s = x.strip()
                    if not s:
                        continue
                    try:
                        items.append(str(int(float(s))))
                    except Exception:
                        continue
                # Dedupe keep order
                seen = set()
                out = []
                for it in items:
                    if it not in seen:
                        seen.add(it)
                        out.append(it)
                return out

            # Preload existing books for faster matching
            cur.execute("SELECT book_id, title, author, total_copies, available_copies, barcode FROM books")
            existing_by_key = {}
            used_accessions = set()
            max_numeric_book_id = 1000

            for row in cur.fetchall() or []:
                ex_book_id = str(row[0])
                ex_title = str(row[1] or '')
                ex_author = str(row[2] or '')
                ex_total = int(row[3] or 0)
                ex_avail = int(row[4] or 0)
                ex_barcode = str(row[5] or '')

                ek = (normalize(ex_title), normalize(ex_author))
                existing_by_key[ek] = (ex_book_id, ex_total, ex_avail, ex_barcode)

                # Track already-used copy IDs (accession numbers) across the whole library
                for acc in _parse_acc_csv(ex_barcode):
                    used_accessions.add(acc)

                # For auto-book_id generation
                try:
                    if ex_book_id.isdigit():
                        max_numeric_book_id = max(max_numeric_book_id, int(ex_book_id))
                except Exception:
                    pass

            # Single transaction for speed
            try:
                cur.execute("BEGIN")
            except Exception:
                pass

            for key, book in books_map.items():
                try:
                    title_val     = safe_str(book.get('title', ''))
                    author_val    = safe_str(book.get('author', ''))
                    if not title_val:
                        continue

                    copies_val    = max(int(book.get('copies', 1) or 1), 1)
                    book_id_v     = safe_str(book.get('book_id', ''))
                    publisher     = safe_str(book.get('publisher', ''))
                    price_val     = float(book.get('price', 0.0) or 0.0)

                    acc_list = _parse_acc_csv(book.get('accession_csv', ''))
                    acc_set = set(acc_list)

                    if key in existing_by_key:
                        ex_book_id, ex_total, ex_avail, ex_barcode = existing_by_key[key]
                        ex_acc_list = _parse_acc_csv(ex_barcode)
                        ex_acc_set = set(ex_acc_list)
                        other_used = used_accessions - ex_acc_set

                        # Allow accession numbers already belonging to this book; reject those used by other books
                        if acc_set:
                            acc_set = {a for a in acc_set if a not in other_used}
                            union_acc = list(ex_acc_set.union(acc_set))
                            # Stable numeric sort when possible
                            try:
                                union_acc.sort(key=lambda x: int(x))
                            except Exception:
                                union_acc.sort()
                            new_barcode = ','.join(union_acc)

                            # BUG-FIX: Trust accession list as ground truth for copy count.
                            # Using max(ex_total, len(union_acc)) would prevent fixing a corrupted ex_total
                            # (e.g. if ex_total was erroneously set to 18000, max() would keep it forever).
                            new_total = len(union_acc)
                        else:
                            # No per-copy accession IDs available: use incoming value as authoritative
                            # (don't use max() here either — if copies_val is sane and ex_total is not, fix it)
                            new_barcode = ex_barcode or ''
                            new_total = copies_val if copies_val > 0 else ex_total

                        issued = max(ex_total - ex_avail, 0)
                        new_avail = max(new_total - issued, 0)

                        # Only write if something changes
                        if new_total != ex_total or new_avail != ex_avail or (acc_set and new_barcode != (ex_barcode or '')):
                            cur.execute(
                                "UPDATE books SET total_copies=?, available_copies=?, barcode=?, price=?, updated_at=CURRENT_TIMESTAMP WHERE book_id=?",
                                (new_total, new_avail, new_barcode or None, price_val, ex_book_id)
                            )
                        existing_by_key[key] = (ex_book_id, new_total, new_avail, new_barcode)
                        used_accessions.update(_parse_acc_csv(new_barcode))
                        merged_count += 1
                        continue

                    # New book insert
                    if acc_set:
                        # Reject accessions that already exist anywhere
                        acc_set = {a for a in acc_set if a not in used_accessions}
                        acc_list = [a for a in acc_list if a in acc_set]
                        try:
                            acc_list.sort(key=lambda x: int(x))
                        except Exception:
                            acc_list.sort()
                        accession_csv = ','.join(acc_list)
                        copies_val = max(1, len(acc_list))
                    else:
                        accession_csv = ''

                    # Ensure unique book_id
                    if not book_id_v:
                        max_numeric_book_id += 1
                        book_id_v = str(max_numeric_book_id)

                    orig_id = book_id_v
                    suffix  = 0
                    while True:
                        cur.execute("SELECT 1 FROM books WHERE book_id=? LIMIT 1", (book_id_v,))
                        if not cur.fetchone():
                            break
                        suffix += 1
                        book_id_v = f"{orig_id}-{suffix}"

                    cur.execute(
                        "INSERT INTO books (book_id,title,author,isbn,category,total_copies,available_copies,barcode,price,updated_at) "
                        "VALUES (?,?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP)",
                        (book_id_v, title_val, author_val, publisher, 'Technology', copies_val, copies_val, accession_csv or None, price_val)
                    )
                    used_accessions.update(_parse_acc_csv(accession_csv))
                    existing_by_key[key] = (book_id_v, copies_val, copies_val, accession_csv)
                    success_count += 1

                except Exception as e:
                    error_count += 1
                    errors.append(f"'{book.get('title','?')}': {e}")

            conn.commit()
            conn.close()

            # ── Summary ──────────────────────────────────────────────────────
            sheets_str = ', '.join(sheets_processed) if sheets_processed else 'auto-detected'
            result_msg = (
                f"Import completed!\n\n"
                f"📄  Sheets processed : {len(sheets_processed)}\n"
                f"📚  Unique books found: {len(books_map)}\n\n"
                f"✅  New books added   : {success_count}\n"
                f"🔁  Copies merged     : {merged_count}  (same title+author already in DB)\n"
                f"❌  Errors            : {error_count}\n\n"
                f"Sheets: {sheets_str}"
            )
            if errors:
                result_msg += "\n\nErrors:\n" + "\n".join(errors[:5])
                if len(errors) > 5:
                    result_msg += f"\n... and {len(errors)-5} more."
            messagebox.showinfo("Import Results", result_msg)

            if success_count > 0 or merged_count > 0:
                self.refresh_books()
                self.refresh_dashboard()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to import Excel file:\n{str(e)}")

    def share_data_dialog(self):
        """Show data sharing dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Share Library Data")
        dialog.geometry("500x400")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (self.root.winfo_rootx() + 200, self.root.winfo_rooty() + 200))
        
        # Title
        title_label = tk.Label(
            dialog,
            text="📤 Share Library Data",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg=self.colors['accent']
        )
        title_label.pack(pady=(20, 30))
        
        # Instructions
        info_label = tk.Label(
            dialog,
            text="Export data to Excel and share via email, WhatsApp, or other platforms:",
            font=('Segoe UI', 11),
            bg='white',
            fg=self.colors['accent'],
            wraplength=400
        )
        info_label.pack(pady=(0, 20))
        
        # Buttons frame
        buttons_frame = tk.Frame(dialog, bg='white')
        buttons_frame.pack(expand=True)
        
        # Data type buttons
        data_types = [
            ("📊 Dashboard Summary", self.export_dashboard_summary),
            ("👥 Students Data", self.export_all_students_direct),
            ("📚 Books Data", self.export_books_to_excel),
            ("📋 Transaction Records", self.export_records_to_excel)
        ]
        
        def create_button_command(cmd, dialog_ref):
            """Create button command that handles dialog closing properly"""
            def button_action():
                dialog_ref.destroy()
                cmd()
            return button_action
        
        for text, command in data_types:
            btn = tk.Button(
                buttons_frame,
                text=text,
                font=('Segoe UI', 12, 'bold'),
                bg=self.colors['secondary'],
                fg='white',
                relief='flat',
                padx=20,
                pady=15,
                command=create_button_command(command, dialog),
                cursor='hand2',
                width=25
            )
            btn.pack(pady=10)
        
        # Close button
        close_btn = tk.Button(
            dialog,
            text="❌ Close",
            font=('Segoe UI', 12, 'bold'),
            bg='#6c757d',
            fg='white',
            relief='flat',
            padx=20,
            pady=10,
            command=dialog.destroy,
            cursor='hand2'
        )
        close_btn.pack(pady=20)
    
    def export_dashboard_summary(self):
        """Export dashboard summary to Excel"""
        try:
            # Get statistics
            stats = self.get_library_statistics()
            
            # Get recent activities
            activities = self.get_recent_activities()
            
            # Create summary data
            summary_data = {
                'Library Statistics': [
                    ['Metric', 'Value'],
                    ['Total Books', stats['total_books']],
                    ['Available Books', stats['available_books']],
                    ['Borrowed Books', stats['borrowed_books']],
                    ['Total Students', stats['total_students']],
                    ['Generated On', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
                ]
            }
            
            # Save to file
            filename = f"dashboard_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                # Correct parameter name
                initialfile=filename
            )
            
            if file_path:
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    # Write statistics
                    stats_df = pd.DataFrame(summary_data['Library Statistics'][1:], columns=summary_data['Library Statistics'][0])
                    header_end_row = 6  # Enhanced header takes 6 rows
                    stats_df.to_excel(writer, sheet_name='Statistics', index=False, startrow=header_end_row)
                    ws_stats = writer.book['Statistics']
                    self._write_excel_header_openpyxl(ws_stats, start_row=1)
                    self._auto_adjust_column_width(ws_stats)
                    
                    # Write recent activities
                    if activities:
                        activities_df = pd.DataFrame(activities, columns=[
                            'Type', 'Student', 'Book', 'Date', 'Status'
                        ])
                        activities_df.to_excel(writer, sheet_name='Recent Activities', index=False, startrow=header_end_row)
                        ws_act = writer.book['Recent Activities']
                        self._write_excel_header_openpyxl(ws_act, start_row=1)
                        self._auto_adjust_column_width(ws_act)
                
                messagebox.showinfo("Success", f"Dashboard summary exported to {file_path}")
                
                # Ask if user wants to open the file
                if messagebox.askyesno("Open File", "Do you want to open the exported file?"):
                    self.open_file(file_path)
                    
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export dashboard summary: {str(e)}")
    
    def open_file(self, file_path):
        """Open file with default application"""
        try:
            if platform.system() == 'Windows':
                os.startfile(file_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.run(['open', file_path])
            else:  # Linux
                subprocess.run(['xdg-open', file_path])
        except Exception as e:
            print(f"Failed to open file: {e}")

    # ---------------------- Excel Helpers ----------------------
    def _write_excel_header_openpyxl(self, worksheet, start_row=1):
        """Write the required header into an openpyxl worksheet with professional formatting like Word documents.
        Creates a beautifully formatted header with logo, merged cells, colors, and borders.
        Returns the next row after the header.
        """
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.utils import get_column_letter
        
        # Set generous row heights for professional appearance
        worksheet.row_dimensions[start_row].height = 45
        worksheet.row_dimensions[start_row + 1].height = 35
        worksheet.row_dimensions[start_row + 2].height = 30
        worksheet.row_dimensions[start_row + 3].height = 8  # Separator row
        worksheet.row_dimensions[start_row + 4].height = 15  # Spacing row
        
        # Add logo if available
        logo_path = os.path.join(os.path.dirname(__file__), 'logo.png')
        if os.path.exists(logo_path):
            try:
                img = XLImage(logo_path)
                # Larger logo for professional appearance
                img.width = 80
                img.height = 80
                # Position logo in column A, centered
                worksheet.add_image(img, f'A{start_row}')
            except Exception as e:
                print(f"Could not add logo: {e}")
        
        # Determine the number of columns to merge (ensure at least 8 columns for better appearance)
        max_col = max(8, worksheet.max_column if worksheet.max_column > 1 else 8)
        
        # Create border styles
        thin_border = Border(
            left=Side(style='thin', color='365F91'),
            right=Side(style='thin', color='365F91'),
            top=Side(style='thin', color='365F91'),
            bottom=Side(style='thin', color='365F91')
        )
        
        # Create background fills for gradient effect
        main_fill = PatternFill(start_color='E8F0FE', end_color='E8F0FE', fill_type='solid')  # Light blue background
        
        # Main heading - VERY LARGE, bold, centered with background
        merge_range_main = f'B{start_row}:{get_column_letter(max_col)}{start_row}'
        worksheet.merge_cells(merge_range_main)
        cell_main = worksheet[f'B{start_row}']
        cell_main.value = "Government Polytechnic Awasari(Kh)"
        cell_main.font = Font(size=22, bold=True, name='Calibri', color='1F4788')  # Dark blue
        cell_main.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell_main.fill = main_fill
        cell_main.border = thin_border
        
        # Subheading - Large, bold, centered with background
        merge_range_sub1 = f'B{start_row + 1}:{get_column_letter(max_col)}{start_row + 1}'
        worksheet.merge_cells(merge_range_sub1)
        cell_sub1 = worksheet[f'B{start_row + 1}']
        cell_sub1.value = "Main Library"
        cell_sub1.font = Font(size=18, bold=True, name='Calibri', color='2E5C8A')  # Medium blue
        cell_sub1.alignment = Alignment(horizontal='center', vertical='center')
        cell_sub1.fill = PatternFill(start_color='F0F6FF', end_color='F0F6FF', fill_type='solid')
        cell_sub1.border = thin_border
        
        # Sub-subheading - Medium, bold, centered with background
        merge_range_sub2 = f'B{start_row + 2}:{get_column_letter(max_col)}{start_row + 2}'
        worksheet.merge_cells(merge_range_sub2)
        cell_sub2 = worksheet[f'B{start_row + 2}']
        cell_sub2.value = "Government Polytechnic Awasari (Kh) - Library Management System"
        cell_sub2.font = Font(size=16, bold=True, name='Calibri', color='365F91')  # Light blue
        cell_sub2.alignment = Alignment(horizontal='center', vertical='center')
        cell_sub2.fill = PatternFill(start_color='F8FAFF', end_color='F8FAFF', fill_type='solid')
        cell_sub2.border = thin_border
        
        # Add a professional separator line across all columns
        thick_border = Border(bottom=Side(style='thick', color='1F4788'))
        for col in range(1, max_col + 1):
            cell = worksheet.cell(row=start_row + 3, column=col)
            cell.border = thick_border
            cell.fill = PatternFill(start_color='1F4788', end_color='1F4788', fill_type='solid')
        
        # Add borders to logo column for consistency
        for row in range(start_row, start_row + 3):
            logo_cell = worksheet.cell(row=row, column=1)
            logo_cell.border = thin_border
            logo_cell.fill = main_fill
        
        # Return next available row (after spacing)
        return start_row + 6

    def _xlsxwriter_write_header(self, worksheet, workbook, start_row=0):
        """Write the required header into an xlsxwriter worksheet with proper formatting and logo."""
        # Main heading - VERY LARGE, bold, centered with color
        fmt_main = workbook.add_format({
            'bold': True, 
            'font_size': 20, 
            'align': 'center',
            'valign': 'vcenter',
            'font_name': 'Arial',
            'font_color': '#1F4788'  # Dark blue
        })
        # Subheading - Large, bold, centered with color
        fmt_sub1 = workbook.add_format({
            'bold': True, 
            'font_size': 16, 
            'align': 'center',
            'valign': 'vcenter',
            'font_name': 'Arial',
            'font_color': '#2E5C8A'  # Medium blue
        })
        # Sub-subheading - Medium, bold, centered with color
        fmt_sub2 = workbook.add_format({
            'bold': True, 
            'font_size': 14, 
            'align': 'center',
            'valign': 'vcenter',
            'font_name': 'Arial',
            'font_color': '#365F91'  # Light blue
        })
        # Separator line format
        fmt_line = workbook.add_format({
            'bottom': 1,
            'bottom_color': '#000000'
        })
        
        # Set row heights for better appearance
        worksheet.set_row(start_row, 30)
        worksheet.set_row(start_row + 1, 25)
        worksheet.set_row(start_row + 2, 20)
        
        # Add logo if available
        logo_path = os.path.join(os.path.dirname(__file__), 'logo.png')
        if os.path.exists(logo_path):
            try:
                worksheet.insert_image(start_row, 0, logo_path, 
                                     {'x_scale': 0.3, 'y_scale': 0.3, 'x_offset': 5, 'y_offset': 5})
            except Exception as e:
                print(f"Could not add logo: {e}")
        
        # Merge cells for headers (assuming 7 columns minimum)
        worksheet.merge_range(start_row, 1, start_row, 6, "Government Polytechnic Awasari(Kh)", fmt_main)
        worksheet.merge_range(start_row + 1, 1, start_row + 1, 6, "Main Library", fmt_sub1)
        worksheet.merge_range(start_row + 2, 1, start_row + 2, 6, "Government Polytechnic Awasari (Kh) - Library", fmt_sub2)
        
        # Add separator line
        for col in range(7):
            worksheet.write(start_row + 3, col, "", fmt_line)
        
        # Add blank row for spacing
        return start_row + 5

    def _auto_adjust_column_width(self, worksheet):
        """Auto-adjust column widths in openpyxl worksheet based on content with enhanced formatting"""
        try:
            # First, apply column formatting to data headers
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            
            # Find data start row (after header)
            data_start_row = None
            for row in range(1, worksheet.max_row + 1):
                cell_value = worksheet.cell(row=row, column=1).value
                if cell_value and isinstance(cell_value, str) and any(header in cell_value.lower() for header in ['enrollment', 'book', 'student', 'name', 'id']):
                    data_start_row = row
                    break
            
            # Style data headers if found
            if data_start_row:
                header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                header_font = Font(color='FFFFFF', bold=True, size=11, name='Calibri')
                header_alignment = Alignment(horizontal='center', vertical='center')
                header_border = Border(
                    left=Side(style='thin', color='FFFFFF'),
                    right=Side(style='thin', color='FFFFFF'),
                    top=Side(style='thin', color='FFFFFF'),
                    bottom=Side(style='thin', color='FFFFFF')
                )
                
                # Apply header styling
                for col in range(1, worksheet.max_column + 1):
                    cell = worksheet.cell(row=data_start_row, column=col)
                    if cell.value:  # Only style cells with content
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.alignment = header_alignment
                        cell.border = header_border
            
            # Auto-adjust column widths based on content
            from openpyxl.utils import get_column_letter as _gcl
            for col_idx in range(1, worksheet.max_column + 1):
                max_length = 0
                column_letter = _gcl(col_idx)
                
                for row_idx in range(1, worksheet.max_row + 1):
                    try:
                        cell = worksheet.cell(row=row_idx, column=col_idx)
                        if hasattr(cell, 'value') and cell.value and col_idx != 1:
                            cell_length = len(str(cell.value))
                            if cell_length > max_length:
                                max_length = cell_length
                    except:
                        pass
                
                # Enhanced width calculation with better minimum and maximum
                if column_letter == 'A':  # Logo column
                    adjusted_width = 12
                elif max_length < 8:  # Short content
                    adjusted_width = 15
                elif max_length < 20:  # Medium content
                    adjusted_width = max_length + 6
                else:  # Long content
                    adjusted_width = min(max_length + 4, 65)
                
                worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # Add alternating row colors for better readability
            if data_start_row:
                light_fill = PatternFill(start_color='F8F9FA', end_color='F8F9FA', fill_type='solid')
                data_border = Border(
                    left=Side(style='thin', color='E0E0E0'),
                    right=Side(style='thin', color='E0E0E0'),
                    top=Side(style='thin', color='E0E0E0'),
                    bottom=Side(style='thin', color='E0E0E0')
                )
                
                for row in range(data_start_row + 1, worksheet.max_row + 1):
                    if (row - data_start_row) % 2 == 0:  # Even rows
                        for col in range(1, worksheet.max_column + 1):
                            cell = worksheet.cell(row=row, column=col)
                            if not cell.coordinate.startswith('A'):  # Skip logo column
                                cell.fill = light_fill
                                cell.border = data_border
                                cell.alignment = Alignment(vertical='center')
                    else:  # Odd rows
                        for col in range(1, worksheet.max_column + 1):
                            cell = worksheet.cell(row=row, column=col)
                            if not cell.coordinate.startswith('A'):  # Skip logo column
                                cell.border = data_border
                                cell.alignment = Alignment(vertical='center')
                                
        except Exception as e:
            print(f"Could not auto-adjust columns: {e}")

    # =====================================================================
    # ANALYSIS TAB - Charts, Graphs, and Data Visualization
    # =====================================================================
    
    def create_analysis_tab(self):
        """Create comprehensive analysis tab using pie/donut charts only, with smooth scrolling"""
        analysis_frame = tk.Frame(self.notebook, bg=self.colors['primary'])
        self.notebook.add(analysis_frame, text="📊 Analysis")
        
        if not MATPLOTLIB_AVAILABLE:
            self.create_analysis_unavailable_message(analysis_frame)
            return
        
        # Main container with scrollable area (both directions)
        main_canvas = tk.Canvas(analysis_frame, bg=self.colors['primary'], highlightthickness=0)
        vscroll = ttk.Scrollbar(analysis_frame, orient="vertical", command=main_canvas.yview)
        hscroll = ttk.Scrollbar(analysis_frame, orient="horizontal", command=main_canvas.xview)
        scrollable_frame = tk.Frame(main_canvas, bg=self.colors['primary'])

        # Create a window inside the canvas and keep its ID to sync width
        window_id = main_canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

        def _on_frame_configure(event=None):
            # Update scroll region to include the whole inner frame
            main_canvas.configure(scrollregion=main_canvas.bbox("all"))
        scrollable_frame.bind("<Configure>", _on_frame_configure)

        def _on_canvas_configure(event=None):
            # Make inner frame match canvas width for pleasant layout
            try:
                main_canvas.itemconfig(window_id, width=event.width)
            except Exception:
                pass
        main_canvas.bind("<Configure>", _on_canvas_configure)

        main_canvas.configure(yscrollcommand=vscroll.set, xscrollcommand=hscroll.set)

        # Mousewheel scrolling (Windows/Linux)
        def _wheel_units(delta):
            # Smooth scrolling with moderate speed
            units = int(-delta / 120)  # Standard Windows mouse wheel delta is 120
            if units == 0:
                units = -1 if delta > 0 else 1
            return units * 2  # Multiply by 2 for balanced smooth scrolling
        
        def _on_vwheel(event):
            units = _wheel_units(event.delta)
            if units:
                main_canvas.yview_scroll(units, 'units')
            return 'break'
        
        def _on_hwheel(event):
            units = _wheel_units(event.delta)
            if units:
                main_canvas.xview_scroll(units, 'units')
            return 'break'
        # Bind locally to inner frame and canvas; avoid global bind/unbind to not break other tabs
        scrollable_frame.bind('<MouseWheel>', _on_vwheel)
        scrollable_frame.bind('<Shift-MouseWheel>', _on_hwheel)
        main_canvas.bind('<MouseWheel>', _on_vwheel)
        main_canvas.bind('<Shift-MouseWheel>', _on_hwheel)
        # Focus canvas when pointer enters so it receives wheel events
        main_canvas.bind('<Enter>', lambda e: main_canvas.focus_set())
        scrollable_frame.bind('<Enter>', lambda e: main_canvas.focus_set())
        # Linux fallbacks
        def _an_btn4(_e=None):
            main_canvas.yview_scroll(-1, 'units'); return 'break'
        def _an_btn5(_e=None):
            main_canvas.yview_scroll(1, 'units'); return 'break'
        main_canvas.bind('<Button-4>', _an_btn4)
        main_canvas.bind('<Button-5>', _an_btn5)
        # Additionally, pointer-scoped global binds when mouse enters/leaves this analysis tab
        def _enter_analysis(_e=None):
            try:
                analysis_frame.bind_all('<MouseWheel>', _on_vwheel, add='+')
                analysis_frame.bind_all('<Shift-MouseWheel>', _on_hwheel, add='+')
            except Exception:
                pass
        def _leave_analysis(_e=None):
            try:
                analysis_frame.unbind_all('<MouseWheel>')
                analysis_frame.unbind_all('<Shift-MouseWheel>')
            except Exception:
                pass
        analysis_frame.bind('<Enter>', _enter_analysis)
        analysis_frame.bind('<Leave>', _leave_analysis)
        # Helper so figures can hook wheel events too
        def _bind_analysis_wheel(widget):
            try:
                widget.bind('<MouseWheel>', _on_vwheel)
                widget.bind('<Shift-MouseWheel>', _on_hwheel)
            except Exception:
                pass
        self._analysis_bind_wheel = _bind_analysis_wheel

        # Pack scrollable components
        main_canvas.pack(side="top", fill="both", expand=True)
        vscroll.pack(side="right", fill="y")
        hscroll.pack(side="bottom", fill="x")
        
        # Header
        header_frame = tk.Frame(scrollable_frame, bg=self.colors['primary'])
        header_frame.pack(fill=tk.X, padx=20, pady=20)
        
        tk.Label(
            header_frame,
            text="📊 Library Analytics Dashboard",
            font=('Segoe UI', 18, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        ).pack(side=tk.LEFT)
        
        # Time period filter
        period_frame = tk.LabelFrame(
            scrollable_frame,
            text="📅 Analysis Period",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        period_frame.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        period_controls = tk.Frame(period_frame, bg=self.colors['primary'])
        period_controls.pack(padx=15, pady=15)
        
        self.analysis_period = tk.StringVar(value="7")
        
        periods = [
            ("Last 7 Days", "7"),
            ("Last 15 Days", "15"), 
            ("Last 30 Days", "30")
        ]
        
        for text, value in periods:
            rb = tk.Radiobutton(
                period_controls,
                text=text,
                variable=self.analysis_period,
                value=value,
                font=('Segoe UI', 11),
                bg=self.colors['primary'],
                fg=self.colors['accent'],
                selectcolor='white',
                command=self.refresh_analysis
            )
            rb.pack(side=tk.LEFT, padx=(0, 30))
        
        # Export controls
        export_frame = tk.Frame(period_frame, bg=self.colors['primary'])
        export_frame.pack(side=tk.RIGHT, padx=15, pady=15)
        
        export_excel_btn = tk.Button(
            export_frame,
            text="📊 Export to Excel",
            font=('Segoe UI', 10, 'bold'),
            bg='#28a745',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.export_analysis_excel,
            cursor='hand2'
        )
        export_excel_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        export_word_btn = tk.Button(
            export_frame,
            text="📄 Export to Word",
            font=('Segoe UI', 10, 'bold'),
            bg='#6f42c1',
            fg='white',
            relief='flat',
            padx=15,
            pady=8,
            command=self.export_analysis_word,
            cursor='hand2'
        )
        export_word_btn.pack(side=tk.LEFT)
        
        # Filter controls for student and book analytics
        filter_frame = tk.LabelFrame(
            scrollable_frame,
            text="🎯 Focused Analysis (Student / Book)",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        filter_frame.pack(fill=tk.X, padx=20, pady=(0, 20))

        self.analysis_student_var = tk.StringVar()
        self.analysis_book_var = tk.StringVar()
        self.analysis_branch_var = tk.StringVar(value="All")

        ff_row = tk.Frame(filter_frame, bg=self.colors['primary'])
        ff_row.pack(fill=tk.X, padx=15, pady=12)

        tk.Label(ff_row, text="Enrollment No:", font=('Segoe UI', 10, 'bold'), bg=self.colors['primary'], fg=self.colors['accent']).pack(side=tk.LEFT)
        student_entry = tk.Entry(ff_row, textvariable=self.analysis_student_var, width=20)
        student_entry.pack(side=tk.LEFT, padx=(6, 18))

        tk.Label(ff_row, text="Book ID:", font=('Segoe UI', 10, 'bold'), bg=self.colors['primary'], fg=self.colors['accent']).pack(side=tk.LEFT)
        book_entry = tk.Entry(ff_row, textvariable=self.analysis_book_var, width=20)
        book_entry.pack(side=tk.LEFT, padx=(6, 18))

        tk.Button(
            ff_row,
            text="Apply Filter",
            font=('Segoe UI', 10, 'bold'),
            bg='#0FA958', fg='white', relief='flat', cursor='hand2',
            command=self.apply_analysis_filter
        ).pack(side=tk.LEFT, padx=(0, 10))

        tk.Button(
            ff_row,
            text="Clear",
            font=('Segoe UI', 10, 'bold'),
            bg='#9e9e9e', fg='white', relief='flat', cursor='hand2',
            command=self.clear_analysis_filter
        ).pack(side=tk.LEFT)
        
        # Add Refresh Analysis button near the filters for easy access
        tk.Button(
            ff_row,
            text="🔄 Refresh Charts",
            font=('Segoe UI', 10, 'bold'),
            bg=self.colors['secondary'], fg='white', relief='flat', cursor='hand2',
            command=self.refresh_analysis,
            activebackground=self.colors['accent']
        ).pack(side=tk.LEFT, padx=(20, 0))

        # Display current filter summary
        self.analysis_filter_summary = tk.Label(
            filter_frame,
            text="No focused filter applied",
            font=('Segoe UI', 10),
            bg=self.colors['primary'],
            fg='#555'
        )
        self.analysis_filter_summary.pack(anchor='w', padx=15, pady=(0, 10))

        # Focused (conditional) sections - place near the filter inputs
        self.focused_container = tk.Frame(scrollable_frame, bg=self.colors['primary'])
        self.focused_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 10))
        # Student-specific (left) and Book-specific (right)
        self.student_specific_frame = tk.LabelFrame(
            self.focused_container,
            text="👤 Student-specific Insights",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.book_specific_frame = tk.LabelFrame(
            self.focused_container,
            text="📖 Book-specific Insights",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )

    # Charts container
        charts_container = tk.Frame(scrollable_frame, bg=self.colors['primary'])
        charts_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Row 1: Pie charts
        pie_row = tk.Frame(charts_container, bg=self.colors['primary'])
        pie_row.pack(fill=tk.X, pady=(0, 20))
        
        # Borrowing Status Pie Chart
        self.borrow_status_frame = tk.LabelFrame(
            pie_row,
            text="📚 Book Status Distribution",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.borrow_status_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # Student Activity Pie Chart
        self.student_activity_frame = tk.LabelFrame(
            pie_row,
            text="👥 Student Activity Distribution",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.student_activity_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        # Row 2: Branch-wise charts
        branch_row = tk.Frame(charts_container, bg=self.colors['primary'])
        branch_row.pack(fill=tk.X, pady=(0, 20))
        
        # Branch-wise Borrowing Distribution
        self.branch_activity_frame = tk.LabelFrame(
            branch_row,
            text="🏢 Branch-wise Borrowing Distribution",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.branch_activity_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # Branch-wise Overdue Distribution
        self.branch_overdue_frame = tk.LabelFrame(
            branch_row,
            text="⚠️ Branch-wise Overdue Distribution",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.branch_overdue_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))

        # Row 3: Advanced donut pie (Inventory & Overdue breakdown)
        advanced_row = tk.Frame(charts_container, bg=self.colors['primary'])
        advanced_row.pack(fill=tk.X, pady=(0, 20))

        self.inventory_overdue_frame = tk.LabelFrame(
            advanced_row,
            text="🍩 Inventory & Overdue Breakdown",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.inventory_overdue_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # Row 3: Summary stats
        self.stats_summary_frame = tk.LabelFrame(
            charts_container,
            text="📋 Summary Statistics",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.stats_summary_frame.pack(fill=tk.X, pady=(0, 20))

        # Row 4: Popular & Least Popular Books
        books_popularity_row = tk.Frame(charts_container, bg=self.colors['primary'])
        books_popularity_row.pack(fill=tk.X, pady=(0, 20))
        
        self.popular_books_frame = tk.LabelFrame(
            books_popularity_row,
            text="📈 Most Popular Books (High Demand)",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.popular_books_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        self.least_popular_books_frame = tk.LabelFrame(
            books_popularity_row,
            text="📉 Least Popular Books (Low Demand)",
            font=('Segoe UI', 12, 'bold'),
            bg=self.colors['primary'],
            fg=self.colors['accent']
        )
        self.least_popular_books_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))

        # (Focused sections are now placed near the filter controls above)
        
    # Footer: compact mode and maintenance actions
        footer_frame = tk.Frame(scrollable_frame, bg=self.colors['primary'])
        footer_frame.pack(fill=tk.X, padx=20, pady=(10, 20))

        self.analysis_compact_mode = tk.BooleanVar(value=True)
        compact_cb = tk.Checkbutton(
            footer_frame,
            text="Compact Mode (show fewer charts)",
            variable=self.analysis_compact_mode,
            onvalue=True,
            offvalue=False,
            command=self.refresh_analysis,
            bg=self.colors['primary'],
            fg=self.colors['accent'],
            selectcolor='white'
        )
        compact_cb.pack(side=tk.LEFT)

        # Toggle to hide charts entirely (for super clean UI)
        self.analysis_show_charts = tk.BooleanVar(value=True)
        hide_cb = tk.Checkbutton(
            footer_frame,
            text="Show Charts",
            variable=self.analysis_show_charts,
            onvalue=True,
            offvalue=False,
            command=self.refresh_analysis,
            bg=self.colors['primary'],
            fg=self.colors['accent'],
            selectcolor='white'
        )
        hide_cb.pack(side=tk.LEFT, padx=(20, 0))

        # Manual refresh
        tk.Button(
            footer_frame,
            text="Refresh Analysis",
            font=('Segoe UI', 10, 'bold'),
            bg=self.colors['secondary'], fg='white', relief='flat', padx=12, pady=4,
            cursor='hand2', command=self.refresh_analysis
        ).pack(side=tk.LEFT, padx=(20, 0))

        # Removed duplicate Promote Student button from footer; action available in header

        # Store chart references for click handling
        self.current_charts = {}
        # Filter state (None or string values)
        self.analysis_filter = {'enrollment_no': None, 'book_id': None}
        
        # Initial load
        self.refresh_analysis()
    
    def create_analysis_unavailable_message(self, parent):
        """Show message when matplotlib is not available"""
        container = tk.Frame(parent, bg=self.colors['primary'])
        container.pack(expand=True, fill=tk.BOTH)
        
        message_frame = tk.Frame(container, bg='white', relief='solid', bd=1)
        message_frame.place(relx=0.5, rely=0.5, anchor='center', width=600, height=400)
        
        tk.Label(
            message_frame,
            text="📊 Analysis Features Unavailable",
            font=('Segoe UI', 18, 'bold'),
            bg='white',
            fg=self.colors['accent']
        ).pack(pady=(40, 20))
        
        tk.Label(
            message_frame,
            text="The Analysis tab requires additional packages.\nPlease install the missing dependencies:",
            font=('Segoe UI', 12),
            bg='white',
            fg='#666666',
            justify='center'
        ).pack(pady=(0, 20))
        
        code_frame = tk.Frame(message_frame, bg='#f8f9fa', relief='solid', bd=1)
        code_frame.pack(padx=40, pady=20, fill=tk.X)
        
        tk.Label(
            code_frame,
            text="pip install matplotlib xlsxwriter",
            font=('Consolas', 11, 'bold'),
            bg='#f8f9fa',
            fg='#d73502'
        ).pack(pady=15)
        
        tk.Label(
            message_frame,
            text="After installation, restart the application to access charts and graphs.",
            font=('Segoe UI', 10),
            bg='white',
            fg='#666666'
        ).pack(pady=20)
    
    def create_student_portal_tab(self):
        """Create comprehensive student portal management tab with sub-sections"""
        portal_frame = tk.Frame(self.notebook, bg=self.colors['primary'])
        self.notebook.add(portal_frame, text="📱 Portal")
        
        # Initialize server tracking variables FIRST
        self.portal_url_history = []
        self.portal_start_time = None
        self.portal_request_count = 0
        self.portal_url = f"http://127.0.0.1:{self.portal_port}"
        self.health_indicators = {}
        
        # Create internal notebook for sub-tabs
        portal_notebook = ttk.Notebook(portal_frame)
        portal_notebook.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
        
        # Configure notebook style
        style = ttk.Style()
        style.configure('Portal.TNotebook.Tab', padding=[15, 8], font=('Segoe UI', 10, 'bold'))
        
        # Sub-tab 1: QR Access
        qr_tab = tk.Frame(portal_notebook, bg='white')
        portal_notebook.add(qr_tab, text="📱 QR Access")
        self._create_qr_access_section(qr_tab)
        
        # Sub-tab 2: All Requests
        requests_tab = tk.Frame(portal_notebook, bg='white')
        portal_notebook.add(requests_tab, text="📋 Requests")
        self._create_requests_section(requests_tab)
        
        # Sub-tab 3: Deletion Requests
        deletion_tab = tk.Frame(portal_notebook, bg='white')
        portal_notebook.add(deletion_tab, text="🗑️ Deletions")
        self._create_deletion_section(deletion_tab)
        
        # Sub-tab 4: Password Resets
        password_tab = tk.Frame(portal_notebook, bg='white')
        portal_notebook.add(password_tab, text="🔑 Password Reset")
        self._create_password_reset_section(password_tab)
        
        # Sub-tab 5: Broadcasts (New)
        broadcast_tab = tk.Frame(portal_notebook, bg='white')
        portal_notebook.add(broadcast_tab, text="📢 Broadcasts")
        self._create_broadcast_section(broadcast_tab)

        # Sub-tab 6: Observability
        observability_tab = tk.Frame(portal_notebook, bg='white')
        portal_notebook.add(observability_tab, text="📈 Observability")
        self._create_observability_section(observability_tab)
        
        # Sub-tab 7: Study Materials
        materials_tab = tk.Frame(portal_notebook, bg='white')
        portal_notebook.add(materials_tab, text="📚 Study Materials")
        self._create_study_materials_section(materials_tab)

    def _create_broadcast_section(self, parent):
        """Create broadcast notice management section"""
        # Header
        tk.Label(
            parent,
            text="📢 Broadcast System",
            font=('Segoe UI', 18, 'bold'),
            bg='white',
            fg=self.colors['accent']
        ).pack(pady=(20, 5), padx=20, anchor='w')
        
        tk.Label(
            parent,
            text="Post important announcements to all students on the portal dashboard.",
            font=('Segoe UI', 10),
            bg='white',
            fg='#666'
        ).pack(padx=20, anchor='w', pady=(0, 20))
        
        # Main layout: Left (Form), Right (Active List)
        container = tk.Frame(parent, bg='white')
        container.pack(fill=tk.BOTH, expand=True, padx=20)
        container.columnconfigure(0, weight=1)
        container.columnconfigure(1, weight=1)
        
        # --- Left: Post Notice Form ---
        form_card = tk.Frame(container, bg='#f8f9fa', relief='solid', bd=1)
        form_card.grid(row=0, column=0, sticky='nsew', padx=(0, 10), pady=5)
        
        tk.Label(
            form_card,
            text="📝 Post New Notice",
            font=('Segoe UI', 12, 'bold'),
            bg='#f1f3f4',
            fg='#333',
            padx=15,
            pady=10
        ).pack(fill=tk.X)
        
        form_inner = tk.Frame(form_card, bg='#f8f9fa')
        form_inner.pack(padx=20, pady=20, fill=tk.X)
        
        tk.Label(form_inner, text="Title:", font=('Segoe UI', 10, 'bold'), bg='#f8f9fa').pack(anchor='w')
        self.notice_title = tk.Entry(form_inner, font=('Segoe UI', 11), width=30)
        self.notice_title.pack(fill=tk.X, pady=(5, 15))
        
        tk.Label(form_inner, text="Message:", font=('Segoe UI', 10, 'bold'), bg='#f8f9fa').pack(anchor='w')
        self.notice_msg = tk.Text(form_inner, height=4, font=('Segoe UI', 10), width=30)
        self.notice_msg.pack(fill=tk.X, pady=(5, 15))
        
        post_btn = tk.Button(
            form_inner,
            text="📢 Post Notice",
            font=('Segoe UI', 10, 'bold'),
            bg='#007bff', # Hardcoded blue to ensure visibility
            fg='white',
            activebackground='#0056b3',
            activeforeground='white',
            padx=20,
            pady=8,
            cursor='hand2',
            relief='flat',
            command=self._post_notice
        )
        post_btn.pack(anchor='w', pady=(10, 0))
        
        # --- Right: Active Notices ---
        list_card = tk.Frame(container, bg='white', relief='solid', bd=1)
        list_card.grid(row=0, column=1, sticky='nsew', padx=(10, 0), pady=0)
        
        header_frame = tk.Frame(list_card, bg='#f1f3f4')
        header_frame.pack(fill=tk.X)
        
        tk.Label(
            header_frame,
            text="📡 Active Broadcasts",
            font=('Segoe UI', 12, 'bold'),
            bg='#f1f3f4',
            fg='#333',
            padx=15,
            pady=10
        ).pack(side=tk.LEFT)
        
        tk.Button(
            header_frame,
            text="🔄 Refresh",
            font=('Segoe UI', 8),
            bg='white',
            fg='#333',
            relief='solid',
            bd=1,
            command=self._refresh_notices
        ).pack(side=tk.RIGHT, padx=10)
        
        # Scrollable container
        self.notices_canvas = tk.Canvas(list_card, bg='white')
        scrollbar = ttk.Scrollbar(list_card, orient="vertical", command=self.notices_canvas.yview)
        self.notices_container = tk.Frame(self.notices_canvas, bg='white')
        
        self.notices_container.bind(
            "<Configure>",
            lambda e: self.notices_canvas.configure(scrollregion=self.notices_canvas.bbox("all"))
        )
        
        self.notices_canvas.create_window((0, 0), window=self.notices_container, anchor="nw")
        
        self.notices_canvas.configure(yscrollcommand=scrollbar.set)
        
        self.notices_canvas.pack(side="left", fill="both", expand=True, padx=5, pady=5)
        scrollbar.pack(side="right", fill="y", pady=5)
        
        # Initial Load
        self._refresh_notices()

    def _post_notice(self):
        """Send notice to backend"""
        if not WEB_PORTAL_AVAILABLE:
            messagebox.showwarning("Unavailable", "Web portal server is not running.")
            return

        title = self.notice_title.get().strip()
        msg = self.notice_msg.get("1.0", tk.END).strip()
        
        if not title or not msg:
            messagebox.showwarning("Incomplete", "Please enter both title and message.")
            return
        
        try:
            import urllib.request
            import json
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/notices"
            data = json.dumps({"title": title, "message": msg}).encode('utf-8')
            req = urllib.request.Request(url, data=data, headers={'Content-Type': 'application/json'}, method='POST')
            
            with urllib.request.urlopen(req, timeout=1) as response:
                res = json.loads(response.read().decode())
                if res['status'] == 'success':
                    messagebox.showinfo("Success", "Notice posted successfully!")
                    self.notice_title.delete(0, tk.END)
                    self.notice_msg.delete("1.0", tk.END)
                    self._refresh_notices()
                else:
                    messagebox.showerror("Error", res.get('message', 'Failed to post'))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to post notice: {e}")

    def _refresh_notices(self):
        """Fetch active notices"""
        for w in self.notices_container.winfo_children():
            w.destroy()
            
        if not WEB_PORTAL_AVAILABLE:
            tk.Label(self.notices_container, text="Portal server unavailable", bg='white', fg='red').pack(pady=20)
            return
            
        try:
            import urllib.request
            import json
            
            url = f"http://127.0.0.1:{self.portal_port}/api/notices" # Get active only
            req = urllib.request.Request(url)
            
            with urllib.request.urlopen(req, timeout=1) as response:
                data = json.loads(response.read().decode())
                notices = data.get('notices', [])
                
                if not notices:
                    tk.Label(self.notices_container, text="No active notices", bg='white', fg='#999').pack(pady=20)
                    return
                
                for notice in notices:
                    card = tk.Frame(self.notices_container, bg='#fff3cd', relief='solid', bd=1)
                    card.pack(fill=tk.X, padx=10, pady=5)
                    
                    header = tk.Frame(card, bg='#fff3cd')
                    header.pack(fill=tk.X, padx=10, pady=5)
                    
                    tk.Label(header, text=notice['title'], font=('Segoe UI', 10, 'bold'), bg='#fff3cd', fg='#856404').pack(side=tk.LEFT)
                    
                    # Delete Button
                    tk.Button(
                        header,
                        text="✕",
                        font=('Segoe UI', 8, 'bold'),
                        bg='#dc3545',
                        fg='white',
                        relief='flat',
                        width=2,
                        command=lambda nid=notice['id']: self._delete_notice(nid)
                    ).pack(side=tk.RIGHT)
                    
                    tk.Label(card, text=notice['message'], font=('Segoe UI', 9), bg='#fff3cd', wraplength=300, justify='left').pack(anchor='w', padx=10, pady=(0, 10))
                    
                    creation = notice.get('created_at', '')[:16]
                    tk.Label(card, text=creation, font=('Segoe UI', 7), bg='#fff3cd', fg='#856404').pack(anchor='e', padx=10, pady=(0, 5))
                    
        except Exception as e:
            tk.Label(self.notices_container, text=f"Error loading: {e}", bg='white', fg='red').pack()

    def _delete_notice(self, notice_id):
        """Deactivate notice"""
        if not WEB_PORTAL_AVAILABLE:
            return

        if not messagebox.askyesno("Confirm", "Are you sure you want to delete this notice?"):
            return
            
        try:
            import urllib.request
            import json
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/notices/{notice_id}"
            req = urllib.request.Request(url, method='DELETE')
            
            with urllib.request.urlopen(req, timeout=1) as response:
                self._refresh_notices()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete: {e}")
        
        # Store reference for refreshing
        # Line removed to fix NameError
        
    def _create_observability_section(self, parent):
        """Create comprehensive observability dashboard with multiple charts and insights"""
        if not MATPLOTLIB_AVAILABLE:
            self.create_analysis_unavailable_message(parent)
            return

        # Create scrollable container
        main_canvas = tk.Canvas(parent, bg='white', highlightthickness=0)
        v_scrollbar = ttk.Scrollbar(parent, orient='vertical', command=main_canvas.yview)
        scrollable_frame = tk.Frame(main_canvas, bg='white')
        
        # Store window id for width updates
        window_id = main_canvas.create_window((0, 0), window=scrollable_frame, anchor='nw')
        
        def _on_frame_configure(e):
            main_canvas.configure(scrollregion=main_canvas.bbox("all"))
        
        def _on_canvas_configure(e):
            # Make scrollable_frame fill the full canvas width
            main_canvas.itemconfig(window_id, width=e.width)
        
        scrollable_frame.bind("<Configure>", _on_frame_configure)
        main_canvas.bind("<Configure>", _on_canvas_configure)
        main_canvas.configure(yscrollcommand=v_scrollbar.set)
        
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        main_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Mousewheel scrolling with proper enter/leave handling
        def _on_mousewheel(event):
            main_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        def _on_enter(event):
            main_canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        def _on_leave(event):
            main_canvas.unbind_all("<MouseWheel>")
        
        main_canvas.bind("<Enter>", _on_enter)
        main_canvas.bind("<Leave>", _on_leave)
        parent.bind("<Destroy>", lambda e: main_canvas.unbind_all("<MouseWheel>"), add='+')


        # Header Section
        header_frame = tk.Frame(scrollable_frame, bg='white')
        header_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        title_container = tk.Frame(header_frame, bg='white')
        title_container.pack(side=tk.LEFT)
        
        tk.Label(title_container, text="📈 Observability Dashboard",
            font=('Segoe UI', 20, 'bold'), bg='white', fg=self.colors['accent']).pack(anchor='w')
        tk.Label(title_container, text="Real-time server traffic analytics and performance insights",
            font=('Segoe UI', 10), bg='white', fg='#666').pack(anchor='w')
        
        tk.Button(header_frame, text="🔄 Refresh All", font=('Segoe UI', 10, 'bold'),
            bg=self.colors['secondary'], fg='white', relief='flat', padx=15, pady=6, cursor='hand2',
            command=lambda: self._refresh_observability_dashboard(scrollable_frame)).pack(side=tk.RIGHT)
        
        self.obs_last_updated = tk.Label(header_frame, text="", font=('Segoe UI', 9), bg='white', fg='#888')
        self.obs_last_updated.pack(side=tk.RIGHT, padx=15)
        
        ttk.Separator(scrollable_frame, orient='horizontal').pack(fill=tk.X, padx=20, pady=10)
        
        # KPI Cards, Charts, and Insights containers
        self.obs_kpi_container = tk.Frame(scrollable_frame, bg='white')
        self.obs_kpi_container.pack(fill=tk.X, padx=20, pady=(5, 15))
        
        self.obs_charts_row1 = tk.Frame(scrollable_frame, bg='white')
        self.obs_charts_row1.pack(fill=tk.X, padx=20, pady=10)
        
        self.obs_charts_row2 = tk.Frame(scrollable_frame, bg='white')
        self.obs_charts_row2.pack(fill=tk.X, padx=20, pady=10)
        
        self.obs_trend_container = tk.Frame(scrollable_frame, bg='white')
        self.obs_trend_container.pack(fill=tk.X, padx=20, pady=10)
        
        self.obs_insights_container = tk.Frame(scrollable_frame, bg='#f0f8ff', relief='solid', bd=1)
        self.obs_insights_container.pack(fill=tk.X, padx=20, pady=(10, 20))
        
        self.obs_scrollable_frame = scrollable_frame
        self.traffic_graph_container = self.obs_charts_row1
        
        self._refresh_observability_dashboard(scrollable_frame)

    def _refresh_observability_dashboard(self, parent):
        """Refresh all observability components with comprehensive analytics"""
        # Clear all containers
        for container in [self.obs_kpi_container, self.obs_charts_row1, 
                          self.obs_charts_row2, self.obs_trend_container, self.obs_insights_container]:
            for widget in container.winfo_children():
                widget.destroy()
        
        try:
            import urllib.request
            import urllib.error
            
            # Fetch observability data from API
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/observability"
            try:
                response = urllib.request.urlopen(url, timeout=5)
                data = json.loads(response.read().decode())
            except urllib.error.URLError as e:
                tk.Label(self.obs_kpi_container, text="⚠️ Server not running. Start the portal server to view analytics.",
                    font=('Segoe UI', 12), bg='white', fg='#dc3545').pack(pady=30)
                return
            except Exception as e:
                tk.Label(self.obs_kpi_container, text=f"⚠️ Error connecting to server: {e}",
                    font=('Segoe UI', 12), bg='white', fg='#dc3545').pack(pady=30)
                return
            
            if data.get('status') == 'error':
                tk.Label(self.obs_kpi_container, text=f"⚠️ {data.get('message', 'Unknown error')}",
                    font=('Segoe UI', 12), bg='white', fg='#dc3545').pack(pady=30)
                return

            # ═══════════════════════════════════════════════════════════════
            # EXTRACT DATA FROM API RESPONSE
            # ═══════════════════════════════════════════════════════════════
            total_24h = data.get('total_24h', 0)
            total_7d = data.get('total_7d', 0)
            success_24h = data.get('success_24h', 0)
            success_rate = data.get('success_rate', 0)
            errors_24h = data.get('errors_24h', 0)
            peak_hour = data.get('peak_hour', 'N/A')
            peak_count = data.get('peak_count', 0)
            hourly_data = data.get('hourly_data', {})
            endpoint_data = data.get('endpoint_data', [])
            status_data = data.get('status_data', {})
            trend_data = data.get('trend_data', [])

            # ═══════════════════════════════════════════════════════════════
            # BUILD KPI CARDS
            # ═══════════════════════════════════════════════════════════════
            kpis = [
                ("📊 Total Requests", f"{total_24h:,}", "Last 24 Hours", self.colors['secondary']),
                ("✅ Success Rate", f"{success_rate:.1f}%", f"{success_24h:,} successful", '#28a745' if success_rate > 95 else '#ffc107' if success_rate > 80 else '#dc3545'),
                ("⏰ Peak Hour", peak_hour, f"{peak_count:,} requests", '#6f42c1'),
                ("⚠️ Errors", f"{errors_24h:,}", "4xx + 5xx responses", '#dc3545' if errors_24h > 10 else '#28a745'),
                ("📈 Weekly Total", f"{total_7d:,}", "Last 7 Days", '#17a2b8'),
            ]
            
            for i, (title, value, subtitle, color) in enumerate(kpis):
                card = tk.Frame(self.obs_kpi_container, bg=color, relief='flat', bd=0)
                card.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
                
                inner = tk.Frame(card, bg='white', relief='flat')
                inner.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
                
                tk.Label(inner, text=title, font=('Segoe UI', 9), bg='white', fg='#666').pack(pady=(10, 2))
                tk.Label(inner, text=value, font=('Segoe UI', 22, 'bold'), bg='white', fg=color).pack()
                tk.Label(inner, text=subtitle, font=('Segoe UI', 8), bg='white', fg='#888').pack(pady=(2, 10))
            
            # ═══════════════════════════════════════════════════════════════
            # CHART 1: Hourly Traffic (Bar Chart)
            # ═══════════════════════════════════════════════════════════════
            chart1_frame = tk.LabelFrame(self.obs_charts_row1, text=" 📊 Requests by Hour (24h) ", 
                font=('Segoe UI', 10, 'bold'), bg='white', fg=self.colors['accent'])
            chart1_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
            
            fig1 = Figure(figsize=(5, 3), dpi=100)
            fig1.patch.set_facecolor('white')
            ax1 = fig1.add_subplot(111)
            
            hours = [f"{h:02d}" for h in range(24)]
            counts = [hourly_data.get(h, 0) for h in hours]
            colors_bars = [self.colors['secondary'] if c < peak_count else '#28a745' for c in counts]
            
            ax1.bar(hours, counts, color=colors_bars, alpha=0.85, width=0.7)
            ax1.set_xlabel("Hour", fontsize=8)
            ax1.set_ylabel("Requests", fontsize=8)
            ax1.tick_params(axis='both', labelsize=7)
            ax1.spines['top'].set_visible(False)
            ax1.spines['right'].set_visible(False)
            ax1.grid(axis='y', alpha=0.3, linestyle='--')
            fig1.tight_layout()
            
            canvas1 = FigureCanvasTkAgg(fig1, chart1_frame)
            canvas1.draw()
            canvas1.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            
            # ═══════════════════════════════════════════════════════════════
            # CHART 2: Top Endpoints (Horizontal Bar)
            # ═══════════════════════════════════════════════════════════════
            chart2_frame = tk.LabelFrame(self.obs_charts_row1, text=" 🔗 Top Endpoints ", 
                font=('Segoe UI', 10, 'bold'), bg='white', fg=self.colors['accent'])
            chart2_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            
            fig2 = Figure(figsize=(5, 3), dpi=100)
            fig2.patch.set_facecolor('white')
            ax2 = fig2.add_subplot(111)
            
            if endpoint_data:
                endpoints = [e[0][-25:] for e in reversed(endpoint_data)]  # Truncate long names
                ep_counts = [e[1] for e in reversed(endpoint_data)]
                ax2.barh(endpoints, ep_counts, color='#6f42c1', alpha=0.8)
                ax2.set_xlabel("Requests", fontsize=8)
                ax2.tick_params(axis='both', labelsize=7)
            else:
                ax2.text(0.5, 0.5, "No data", ha='center', va='center')
            ax2.spines['top'].set_visible(False)
            ax2.spines['right'].set_visible(False)
            fig2.tight_layout()
            
            canvas2 = FigureCanvasTkAgg(fig2, chart2_frame)
            canvas2.draw()
            canvas2.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            
            # ═══════════════════════════════════════════════════════════════
            # CHART 3: Status Code Pie
            # ═══════════════════════════════════════════════════════════════
            chart3_frame = tk.LabelFrame(self.obs_charts_row2, text=" 🎯 Response Status Distribution ", 
                font=('Segoe UI', 10, 'bold'), bg='white', fg=self.colors['accent'])
            chart3_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
            
            fig3 = Figure(figsize=(4, 3), dpi=100)
            fig3.patch.set_facecolor('white')
            ax3 = fig3.add_subplot(111)
            
            if status_data:
                labels = list(status_data.keys())
                sizes = list(status_data.values())
                # Guard against all-zero data (causes NaN in matplotlib pie)
                if sum(sizes) > 0:
                    colors_pie = ['#28a745', '#17a2b8', '#ffc107', '#dc3545', '#6c757d'][:len(labels)]
                    wedges, texts, autotexts = ax3.pie(sizes, labels=labels, autopct='%1.0f%%', 
                        colors=colors_pie, startangle=90, textprops={'fontsize': 8})
                    for autotext in autotexts:
                        autotext.set_fontsize(8)
                        autotext.set_fontweight('bold')
                else:
                    ax3.text(0.5, 0.5, "No data yet", ha='center', va='center', fontsize=10, color='#888')
            else:
                ax3.text(0.5, 0.5, "No data", ha='center', va='center')
            ax3.axis('equal')
            fig3.tight_layout()
            
            canvas3 = FigureCanvasTkAgg(fig3, chart3_frame)
            canvas3.draw()
            canvas3.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            
            # ═══════════════════════════════════════════════════════════════
            # CHART 4: Live Metrics Panel
            # ═══════════════════════════════════════════════════════════════
            chart4_frame = tk.LabelFrame(self.obs_charts_row2, text=" 📈 Quick Stats ", 
                font=('Segoe UI', 10, 'bold'), bg='white', fg=self.colors['accent'])
            chart4_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            
            stats_inner = tk.Frame(chart4_frame, bg='white')
            stats_inner.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
            
            avg_per_hour = total_24h / 24 if total_24h else 0
            unique_endpoints = len(endpoint_data)
            
            quick_stats = [
                ("📊 Avg Requests/Hour", f"{avg_per_hour:.1f}"),
                ("🔗 Unique Endpoints", f"{unique_endpoints}"),
                ("🟢 2xx Responses", f"{status_data.get('2xx Success', 0):,}"),
                ("🟡 4xx Responses", f"{status_data.get('4xx Client Error', 0):,}"),
                ("🔴 5xx Responses", f"{status_data.get('5xx Server Error', 0):,}"),
            ]
            
            for label, val in quick_stats:
                row = tk.Frame(stats_inner, bg='white')
                row.pack(fill=tk.X, pady=3)
                tk.Label(row, text=label, font=('Segoe UI', 9), bg='white', fg='#555', anchor='w').pack(side=tk.LEFT)
                tk.Label(row, text=val, font=('Segoe UI', 10, 'bold'), bg='white', fg=self.colors['accent'], anchor='e').pack(side=tk.RIGHT)
            
            # ═══════════════════════════════════════════════════════════════
            # CHART 5: 7-Day Trend (Line Chart)
            # ═══════════════════════════════════════════════════════════════
            trend_frame = tk.LabelFrame(self.obs_trend_container, text=" 📅 7-Day Traffic Trend ", 
                font=('Segoe UI', 10, 'bold'), bg='white', fg=self.colors['accent'])
            trend_frame.pack(fill=tk.BOTH, expand=True)
            
            fig5 = Figure(figsize=(10, 2.5), dpi=100)
            fig5.patch.set_facecolor('white')
            ax5 = fig5.add_subplot(111)
            
            if trend_data:
                days = [t[0][-5:] for t in trend_data]  # MM-DD format
                day_counts = [t[1] for t in trend_data]
                ax5.fill_between(range(len(days)), day_counts, alpha=0.3, color=self.colors['secondary'])
                ax5.plot(range(len(days)), day_counts, color=self.colors['secondary'], linewidth=2, marker='o', markersize=6)
                ax5.set_xticks(range(len(days)))
                ax5.set_xticklabels(days, fontsize=8)
                ax5.set_ylabel("Requests", fontsize=9)
                for i, v in enumerate(day_counts):
                    ax5.annotate(str(v), (i, v), textcoords="offset points", xytext=(0, 8), ha='center', fontsize=8, fontweight='bold')
            else:
                ax5.text(0.5, 0.5, "No trend data available", ha='center', va='center')
            ax5.spines['top'].set_visible(False)
            ax5.spines['right'].set_visible(False)
            ax5.grid(axis='y', alpha=0.3, linestyle='--')
            fig5.tight_layout()
            
            canvas5 = FigureCanvasTkAgg(fig5, trend_frame)
            canvas5.draw()
            canvas5.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            
            # ═══════════════════════════════════════════════════════════════
            # INSIGHTS PANEL
            # ═══════════════════════════════════════════════════════════════
            insights_header = tk.Frame(self.obs_insights_container, bg='#f0f8ff')
            insights_header.pack(fill=tk.X, padx=15, pady=(10, 5))
            
            tk.Label(insights_header, text="💡 Intelligent Insights", 
                font=('Segoe UI', 12, 'bold'), bg='#f0f8ff', fg=self.colors['accent']).pack(side=tk.LEFT)
            
            insights_content = tk.Frame(self.obs_insights_container, bg='#f0f8ff')
            insights_content.pack(fill=tk.X, padx=15, pady=(0, 15))
            
            # Generate insights
            insights = []
            if peak_hour != "N/A":
                insights.append(f"⏰ Peak traffic occurs at {peak_hour} with {peak_count:,} requests")
            if success_rate < 95:
                insights.append(f"⚠️ Success rate is {success_rate:.1f}% - investigate error responses")
            elif success_rate >= 99:
                insights.append(f"✅ Excellent success rate of {success_rate:.1f}%!")
            if errors_24h > 0:
                error_pct = (errors_24h / total_24h * 100) if total_24h > 0 else 0
                insights.append(f"🔴 {errors_24h} errors detected ({error_pct:.1f}% of traffic)")
            if avg_per_hour < 5:
                insights.append("📉 Low traffic volume - server is underutilized")
            elif avg_per_hour > 100:
                insights.append("📈 High traffic volume - monitor server performance")
            if total_7d > 0 and total_24h > 0:
                daily_avg = total_7d / 7
                if total_24h > daily_avg * 1.5:
                    insights.append("🚀 Today's traffic is 50%+ higher than the weekly average!")
                elif total_24h < daily_avg * 0.5:
                    insights.append("📉 Today's traffic is below the weekly average")
            
            if not insights:
                insights.append("ℹ️ Collecting more data for actionable insights...")
            
            for insight in insights:
                tk.Label(insights_content, text=f"  • {insight}", 
                    font=('Segoe UI', 10), bg='#f0f8ff', fg='#333', anchor='w').pack(fill=tk.X, pady=2)
            
            # Update timestamp
            self.obs_last_updated.config(text=f"Last updated: {datetime.now().strftime('%H:%M:%S')}")
            
        except Exception as e:
            tk.Label(self.obs_kpi_container, text=f"Error loading observability data: {e}",
                font=('Segoe UI', 11), bg='white', fg='#dc3545').pack(pady=30)
            import traceback
            traceback.print_exc()

    def _refresh_traffic_graph(self, parent):

        """Fetch logs and plot traffic"""
        # Clear previous
        for widget in self.traffic_graph_container.winfo_children():
            widget.destroy()
            
        try:
            import urllib.request
            import urllib.error
            
            # Fetch observability data from API
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/observability"
            try:
                response = urllib.request.urlopen(url, timeout=5)
                data = json.loads(response.read().decode())
            except urllib.error.URLError:
                tk.Label(self.traffic_graph_container, text="Server not running.\nStart the portal server to view traffic.", font=('Segoe UI', 12), bg='white', fg='#dc3545').pack(pady=50)
                return
            except Exception as e:
                tk.Label(self.traffic_graph_container, text=f"Error connecting to server: {e}", font=('Segoe UI', 12), bg='white', fg='#dc3545').pack(pady=50)
                return
            
            if data.get('status') == 'error':
                tk.Label(self.traffic_graph_container, text=f"{data.get('message', 'Unknown error')}", font=('Segoe UI', 12), bg='white', fg='#dc3545').pack(pady=50)
                return
            
            # Get hourly data from API response
            hours_map = data.get('hourly_data', {})
            
            # Generate full 24h timeline (last 24 hours from now) or just 00-23?
            # Let's do a simple 00-23 plot for "Today" or just mapped to hours which is simpler
            # A 24-hour rolling window might be complex to label.
            # Let's stick to 00-23 clock hours for simplicity as typically used in such dashboards
            
            full_hours = [f"{h:02d}" for h in range(24)]
            full_counts = [hours_map.get(h, 0) for h in full_hours]
            
            # Create Plot
            fig = Figure(figsize=(8, 4), dpi=100)
            fig.patch.set_facecolor('white')
            ax = fig.add_subplot(111)
            
            # Plot bar chart
            # Use a nice blue color
            bars = ax.bar(full_hours, full_counts, color=self.colors['secondary'], alpha=0.8, width=0.6)
            
            # Styling
            ax.set_title("Request Volume by Hour (Last 24 Hours)", fontsize=12, fontweight='bold', color=self.colors['accent'], pad=15)
            ax.set_xlabel("Hour of Day (00-23)", fontsize=10, color='#555')
            ax.set_ylabel("Number of Requests", fontsize=10, color='#555')
            
            # Ticks styling
            ax.tick_params(axis='x', colors='#666', labelsize=8)
            ax.tick_params(axis='y', colors='#666', labelsize=8)
            
            # Grid
            ax.grid(True, axis='y', linestyle='--', alpha=0.2, color='#999')
            
            # Remove top and right spines
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#ddd')
            ax.spines['bottom'].set_color('#ddd')
            
            # Annotate non-zero bars
            for bar in bars:
                height = bar.get_height()
                if height > 0:
                    ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                            f'{int(height)}',
                            ha='center', va='bottom', fontsize=8, color=self.colors['accent'])

            # Embed in Tkinter
            canvas = FigureCanvasTkAgg(fig, master=self.traffic_graph_container)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
            
        except Exception as e:
            tk.Label(self.traffic_graph_container, text=f"Error loading traffic graph: {e}", bg='white', fg='red').pack(pady=50)
    
    def _create_study_materials_section(self, parent):
        """Create study materials management section for uploading files"""
        # Header
        header_frame = tk.Frame(parent, bg='white')
        header_frame.pack(fill=tk.X, padx=20, pady=(20, 10))
        
        tk.Label(
            header_frame,
            text="📚 Study Materials Manager",
            font=('Segoe UI', 18, 'bold'),
            bg='white',
            fg=self.colors['accent']
        ).pack(side=tk.LEFT)
        
        tk.Button(
            header_frame,
            text="🔄 Refresh",
            font=('Segoe UI', 9, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            relief='flat',
            padx=12,
            pady=4,
            cursor='hand2',
            command=lambda: self._refresh_study_materials()
        ).pack(side=tk.RIGHT)
        
        tk.Label(
            parent,
            text="Upload study materials (PDFs, docs, images, etc.) and share with students.",
            font=('Segoe UI', 10),
            bg='white',
            fg='#666'
        ).pack(padx=20, anchor='w', pady=(0, 20))
        
        # Main container with two columns
        main_container = tk.Frame(parent, bg='white')
        main_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # LEFT: Upload Form
        form_frame = tk.Frame(main_container, bg='white', relief='solid', bd=1)
        form_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        tk.Label(form_frame, text="📤 Upload New Material", font=('Segoe UI', 13, 'bold'), 
                bg='white', fg=self.colors['accent']).pack(pady=15, padx=15, anchor='w')
        
        # Form fields
        form_inner = tk.Frame(form_frame, bg='white')
        form_inner.pack(fill=tk.BOTH, expand=True, padx=20, pady=(5, 10))
        
        # Title
        tk.Label(form_inner, text="Title *", font=('Segoe UI', 9, 'bold'), bg='white').grid(row=0, column=0, sticky='w', pady=(0,2))
        self.material_title = tk.Entry(form_inner, font=('Segoe UI', 10), width=35)
        self.material_title.grid(row=1, column=0, sticky='ew', pady=(0,5))
        
        # Description
        tk.Label(form_inner, text="Description", font=('Segoe UI', 9, 'bold'), bg='white').grid(row=2, column=0, sticky='w', pady=(0,2))
        self.material_desc = tk.Text(form_inner, font=('Segoe UI', 10), height=2, width=35, wrap=tk.WORD)
        self.material_desc.grid(row=3, column=0, sticky='ew', pady=(0,5))
        
        # File Selection
        tk.Label(form_inner, text="Select File *", font=('Segoe UI', 9, 'bold'), bg='white').grid(row=4, column=0, sticky='w', pady=(0,2))
        
        self.selected_file_path = tk.StringVar(value="No file selected")
        self.selected_file_full_path = None  # Initialize file path variable
        
        # File row: entry + browse button side by side
        file_row = tk.Frame(form_inner, bg='white')
        file_row.grid(row=5, column=0, sticky='ew', pady=(0,5))
        
        file_display = tk.Entry(file_row, textvariable=self.selected_file_path, font=('Segoe UI', 9), 
                               state='readonly', fg='#666', relief='solid', bd=1)
        file_display.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        
        browse_btn = tk.Button(file_row, text="📁 Browse", font=('Segoe UI', 9, 'bold'), 
                              bg='#0078d4', fg='white', relief='flat', 
                              padx=10, pady=3, cursor='hand2', 
                              command=self._browse_study_material_file,
                              activebackground='#106ebe')
        browse_btn.pack(side=tk.RIGHT)
        
        # Semester + Category row (side by side)
        sem_cat_frame = tk.Frame(form_inner, bg='white')
        sem_cat_frame.grid(row=6, column=0, sticky='ew', pady=(0,5))
        
        tk.Label(sem_cat_frame, text="Semester *", font=('Segoe UI', 9, 'bold'), bg='white').pack(side=tk.LEFT, padx=(0,5))
        self.material_year = ttk.Combobox(sem_cat_frame, values=['1st', '2nd', '3rd', '4th', '5th', '6th'], state='readonly', width=8)
        self.material_year.set('1st')
        self.material_year.pack(side=tk.LEFT, padx=(0,15))
        
        tk.Label(sem_cat_frame, text="Category", font=('Segoe UI', 9, 'bold'), bg='white').pack(side=tk.LEFT, padx=(0,5))
        self.material_category = ttk.Combobox(sem_cat_frame, values=['Notes', 'PYQ', 'Study Material', 'Syllabus', 'Other'], state='readonly', width=12)
        self.material_category.set('Notes')
        self.material_category.pack(side=tk.LEFT)
        
        # Branch selection
        branch_frame = tk.Frame(form_inner, bg='white')
        branch_frame.grid(row=7, column=0, sticky='ew', pady=(0,5))
        tk.Label(branch_frame, text="Branch *", font=('Segoe UI', 9, 'bold'), bg='white').pack(side=tk.LEFT, padx=(0,5))
        branch_values = ['All Branches', 'Computer Engineering', 'Information Technology', 'Civil Engineering',
                         'Electronics & Telecommunication', 'Mechanical Engineering', 'Automobile Engineering', 'Electrical Engineering']
        self.material_branch = ttk.Combobox(branch_frame, values=branch_values, state='readonly', width=28)
        self.material_branch.set('All Branches')
        self.material_branch.pack(side=tk.LEFT)
        
        # Upload Button
        upload_btn = tk.Button(
            form_inner,
            text="📤 UPLOAD MATERIAL",
            font=('Segoe UI', 11, 'bold'),
            bg='#28a745',
            fg='white',
            relief='flat',
            padx=25,
            pady=8,
            cursor='hand2',
            command=self._upload_study_material,
            activebackground='#218838'
        )
        upload_btn.grid(row=8, column=0, sticky='ew', pady=(8,5))
        
        form_inner.grid_columnconfigure(0, weight=1)
        
        # RIGHT: Materials List
        list_frame = tk.Frame(main_container, bg='white', relief='solid', bd=1)
        list_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        tk.Label(list_frame, text="📋 Uploaded Materials", font=('Segoe UI', 13, 'bold'), 
                bg='white', fg=self.colors['accent']).pack(pady=15, padx=15, anchor='w')
        
        # Branch filter for material list
        mat_filter_frame = tk.Frame(list_frame, bg='#f8f9fa')
        mat_filter_frame.pack(fill=tk.X, padx=15, pady=(0, 5))
        
        tk.Label(mat_filter_frame, text="🏫 Branch:", font=('Segoe UI', 9, 'bold'), bg='#f8f9fa').pack(side=tk.LEFT, padx=(5, 5))
        self.material_filter_branch = ttk.Combobox(mat_filter_frame, state='readonly', width=28, font=('Segoe UI', 9))
        self.material_filter_branch['values'] = ['All Branches', 'Computer Engineering', 'Information Technology', 'Civil Engineering',
                                                   'Electronics & Telecommunication', 'Mechanical Engineering', 'Automobile Engineering', 'Electrical Engineering']
        self.material_filter_branch.set('All Branches')
        self.material_filter_branch.pack(side=tk.LEFT, padx=5)
        self.material_filter_branch.bind("<<ComboboxSelected>>", lambda e: self._refresh_study_materials())
        
        # Materials tree
        tree_container = tk.Frame(list_frame, bg='white')
        tree_container.pack(fill=tk.BOTH, expand=True, padx=15, pady=(10, 0))
        
        columns = ('ID', 'Title', 'Branch', 'Semester', 'Category', 'Date')
        self.materials_tree = ttk.Treeview(tree_container, columns=columns, show='headings', height=12)
        
        self.materials_tree.heading('ID', text='ID')
        self.materials_tree.heading('Title', text='Title')
        self.materials_tree.heading('Branch', text='Branch')
        self.materials_tree.heading('Semester', text='Semester')
        self.materials_tree.heading('Category', text='Category')
        self.materials_tree.heading('Date', text='Upload Date')
        
        self.materials_tree.column('ID', width=40)
        self.materials_tree.column('Title', width=150)
        self.materials_tree.column('Branch', width=130)
        self.materials_tree.column('Semester', width=65)
        self.materials_tree.column('Category', width=85)
        self.materials_tree.column('Date', width=100)
        
        scrollbar = ttk.Scrollbar(tree_container, orient='vertical', command=self.materials_tree.yview)
        self.materials_tree.configure(yscrollcommand=scrollbar.set)
        
        # Bind double-click to edit
        self.materials_tree.bind("<Double-1>", self._on_material_double_click)
        
        self.materials_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Action buttons
        btn_frame = tk.Frame(list_frame, bg='white')
        btn_frame.pack(pady=10)
        
        tk.Button(btn_frame, text="🗑️ Delete", font=('Segoe UI', 9, 'bold'), bg='#dc3545', fg='white',
                 relief='flat', padx=12, pady=5, cursor='hand2', command=self._delete_study_material).pack(side=tk.LEFT, padx=5)
        
        # Load materials
        self._refresh_study_materials()
    
    def _browse_study_material_file(self):
        """Open file dialog to select study material file"""
        from tkinter import filedialog
        
        filetypes = [
            ("All Files", "*.*"),
            ("PDF Files", "*.pdf"),
            ("Word Documents", "*.doc;*.docx"),
            ("PowerPoint", "*.ppt;*.pptx"),
            ("Text Files", "*.txt"),
            ("Images", "*.jpg;*.jpeg;*.png"),
            ("Archives", "*.zip;*.rar")
        ]
        
        filepath = filedialog.askopenfilename(
            title="Select Study Material File",
            filetypes=filetypes
        )
        
        if filepath:
            import os
            filename = os.path.basename(filepath)
            self.selected_file_path.set(filename)
            self.selected_file_full_path = filepath
    
    def _upload_study_material(self):
        """Upload study material file to server"""
        if not WEB_PORTAL_AVAILABLE:
            messagebox.showwarning("Unavailable", "Web portal is not running.")
            return
        
        title = self.material_title.get().strip()
        desc = self.material_desc.get("1.0", tk.END).strip()
        year = self.material_year.get()
        category = self.material_category.get()
        
        # Check if file is selected
        if not hasattr(self, 'selected_file_full_path') or not self.selected_file_full_path:
            messagebox.showwarning("No File", "Please select a file to upload.")
            return
        
        if not title or not year:
            messagebox.showwarning("Incomplete", "Please fill in Title and Year.")
            return
        
        try:
            import requests
            import os
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/study-materials"
            
            # Prepare multipart form data
            with open(self.selected_file_full_path, 'rb') as f:
                files = {'file': (os.path.basename(self.selected_file_full_path), f)}
                branch = self.material_branch.get() if hasattr(self, 'material_branch') else 'All Branches'
                data = {
                    "title": title,
                    "description": desc,
                    "year": year,
                    "category": category,
                    "branch": branch
                }
                
                response = requests.post(url, files=files, data=data, timeout=30)
                res = response.json()
                
                if res['status'] == 'success':
                    messagebox.showinfo("Success", "Study material uploaded successfully!")
                    # Clear form
                    self.material_title.delete(0, tk.END)
                    self.material_desc.delete("1.0", tk.END)
                    self.material_year.set('1st')
                    self.material_category.set('Notes')
                    if hasattr(self, 'material_branch'):
                        self.material_branch.set('All Branches')
                    self.selected_file_path.set("No file selected")
                    self.selected_file_full_path = None
                    self._refresh_study_materials()
                else:
                    messagebox.showerror("Error", res.get('message', 'Failed to upload'))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to upload: {e}")
    
    def _refresh_study_materials(self):
        """Refresh materials list from database"""
        if not WEB_PORTAL_AVAILABLE:
            return
        
        try:
            import urllib.request
            import json
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/study-materials"
            req = urllib.request.Request(url)
            
            with urllib.request.urlopen(req, timeout=2) as response:
                res = json.loads(response.read().decode())
                materials = res.get('materials', [])
                
                # Clear tree
                for item in self.materials_tree.get_children():
                    self.materials_tree.delete(item)
                
                # Get selected branch filter
                selected_branch = self.material_filter_branch.get() if hasattr(self, 'material_filter_branch') else 'All Branches'
                
                # Populate tree
                for mat in materials:
                    if mat.get('active', 1) == 1:  # Only show active materials
                        mat_branch = mat.get('branch', 'All Branches')
                        # Apply branch filter
                        if selected_branch != 'All Branches' and mat_branch != 'All Branches' and mat_branch != selected_branch:
                            continue
                        
                        date_str = mat['upload_date'].split()[0] if mat.get('upload_date') else 'N/A'
                        # Format file size
                        file_size = mat.get('file_size', 0)
                        size_str = f"{file_size / 1024:.1f} KB" if file_size < 1024*1024 else f"{file_size / (1024*1024):.1f} MB"
                        
                        # Shorten branch name for display
                        branch_display = mat_branch[:20] + '..' if len(mat_branch) > 20 else mat_branch
                        
                        self.materials_tree.insert('', 'end', values=(
                            mat['id'],
                            mat['title'][:35] + '...' if len(mat['title']) > 35 else mat['title'],
                            branch_display,
                            mat['year'],
                            mat.get('category', 'N/A'),
                            date_str
                        ), tags=(mat['id'],))
                        
        except Exception as e:
            print(f"Error refreshing materials: {e}")
    
    def _delete_study_material(self):
        """Delete selected study material"""
        selected = self.materials_tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select a material to delete.")
            return
        
        if not messagebox.askyesno("Confirm", "Delete this study material?"):
            return
        
        try:
            import urllib.request
            import json
            
            material_id = self.materials_tree.item(selected[0])['values'][0]
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/study-materials/{material_id}"
            req = urllib.request.Request(url, method='DELETE')
            
            with urllib.request.urlopen(req, timeout=2) as response:
                res = json.loads(response.read().decode())
                if res['status'] == 'success':
                    messagebox.showinfo("Success", "Material deleted successfully!")
                    self._refresh_study_materials()
                else:
                    messagebox.showerror("Error", res.get('message', 'Failed to delete'))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete: {e}")


    def _on_material_double_click(self, event):
        """Handle double-click on material to edit"""
        item = self.materials_tree.identify_row(event.y)
        if not item:
            return
            
        # Get material data
        values = self.materials_tree.item(item)['values']
        material_id = values[0]
        
        # Show edit dialog
        self._edit_study_material_dialog(material_id)
        
    def _edit_study_material_dialog(self, material_id):
        """Show dialog to edit study material"""
        current_item = None
        for item in self.materials_tree.get_children():
            if str(self.materials_tree.item(item)['values'][0]) == str(material_id):
                current_item = self.materials_tree.item(item)
                break
        
        if not current_item:
            return

        vals = current_item['values']
        current_title = vals[1]
        current_branch = vals[2]
        current_year = vals[3]
        current_cat = vals[4]
        
        # Remove ellipsis from title if present (rough restore)
        if str(current_title).endswith('...'):
            current_title = str(current_title)[:-3]
        if str(current_branch).endswith('..'):
            current_branch = str(current_branch)[:-2]

        dialog = tk.Toplevel(self.root)
        dialog.title(f"Edit Material #{material_id}")
        dialog.geometry("500x520")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (500 // 2)
        y = (dialog.winfo_screenheight() // 2) - (520 // 2)
        dialog.geometry(f"+{x}+{y}")
        
        tk.Label(dialog, text="✏️ Edit Material", font=('Segoe UI', 14, 'bold'), bg='white', fg=self.colors['accent']).pack(fill=tk.X, pady=10,padx=20)
        
        form_frame = tk.Frame(dialog, padx=20, pady=10)
        form_frame.pack(fill=tk.BOTH, expand=True)
        
        tk.Label(form_frame, text="Title", font=('Segoe UI', 10, 'bold')).pack(anchor='w')
        title_entry = tk.Entry(form_frame, font=('Segoe UI', 10), width=40)
        title_entry.insert(0, current_title)
        title_entry.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(form_frame, text="Description (Update optional)", font=('Segoe UI', 10, 'bold')).pack(anchor='w')
        desc_entry = tk.Text(form_frame, font=('Segoe UI', 10), height=3, width=40)
        desc_entry.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(form_frame, text="Semester", font=('Segoe UI', 10, 'bold')).pack(anchor='w')
        year_combo = ttk.Combobox(form_frame, values=['1st', '2nd', '3rd', '4th', '5th', '6th'], state='readonly')
        year_combo.set(current_year)
        year_combo.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(form_frame, text="Category", font=('Segoe UI', 10, 'bold')).pack(anchor='w')
        cat_combo = ttk.Combobox(form_frame, values=['Notes', 'PYQ', 'Study Material', 'Syllabus', 'Other'], state='readonly')
        cat_combo.set(current_cat)
        cat_combo.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(form_frame, text="Branch", font=('Segoe UI', 10, 'bold')).pack(anchor='w')
        branch_values = ['All Branches', 'Computer Engineering', 'Information Technology', 'Civil Engineering',
                         'Electronics & Telecommunication', 'Mechanical Engineering', 'Automobile Engineering', 'Electrical Engineering']
        branch_combo = ttk.Combobox(form_frame, values=branch_values, state='readonly')
        branch_combo.set(current_branch if current_branch in branch_values else 'All Branches')
        branch_combo.pack(fill=tk.X, pady=(0, 10))
        
        btn_frame = tk.Frame(dialog, pady=20)
        btn_frame.pack(fill=tk.X, padx=20)
        
        def save_changes():
            import urllib.request
            import json
            
            new_title = title_entry.get().strip()
            if not new_title:
                messagebox.showerror("Error", "Title is required", parent=dialog)
                return
                
            data = {
                "title": new_title,
                "description": desc_entry.get("1.0", tk.END).strip(),
                "year": year_combo.get(),
                "category": cat_combo.get(),
                "branch": branch_combo.get(),
                "drive_link": "" 
            }
            
            try:
                url = f"http://127.0.0.1:{self.portal_port}/api/admin/study-materials/{material_id}"
                req = urllib.request.Request(url, method='PUT')
                req.add_header('Content-Type', 'application/json')
                
                jsondata = json.dumps(data).encode('utf-8')
                
                with urllib.request.urlopen(req, data=jsondata, timeout=5) as response:
                    res = json.loads(response.read().decode())
                    if res.get('status') == 'success':
                        messagebox.showinfo("Success", "Material updated!", parent=dialog)
                        dialog.destroy()
                        self._refresh_study_materials()
                    else:
                        messagebox.showerror("Error", res.get('message', 'Update failed'), parent=dialog)
            except Exception as e:
                messagebox.showerror("Error", f"Update failed: {e}", parent=dialog)

        def delete_material():
            if messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this material?", parent=dialog):
                try:
                    import urllib.request
                    import json
                    url = f"http://127.0.0.1:{self.portal_port}/api/admin/study-materials/{material_id}"
                    req = urllib.request.Request(url, method='DELETE')
                    with urllib.request.urlopen(req, timeout=5) as response:
                        res = json.loads(response.read().decode())
                        if res.get('status') == 'success':
                            messagebox.showinfo("Deleted", "Material deleted successfully", parent=dialog)
                            dialog.destroy()
                            self._refresh_study_materials()
                        else:
                            messagebox.showerror("Error", res.get('message', 'Deletion failed'), parent=dialog)
                except Exception as e:
                    messagebox.showerror("Error", f"Deletion failed: {e}", parent=dialog)

        tk.Button(btn_frame, text="💾 Save Changes", command=save_changes, 
                 bg='#28a745', fg='white', font=('Segoe UI', 10, 'bold'), padx=15).pack(side=tk.RIGHT)
                 
        tk.Button(btn_frame, text="🗑️ Delete", command=delete_material, 
                 bg='#dc3545', fg='white', font=('Segoe UI', 10, 'bold'), padx=15).pack(side=tk.LEFT)

    def _create_qr_access_section(self, parent):
        """Create comprehensive QR code access section with server dashboard"""
        # Main container - fill entire space
        container = tk.Frame(parent, bg='white')
        container.pack(fill=tk.BOTH, expand=True, padx=25, pady=15)
        
        # ═══════════════════════════════════════════════════════════════
        # HEADER SECTION
        # ═══════════════════════════════════════════════════════════════
        header_frame = tk.Frame(container, bg='white')
        header_frame.pack(fill=tk.X, pady=(0, 15))
        
        tk.Label(
            header_frame,
            text="📱 Student Portal Server",
            font=('Segoe UI', 18, 'bold'),
            bg='white',
            fg=self.colors['accent']
        ).pack(side=tk.LEFT)
        
        # Server status badge
        self.server_status_badge = tk.Label(
            header_frame,
            text="● ONLINE",
            font=('Segoe UI', 10, 'bold'),
            bg='#d4edda',
            fg='#155724',
            padx=12,
            pady=4
        )
        self.server_status_badge.pack(side=tk.RIGHT, padx=5)
        
        # Refresh all button
        refresh_all_btn = tk.Button(
            header_frame,
            text="🔄 Refresh All",
            font=('Segoe UI', 9, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            padx=15,
            pady=4,
            cursor='hand2',
            relief='flat',
            command=self._refresh_server_dashboard
        )
        refresh_all_btn.pack(side=tk.RIGHT, padx=5)
        
        # Separator
        tk.Frame(container, height=2, bg='#e9ecef').pack(fill=tk.X, pady=(0, 20))
        
        # ═══════════════════════════════════════════════════════════════
        # MAIN CONTENT: Two-column layout using grid
        # ═══════════════════════════════════════════════════════════════
        main_content = tk.Frame(container, bg='white')
        main_content.pack(fill=tk.BOTH, expand=True)
        
        # Configure grid columns with weights for proper expansion
        main_content.columnconfigure(0, weight=2, minsize=280)  # QR Column
        main_content.columnconfigure(1, weight=3, minsize=400)  # Status Column
        main_content.rowconfigure(0, weight=1)
        
        # ═══════════════════════════════════════════════════════════════
        # LEFT COLUMN: QR Code Card
        # ═══════════════════════════════════════════════════════════════
        left_frame = tk.Frame(main_content, bg='white')
        left_frame.grid(row=0, column=0, sticky='nsew', padx=(0, 15))
        
        # QR Card
        qr_card = tk.Frame(left_frame, bg='#f8f9fa', relief='solid', bd=1)
        qr_card.pack(fill=tk.X)  # Don't expand, just fit content
        
        qr_inner = tk.Frame(qr_card, bg='#f8f9fa')
        qr_inner.pack(pady=15, padx=15)

        # QR Card Header
        tk.Label(
            qr_inner,
            text="📲 Scan to Access",
            font=('Segoe UI', 12, 'bold'),
            bg='#f8f9fa',
            fg='#333'
        ).pack()
        
        tk.Label(
            qr_inner,
            text="Point your phone camera at this QR code",
            font=('Segoe UI', 9),
            bg='#f8f9fa',
            fg='#666'
        ).pack(pady=(2, 10))
        
        # QR Code Container
        self.portal_qr_container = tk.Frame(qr_inner, bg='white', relief='solid', bd=2)
        self.portal_qr_container.pack(pady=5)
        
        # URL and Buttons - OUTSIDE the card for visibility
        url_btn_frame = tk.Frame(left_frame, bg='white')
        url_btn_frame.pack(fill=tk.X, pady=(10, 0))
        
        # Current URL Display
        self.portal_current_url_label = tk.Label(
            url_btn_frame,
            text="Loading...",
            font=('Consolas', 10, 'bold'),
            bg='white',
            fg=self.colors['secondary']
        )
        self.portal_current_url_label.pack(pady=(0, 8))
        
        # Button frame for Copy and Save buttons
        btn_frame = tk.Frame(url_btn_frame, bg='white')
        btn_frame.pack()
        
        # Copy URL button
        copy_btn = tk.Button(
            btn_frame,
            text="📋 Copy URL",
            font=('Segoe UI', 9, 'bold'),
            bg='#28a745',
            fg='white',
            padx=15,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._copy_portal_url
        )
        copy_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        # Save QR as PNG button
        save_qr_btn = tk.Button(
            btn_frame,
            text="💾 Save QR",
            font=('Segoe UI', 9, 'bold'),
            bg='#17a2b8',
            fg='white',
            padx=15,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._save_qr_as_image
        )
        save_qr_btn.pack(side=tk.LEFT)
        
        # WiFi Note
        tk.Label(
            url_btn_frame,
            text="⚡ Both devices must be on the same network",
            font=('Segoe UI', 8, 'italic'),
            bg='white',
            fg='#888'
        ).pack(pady=(8, 5))



        
        # ═══════════════════════════════════════════════════════════════
        # RIGHT COLUMN: Status Dashboard
        # ═══════════════════════════════════════════════════════════════
        right_frame = tk.Frame(main_content, bg='white')
        right_frame.grid(row=0, column=1, sticky='nsew')
        
        # STATUS CARD 1: Server Info
        self._create_status_card(right_frame, "🖥️ Server Information", self._build_server_info_content)
        
        # STATUS CARD 2: URL History
        self._create_status_card(right_frame, "🔗 URL History", self._build_url_history_content)
        
        # STATUS CARD 3: Health & Diagnostics
        self._create_status_card(right_frame, "🩺 Health & Diagnostics", self._build_health_content)
        
        # STATUS CARD 4: Quick Actions
        self._create_status_card(right_frame, "⚡ Quick Actions", self._build_quick_actions_content)
        
        # Initialize server and generate QR
        self._initialize_portal_server()
        
    def _create_status_card(self, parent, title, content_builder):
        """Create a status card with title and dynamic content"""
        card = tk.Frame(parent, bg='white', relief='solid', bd=1)
        card.pack(fill=tk.X, pady=(0, 12))
        
        # Header
        header = tk.Frame(card, bg='#f1f3f4')
        header.pack(fill=tk.X)
        
        tk.Label(
            header,
            text=title,
            font=('Segoe UI', 10, 'bold'),
            bg='#f1f3f4',
            fg='#333',
            padx=15,
            pady=8
        ).pack(side=tk.LEFT)
        
        # Content
        content_frame = tk.Frame(card, bg='white')
        content_frame.pack(fill=tk.X, padx=15, pady=12)
        
        # Build content
        content_builder(content_frame)
        
    def _build_server_info_content(self, parent):
        """Build server information content"""
        # Store reference for updates
        self.server_info_frame = parent
        
        # Row 1: Status + Port
        row1 = tk.Frame(parent, bg='white')
        row1.pack(fill=tk.X, pady=3)
        
        tk.Label(row1, text="Status:", font=('Segoe UI', 9, 'bold'), bg='white', fg='#555', width=12, anchor='w').pack(side=tk.LEFT)
        self.server_status_label = tk.Label(row1, text="● Running", font=('Segoe UI', 9), bg='white', fg='#28a745')
        self.server_status_label.pack(side=tk.LEFT)
        
        tk.Label(row1, text="Port:", font=('Segoe UI', 9, 'bold'), bg='white', fg='#555').pack(side=tk.LEFT, padx=20)
        tk.Label(row1, text=str(self.portal_port), font=('Consolas', 9), bg='white', fg='#333').pack(side=tk.LEFT, padx=5)
        
        # Row 2: Runtime
        row2 = tk.Frame(parent, bg='white')
        row2.pack(fill=tk.X, pady=3)
        
        tk.Label(row2, text="Runtime:", font=('Segoe UI', 9, 'bold'), bg='white', fg='#555', width=12, anchor='w').pack(side=tk.LEFT)
        self.server_runtime_label = tk.Label(row2, text="0h 0m 0s", font=('Consolas', 9), bg='white', fg='#333')
        self.server_runtime_label.pack(side=tk.LEFT)
        
        # Row 3: IP Address
        row3 = tk.Frame(parent, bg='white')
        row3.pack(fill=tk.X, pady=3)
        
        tk.Label(row3, text="IP Address:", font=('Segoe UI', 9, 'bold'), bg='white', fg='#555', width=12, anchor='w').pack(side=tk.LEFT)
        self.server_ip_label = tk.Label(row3, text="Detecting...", font=('Consolas', 9), bg='white', fg='#333')
        self.server_ip_label.pack(side=tk.LEFT)
        
        # Row 4: Traffic Estimate
        row4 = tk.Frame(parent, bg='white')
        row4.pack(fill=tk.X, pady=3)
        
        tk.Label(row4, text="Est. Traffic:", font=('Segoe UI', 9, 'bold'), bg='white', fg='#555', width=12, anchor='w').pack(side=tk.LEFT)
        self.server_traffic_label = tk.Label(row4, text="—", font=('Segoe UI', 9), bg='white', fg='#666')
        self.server_traffic_label.pack(side=tk.LEFT)
        
    def _build_url_history_content(self, parent):
        """Build URL history content with toggle"""
        self.url_history_frame = parent
        
        # Toggle for showing history
        toggle_frame = tk.Frame(parent, bg='white')
        toggle_frame.pack(fill=tk.X, pady=(0, 8))
        
        self.show_url_history = tk.BooleanVar(value=False)
        
        tk.Label(toggle_frame, text="Current:", font=('Segoe UI', 9, 'bold'), bg='white', fg='#555').pack(side=tk.LEFT)
        self.current_url_display = tk.Label(toggle_frame, text="Loading...", font=('Consolas', 9), bg='white', fg=self.colors['secondary'])
        self.current_url_display.pack(side=tk.LEFT, padx=8)
        
        # History toggle button
        self.history_toggle_btn = tk.Button(
            toggle_frame,
            text="▼ Show History",
            font=('Segoe UI', 8),
            bg='white',
            fg='#666',
            relief='flat',
            cursor='hand2',
            command=self._toggle_url_history
        )
        self.history_toggle_btn.pack(side=tk.RIGHT)
        
        # History container (hidden by default)
        self.url_history_container = tk.Frame(parent, bg='#f8f9fa')
        # Don't pack - will be toggled
        
    def _toggle_url_history(self):
        """Toggle URL history visibility"""
        if self.show_url_history.get():
            self.show_url_history.set(False)
            self.url_history_container.pack_forget()
            self.history_toggle_btn.config(text="▼ Show History")
        else:
            self.show_url_history.set(True)
            self._populate_url_history()
            self.url_history_container.pack(fill=tk.X, pady=(8, 0))
            self.history_toggle_btn.config(text="▲ Hide History")
    
    def _populate_url_history(self):
        """Populate URL history list"""
        for w in self.url_history_container.winfo_children():
            w.destroy()
        
        if not self.portal_url_history:
            tk.Label(
                self.url_history_container,
                text="No previous URLs recorded",
                font=('Segoe UI', 8, 'italic'),
                bg='#f8f9fa',
                fg='#999',
                pady=8,
                padx=10
            ).pack(anchor='w')
            return
        
        tk.Label(
            self.url_history_container,
            text="Previous Sessions (Last 3):",
            font=('Segoe UI', 8, 'bold'),
            bg='#f8f9fa',
            fg='#555',
            pady=5,
            padx=10
        ).pack(anchor='w')
        
        for i, url_entry in enumerate(self.portal_url_history[-3:]):
            row = tk.Frame(self.url_history_container, bg='#f8f9fa')
            row.pack(fill=tk.X, padx=10, pady=2)
            
            tk.Label(row, text=f"#{i+1}", font=('Segoe UI', 8), bg='#f8f9fa', fg='#888', width=3).pack(side=tk.LEFT)
            tk.Label(row, text=url_entry.get('url', ''), font=('Consolas', 8), bg='#f8f9fa', fg='#333').pack(side=tk.LEFT, padx=5)
            tk.Label(row, text=url_entry.get('time', ''), font=('Segoe UI', 7), bg='#f8f9fa', fg='#999').pack(side=tk.RIGHT)
        
    def _build_health_content(self, parent):
        """Build health diagnostics content"""
        self.health_frame = parent
        
        # Health indicators
        indicators = [
            ("Portal Server", "checking"),
            ("Database Connection", "checking"),
            ("Network Access", "checking")
        ]
        
        self.health_indicators = {}
        
        for name, status in indicators:
            row = tk.Frame(parent, bg='white')
            row.pack(fill=tk.X, pady=3)
            
            tk.Label(row, text=name + ":", font=('Segoe UI', 9), bg='white', fg='#555', width=18, anchor='w').pack(side=tk.LEFT)
            
            indicator = tk.Label(row, text="◯ Checking...", font=('Segoe UI', 9), bg='white', fg='#ffc107')
            indicator.pack(side=tk.LEFT)
            self.health_indicators[name] = indicator
        
        # Issues section
        tk.Frame(parent, height=1, bg='#e9ecef').pack(fill=tk.X, pady=(10, 8))
        
        self.issues_label = tk.Label(
            parent,
            text="✓ No issues detected",
            font=('Segoe UI', 9),
            bg='white',
            fg='#28a745'
        )
        self.issues_label.pack(anchor='w')
    
    def _build_quick_actions_content(self, parent):
        """Build quick actions content"""
        # Button container
        btn_frame = tk.Frame(parent, bg='white')
        btn_frame.pack(fill=tk.X, pady=5)
        
        # Restart Server Button
        restart_btn = tk.Button(
            btn_frame,
            text="🔄 Restart Server",
            font=('Segoe UI', 9),
            bg='#e9ecef',
            fg='#333',
            padx=12,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._restart_portal_server
        )
        restart_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        # View Logs Button
        logs_btn = tk.Button(
            btn_frame,
            text="📋 View Stats",
            font=('Segoe UI', 9),
            bg='#e9ecef',
            fg='#333',
            padx=12,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._show_portal_stats
        )
        logs_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        # Open in Browser Button
        browser_btn = tk.Button(
            btn_frame,
            text="🌐 Open in Browser",
            font=('Segoe UI', 9),
            bg='#e9ecef',
            fg='#333',
            padx=12,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._open_portal_in_browser
        )
        browser_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        # Save QR Button
        save_qr_btn = tk.Button(
            btn_frame,
            text="💾 Save QR",
            font=('Segoe UI', 9),
            bg='#17a2b8',
            fg='white',
            padx=12,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._save_qr_as_image
        )
        save_qr_btn.pack(side=tk.LEFT)
        
        # Info text
        tk.Label(
            parent,
            text="Quick access to common portal management tasks",
            font=('Segoe UI', 8, 'italic'),
            bg='white',
            fg='#888'
        ).pack(anchor='w', pady=(8, 0))
    
    def _restart_portal_server(self):
        """Restart the portal server"""
        self.start_student_portal()
        self._run_health_checks()
        messagebox.showinfo("Server", "Portal server has been restarted.")
    
    def _show_portal_stats(self):
        """Show portal statistics"""
        try:
            import urllib.request
            import json
            response = urllib.request.urlopen(f"http://127.0.0.1:{self.portal_port}/api/admin/stats", timeout=3)
            stats = json.loads(response.read().decode())
            
            # Parse request stats
            requests = stats.get('requests', {})
            deletions = stats.get('deletions', {})
            
            pending_requests = requests.get('pending', 0)
            approved_requests = requests.get('approved', 0)
            rejected_requests = requests.get('rejected', 0)
            
            pending_deletions = deletions.get('pending', 0)
            
            stats_msg = f"""📊 Portal Statistics

👥 Portal Users: {stats.get('portal_users', 0)}
🔐 Pending Password Change: {stats.get('pending_password_change', 0)}

📝 Requests:
   • Pending: {pending_requests}
   • Approved: {approved_requests}
   • Rejected: {rejected_requests}

🗑️ Deletion Requests:
   • Pending: {pending_deletions}
"""
            messagebox.showinfo("Portal Stats", stats_msg)
        except Exception as e:
            messagebox.showinfo("Portal Stats", f"Statistics not available.\n{str(e)}")
    
    def _open_portal_in_browser(self):
        """Open portal in default browser"""
        import webbrowser
        webbrowser.open(self.portal_url)
        
    def _initialize_portal_server(self):
        """Initialize portal server and update dashboard"""
        self.start_student_portal()
        self.portal_start_time = datetime.now()
        
        # Get current IP
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            local_ip = s.getsockname()[0]
            s.close()
        except:
            local_ip = "127.0.0.1"
        
        self.portal_url = f"http://{local_ip}:{self.portal_port}"
        
        # Add to history
        self.portal_url_history.append({
            'url': self.portal_url,
            'time': datetime.now().strftime('%H:%M:%S'),
            'ip': local_ip
        })
        
        # Keep only last 3
        if len(self.portal_url_history) > 3:
            self.portal_url_history = self.portal_url_history[-3:]
        
        # Update displays
        self._update_dashboard_displays(local_ip)
        
        # Generate QR
        self._generate_portal_qr_code()
        
        # Run health checks
        self._run_health_checks()
        
        # Start runtime updater
        self._start_runtime_updater()
        
    def _update_dashboard_displays(self, local_ip):
        """Update all dashboard displays"""
        if hasattr(self, 'portal_current_url_label'):
            self.portal_current_url_label.config(text=self.portal_url)
        if hasattr(self, 'current_url_display'):
            self.current_url_display.config(text=self.portal_url)
        if hasattr(self, 'server_ip_label'):
            self.server_ip_label.config(text=local_ip)
        if hasattr(self, 'server_traffic_label'):
            self.server_traffic_label.config(text="Low (< 10 active)")
        
    def _generate_portal_qr_code(self):
        """Generate and display QR code"""
        # Clear existing
        for w in self.portal_qr_container.winfo_children():
            w.destroy()
        
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            local_ip = s.getsockname()[0]
            s.close()
        except:
            local_ip = "127.0.0.1"
        
        url = f"http://{local_ip}:{self.portal_port}"
        
        try:
            qr = qrcode.QRCode(box_size=8, border=2)
            qr.add_data(url)
            qr.make(fit=True)
            img = qr.make_image(fill_color="#0F3460", back_color="white")
            
            img_tk = ImageTk.PhotoImage(img)
            qr_label = tk.Label(self.portal_qr_container, image=img_tk, bg='white', padx=10, pady=10)
            qr_label.image = img_tk
            qr_label.pack()
        except Exception as e:
            tk.Label(
                self.portal_qr_container,
                text="⚠️ QR Generation Failed",
                font=('Segoe UI', 10),
                bg='white',
                fg='#dc3545',
                padx=30,
                pady=30
            ).pack()
    
    def _copy_portal_url(self):
        """Copy portal URL to clipboard"""
        self.root.clipboard_clear()
        self.root.clipboard_append(self.portal_url)
        messagebox.showinfo("Copied", "Portal URL copied to clipboard!")
    
    def _save_qr_as_image(self):
        """Save QR code as PNG image file"""
        from tkinter import filedialog
        
        if not hasattr(self, 'portal_url') or not self.portal_url:
            messagebox.showerror("Error", "No portal URL available. Please start the server first.")
            return
        
        # Ask user for save location
        file_path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG files", "*.png"), ("All files", "*.*")],
            initialfile="student_portal_qr.png",
            title="Save QR Code"
        )
        
        if not file_path:
            return  # User cancelled
        
        try:
            # Generate QR code
            qr = qrcode.QRCode(box_size=10, border=2)
            qr.add_data(self.portal_url)
            qr.make(fit=True)
            img = qr.make_image(fill_color="#0F3460", back_color="white")
            
            # Save to file
            img.save(file_path)
            
            messagebox.showinfo("Success", f"QR code saved successfully!\n\nLocation: {file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save QR code: {str(e)}")
    
    def _run_health_checks(self):
        """Run health diagnostics"""
        # Skip if health indicators not yet created
        if not hasattr(self, 'health_indicators') or not self.health_indicators:
            return
        
        issues = []
        
        # Check 1: Portal Server
        try:
            import urllib.request
            urllib.request.urlopen(f"http://127.0.0.1:{self.portal_port}/api/me", timeout=2)
            if "Portal Server" in self.health_indicators:
                self.health_indicators["Portal Server"].config(text="● Healthy", fg='#28a745')
        except:
            if "Portal Server" in self.health_indicators:
                self.health_indicators["Portal Server"].config(text="● Unavailable", fg='#dc3545')
            issues.append("Portal server not responding")
        
        # Check 2: Database
        try:
            portal_db_path = os.path.join(os.path.dirname(__file__), 'Web-Extension', 'portal.db')
            if os.path.exists(portal_db_path):
                if "Database Connection" in self.health_indicators:
                    self.health_indicators["Database Connection"].config(text="● Connected", fg='#28a745')
            else:
                if "Database Connection" in self.health_indicators:
                    self.health_indicators["Database Connection"].config(text="● Not Found", fg='#ffc107')
                issues.append("Portal database not found")
        except:
            if "Database Connection" in self.health_indicators:
                self.health_indicators["Database Connection"].config(text="● Error", fg='#dc3545')
            issues.append("Database connection error")
        
        # Check 3: Network
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            s.close()
            if "Network Access" in self.health_indicators:
                self.health_indicators["Network Access"].config(text="● Available", fg='#28a745')
        except:
            if "Network Access" in self.health_indicators:
                self.health_indicators["Network Access"].config(text="● Limited", fg='#ffc107')
            issues.append("Limited network access")
        
        # Update issues label
        if hasattr(self, 'issues_label') and hasattr(self, 'server_status_badge'):
            if issues:
                self.issues_label.config(text=f"⚠️ {len(issues)} issue(s): " + "; ".join(issues), fg='#dc3545')
                self.server_status_badge.config(text="● ISSUES", bg='#fff3cd', fg='#856404')
            else:
                self.issues_label.config(text="✓ All systems operational", fg='#28a745')
                self.server_status_badge.config(text="● ONLINE", bg='#d4edda', fg='#155724')
    
    def _start_runtime_updater(self):
        """Start background runtime updater"""
        def update_runtime():
            if self.portal_start_time and hasattr(self, 'server_runtime_label'):
                delta = datetime.now() - self.portal_start_time
                hours, remainder = divmod(int(delta.total_seconds()), 3600)
                minutes, seconds = divmod(remainder, 60)
                self.server_runtime_label.config(text=f"{hours}h {minutes}m {seconds}s")
            self.root.after(1000, update_runtime)
        
        update_runtime()
    
    def _refresh_server_dashboard(self):
        """Refresh all server dashboard data"""
        self._run_health_checks()
        self._generate_portal_qr_code()
        
        # Re-detect IP
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            local_ip = s.getsockname()[0]
            s.close()
        except:
            local_ip = "127.0.0.1"
        
        new_url = f"http://{local_ip}:{self.portal_port}"
        
        if new_url != self.portal_url:
            self.portal_url = new_url
            self.portal_url_history.append({
                'url': new_url,
                'time': datetime.now().strftime('%H:%M:%S'),
                'ip': local_ip
            })
            if len(self.portal_url_history) > 3:
                self.portal_url_history = self.portal_url_history[-3:]
        
        self._update_dashboard_displays(local_ip)
        
        if hasattr(self, 'show_url_history') and self.show_url_history.get():
            self._populate_url_history()
            

    
    def _create_requests_section(self, parent):
        """Create requests management section with sub-tabs for different request types"""
        # Main container
        container = tk.Frame(parent, bg='white')
        container.pack(fill=tk.BOTH, expand=True)
        
        # Header frame with title and buttons
        header_frame = tk.Frame(container, bg='white')
        header_frame.pack(fill=tk.X, padx=20, pady=(20, 15))
        
        title_frame = tk.Frame(header_frame, bg='white')
        title_frame.pack(side=tk.LEFT)
        
        tk.Label(
            title_frame,
            text="📋 Student Portal Requests",
            font=('Segoe UI', 18, 'bold'),
            bg='white',
            fg=self.colors['accent']
        ).pack(side=tk.LEFT)
        
        # Count badge
        self.requests_count_badge = tk.Label(
            title_frame,
            text="0",
            font=('Segoe UI', 10, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            padx=10,
            pady=2
        )
        self.requests_count_badge.pack(side=tk.LEFT, padx=(10, 0))
        
        # Buttons frame on the right
        buttons_frame = tk.Frame(header_frame, bg='white')
        buttons_frame.pack(side=tk.RIGHT)
        
        # History toggle state
        self.show_request_history = tk.BooleanVar(value=False)
        
        self.history_toggle_btn = tk.Button(
            buttons_frame,
            text="📜 View History",
            font=('Segoe UI', 10, 'bold'),
            bg='#6c757d',
            fg='white',
            padx=12,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._toggle_request_history
        )
        self.history_toggle_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        refresh_btn = tk.Button(
            buttons_frame,
            text="🔄 Refresh",
            font=('Segoe UI', 10, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            padx=15,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._refresh_portal_requests
        )
        refresh_btn.pack(side=tk.LEFT)
        
        # ═══════════════════════════════════════════════════════════════
        # SEARCH AND FILTER BAR
        # ═══════════════════════════════════════════════════════════════
        search_filter_frame = tk.Frame(container, bg='#f8f9fa', relief='solid', bd=1)
        search_filter_frame.pack(fill=tk.X, padx=20, pady=(0, 10))
        
        search_inner = tk.Frame(search_filter_frame, bg='#f8f9fa')
        search_inner.pack(fill=tk.X, padx=15, pady=10)
        
        # Search entry
        tk.Label(search_inner, text="🔍 Search:", font=('Segoe UI', 9, 'bold'), bg='#f8f9fa').pack(side=tk.LEFT, padx=(0, 5))
        
        self.request_search_var = tk.StringVar()
        self.request_search_var.trace("w", lambda *args: self.root.after(400, self._apply_request_filters))
        search_entry = tk.Entry(search_inner, textvariable=self.request_search_var, width=25, font=('Segoe UI', 9))
        search_entry.pack(side=tk.LEFT, padx=5, ipady=3)
        
        # Date range filter
        tk.Label(search_inner, text="📅 Period:", font=('Segoe UI', 9), bg='#f8f9fa').pack(side=tk.LEFT, padx=(20, 5))
        
        self.request_days_var = tk.StringVar(value="All Time")
        days_combo = ttk.Combobox(search_inner, textvariable=self.request_days_var, state="readonly", width=12, font=('Segoe UI', 9))
        days_combo['values'] = ("Today", "Last 7 Days", "Last 30 Days", "All Time")
        days_combo.pack(side=tk.LEFT, padx=5)
        days_combo.bind("<<ComboboxSelected>>", lambda e: self._apply_request_filters())
        
        # Branch filter
        tk.Label(search_inner, text="🏫 Branch:", font=('Segoe UI', 9), bg='#f8f9fa').pack(side=tk.LEFT, padx=(20, 5))
        
        self.request_branch_var = tk.StringVar(value="All Branches")
        branch_combo = ttk.Combobox(search_inner, textvariable=self.request_branch_var, state="readonly", width=22, font=('Segoe UI', 9))
        branch_combo['values'] = ('All Branches', 'Computer Engineering', 'Information Technology', 'Civil Engineering',
                                   'Electronics & Telecommunication', 'Mechanical Engineering', 'Automobile Engineering', 'Electrical Engineering')
        branch_combo.pack(side=tk.LEFT, padx=5)
        branch_combo.bind("<<ComboboxSelected>>", lambda e: self._apply_request_filters())
        
        # Clear filters button
        clear_btn = tk.Button(
            search_inner,
            text="✕ Clear",
            font=('Segoe UI', 8),
            bg='#dc3545',
            fg='white',
            padx=8,
            pady=2,
            cursor='hand2',
            relief='flat',
            command=self._clear_request_filters
        )
        clear_btn.pack(side=tk.RIGHT)
        
        # ═══════════════════════════════════════════════════════════════
        # REQUEST TYPE SUB-TABS (Filter Tabs)
        # ═══════════════════════════════════════════════════════════════
        tabs_frame = tk.Frame(container, bg='white')
        tabs_frame.pack(fill=tk.X, padx=20, pady=(0, 10))
        
        # Current filter type
        self.request_type_filter = tk.StringVar(value="all")
        
        # Define request type tabs with colors and icons
        self.request_tab_configs = [
            ("all", "📋 All", "#1a1a2e", "#e2e8f0"),
            ("student_registration", "📝 Registration", "#0ea5e9", "#e0f2fe"),
            ("password_reset", "🔑 Password", "#8b5cf6", "#ede9fe"),
            ("profile_update", "👤 Profile", "#17a2b8", "#cffafe"),
            ("book_request", "📚 Books", "#6f42c1", "#f3e8ff"),
            ("renewal", "🔄 Renewal", "#28a745", "#dcfce7"),
            ("extension", "⏰ Extension", "#fd7e14", "#ffedd5"),
            ("service", "🛎️ Services", "#ec4899", "#fce7f3")
        ]
        
        self.request_tab_buttons = {}
        
        for type_key, label, active_color, bg_color in self.request_tab_configs:
            btn = tk.Button(
                tabs_frame,
                text=label,
                font=('Segoe UI', 9, 'bold'),
                bg=bg_color if type_key != "all" else active_color,
                fg='white' if type_key == "all" else '#333',
                padx=12,
                pady=6,
                cursor='hand2',
                relief='flat',
                command=lambda t=type_key: self._set_request_type_filter(t)
            )
            btn.pack(side=tk.LEFT, padx=(0, 6))
            self.request_tab_buttons[type_key] = {
                'button': btn,
                'active_color': active_color,
                'bg_color': bg_color
            }
        
        # ═══════════════════════════════════════════════════════════════
        # STATS SUMMARY BAR (Counts for each type)
        # ═══════════════════════════════════════════════════════════════
        stats_frame = tk.Frame(container, bg='#f8f9fa', relief='solid', bd=1)
        stats_frame.pack(fill=tk.X, padx=20, pady=(0, 10))
        
        stats_inner = tk.Frame(stats_frame, bg='#f8f9fa')
        stats_inner.pack(padx=15, pady=8)
        
        self.request_stats_labels = {}
        stat_types = [
            ("Reg", "#0ea5e9", "student_registration"),
            ("Password", "#8b5cf6", "password_reset"),
            ("Profile", "#17a2b8", "profile_update"),
            ("Books", "#6f42c1", "book_request"),
            ("Renew", "#28a745", "renewal"),
            ("Extend", "#fd7e14", "extension"),
            ("Services", "#ec4899", "service"),
            ("Rejected", "#dc3545", "rejected")
        ]
        
        for label, color, key in stat_types:
            stat_item = tk.Frame(stats_inner, bg='#f8f9fa')
            stat_item.pack(side=tk.LEFT, padx=10)
            
            count_label = tk.Label(
                stat_item,
                text="0",
                font=('Segoe UI', 12, 'bold'),
                bg='#f8f9fa',
                fg=color
            )
            count_label.pack()
            self.request_stats_labels[key] = count_label
            
            tk.Label(
                stat_item,
                text=label,
                font=('Segoe UI', 7),
                bg='#f8f9fa',
                fg='#666'
            ).pack()
        
        # ═══════════════════════════════════════════════════════════════
        # HISTORY FILTER FRAME (Hidden by default, shown in history mode)
        # ═══════════════════════════════════════════════════════════════
        self.request_filters_frame = tk.Frame(container, bg='white')
        # (Not packed by default - shown when history mode is active)
        
        # ═══════════════════════════════════════════════════════════════
        # SCROLLABLE REQUEST LIST
        # ═══════════════════════════════════════════════════════════════
        self.requests_list_frame = tk.Frame(container, bg='white', relief='solid', bd=1)
        self.requests_list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        list_frame = self.requests_list_frame
        
        # Canvas for scrolling
        self.requests_canvas = tk.Canvas(list_frame, bg='white', highlightthickness=0)
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.requests_canvas.yview)
        self.requests_container = tk.Frame(self.requests_canvas, bg='white')
        
        self.requests_container.bind(
            "<Configure>",
            lambda e: self.requests_canvas.configure(scrollregion=self.requests_canvas.bbox("all"))
        )
        
        # Create window and bind width updates
        self.requests_canvas_window = self.requests_canvas.create_window((0, 0), window=self.requests_container, anchor="nw")
        
        def on_canvas_configure(event):
            self.requests_canvas.itemconfig(self.requests_canvas_window, width=event.width)
        self.requests_canvas.bind('<Configure>', on_canvas_configure)
        
        self.requests_canvas.configure(yscrollcommand=scrollbar.set)
        
        self.requests_canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Mousewheel scrolling - only scroll when content overflows
        def _on_mousewheel(event):
            # Only scroll if content is larger than canvas
            canvas_height = self.requests_canvas.winfo_height()
            content_height = self.requests_container.winfo_reqheight()
            if content_height > canvas_height:
                self.requests_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
                return "break"  # Prevent event propagation
            return None
        
        # Bind to canvas and container for local scrolling
        self.requests_canvas.bind("<MouseWheel>", _on_mousewheel)
        self.requests_container.bind("<MouseWheel>", _on_mousewheel)
        
        # Recursive function to bind mousewheel to all children
        def bind_mousewheel_recursive(widget):
            widget.bind("<MouseWheel>", _on_mousewheel)
            for child in widget.winfo_children():
                bind_mousewheel_recursive(child)
        
        # Store the bind function for later use when adding cards
        self._bind_requests_mousewheel = bind_mousewheel_recursive
        
        # Store all requests for filtering
        self.all_pending_requests = []
        
        # Delay initial load to allow Waitress server to start (race condition fix)
        self.root.after(3000, self._refresh_portal_requests)
        
        # Auto-refresh every 30 seconds so new web-portal requests appear
        self.root.after(30000, self._start_requests_auto_refresh)
    
    def _set_request_type_filter(self, type_key):
        """Set the active request type filter and update tab styling"""
        self.request_type_filter.set(type_key)
        
        # Update tab button styles
        for key, config in self.request_tab_buttons.items():
            btn = config['button']
            if key == type_key:
                # Active state
                btn.config(bg=config['active_color'], fg='white')
            else:
                # Inactive state
                btn.config(bg=config['bg_color'], fg='#333')
        
        # Apply filter
        self._apply_request_filters()
    
    def _clear_request_filters(self):
        """Clear all filters and reset to default"""
        self.request_search_var.set("")
        self.request_days_var.set("All Time")
        self.request_type_filter.set("all")
        if hasattr(self, 'request_branch_var'):
            self.request_branch_var.set("All Branches")
        
        # Reset tab styles
        for key, config in self.request_tab_buttons.items():
            btn = config['button']
            if key == "all":
                btn.config(bg=config['active_color'], fg='white')
            else:
                btn.config(bg=config['bg_color'], fg='#333')
        
        self._apply_request_filters()
    
    def _apply_request_filters(self):
        """Apply search, date, and type filters to requests"""
        if self.show_request_history.get():
            self._refresh_request_history()
            return
            
        # Filter the cached requests
        filtered = self.all_pending_requests.copy()
        
        # Type filter
        type_filter = self.request_type_filter.get()
        if type_filter != "all":
            if type_filter == "service":
                # Service filter matches multiple service request types
                service_types = ['Extended Hours Permission', 'Library Opening Request', 
                                'Study Materials Request', 'General Request', 'service_request',
                                'extended_hours', 'library_opening', 'study_materials', 'general']
                filtered = [r for r in filtered if r.get('request_type') in service_types or 
                           'request' in r.get('request_type', '').lower() and 
                           r.get('request_type') not in ['book_request', 'password_reset', 'profile_update', 
                                                          'renewal', 'extension', 'student_registration']]
            else:
                filtered = [r for r in filtered if r.get('request_type') == type_filter]
        
        # Search filter
        search_term = self.request_search_var.get().strip().lower()
        if search_term:
            filtered = [r for r in filtered if 
                search_term in r.get('student_name', '').lower() or
                search_term in r.get('enrollment_no', '').lower() or
                search_term in r.get('request_type', '').lower() or
                search_term in str(r.get('details', '')).lower()
            ]
        
        # Branch filter
        if hasattr(self, 'request_branch_var'):
            branch_filter = self.request_branch_var.get()
            if branch_filter and branch_filter != 'All Branches':
                filtered = [r for r in filtered if r.get('department', '') == branch_filter]
        
        # Date filter
        days_filter = self.request_days_var.get()
        if days_filter != "All Time":
            from datetime import datetime, timedelta
            try:
                if days_filter == "Today":
                    cutoff = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
                elif days_filter == "Last 7 Days":
                    cutoff = datetime.now() - timedelta(days=7)
                elif days_filter == "Last 30 Days":
                    cutoff = datetime.now() - timedelta(days=30)
                else:
                    cutoff = None
                
                if cutoff:
                    def parse_date(date_str):
                        try:
                            for fmt in ['%Y-%m-%d %H:%M:%S', '%Y-%m-%d', '%d/%m/%Y']:
                                try:
                                    return datetime.strptime(str(date_str)[:19], fmt)
                                except:
                                    continue
                            return datetime.now()
                        except:
                            return datetime.now()
                    
                    filtered = [r for r in filtered if parse_date(r.get('created_at', '')) >= cutoff]
            except Exception as e:
                print(f"Date filter error: {e}")
        
        # Update display
        self._display_filtered_requests(filtered)
    
    def _display_filtered_requests(self, requests_list):
        """Display filtered requests in the container"""
        # Clear existing
        for w in self.requests_container.winfo_children():
            w.destroy()
        
        # Update count badge
        if hasattr(self, 'requests_count_badge'):
            self.requests_count_badge.config(text=str(len(requests_list)))
        
        if not requests_list:
            # Check if we have any requests at all
            if not self.all_pending_requests:
                self._show_empty_message(self.requests_container, "No pending requests", 
                    "All student requests have been processed! 🎉")
            else:
                self._show_empty_message(self.requests_container, "No matching requests", 
                    "No requests match your current filters. Try adjusting your search or filters.")
            return
        
        # Create two-column grid layout
        self.requests_container.columnconfigure(0, weight=1)
        self.requests_container.columnconfigure(1, weight=1)
        
        # Place cards in grid
        for idx, req_data in enumerate(requests_list):
            row = idx // 2
            col = idx % 2
            self._create_request_card(self.requests_container, req_data, row, col, idx + 1)
        
        # Bind mousewheel to all new cards
        if hasattr(self, '_bind_requests_mousewheel'):
            self._bind_requests_mousewheel(self.requests_container)
    
    def _refresh_portal_requests(self):
        """Fetch and display pending requests in two-column layout"""
        # Clear existing
        for w in self.requests_container.winfo_children():
            w.destroy()
            
        if not WEB_PORTAL_AVAILABLE:
            self._show_empty_message(self.requests_container, "Portal Unavailable", "Web portal server is not running.")
            return
        
        try:
            import urllib.request
            import urllib.error
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/all-requests"
            req = urllib.request.Request(url)
            
            with urllib.request.urlopen(req, timeout=10) as response:
                data = json.loads(response.read().decode())
                requests_list = data.get('requests', [])
                deletion_requests = data.get('deletion_requests', [])
                
                # Debug: Log what we received
                print(f"[Requests Tab] Loaded {len(requests_list)} pending requests, {len(deletion_requests)} deletion requests")
                
                # Store all requests for filtering
                self.all_pending_requests = requests_list.copy()
                
                # Update stats by type (always show full stats)
                if hasattr(self, 'request_stats_labels'):
                    type_counts = {}
                    service_count = 0
                    known_types = ['student_registration', 'password_reset', 'profile_update', 
                                   'book_request', 'renewal', 'extension']
                    
                    for req_data in requests_list:
                        req_type = req_data.get('request_type', 'other')
                        if req_type in known_types:
                            type_counts[req_type] = type_counts.get(req_type, 0) + 1
                        else:
                            # Count as service request
                            service_count += 1
                    
                    type_counts['service'] = service_count
                    
                    # Get rejected count from API response
                    rejected_count = data.get('rejected_count', 0)
                    type_counts['rejected'] = rejected_count
                    
                    for key, label in self.request_stats_labels.items():
                        label.config(text=str(type_counts.get(key, 0)))
                
                # Apply current filters
                self._apply_request_filters()
                    
        except Exception as e:
            self._show_empty_message(self.requests_container, "Could not load requests", f"Portal server may not be running.\n{str(e)}")

    def _start_requests_auto_refresh(self):
        """Periodically poll the portal API so new requests appear automatically.
        Bug 10 Fix: HTTP call runs in a daemon thread; Tkinter UI never blocked."""
        if not WEB_PORTAL_AVAILABLE:
            return
        # Only refresh if we're NOT in history mode
        if hasattr(self, 'show_request_history') and not self.show_request_history.get():
            import threading as _threading
            def _bg_refresh():
                try:
                    self._refresh_portal_requests()
                except Exception:
                    pass
            _threading.Thread(target=_bg_refresh, daemon=True).start()
        # Schedule next poll in 30 seconds
        self.root.after(30000, self._start_requests_auto_refresh)

    def _toggle_request_history(self):
        """Toggle between pending requests and history view"""
        current = self.show_request_history.get()
        self.show_request_history.set(not current)
        
        if self.show_request_history.get():
            # Switch to history mode
            self.history_toggle_btn.config(text="📋 View Pending", bg='#28a745')
            if hasattr(self, 'request_filters_frame') and hasattr(self, 'requests_list_frame'):
                self.request_filters_frame.pack(fill=tk.X, padx=20, pady=(0, 10), before=self.requests_list_frame)
            self._refresh_request_history()
        else:
            # Switch to pending mode
            self.history_toggle_btn.config(text="📜 View History", bg='#6c757d')
            if hasattr(self, 'request_filters_frame'):
                self.request_filters_frame.pack_forget()
            self._refresh_portal_requests()
    
    def _refresh_request_history(self):
        """Fetch and display processed request history with filters"""
        # Clear existing
        for w in self.requests_container.winfo_children():
            w.destroy()
            
        if not WEB_PORTAL_AVAILABLE:
            self._show_empty_message(self.requests_container, "Portal Unavailable", "Web portal server is not running.")
            return
        
        try:
            import urllib.request
            import urllib.error
            import urllib.parse
            
            # Prepare params
            days_map = {
                "Last 7 Days": "7",
                "Last 30 Days": "30",
                "Last 90 Days": "90",
                "All Time": ""
            }
            days = days_map.get(self.request_days_var.get(), "") if hasattr(self, 'request_days_var') else ""
            q = urllib.parse.quote(self.request_search_var.get().strip()) if hasattr(self, 'request_search_var') else ""
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/request-history?q={q}&days={days}"
            req = urllib.request.Request(url)
            
            with urllib.request.urlopen(req, timeout=5) as response:
                data = json.loads(response.read().decode())
                history_list = data.get('history', [])
                counts = data.get('counts', {})
                
                # Update count badge to show total history
                if hasattr(self, 'requests_count_badge'):
                    self.requests_count_badge.config(text=str(counts.get('total', 0)), bg='#6c757d')
                
                # Update stats labels
                if hasattr(self, 'request_stats_labels'):
                    # Reset all to 0, then show approved/rejected in appropriate slots
                    for key in self.request_stats_labels:
                        self.request_stats_labels[key].config(text="0")
                    if 'rejected' in self.request_stats_labels:
                        self.request_stats_labels['rejected'].config(text=str(counts.get('rejected', 0)))
                
                if not history_list:
                    self._show_empty_message(self.requests_container, "No request history", "No processed requests found.")
                    return
                
                # Create two-column grid layout
                self.requests_container.columnconfigure(0, weight=1)
                self.requests_container.columnconfigure(1, weight=1)
                
                # Place history cards
                for idx, req_data in enumerate(history_list):
                    row = idx // 2
                    col = idx % 2
                    self._create_history_card(self.requests_container, req_data, row, col, idx + 1)
                    
        except Exception as e:
            self._show_empty_message(self.requests_container, "Could not load history", str(e))
    
    def _create_history_card(self, parent, req_data, row=0, col=0, index=1):
        """Create a card for displaying processed request in history view"""
        status = req_data.get('status', 'unknown')
        
        # Color scheme based on status
        if status == 'approved':
            bg_color = '#d4edda'  # Light green
            border_color = '#28a745'
            status_text = '✓ Approved'
            status_bg = '#28a745'
        else:  # rejected
            bg_color = '#f8d7da'  # Light red
            border_color = '#dc3545'
            status_text = '✕ Rejected'
            status_bg = '#dc3545'
        
        card = tk.Frame(parent, bg=bg_color, relief='solid', bd=1)
        card.grid(row=row, column=col, sticky='nsew', padx=8, pady=8)
        
        # Bind mousewheel
        def _on_card_mousewheel(event):
            self.requests_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        card.bind("<MouseWheel>", _on_card_mousewheel)
        
        inner = tk.Frame(card, bg=bg_color)
        inner.pack(fill=tk.BOTH, expand=True, padx=18, pady=15)
        inner.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Header row with index and status badge
        header_row = tk.Frame(inner, bg=bg_color)
        header_row.pack(fill=tk.X)
        
        # Index badge
        tk.Label(
            header_row,
            text=f"#{index}",
            font=('Segoe UI', 11, 'bold'),
            bg='#6c757d',
            fg='white',
            padx=8,
            pady=2
        ).pack(side=tk.LEFT)
        
        # Status badge
        tk.Label(
            header_row,
            text=status_text,
            font=('Segoe UI', 9, 'bold'),
            bg=status_bg,
            fg='white',
            padx=10,
            pady=3
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Type badge
        type_colors = {
            'profile_update': '#17a2b8',
            'renewal': '#28a745',
            'book_reservation': '#6f42c1',
            'book_request': '#6f42c1',
            'extension': '#fd7e14'
        }
        req_type = req_data.get('request_type', 'request')
        type_color = type_colors.get(req_type, '#6c757d')
        
        tk.Label(
            header_row,
            text=req_type.replace('_', ' ').title(),
            font=('Segoe UI', 8, 'bold'),
            bg=type_color,
            fg='white',
            padx=8,
            pady=2
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Student info
        student_row = tk.Frame(inner, bg=bg_color)
        student_row.pack(fill=tk.X, pady=(8, 0))
        
        tk.Label(
            student_row,
            text=f"{req_data.get('student_name', 'Unknown')}",
            font=('Segoe UI', 11, 'bold'),
            bg=bg_color,
            fg='#333'
        ).pack(side=tk.LEFT)
        
        tk.Label(
            student_row,
            text=f"({req_data.get('enrollment_no', '')})",
            font=('Segoe UI', 9),
            bg=bg_color,
            fg='#666'
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Details
        details = req_data.get('details', {})
        req_type = req_data.get('request_type', '')
        if isinstance(details, dict):
            if req_type == 'renewal':
                # Show clean format for renewal: "Book Title (Copy: 708)"
                title_val = details.get('title', '')
                accession = details.get('accession_no', details.get('book_id', ''))
                detail_text = f"Book: {title_val}" if title_val else ''
                if accession and accession != details.get('book_id', ''):
                    detail_text += f" (Copy: {accession})"
                elif accession:
                    detail_text += f" (ID: {accession})"
            else:
                detail_text = ', '.join([f"{k}: {v}" for k, v in details.items() if v])
        else:
            detail_text = str(details)
        
        if detail_text:
            tk.Label(
                inner,
                text=detail_text[:100] + ('...' if len(detail_text) > 100 else ''),
                font=('Segoe UI', 9),
                bg=bg_color,
                fg='#555',
                wraplength=350,
                justify='left'
            ).pack(anchor='w', pady=(6, 0))
        
        # Timestamp
        tk.Label(
            inner,
            text=f"📅 {req_data.get('created_at', 'N/A')}",
            font=('Segoe UI', 8),
            bg=bg_color,
            fg='#888'
        ).pack(anchor='w', pady=(6, 0))
    
    def _create_request_card(self, parent, req_data, row=0, col=0, index=1):
        """Create a card for displaying a request in grid layout with index"""
        card = tk.Frame(parent, bg='#f8f9fa', relief='solid', bd=1)
        card.grid(row=row, column=col, sticky='nsew', padx=8, pady=8)
        
        # Bind mousewheel to card for scrolling
        def _on_card_mousewheel(event):
            self.requests_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        card.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Inner padding - larger for better visibility
        inner = tk.Frame(card, bg='#f8f9fa')
        inner.pack(fill=tk.BOTH, expand=True, padx=18, pady=15)
        inner.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Header row with index and type badge
        header_row = tk.Frame(inner, bg='#f8f9fa')
        header_row.pack(fill=tk.X)
        header_row.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Index badge - prominent number
        index_badge = tk.Label(
            header_row,
            text=f"#{index}",
            font=('Segoe UI', 11, 'bold'),
            bg=self.colors['accent'],
            fg='white',
            padx=8,
            pady=2
        )
        index_badge.pack(side=tk.LEFT)
        index_badge.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Type badge
        type_colors = {
            'student_registration': '#0ea5e9',
            'password_reset': '#8b5cf6',
            'profile_update': '#17a2b8',
            'renewal': '#28a745',
            'book_reservation': '#6f42c1',
            'book_request': '#6f42c1',
            'extension': '#fd7e14',
            'Extended Hours Permission': '#ec4899',
            'Library Opening Request': '#ec4899',
            'Study Materials Request': '#ec4899',
            'General Request': '#ec4899'
        }
        req_type = req_data.get('request_type', 'request')
        type_color = type_colors.get(req_type, '#ec4899')  # Default to pink for service requests
        
        type_badge = tk.Label(
            header_row,
            text=req_type.replace('_', ' ').title(),
            font=('Segoe UI', 9, 'bold'),
            bg=type_color,
            fg='white',
            padx=10,
            pady=3
        )
        type_badge.pack(side=tk.LEFT, padx=(8, 0))
        type_badge.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Student info row
        student_row = tk.Frame(inner, bg='#f8f9fa')
        student_row.pack(fill=tk.X, pady=(8, 0))
        
        tk.Label(
            student_row,
            text=f"{req_data.get('student_name', 'Unknown')}",
            font=('Segoe UI', 11, 'bold'),
            bg='#f8f9fa',
            fg='#333'
        ).pack(side=tk.LEFT)
        
        tk.Label(
            student_row,
            text=f"({req_data.get('enrollment_no', '')})",
            font=('Segoe UI', 9),
            bg='#f8f9fa',
            fg='#666'
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Department badge
        dept = req_data.get('department', '')
        if dept:
            tk.Label(
                student_row,
                text=dept,
                font=('Segoe UI', 8),
                bg='#e2e8f0',
                fg='#475569',
                padx=6,
                pady=1
            ).pack(side=tk.RIGHT)
        
        # Details
        details = req_data.get('details', {})
        req_type = req_data.get('request_type', '')
        if isinstance(details, dict):
            if req_type == 'renewal':
                title_val = details.get('title', '')
                accession = details.get('accession_no', details.get('book_id', ''))
                detail_text = f"Book: {title_val}" if title_val else ''
                if accession and accession != details.get('book_id', ''):
                    detail_text += f" (Copy: {accession})"
                elif accession:
                    detail_text += f" (ID: {accession})"
            else:
                detail_text = ', '.join([f"{k}: {v}" for k, v in details.items() if v])
        else:
            detail_text = str(details)
        
        if detail_text:
            tk.Label(
                inner,
                text=detail_text[:120] + ('...' if len(detail_text) > 120 else ''),
                font=('Segoe UI', 9),
                bg='#f8f9fa',
                fg='#555',
                anchor='w',
                justify='left',
                wraplength=350
            ).pack(fill=tk.X, pady=(6, 0), anchor='w')
        
        # Timestamp
        tk.Label(
            inner,
            text=f"📅 {req_data.get('created_at', 'N/A')}",
            font=('Segoe UI', 8),
            bg='#f8f9fa',
            fg='#888'
        ).pack(anchor='w', pady=(6, 0))
        
        # Actions row at bottom
        actions_frame = tk.Frame(inner, bg='#f8f9fa')
        actions_frame.pack(fill=tk.X, pady=(12, 0))
        
        req_id = req_data.get('req_id')
        
        approve_btn = tk.Button(
            actions_frame,
            text="✓ Approve",
            font=('Segoe UI', 9, 'bold'),
            bg='#28a745',
            fg='white',
            padx=15,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=lambda rid=req_id: self._handle_request_action(rid, 'approve')
        )
        approve_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        reject_btn = tk.Button(
            actions_frame,
            text="✕ Reject",
            font=('Segoe UI', 9, 'bold'),
            bg='#dc3545',
            fg='white',
            padx=15,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=lambda rid=req_id: self._handle_request_action(rid, 'reject')
        )
        reject_btn.pack(side=tk.LEFT)
    
    def _handle_request_action(self, req_id, action):
        """Handle approve/reject action for a request"""
        try:
            import urllib.request
            import urllib.error
            import json
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/requests/{req_id}/{action}"

            # Bug 3 Fix: Use POST for state-changing operations (GET is semantically
            # wrong and can be cached/replayed by proxies, causing double-actions).
            # The portal's csrf_protect() skips /api/admin/ routes, so POST works fine.
            req = urllib.request.Request(url, data=b'', method='POST')
            req.add_header('Content-Type', 'application/json')

            with urllib.request.urlopen(req, timeout=5) as response:
                result = json.loads(response.read().decode())
                if result.get('status') == 'success':
                    messagebox.showinfo("Success", f"Request {action}d successfully!")
                    self._refresh_portal_requests()
                    # Refresh students list immediately if approved
                    if action == 'approve':
                        self.refresh_students()
                else:
                    messagebox.showerror("Error", result.get('message', 'Action failed'))
        except urllib.error.HTTPError as e:
            # Detailed error logging for 405/500 errors
            err_msg = f"HTTP {e.code}: {e.reason}"
            try:
                err_body = e.read().decode()
                print(f"[API Error] Body: {err_body}")
            except:
                pass
            messagebox.showerror("API Error", f"Failed to {action} request.\n\n{err_msg}")
        except Exception as e:
            print(f"[Request Action Error] {e}")
            messagebox.showerror("Error", f"Failed to {action} request: {str(e)}")
    
    def _create_deletion_section(self, parent):
        """Create deletion requests management section with enhanced UI"""
        # Main container
        container = tk.Frame(parent, bg='white')
        container.pack(fill=tk.BOTH, expand=True)
        
        # Header frame with count badge
        header_frame = tk.Frame(container, bg='white')
        header_frame.pack(fill=tk.X, padx=20, pady=(20, 15))
        
        title_frame = tk.Frame(header_frame, bg='white')
        title_frame.pack(side=tk.LEFT)
        
        tk.Label(
            title_frame,
            text="🗑️ Account Deletion Requests",
            font=('Segoe UI', 18, 'bold'),
            bg='white',
            fg=self.colors['accent']
        ).pack(side=tk.LEFT)
        
        # Count badge
        self.deletion_count_badge = tk.Label(
            title_frame,
            text="0",
            font=('Segoe UI', 10, 'bold'),
            bg='#dc3545',
            fg='white',
            padx=10,
            pady=2
        )
        self.deletion_count_badge.pack(side=tk.LEFT, padx=(10, 0))
        
        # Buttons frame on the right
        buttons_frame = tk.Frame(header_frame, bg='white')
        buttons_frame.pack(side=tk.RIGHT)
        
        # History toggle state
        self.show_deletion_history = tk.BooleanVar(value=False)
        
        self.deletion_history_toggle_btn = tk.Button(
            buttons_frame,
            text="📜 View History",
            font=('Segoe UI', 10, 'bold'),
            bg='#6c757d',
            fg='white',
            padx=12,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._toggle_deletion_history
        )
        self.deletion_history_toggle_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        refresh_btn = tk.Button(
            buttons_frame,
            text="🔄 Refresh",
            font=('Segoe UI', 10, 'bold'),
            bg=self.colors['secondary'],
            fg='white',
            padx=15,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=self._refresh_deletion_requests
        )
        refresh_btn.pack(side=tk.LEFT)
        
        # Stats summary bar
        stats_frame = tk.Frame(container, bg='#f8f9fa', relief='solid', bd=1)
        stats_frame.pack(fill=tk.X, padx=20, pady=(0, 15))
        
        stats_inner = tk.Frame(stats_frame, bg='#f8f9fa')
        stats_inner.pack(padx=15, pady=10)
        
        self.deletion_stats_labels = {}
        stat_types = [
            ("Pending", "#fd7e14", "pending"),
            ("Approved", "#28a745", "approved"),
            ("Rejected", "#dc3545", "rejected")
        ]
        
        for label, color, key in stat_types:
            stat_item = tk.Frame(stats_inner, bg='#f8f9fa')
            stat_item.pack(side=tk.LEFT, padx=20)
            
            count_label = tk.Label(
                stat_item,
                text="0",
                font=('Segoe UI', 16, 'bold'),
                bg='#f8f9fa',
                fg=color
            )
            count_label.pack()
            self.deletion_stats_labels[key] = count_label
            
            tk.Label(
                stat_item,
                text=label,
                font=('Segoe UI', 9),
                bg='#f8f9fa',
                fg='#666'
            ).pack()
        
        # Warning banner
        warning_frame = tk.Frame(container, bg='#fff3cd', relief='solid', bd=1)
        warning_frame.pack(fill=tk.X, padx=20, pady=(0, 15))
        
        tk.Label(
            warning_frame,
            text="⚠️ Approving deletion will remove the student from the portal authentication system.",
            font=('Segoe UI', 9),
            bg='#fff3cd',
            fg='#856404',
            padx=15,
            pady=8
        ).pack(anchor='w')
        
        # History Filter Frame (Hidden by default)
        self.deletion_filters_frame = tk.Frame(container, bg='white')
        
        filter_inner = tk.Frame(self.deletion_filters_frame, bg='#f8f9fa', relief='solid', bd=1)
        filter_inner.pack(fill=tk.X, padx=1)
        
        tk.Label(filter_inner, text="🔍 Search:", font=('Segoe UI', 9), bg='#f8f9fa').pack(side=tk.LEFT, padx=(10, 5), pady=8)
        
        self.deletion_search_var = tk.StringVar()
        self.deletion_search_var.trace("w", lambda *args: self.root.after(500, self._refresh_deletion_history))
        tk.Entry(filter_inner, textvariable=self.deletion_search_var, width=25, font=('Segoe UI', 9)).pack(side=tk.LEFT, padx=5)
        
        tk.Label(filter_inner, text="📅 Period:", font=('Segoe UI', 9), bg='#f8f9fa').pack(side=tk.LEFT, padx=(15, 5))
        
        self.deletion_days_var = tk.StringVar(value="All Time")
        days_combo = ttk.Combobox(filter_inner, textvariable=self.deletion_days_var, state="readonly", width=12, font=('Segoe UI', 9))
        days_combo['values'] = ("Last 7 Days", "Last 30 Days", "Last 90 Days", "All Time")
        days_combo.pack(side=tk.LEFT, padx=5)
        days_combo.bind("<<ComboboxSelected>>", lambda e: self._refresh_deletion_history())
        
        # Scrollable list with proper width binding
        self.deletion_list_frame = tk.Frame(container, bg='white', relief='solid', bd=1)
        self.deletion_list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        list_frame = self.deletion_list_frame
        
        self.deletion_canvas = tk.Canvas(list_frame, bg='white', highlightthickness=0)
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.deletion_canvas.yview)
        self.deletion_container = tk.Frame(self.deletion_canvas, bg='white')
        
        self.deletion_container.bind(
            "<Configure>",
            lambda e: self.deletion_canvas.configure(scrollregion=self.deletion_canvas.bbox("all"))
        )
        
        # Create window and bind width updates
        self.deletion_canvas_window = self.deletion_canvas.create_window((0, 0), window=self.deletion_container, anchor="nw")
        
        def on_canvas_configure(event):
            self.deletion_canvas.itemconfig(self.deletion_canvas_window, width=event.width)
        self.deletion_canvas.bind('<Configure>', on_canvas_configure)
        
        self.deletion_canvas.configure(yscrollcommand=scrollbar.set)
        
        self.deletion_canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Bind mousewheel - only scroll when content overflows
        def _on_mousewheel(event):
            # Only scroll if content is larger than canvas
            canvas_height = self.deletion_canvas.winfo_height()
            content_height = self.deletion_container.winfo_reqheight()
            if content_height > canvas_height:
                self.deletion_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
                return "break"  # Prevent event propagation
            return None
        
        # Bind to canvas and container for local scrolling
        self.deletion_canvas.bind("<MouseWheel>", _on_mousewheel)
        self.deletion_container.bind("<MouseWheel>", _on_mousewheel)
        
        # Recursive function to bind mousewheel to all children
        def bind_mousewheel_recursive(widget):
            widget.bind("<MouseWheel>", _on_mousewheel)
            for child in widget.winfo_children():
                bind_mousewheel_recursive(child)
        
        # Store the bind function for later use when adding cards
        self._bind_deletion_mousewheel = bind_mousewheel_recursive
        
        # Delay initial load to allow Waitress server to start
        self.root.after(3500, self._refresh_deletion_requests)
    
    def _refresh_deletion_requests(self):
        """Fetch and display deletion requests in two-column layout"""
        for w in self.deletion_container.winfo_children():
            w.destroy()
            
        if not WEB_PORTAL_AVAILABLE:
            self._show_empty_message(self.deletion_container, "Portal Unavailable", "Web portal server is not running.")
            return
        
        try:
            import urllib.request
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/all-requests"
            req = urllib.request.Request(url)
            
            with urllib.request.urlopen(req, timeout=5) as response:
                data = json.loads(response.read().decode())
                deletions = data.get('deletion_requests', [])
                
                # Update count badge
                if hasattr(self, 'deletion_count_badge'):
                    self.deletion_count_badge.config(text=str(len(deletions)))
                
                # Get deletion stats from separate query
                if hasattr(self, 'deletion_stats_labels'):
                    deletion_counts = data.get('deletion_counts', {})
                    self.deletion_stats_labels['pending'].config(text=str(len(deletions)))
                    self.deletion_stats_labels['approved'].config(text=str(deletion_counts.get('approved', 0)))
                    self.deletion_stats_labels['rejected'].config(text=str(deletion_counts.get('rejected', 0)))
                
                if not deletions:
                    self._show_empty_message(self.deletion_container, "No deletion requests", "No students have requested account deletion. 🎉")
                    return
                
                # Create two-column grid layout
                self.deletion_container.columnconfigure(0, weight=1)
                self.deletion_container.columnconfigure(1, weight=1)
                
                # Place cards in grid
                for idx, del_data in enumerate(deletions):
                    row = idx // 2
                    col = idx % 2
                    self._create_deletion_card(self.deletion_container, del_data, row, col, idx + 1)
                
                # Bind mousewheel to all new cards
                if hasattr(self, '_bind_deletion_mousewheel'):
                    self._bind_deletion_mousewheel(self.deletion_container)
                    
        except Exception as e:
            self._show_empty_message(self.deletion_container, "Could not load requests", str(e))
    def _toggle_deletion_history(self):
        """Toggle between pending deletions and history view"""
        current = self.show_deletion_history.get()
        self.show_deletion_history.set(not current)
        
        if self.show_deletion_history.get():
            # Switch to history mode
            self.deletion_history_toggle_btn.config(text="📋 View Pending", bg='#28a745')
            if hasattr(self, 'deletion_filters_frame') and hasattr(self, 'deletion_list_frame'):
                self.deletion_filters_frame.pack(fill=tk.X, padx=20, pady=(0, 10), before=self.deletion_list_frame)
            self._refresh_deletion_history()
        else:
            # Switch to pending mode
            self.deletion_history_toggle_btn.config(text="📜 View History", bg='#6c757d')
            if hasattr(self, 'deletion_filters_frame'):
                self.deletion_filters_frame.pack_forget()
            self._refresh_deletion_requests()
    
    def _refresh_deletion_history(self):
        """Fetch and display processed deletion history with filters"""
        # Clear existing
        for w in self.deletion_container.winfo_children():
            w.destroy()
            
        if not WEB_PORTAL_AVAILABLE:
            self._show_empty_message(self.deletion_container, "Portal Unavailable", "Web portal server is not running.")
            return
        
        try:
            import urllib.request
            import urllib.error
            import urllib.parse
            
            # Prepare params
            days_map = {
                "Last 7 Days": "7",
                "Last 30 Days": "30",
                "Last 90 Days": "90",
                "All Time": ""
            }
            days = days_map.get(self.deletion_days_var.get(), "") if hasattr(self, 'deletion_days_var') else ""
            q = urllib.parse.quote(self.deletion_search_var.get().strip()) if hasattr(self, 'deletion_search_var') else ""
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/deletion-history?q={q}&days={days}"
            req = urllib.request.Request(url)
            
            with urllib.request.urlopen(req, timeout=5) as response:
                data = json.loads(response.read().decode())
                history_list = data.get('history', [])
                counts = data.get('counts', {})
                
                # Update count badge to show total history
                if hasattr(self, 'deletion_count_badge'):
                    self.deletion_count_badge.config(text=str(counts.get('total', 0)), bg='#6c757d')
                
                # Update stats labels
                if hasattr(self, 'deletion_stats_labels'):
                    self.deletion_stats_labels['pending'].config(text="0")
                    self.deletion_stats_labels['approved'].config(text=str(counts.get('approved', 0)))
                    self.deletion_stats_labels['rejected'].config(text=str(counts.get('rejected', 0)))
                
                if not history_list:
                    self._show_empty_message(self.deletion_container, "No deletion history", "No processed deletion requests found.")
                    return
                
                # Create two-column grid layout
                self.deletion_container.columnconfigure(0, weight=1)
                self.deletion_container.columnconfigure(1, weight=1)
                
                # Place history cards
                for idx, del_data in enumerate(history_list):
                    row = idx // 2
                    col = idx % 2
                    self._create_deletion_history_card(self.deletion_container, del_data, row, col, idx + 1)
                
                # Bind mousewheel to all new cards
                if hasattr(self, '_bind_deletion_mousewheel'):
                    self._bind_deletion_mousewheel(self.deletion_container)
                    
        except Exception as e:
            self._show_empty_message(self.deletion_container, "Could not load history", str(e))
    
    def _create_deletion_history_card(self, parent, del_data, row=0, col=0, index=1):
        """Create a card for processed deletion in history view"""
        status = del_data.get('status', 'unknown')
        
        # Color scheme based on status
        if status == 'approved':
            bg_color = '#d4edda'  # Light green
            status_text = '✓ Approved'
            status_bg = '#28a745'
        else:  # rejected
            bg_color = '#f8d7da'  # Light red
            status_text = '✕ Rejected'
            status_bg = '#dc3545'
        
        card = tk.Frame(parent, bg=bg_color, relief='solid', bd=1)
        card.grid(row=row, column=col, sticky='nsew', padx=8, pady=8)
        
        # Bind mousewheel
        def _on_card_mousewheel(event):
            self.deletion_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        card.bind("<MouseWheel>", _on_card_mousewheel)
        
        inner = tk.Frame(card, bg=bg_color)
        inner.pack(fill=tk.BOTH, expand=True, padx=18, pady=15)
        inner.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Header row with index and status badge
        header_row = tk.Frame(inner, bg=bg_color)
        header_row.pack(fill=tk.X)
        
        # Index badge
        tk.Label(
            header_row,
            text=f"#{index}",
            font=('Segoe UI', 11, 'bold'),
            bg='#6c757d',
            fg='white',
            padx=8,
            pady=2
        ).pack(side=tk.LEFT)
        
        # Status badge
        tk.Label(
            header_row,
            text=status_text,
            font=('Segoe UI', 9, 'bold'),
            bg=status_bg,
            fg='white',
            padx=10,
            pady=3
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Deletion badge
        tk.Label(
            header_row,
            text="🗑️ Deletion",
            font=('Segoe UI', 8, 'bold'),
            bg='#c0392b',
            fg='white',
            padx=8,
            pady=2
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Student info
        student_row = tk.Frame(inner, bg=bg_color)
        student_row.pack(fill=tk.X, pady=(8, 0))
        
        tk.Label(
            student_row,
            text=f"{del_data.get('student_name', 'Deleted Account')}",
            font=('Segoe UI', 11, 'bold'),
            bg=bg_color,
            fg='#333'
        ).pack(side=tk.LEFT)
        
        tk.Label(
            student_row,
            text=f"({del_data.get('student_id', '')})",
            font=('Segoe UI', 9),
            bg=bg_color,
            fg='#666'
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Reason
        reason = del_data.get('reason', 'No reason provided')
        tk.Label(
            inner,
            text=f"Reason: {reason[:80]}{'...' if len(reason) > 80 else ''}",
            font=('Segoe UI', 9, 'italic'),
            bg=bg_color,
            fg='#555',
            wraplength=350,
            justify='left'
        ).pack(anchor='w', pady=(6, 0))
        
        # Timestamp
        tk.Label(
            inner,
            text=f"📅 {del_data.get('timestamp', 'N/A')}",
            font=('Segoe UI', 8),
            bg=bg_color,
            fg='#888'
        ).pack(anchor='w', pady=(6, 0))
    
    def _create_deletion_card(self, parent, del_data, row=0, col=0, index=1):
        """Create a card for deletion request in grid layout with index"""
        card = tk.Frame(parent, bg='#ffe6e6', relief='solid', bd=1)
        card.grid(row=row, column=col, sticky='nsew', padx=8, pady=8)
        
        # Bind mousewheel to card for scrolling
        def _on_card_mousewheel(event):
            self.deletion_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        card.bind("<MouseWheel>", _on_card_mousewheel)
        
        inner = tk.Frame(card, bg='#ffe6e6')
        inner.pack(fill=tk.BOTH, expand=True, padx=18, pady=15)
        inner.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Header row with index badge
        header_row = tk.Frame(inner, bg='#ffe6e6')
        header_row.pack(fill=tk.X)
        header_row.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Index badge
        index_badge = tk.Label(
            header_row,
            text=f"#{index}",
            font=('Segoe UI', 11, 'bold'),
            bg='#c0392b',
            fg='white',
            padx=8,
            pady=2
        )
        index_badge.pack(side=tk.LEFT)
        index_badge.bind("<MouseWheel>", _on_card_mousewheel)
        
        # Deletion icon badge
        tk.Label(
            header_row,
            text="🗑️ Deletion Request",
            font=('Segoe UI', 9, 'bold'),
            bg='#dc3545',
            fg='white',
            padx=10,
            pady=3
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Student name row
        name_row = tk.Frame(inner, bg='#ffe6e6')
        name_row.pack(fill=tk.X, pady=(8, 0))
        
        tk.Label(
            name_row,
            text=f"{del_data.get('student_name', 'Unknown')}",
            font=('Segoe UI', 11, 'bold'),
            bg='#ffe6e6',
            fg='#c0392b'
        ).pack(side=tk.LEFT)
        
        tk.Label(
            name_row,
            text=f"({del_data.get('student_id', '')})",
            font=('Segoe UI', 9),
            bg='#ffe6e6',
            fg='#666'
        ).pack(side=tk.LEFT, padx=(8, 0))
        
        # Reason
        reason = del_data.get('reason', 'No reason provided')
        tk.Label(
            inner,
            text=f"Reason: {reason[:80]}{'...' if len(reason) > 80 else ''}",
            font=('Segoe UI', 9, 'italic'),
            bg='#ffe6e6',
            fg='#555',
            wraplength=350,
            justify='left'
        ).pack(anchor='w', pady=(6, 0))
        
        # Timestamp
        tk.Label(
            inner,
            text=f"📅 {del_data.get('timestamp', 'N/A')}",
            font=('Segoe UI', 8),
            bg='#ffe6e6',
            fg='#888'
        ).pack(anchor='w', pady=(6, 0))
        
        # Actions row at bottom
        actions_frame = tk.Frame(inner, bg='#ffe6e6')
        actions_frame.pack(fill=tk.X, pady=(12, 0))
        
        del_id = del_data.get('id')
        
        approve_btn = tk.Button(
            actions_frame,
            text="✓ Approve Deletion",
            font=('Segoe UI', 9, 'bold'),
            bg='#dc3545',
            fg='white',
            padx=12,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=lambda did=del_id, name=del_data.get('student_name'): self._handle_deletion_action(did, 'approve', name)
        )
        approve_btn.pack(side=tk.LEFT, padx=(0, 8))
        
        reject_btn = tk.Button(
            actions_frame,
            text="✕ Reject",
            font=('Segoe UI', 9, 'bold'),
            bg='#6c757d',
            fg='white',
            padx=12,
            pady=5,
            cursor='hand2',
            relief='flat',
            command=lambda did=del_id: self._handle_deletion_action(did, 'reject')
        )
        reject_btn.pack(side=tk.LEFT)
    
    def _handle_deletion_action(self, del_id, action, student_name=None):
        """Handle deletion request action"""
        if action == 'approve':
            if not messagebox.askyesno("Confirm Deletion", 
                f"Are you sure you want to approve the deletion request for {student_name}?\n\n"
                "This will remove the student from the library system,\n"
                "return any borrowed books, and delete their portal access."):
                return
        
        try:
            import urllib.request
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/deletion/{del_id}/{action}"
            req = urllib.request.Request(url, method='POST', data=b'')
            
            with urllib.request.urlopen(req, timeout=5) as response:
                result = json.loads(response.read().decode())
                if result.get('status') == 'success':
                    messagebox.showinfo("Success", result.get('message', f"Deletion {action}d!"))
                    self._refresh_deletion_requests()
                    # Refresh students list since student was removed from library DB
                    if action == 'approve':
                        self.refresh_students()
                else:
                    messagebox.showerror("Error", result.get('message', 'Action failed'))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to {action} deletion: {str(e)}")
    
    def _create_password_reset_section(self, parent):
        """Create password reset section with auth dashboard - PREMIUM UI"""
        # Main container with subtle gradient effect
        container = tk.Frame(parent, bg='white')
        container.pack(fill=tk.BOTH, expand=True)
        
        # ════════════════════════════════════════════════════════════════
        # HEADER SECTION - Clean and Modern
        # ════════════════════════════════════════════════════════════════
        header_frame = tk.Frame(container, bg='white')
        header_frame.pack(fill=tk.X, padx=25, pady=(20, 20))
        
        # Title with icon
        title_frame = tk.Frame(header_frame, bg='white')
        title_frame.pack(side=tk.LEFT)
        
        tk.Label(
            title_frame,
            text="🔐",
            font=('Segoe UI', 22),
            bg='white'
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        title_text_frame = tk.Frame(title_frame, bg='white')
        title_text_frame.pack(side=tk.LEFT)
        
        tk.Label(
            title_text_frame,
            text="Password & Authentication",
            font=('Segoe UI', 16, 'bold'),
            bg='white',
            fg='#1a1a2e'
        ).pack(anchor='w')
        
        tk.Label(
            title_text_frame,
            text="Manage student credentials and access",
            font=('Segoe UI', 9),
            bg='white',
            fg='#888'
        ).pack(anchor='w')
        
        # Refresh button - Premium styling
        refresh_btn = tk.Button(
            header_frame,
            text="⟳ Sync Data",
            font=('Segoe UI', 10, 'bold'),
            bg='#4f46e5',
            fg='white',
            padx=18,
            pady=8,
            cursor='hand2',
            relief='flat',
            activebackground='#4338ca',
            activeforeground='white',
            command=self._refresh_auth_stats
        )
        refresh_btn.pack(side=tk.RIGHT)
        
        # ════════════════════════════════════════════════════════════════
        # STATS DASHBOARD - Premium Cards
        # ════════════════════════════════════════════════════════════════
        stats_frame = tk.Frame(container, bg='white')
        stats_frame.pack(fill=tk.X, padx=25, pady=(0, 20))
        
        self.auth_stats_labels = {}
        stat_configs = [
            ("Total Registered", "#0ea5e9", "#e0f2fe", "total_registered", "📊"),
            ("Active Users", "#10b981", "#d1fae5", "active_users", "✅"),
            ("Pending Change", "#f59e0b", "#fef3c7", "pending_change", "⏳")
        ]
        
        for i, (label, accent, bg_color, key, icon) in enumerate(stat_configs):
            card = tk.Frame(stats_frame, bg=bg_color, relief='flat', bd=0)
            card.pack(side=tk.LEFT, padx=(0 if i == 0 else 12, 0), ipadx=25, ipady=12)
            
            # Icon
            tk.Label(card, text=icon, font=('Segoe UI', 16), bg=bg_color).pack(anchor='w', padx=(10, 0))
            
            # Value
            count_label = tk.Label(
                card,
                text="0",
                font=('Segoe UI', 28, 'bold'),
                bg=bg_color,
                fg=accent
            )
            count_label.pack(anchor='w', padx=(10, 0))
            self.auth_stats_labels[key] = count_label
            
            # Label
            tk.Label(
                card,
                text=label,
                font=('Segoe UI', 9, 'bold'),
                bg=bg_color,
                fg='#555'
            ).pack(anchor='w', padx=(10, 10))
        
        # ════════════════════════════════════════════════════════════════
        # TWO COLUMN LAYOUT
        # ════════════════════════════════════════════════════════════════
        columns_frame = tk.Frame(container, bg='white')
        columns_frame.pack(fill=tk.BOTH, expand=True, padx=25, pady=(0, 20))
        columns_frame.columnconfigure(0, weight=1, minsize=350)
        columns_frame.columnconfigure(1, weight=1, minsize=350)
        columns_frame.rowconfigure(0, weight=1)
        
        # ════════════════════════════════════════════════════════════════
        # LEFT COLUMN - Password Tools
        # ════════════════════════════════════════════════════════════════
        left_card = tk.Frame(columns_frame, bg='#f8fafc', relief='flat', bd=0)
        left_card.grid(row=0, column=0, sticky='nsew', padx=(0, 12), pady=0)
        
        # Card Header
        left_header = tk.Frame(left_card, bg='#0f172a')
        left_header.pack(fill=tk.X)
        tk.Label(
            left_header,
            text="🛠️ Password Management",
            font=('Segoe UI', 11, 'bold'),
            bg='#0f172a',
            fg='white',
            padx=15,
            pady=12
        ).pack(anchor='w')
        
        left_inner = tk.Frame(left_card, bg='#f8fafc')
        left_inner.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # ─────────────────────────────────────────
        # TOOL 1: Student Search
        # ─────────────────────────────────────────
        search_card = tk.Frame(left_inner, bg='white', relief='flat', bd=0)
        search_card.pack(fill=tk.X, pady=(0, 12))
        
        # Search header with accent bar
        search_accent = tk.Frame(search_card, bg='#3b82f6', height=4)
        search_accent.pack(fill=tk.X)
        
        search_content = tk.Frame(search_card, bg='white')
        search_content.pack(fill=tk.X, padx=12, pady=12)
        
        tk.Label(
            search_content,
            text="🔍 Find Student",
            font=('Segoe UI', 10, 'bold'),
            bg='white',
            fg='#1e40af'
        ).pack(anchor='w')
        
        tk.Label(
            search_content,
            text="Search by name or enrollment to reset password",
            font=('Segoe UI', 8),
            bg='white',
            fg='#888'
        ).pack(anchor='w', pady=(0, 8))
        
        search_row = tk.Frame(search_content, bg='white')
        search_row.pack(fill=tk.X)
        
        self.student_search_entry = tk.Entry(
            search_row,
            font=('Segoe UI', 10),
            relief='solid',
            bd=1,
            fg='#666'
        )
        self.student_search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=6)
        self.student_search_entry.insert(0, "Type name or ID...")
        self.student_search_entry.bind("<FocusIn>", lambda e: (self.student_search_entry.delete(0, tk.END), self.student_search_entry.config(fg='#333')) if self.student_search_entry.get() == "Type name or ID..." else None)
        self.student_search_entry.bind("<Return>", lambda e: self._search_student_for_reset())
        
        search_btn = tk.Button(
            search_row,
            text="Search",
            font=('Segoe UI', 9, 'bold'),
            bg='#3b82f6',
            fg='white',
            padx=14,
            pady=4,
            cursor='hand2',
            relief='flat',
            activebackground='#2563eb',
            command=self._search_student_for_reset
        )
        search_btn.pack(side=tk.LEFT, padx=(8, 0))
        
        # Search results container
        self.search_results_container = tk.Frame(search_content, bg='white')
        self.search_results_container.pack(fill=tk.X, pady=(8, 0))
        
        # ─────────────────────────────────────────
        # TOOL 2: Quick Reset
        # ─────────────────────────────────────────
        reset_card = tk.Frame(left_inner, bg='white', relief='flat', bd=0)
        reset_card.pack(fill=tk.X, pady=(0, 12))
        
        reset_accent = tk.Frame(reset_card, bg='#f59e0b', height=4)
        reset_accent.pack(fill=tk.X)
        
        reset_content = tk.Frame(reset_card, bg='white')
        reset_content.pack(fill=tk.X, padx=12, pady=12)
        
        tk.Label(
            reset_content,
            text="⚡ Quick Reset",
            font=('Segoe UI', 10, 'bold'),
            bg='white',
            fg='#b45309'
        ).pack(anchor='w')
        
        tk.Label(
            reset_content,
            text="Directly reset using enrollment number",
            font=('Segoe UI', 8),
            bg='white',
            fg='#888'
        ).pack(anchor='w', pady=(0, 8))
        
        reset_row = tk.Frame(reset_content, bg='white')
        reset_row.pack(fill=tk.X)
        
        self.password_reset_enrollment = tk.Entry(
            reset_row,
            font=('Segoe UI', 10),
            relief='solid',
            bd=1
        )
        self.password_reset_enrollment.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=6)
        
        reset_btn = tk.Button(
            reset_row,
            text="Reset",
            font=('Segoe UI', 9, 'bold'),
            bg='#f59e0b',
            fg='white',
            padx=14,
            pady=4,
            cursor='hand2',
            relief='flat',
            activebackground='#d97706',
            command=self._handle_password_reset
        )
        reset_btn.pack(side=tk.LEFT, padx=(8, 0))
        
        # ─────────────────────────────────────────
        # TOOL 3: Bulk Reset (Danger Zone)
        # ─────────────────────────────────────────
        danger_card = tk.Frame(left_inner, bg='#fef2f2', relief='flat', bd=0)
        danger_card.pack(fill=tk.X)
        
        danger_accent = tk.Frame(danger_card, bg='#dc2626', height=4)
        danger_accent.pack(fill=tk.X)
        
        danger_content = tk.Frame(danger_card, bg='#fef2f2')
        danger_content.pack(fill=tk.X, padx=12, pady=12)
        
        tk.Label(
            danger_content,
            text="⚠️ Batch Reset",
            font=('Segoe UI', 10, 'bold'),
            bg='#fef2f2',
            fg='#b91c1c'
        ).pack(anchor='w')
        
        tk.Label(
            danger_content,
            text="Reset ALL students in a year. Use with extreme caution!",
            font=('Segoe UI', 8),
            bg='#fef2f2',
            fg='#888'
        ).pack(anchor='w', pady=(0, 8))
        
        danger_row = tk.Frame(danger_content, bg='#fef2f2')
        danger_row.pack(fill=tk.X)
        
        self.bulk_reset_year_var = tk.StringVar(value="Select Year...")
        year_combo = ttk.Combobox(danger_row, textvariable=self.bulk_reset_year_var, state="readonly", width=18, font=('Segoe UI', 10))
        year_combo['values'] = ("1st Year", "2nd Year", "3rd Year", "All Years")
        year_combo.pack(side=tk.LEFT, ipady=3)
        
        danger_btn = tk.Button(
            danger_row,
            text="⚠ Reset All",
            font=('Segoe UI', 9, 'bold'),
            bg='#dc2626',
            fg='white',
            padx=12,
            pady=4,
            cursor='hand2',
            relief='flat',
            activebackground='#b91c1c',
            command=self._handle_bulk_reset
        )
        danger_btn.pack(side=tk.LEFT, padx=(8, 0))
        
        # ════════════════════════════════════════════════════════════════
        # RIGHT COLUMN - Recent Activity
        # ════════════════════════════════════════════════════════════════
        right_card = tk.Frame(columns_frame, bg='#f8fafc', relief='flat', bd=0)
        right_card.grid(row=0, column=1, sticky='nsew', padx=(12, 0), pady=0)
        
        # Card Header
        right_header = tk.Frame(right_card, bg='#0f172a')
        right_header.pack(fill=tk.X)
        tk.Label(
            right_header,
            text="📜 Recent Password Resets",
            font=('Segoe UI', 11, 'bold'),
            bg='#0f172a',
            fg='white',
            padx=15,
            pady=12
        ).pack(anchor='w')
        
        self.recent_resets_container = tk.Frame(right_card, bg='#f8fafc')
        self.recent_resets_container.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # Delay initial stats load for server startup
        self.root.after(4000, self._refresh_auth_stats)
    
    def _handle_password_reset(self):
        """Handle password reset action"""

        enrollment = self.password_reset_enrollment.get().strip()
        
        if not enrollment:
            messagebox.showwarning("Missing Information", "Please enter an enrollment number.")
            return
        
        if not messagebox.askyesno("Confirm Reset", 
            f"Reset password for enrollment: {enrollment}?\n\n"
            "Their password will be set to their enrollment number."):
            return
        
        try:
            import urllib.request
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/password-reset/{enrollment}"
            req = urllib.request.Request(url, method='POST', data=b'')
            
            with urllib.request.urlopen(req, timeout=5) as response:
                result = json.loads(response.read().decode())
                if result.get('status') == 'success':
                    messagebox.showinfo("Success", result.get('message', 'Password reset successfully!'))
                    self.password_reset_enrollment.delete(0, tk.END)
                else:
                    messagebox.showerror("Error", result.get('message', 'Reset failed'))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to reset password: {str(e)}")
        
        # Refresh stats after reset
        self._refresh_auth_stats()
    
    def _search_student_for_reset(self):
        """Search for student by name or enrollment number"""
        query = self.student_search_entry.get().strip()
        
        if not query or query == "Name or Enrollment...":
            messagebox.showwarning("Search", "Please enter a name or enrollment number to search.")
            return
        
        # Clear previous results
        for w in self.search_results_container.winfo_children():
            w.destroy()
        
        try:
            # Search in local database
            conn = self.db.get_connection()
            cursor = conn.cursor()
            
            # Search by name or enrollment (case-insensitive)
            cursor.execute("""
                SELECT enrollment_no, name, year, department 
                FROM students 
                WHERE LOWER(name) LIKE ? OR enrollment_no LIKE ?
                LIMIT 5
            """, (f"%{query.lower()}%", f"%{query}%"))
            
            results = cursor.fetchall()
            conn.close()
            
            if not results:
                tk.Label(
                    self.search_results_container,
                    text="No students found",
                    font=('Segoe UI', 9, 'italic'),
                    bg='#e8f4fd',
                    fg='#888'
                ).pack(anchor='w')
                return
            
            # Display results as clickable items
            for student in results:
                row = tk.Frame(self.search_results_container, bg='white', relief='solid', bd=1)
                row.pack(fill=tk.X, pady=2)
                
                info_frame = tk.Frame(row, bg='white')
                info_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5, pady=3)
                
                tk.Label(
                    info_frame,
                    text=f"{student[1]}",  # Name
                    font=('Segoe UI', 9, 'bold'),
                    bg='white',
                    fg='#333'
                ).pack(anchor='w')
                
                tk.Label(
                    info_frame,
                    text=f"{student[0]} | {student[2] or 'N/A'} | {student[3] or 'N/A'}",
                    font=('Segoe UI', 7),
                    bg='white',
                    fg='#888'
                ).pack(anchor='w')
                
                # Reset button for this student
                reset_btn = tk.Button(
                    row,
                    text="Reset",
                    font=('Segoe UI', 8, 'bold'),
                    bg='#fd7e14',
                    fg='white',
                    padx=6,
                    pady=1,
                    cursor='hand2',
                    relief='flat',
                    command=lambda e=student[0]: self._reset_student_password(e)
                )
                reset_btn.pack(side=tk.RIGHT, padx=5, pady=3)
                
        except Exception as e:
            tk.Label(
                self.search_results_container,
                text=f"Search error: {str(e)[:30]}",
                font=('Segoe UI', 8),
                bg='#e8f4fd',
                fg='#dc3545'
            ).pack(anchor='w')
    
    def _reset_student_password(self, enrollment):
        """Reset a specific student's password (from search results)"""
        if not messagebox.askyesno("Confirm Reset", 
            f"Reset password for: {enrollment}?\n\nPassword will be set to enrollment number."):
            return
        
        try:
            import urllib.request
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/password-reset/{enrollment}"
            req = urllib.request.Request(url, method='POST', data=b'')
            
            with urllib.request.urlopen(req, timeout=5) as response:
                result = json.loads(response.read().decode())
                if result.get('status') == 'success':
                    messagebox.showinfo("Success", f"Password reset for {enrollment}!")
                    self._refresh_auth_stats()
                else:
                    messagebox.showerror("Error", result.get('message', 'Reset failed'))
        except Exception as e:
            messagebox.showerror("Error", f"Failed: {str(e)}")
    
    def _handle_bulk_reset(self):
        """Handle bulk password reset for a year group"""
        year = self.bulk_reset_year_var.get()
        
        if year == "Select Year":
            messagebox.showwarning("Select Year", "Please select a year group first.")
            return
        
        # Map display value to database year value
        year_map = {
            "1st Year": "1st",
            "2nd Year": "2nd", 
            "3rd Year": "3rd",
            "All Years": None  # None means all years
        }
        db_year = year_map.get(year)
        
        # Get count of students affected
        try:
            conn = self.db.get_connection()
            cursor = conn.cursor()
            
            if db_year:
                cursor.execute("SELECT COUNT(*) FROM students WHERE year = ?", (db_year,))
            else:
                cursor.execute("SELECT COUNT(*) FROM students")
            
            count = cursor.fetchone()[0]
            conn.close()
            
            if count == 0:
                messagebox.showinfo("No Students", f"No students found for {year}.")
                return
            
            # Confirmation with student count
            if not messagebox.askyesno("⚠️ Bulk Reset Warning", 
                f"This will reset passwords for {count} students in {year}!\n\n"
                "All their passwords will be set to their enrollment numbers.\n\n"
                "Are you absolutely sure?"):
                return
            
            # Second confirmation
            if not messagebox.askyesno("Final Confirmation", 
                f"LAST CHANCE: Reset {count} passwords?\n\nThis action cannot be undone!"):
                return
            
            # Perform bulk reset via API
            import urllib.request
            import urllib.parse
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/bulk-password-reset"
            data = json.dumps({'year': db_year}).encode('utf-8')
            req = urllib.request.Request(url, method='POST', data=data, 
                                         headers={'Content-Type': 'application/json'})
            
            with urllib.request.urlopen(req, timeout=30) as response:
                result = json.loads(response.read().decode())
                if result.get('status') == 'success':
                    reset_count = result.get('count', count)
                    messagebox.showinfo("✅ Bulk Reset Complete", 
                        f"Successfully reset {reset_count} passwords!\n\n"
                        "Students will be prompted to change on next login.")
                    self._refresh_auth_stats()
                else:
                    messagebox.showerror("Error", result.get('message', 'Bulk reset failed'))
                    
        except Exception as e:
            messagebox.showerror("Error", f"Failed to perform bulk reset: {str(e)}")
    
    def _refresh_auth_stats(self):
        """Fetch and display auth statistics and recent password resets"""
        if not WEB_PORTAL_AVAILABLE:
            return


        try:
            import urllib.request
            
            url = f"http://127.0.0.1:{self.portal_port}/api/admin/auth-stats"
            req = urllib.request.Request(url)
            
            with urllib.request.urlopen(req, timeout=5) as response:
                data = json.loads(response.read().decode())
                stats = data.get('stats', {})
                resets = data.get('recent_resets', [])
                
                # Update stats labels
                if hasattr(self, 'auth_stats_labels'):
                    self.auth_stats_labels['total_registered'].config(text=str(stats.get('total_registered', 0)))
                    self.auth_stats_labels['active_users'].config(text=str(stats.get('active_users', 0)))
                    self.auth_stats_labels['pending_change'].config(text=str(stats.get('pending_change', 0)))
                
                # Update recent resets list
                if hasattr(self, 'recent_resets_container'):
                    for w in self.recent_resets_container.winfo_children():
                        w.destroy()
                    
                    if not resets:
                        tk.Label(
                            self.recent_resets_container,
                            text="No recent password resets",
                            font=('Segoe UI', 10),
                            bg='#f8f9fa',
                            fg='#888'
                        ).pack(pady=20)
                    else:
                        for reset in resets[:8]:  # Show max 8
                            row = tk.Frame(self.recent_resets_container, bg='#f8f9fa')
                            row.pack(fill=tk.X, pady=3)
                            
                            tk.Label(
                                row,
                                text=f"👤 {reset.get('student_name', 'Unknown')}",
                                font=('Segoe UI', 9),
                                bg='#f8f9fa',
                                fg='#333'
                            ).pack(side=tk.LEFT)
                            
                            tk.Label(
                                row,
                                text=f"({reset.get('enrollment_no', '')})",
                                font=('Segoe UI', 8),
                                bg='#f8f9fa',
                                fg='#666'
                            ).pack(side=tk.LEFT, padx=5)
                            
                            timestamp = reset.get('last_changed', '')
                            if timestamp:
                                tk.Label(
                                    row,
                                    text=timestamp[:16],  # Trim to date/time
                                    font=('Segoe UI', 8),
                                    bg='#f8f9fa',
                                    fg='#888'
                                ).pack(side=tk.RIGHT)
                    
        except Exception as e:
            # Silently fail - stats will just show 0
            pass
    
    def _show_empty_message(self, container, title, message):
        """Show empty state message"""
        frame = tk.Frame(container, bg='white')
        frame.pack(fill=tk.BOTH, expand=True, pady=40)
        
        tk.Label(
            frame,
            text=title,
            font=('Segoe UI', 14, 'bold'),
            bg='white',
            fg='#666'
        ).pack()
        
        tk.Label(
            frame,
            text=message,
            font=('Segoe UI', 10),
            bg='white',
            fg='#999'
        ).pack(pady=(8, 0))

    
    def refresh_analysis(self):
        """Refresh all analysis charts based on selected time period (Threaded implementation)"""
        if not MATPLOTLIB_AVAILABLE:
            return
        
        # If charts are hidden, clear frames and show a small placeholder
        if hasattr(self, 'analysis_show_charts') and not self.analysis_show_charts.get():
            def _clear_and_placeholder(frame, text):
                for w in frame.winfo_children():
                    w.destroy()
                tk.Label(frame, text=text, font=('Segoe UI', 11), bg=self.colors['primary'], fg='#666').pack(expand=True, pady=20)
            _clear_and_placeholder(self.borrow_status_frame, "Charts hidden")
            _clear_and_placeholder(self.student_activity_frame, "Charts hidden")
            if hasattr(self, 'inventory_overdue_frame'):
                _clear_and_placeholder(self.inventory_overdue_frame, "Charts hidden")
            if hasattr(self, 'branch_activity_frame'):
                _clear_and_placeholder(self.branch_activity_frame, "Charts hidden")
            if hasattr(self, 'branch_overdue_frame'):
                _clear_and_placeholder(self.branch_overdue_frame, "Charts hidden")
            # Also clear summary
            for w in self.stats_summary_frame.winfo_children():
                w.destroy()
            tk.Label(self.stats_summary_frame, text="Charts are hidden. Enable 'Show Charts' to view.", font=('Segoe UI', 11), bg=self.colors['primary'], fg='#666').pack(pady=10)
            return

        # Clear existing widgets to avoid duplicates (charts, labels, legends)
        def _clear(frame):
            try:
                for w in frame.winfo_children():
                    w.destroy()
            except Exception:
                pass
        _clear(self.borrow_status_frame)
        _clear(self.student_activity_frame)
        if hasattr(self, 'daily_trend_frame'): _clear(self.daily_trend_frame)
        if hasattr(self, 'popular_books_frame'): _clear(self.popular_books_frame)
        if hasattr(self, 'least_popular_books_frame'): _clear(self.least_popular_books_frame)
        if hasattr(self, 'inventory_overdue_frame'): _clear(self.inventory_overdue_frame)
        if hasattr(self, 'branch_activity_frame'): _clear(self.branch_activity_frame)
        if hasattr(self, 'branch_overdue_frame'): _clear(self.branch_overdue_frame)
        _clear(self.stats_summary_frame)
        
        # Clear focused frames
        for w in self.student_specific_frame.winfo_children():
            if isinstance(w, FigureCanvasTkAgg): w.get_tk_widget().destroy()
            else: w.destroy()
        for w in self.book_specific_frame.winfo_children():
            if isinstance(w, FigureCanvasTkAgg): w.get_tk_widget().destroy()
            else: w.destroy()

        # Show Loading Indicator
        tk.Label(self.stats_summary_frame, text="⏳ Loading analysis data...", font=('Segoe UI', 12), bg=self.colors['primary'], fg='#666').pack(pady=20)
        
        days = int(self.analysis_period.get())
        en = self.analysis_filter.get('enrollment_no')
        bk = self.analysis_filter.get('book_id')
        br = self.analysis_branch_var.get() if hasattr(self, 'analysis_branch_var') else "All"
        
        self.analysis_filter_summary.config(text="Loading filter details...")
        
        # Run in background
        self.run_in_background_thread(
            self._fetch_analysis_data, 
            self._on_analysis_data_ready, 
            days=days, enrollment_no=en, book_id=bk, branch=br
        )

    def _on_analysis_data_ready(self, result):
        try:
            # Clear loading indicator
            for w in self.stats_summary_frame.winfo_children():
                w.destroy()
                
            if isinstance(result, Exception):
                tk.Label(self.stats_summary_frame, text=f"Error loading analysis: {result}", fg="red", bg=self.colors['primary']).pack()
                return

            data = result
            
            # Update filter summary with fetched names
            names = data.get('filter_names', {})
            en = self.analysis_filter.get('enrollment_no')
            bk = self.analysis_filter.get('book_id')
            br = self.analysis_branch_var.get() if hasattr(self, 'analysis_branch_var') else "All"
            parts = []
            if en: parts.append(f"Student: {en} ({names.get('student_name', 'Unknown') or 'Unknown'})")
            if bk: parts.append(f"Book: {bk} ({names.get('book_title', 'Unknown') or 'Unknown'})")
            if br and br != "All": parts.append(f"Branch: {br}")
            if parts:
                self.analysis_filter_summary.config(text=' | '.join(parts))
            else:
                self.analysis_filter_summary.config(text="No focused filter applied")

            days = int(self.analysis_period.get())

            # Render Charts with injected Data
            self.create_borrow_status_pie(data=data.get('borrow_status'))
            self.create_student_activity_pie(days, data=data.get('student_activity'))
            self.create_branch_activity_pie(days, data=data.get('branch_activity'))
            self.create_branch_overdue_pie(data=data.get('branch_overdue'))
            self.create_inventory_overdue_donut(data=data.get('inventory'))
            
            self.create_popular_books_chart(days, data=data.get('popular_books'))
            self.create_least_popular_books_chart(days, data=data.get('least_popular'))
            
            if hasattr(self, 'analysis_compact_mode') and not self.analysis_compact_mode.get():
                self.create_daily_trend_chart(days, data=data.get('daily_trend'))
            
            self.create_summary_stats(days, data=data.get('summary'))
            
            # Focused Insights
            if 'student_specific' in data and data['student_specific']:
                if not self.student_specific_frame.winfo_manager():
                    self.student_specific_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
                self.create_student_specific_pie(days, self.analysis_filter.get('enrollment_no'), data=data.get('student_specific'))

            if 'book_specific' in data and data['book_specific']:
                if not self.book_specific_frame.winfo_manager():
                    self.book_specific_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(10, 0))
                self.create_book_specific_pie(days, self.analysis_filter.get('book_id'), data=data.get('book_specific'))
        except Exception as e:
            print(f"Error in _on_analysis_data_ready: {e}")
    
    def create_borrow_status_pie(self, data=None):
        """Create pie chart showing book status distribution"""
        try:
            # Get data
            if data is not None:
                results = data
            else:
                conn = self.db.get_connection()
                cursor = conn.cursor()
                
                # Count books by status
                cursor.execute("""
                    SELECT 
                        CASE 
                            WHEN br.status = 'borrowed' THEN 'Currently Issued'
                            ELSE 'Available'
                        END as status,
                        COUNT(DISTINCT b.book_id) as count
                    FROM books b
                    LEFT JOIN borrow_records br ON b.book_id = br.book_id AND br.status = 'borrowed'
                    GROUP BY status
                """)
                
                results = cursor.fetchall()
                conn.close()
            
            if not results:
                # Clear frame then show no-data message
                for w in self.borrow_status_frame.winfo_children():
                    w.destroy()
                tk.Label(
                    self.borrow_status_frame,
                    text="No books or borrow data to display",
                    font=('Segoe UI', 12), bg=self.colors['primary'], fg='#666'
                ).pack(expand=True, pady=20)
                return
            
            raw_labels = [row[0] for row in results]
            sizes = [row[1] for row in results]
            labels = [f"{name} ({cnt})" for name, cnt in zip(raw_labels, sizes)]
            # Colorblind-friendly palette
            colors = ['#0072B2', '#D55E00', '#F0E442', '#009E73', '#CC79A7', '#56B4E9']
            
            # Larger, modern figure with tight layout for legends
            fig = Figure(figsize=(6, 4), dpi=100)
            ax = fig.add_subplot(111)
            fig.patch.set_facecolor('white')
            fig.subplots_adjust(left=0.1, right=0.78)  # Make room for legend
            
            def on_pie_click(event):
                if event.inaxes == ax:
                    # Find which wedge was clicked
                    for i, (wedge, label, size) in enumerate(zip(ax.patches, labels, sizes)):
                        contains, info = wedge.contains(event)
                        if contains:
                            # Check the raw status name (not the formatted label with count)
                            status_name = raw_labels[i]
                            if status_name == 'Currently Issued':
                                self.show_currently_borrowed_dialog()
                            else:
                                self.show_available_books_dialog()
                            break
            
            def _autopct(pct, allvals=sizes):
                total = sum(allvals)
                if total == 0:
                    return " "
                val = int(round(pct*total/100.0))
                return f"{pct:.1f}%\n({val})"


            wedges, texts, autotexts = ax.pie(
                sizes, 
                labels=labels, 
                colors=colors[:len(sizes)],
                autopct=_autopct,
                startangle=90,
                wedgeprops=dict(width=0.32, edgecolor='white')  # donut style, slightly thinner
            )
            # Center label with total
            ax.text(0, 0, f"Total\n{sum(sizes)}", ha='center', va='center', fontsize=14, fontweight='bold', color='#333')
            ax.axis('equal')
            ax.set_title('Book Status Distribution', fontsize=16, fontweight='bold', color='#0072B2')
            # Add legend for clarity
            ax.legend(wedges, labels, title="Status", loc='center left', bbox_to_anchor=(1.0, 0.5), fontsize=10)
            
            # Clear any existing widgets in the frame
            for widget in self.borrow_status_frame.winfo_children():
                widget.destroy()
            
            # Add to GUI
            canvas = FigureCanvasTkAgg(fig, self.borrow_status_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            canvas.mpl_connect('button_press_event', on_pie_click)
            # Ensure wheel scrolling works when cursor is over the figure
            try:
                self._analysis_bind_wheel(canvas.get_tk_widget())
            except Exception:
                pass
            
            self.current_charts['borrow_status'] = (fig, labels, sizes)
            
        except Exception as e:
            print(f"Error creating borrow status pie chart: {e}")
    
    def create_student_activity_pie(self, days, data=None):
        """Create pie chart showing student activity levels"""
        try:
            # Get data
            if data is not None:
                results = data
            else:
                conn = self.db.get_connection()
                cursor = conn.cursor()
                
                start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
                
                # Count students by activity level
                cursor.execute("""
                    SELECT 
                        s.year,
                        COUNT(br.id) as borrow_count
                    FROM students s
                    LEFT JOIN borrow_records br ON s.enrollment_no = br.enrollment_no 
                        AND br.borrow_date >= ?
                    GROUP BY s.year
                    HAVING COUNT(br.id) > 0
                    ORDER BY borrow_count DESC
                """, (start_date,))
                
                results = cursor.fetchall()
                conn.close()
            
            if not results:
                # Always render a placeholder donut so the chart area is not blank
                fig = Figure(figsize=(6, 4), dpi=100)
                ax = fig.add_subplot(111)
                fig.patch.set_facecolor('white')
                fig.subplots_adjust(left=0.1, right=0.78)

                sizes = [1]
                labels = ["No Activity"]
                colors = ['#d0d7de']  # light gray
                wedges, texts = ax.pie(
                    sizes,
                    labels=None,
                    colors=colors,
                    startangle=90,
                    wedgeprops=dict(width=0.32, edgecolor='white')
                )
                ax.text(0, 0, "No Data", ha='center', va='center', fontsize=14, fontweight='bold', color='#666')
                ax.axis('equal')
                ax.set_title(f'Student Activity (Last {days} Days)', fontsize=16, fontweight='bold', color='#D55E00')
                ax.legend(wedges, labels, title="Year", loc='center left', bbox_to_anchor=(1.0, 0.5), fontsize=10)

                for w in self.student_activity_frame.winfo_children():
                    w.destroy()
                canvas = FigureCanvasTkAgg(fig, self.student_activity_frame)
                canvas.draw()
                canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
                try:
                    self._analysis_bind_wheel(canvas.get_tk_widget())
                except Exception:
                    pass
                self.current_charts['student_activity'] = (fig, labels, sizes)
                return
            
            def _format_year_label(y):
                if y is None:
                    return "Unknown"
                s = str(y).strip()
                low = s.lower()
                # Normalize any variant that mentions pass + out to 'Pass Out'
                if ("pass" in low) and ("out" in low):
                    return "Pass Out"
                # If it already contains the word 'year', keep as-is with title casing
                if "year" in low:
                    return s.title()
                # Otherwise append 'Year' to numeric/ordinal
                return f"{s} Year"

            raw_labels = [_format_year_label(row[0]) for row in results]
            sizes = [row[1] for row in results]
            # Keep legend labels short to avoid truncation
            labels = list(raw_labels)
            # Colorblind-friendly palette
            colors = ['#0072B2', '#D55E00', '#F0E442', '#009E73', '#CC79A7', '#56B4E9']
            
            # Larger, modern figure with tight layout for legends
            fig = Figure(figsize=(6, 4), dpi=100)
            ax = fig.add_subplot(111)
            fig.patch.set_facecolor('white')
            fig.subplots_adjust(left=0.1, right=0.78)  # Make room for legend
            
            def on_activity_click(event):
                if event.inaxes == ax:
                    for i, (wedge, label, size) in enumerate(zip(ax.patches, labels, sizes)):
                        contains, info = wedge.contains(event)
                        if contains:
                            year = results[i][0]
                            self.show_students_by_year_dialog(year, days)
                            break
            
            # Explode the largest slice slightly for emphasis
            if sizes:
                max_idx = sizes.index(max(sizes))
                explode = [0.08 if i == max_idx else 0 for i in range(len(sizes))]
            else:
                explode = None

            # Build pie with labels and percentages
            def autopct_format(pct):
                return f'{pct:.1f}%' if pct > 5 else ''
            
            wedges, texts, autotexts = ax.pie(
                sizes,
                labels=labels,  # Show labels on slices
                colors=colors[:len(sizes)],
                autopct=autopct_format,
                startangle=90,
                explode=explode,
                wedgeprops=dict(width=0.32, edgecolor='white'),  # donut style, thinner
                textprops={'fontsize': 10, 'weight': 'bold'}
            )
            # Make percentage text black for better visibility
            for autotext in autotexts:
                autotext.set_color('black')
            
            ax.text(0, 0, f"Total\n{sum(sizes)}", ha='center', va='center', fontsize=14, fontweight='bold', color='#333')
            ax.axis('equal')
            ax.set_title(f'Student Activity (Last {days} Days)', fontsize=16, fontweight='bold', color='#D55E00')
            
            # Clear any existing widgets in the frame
            for widget in self.student_activity_frame.winfo_children():
                widget.destroy()
            
            # Add to GUI
            canvas = FigureCanvasTkAgg(fig, self.student_activity_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            canvas.mpl_connect('button_press_event', on_activity_click)
            try:
                self._analysis_bind_wheel(canvas.get_tk_widget())
            except Exception:
                pass
            
            self.current_charts['student_activity'] = (fig, labels, sizes)
            
        except Exception as e:
            try:
                for w in self.student_activity_frame.winfo_children():
                    w.destroy()
                tk.Label(
                    self.student_activity_frame,
                    text="Unable to render chart",
                    font=('Segoe UI', 12),
                    bg=self.colors['primary'],
                    fg='#b00020'
                ).pack(expand=True, pady=20)
            except Exception:
                pass
            print(f"Error creating student activity pie chart: {e}")

    def create_inventory_overdue_donut(self, data=None):
        """Create a nested donut pie showing Available vs Issued (outer), and inner ring splitting Issued into On-time vs Overdue."""
        try:
            if data:
                total_copies = data.get('total_copies', 0)
                total_available = data.get('total_available', 0)
                overdue = data.get('overdue', 0)
            else:
                conn = self.db.get_connection()
                cur = conn.cursor()
                # Sum copies
                cur.execute("SELECT COALESCE(SUM(total_copies),0), COALESCE(SUM(available_copies),0) FROM books")
                total_copies, total_available = cur.fetchone()
                # Overdue issued count (by transactions)
                today = datetime.now().strftime('%Y-%m-%d')
                cur.execute("SELECT COUNT(*) FROM borrow_records WHERE status='borrowed' AND due_date < ?", (today,))
                overdue = cur.fetchone()[0] or 0
                conn.close()

            total_issued = max((total_copies or 0) - (total_available or 0), 0)

            on_time = max(total_issued - overdue, 0)

            # Check if there's any data to display - prevent NaN division
            if (total_copies or 0) == 0:
                # Show a placeholder message instead of empty chart
                placeholder = tk.Label(
                    self.inventory_overdue_frame,
                    text="📊 No book inventory data available.\nAdd books to see inventory breakdown.",
                    font=('Segoe UI', 12),
                    bg=self.colors['primary'],
                    fg='#666',
                    justify='center'
                )
                placeholder.pack(expand=True, fill='both', pady=40)
                return

            outer_labels = ["Available", "Issued"]
            outer_sizes = [total_available, total_issued]
            inner_labels = ["Available", "On-time", "Overdue"]
            inner_sizes = [total_available, on_time, overdue]


            # Colors
            outer_colors = ['#2ed573', '#ff9f43']
            inner_colors = ['#7bed9f', '#ffa502', '#ff4757']

            # Build figure
            fig = Figure(figsize=(6, 4), dpi=100)
            ax = fig.add_subplot(111)
            fig.patch.set_facecolor('white')

            # Outer ring
            res1 = ax.pie(outer_sizes, radius=1.0, labels=outer_labels, labeldistance=1.05,
                                colors=outer_colors, startangle=90, wedgeprops=dict(width=0.3, edgecolor='white'))
            wedges1 = res1[0]

            # Inner ring
            def _autopct(pct, allvals=inner_sizes):
                total = sum(allvals)
                if total == 0:
                    return ""
                try:
                    # Handle cases where pct might be NaN
                    if pct != pct: # Check for NaN
                        return ""
                    val = int(round(pct*total/100.0))
                    return f"{pct:.1f}%\n({val})"
                except:
                    return ""
            res2 = ax.pie(inner_sizes, radius=1.0-0.3, labels=None,
                                   colors=inner_colors, startangle=90,
                                   autopct=_autopct,
                                   wedgeprops=dict(width=0.3, edgecolor='white'))
            wedges2 = res2[0]
            # Center text
            ax.text(0, 0, f"Total\n{int(total_copies or 0)}", ha='center', va='center', fontsize=11, fontweight='bold')
            ax.set_title('Inventory & Overdue Breakdown', fontsize=12, fontweight='bold')

            # Legend shows inner ring details
            ax.legend(wedges2, inner_labels, title="Details", loc='center left', bbox_to_anchor=(1.0, 0.5))

            canvas = FigureCanvasTkAgg(fig, self.inventory_overdue_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            try:
                self._analysis_bind_wheel(canvas.get_tk_widget())
            except Exception:
                pass

            self.current_charts['inventory_overdue_donut'] = (fig, inner_labels, inner_sizes)
        except Exception as e:
            print(f"Error creating inventory/overdue donut: {e}")

    def create_branch_activity_pie(self, days, data=None):
        """Create pie/donut chart showing branch-wise borrowing distribution"""
        try:
            if data is not None:
                results = data
            else:
                conn = self.db.get_connection()
                cursor = conn.cursor()
                start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
                cursor.execute(
                    "SELECT s.department, COUNT(br.id) as borrow_count FROM students s "
                    "JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
                    "WHERE br.borrow_date >= ? AND s.department IS NOT NULL AND s.department != '' "
                    "GROUP BY s.department ORDER BY borrow_count DESC", (start_date,))
                results = cursor.fetchall()
                conn.close()

            if not results:
                fig = Figure(figsize=(6, 4), dpi=100)
                ax = fig.add_subplot(111)
                fig.patch.set_facecolor('white')
                ax.pie([1], labels=None, colors=['#d0d7de'], startangle=90,
                       wedgeprops=dict(width=0.32, edgecolor='white'))
                ax.text(0, 0, "No Data", ha='center', va='center', fontsize=14, fontweight='bold', color='#666')
                ax.axis('equal')
                ax.set_title(f'Branch-wise Borrowing (Last {days} Days)', fontsize=14, fontweight='bold', color='#0072B2')
                for w in self.branch_activity_frame.winfo_children():
                    w.destroy()
                canvas = FigureCanvasTkAgg(fig, self.branch_activity_frame)
                canvas.draw()
                canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
                try: self._analysis_bind_wheel(canvas.get_tk_widget())
                except: pass
                self.current_charts['branch_activity'] = (fig, ["No Data"], [1])
                return

            def _short_branch(name):
                """Shorten branch names for chart labels"""
                name = str(name).strip()
                mapping = {
                    'Computer Engineering': 'CO',
                    'Information Technology': 'IT',
                    'Civil Engineering': 'CE',
                    'Electronics & Telecommunication': 'E&TC',
                    'Mechanical Engineering': 'ME',
                    'Automobile Engineering': 'AE',
                    'Electrical Engineering': 'EE',
                }
                return mapping.get(name, name[:6])

            raw_labels = [str(row[0]) for row in results]
            short_labels = [_short_branch(row[0]) for row in results]
            sizes = [row[1] for row in results]
            legend_labels = [f"{_short_branch(r[0])} - {r[0]} ({r[1]})" for r in results]
            colors = ['#0072B2', '#D55E00', '#F0E442', '#009E73', '#CC79A7', '#56B4E9', '#E69F00', '#999999']

            fig = Figure(figsize=(6, 4), dpi=100)
            ax = fig.add_subplot(111)
            fig.patch.set_facecolor('white')
            fig.subplots_adjust(left=0.05, right=0.65)

            def _autopct(pct, allvals=sizes):
                total = sum(allvals)
                if total == 0: return ""
                val = int(round(pct * total / 100.0))
                return f"{pct:.1f}%\n({val})"

            wedges, texts, autotexts = ax.pie(
                sizes, labels=short_labels, colors=colors[:len(sizes)],
                autopct=_autopct, startangle=90,
                wedgeprops=dict(width=0.35, edgecolor='white'),
                textprops={'fontsize': 9, 'weight': 'bold'}
            )
            for at in autotexts:
                at.set_color('black')
            ax.text(0, 0, f"Total\n{sum(sizes)}", ha='center', va='center', fontsize=13, fontweight='bold', color='#333')
            ax.axis('equal')
            ax.set_title(f'Branch-wise Borrowing (Last {days} Days)', fontsize=14, fontweight='bold', color='#0072B2')
            ax.legend(wedges, legend_labels, title="Branch", loc='center left', bbox_to_anchor=(1.0, 0.5), fontsize=8)

            def on_branch_click(event):
                if event.inaxes == ax:
                    for i, wedge in enumerate(wedges):
                        contains, _ = wedge.contains(event)
                        if contains:
                            self.show_branch_students_dialog(raw_labels[i], days)
                            break

            for w in self.branch_activity_frame.winfo_children():
                w.destroy()
            canvas = FigureCanvasTkAgg(fig, self.branch_activity_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            canvas.mpl_connect('button_press_event', on_branch_click)
            try: self._analysis_bind_wheel(canvas.get_tk_widget())
            except: pass
            self.current_charts['branch_activity'] = (fig, raw_labels, sizes)
        except Exception as e:
            print(f"Error creating branch activity pie: {e}")

    def create_branch_overdue_pie(self, data=None):
        """Create pie/donut chart showing branch-wise overdue distribution"""
        try:
            if data is not None:
                results = data
            else:
                conn = self.db.get_connection()
                cursor = conn.cursor()
                today = datetime.now().strftime('%Y-%m-%d')
                cursor.execute(
                    "SELECT s.department, COUNT(br.id) as overdue_count FROM students s "
                    "JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
                    "WHERE br.status = 'borrowed' AND br.due_date < ? AND s.department IS NOT NULL AND s.department != '' "
                    "GROUP BY s.department ORDER BY overdue_count DESC", (today,))
                results = cursor.fetchall()
                conn.close()

            if not results:
                fig = Figure(figsize=(6, 4), dpi=100)
                ax = fig.add_subplot(111)
                fig.patch.set_facecolor('white')
                ax.pie([1], labels=None, colors=['#7bed9f'], startangle=90,
                       wedgeprops=dict(width=0.32, edgecolor='white'))
                ax.text(0, 0, "No\nOverdue", ha='center', va='center', fontsize=14, fontweight='bold', color='#2ed573')
                ax.axis('equal')
                ax.set_title('Branch-wise Overdue', fontsize=14, fontweight='bold', color='#D55E00')
                for w in self.branch_overdue_frame.winfo_children():
                    w.destroy()
                canvas = FigureCanvasTkAgg(fig, self.branch_overdue_frame)
                canvas.draw()
                canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
                try: self._analysis_bind_wheel(canvas.get_tk_widget())
                except: pass
                self.current_charts['branch_overdue'] = (fig, ["No Overdue"], [0])
                return

            def _short_branch(name):
                name = str(name).strip()
                mapping = {
                    'Computer Engineering': 'CO',
                    'Information Technology': 'IT',
                    'Civil Engineering': 'CE',
                    'Electronics & Telecommunication': 'E&TC',
                    'Mechanical Engineering': 'ME',
                    'Automobile Engineering': 'AE',
                    'Electrical Engineering': 'EE',
                }
                return mapping.get(name, name[:6])

            raw_labels = [str(row[0]) for row in results]
            short_labels = [_short_branch(row[0]) for row in results]
            sizes = [row[1] for row in results]
            legend_labels = [f"{_short_branch(r[0])} - {r[0]} ({r[1]})" for r in results]
            colors = ['#ff4757', '#ff6b81', '#ff9f43', '#ffa502', '#e74c3c', '#c0392b', '#e84118', '#d63031']

            fig = Figure(figsize=(6, 4), dpi=100)
            ax = fig.add_subplot(111)
            fig.patch.set_facecolor('white')
            fig.subplots_adjust(left=0.05, right=0.65)

            def _autopct(pct, allvals=sizes):
                total = sum(allvals)
                if total == 0: return ""
                val = int(round(pct * total / 100.0))
                return f"{pct:.1f}%\n({val})"

            wedges, texts, autotexts = ax.pie(
                sizes, labels=short_labels, colors=colors[:len(sizes)],
                autopct=_autopct, startangle=90,
                wedgeprops=dict(width=0.35, edgecolor='white'),
                textprops={'fontsize': 9, 'weight': 'bold'}
            )
            for at in autotexts:
                at.set_color('black')
            ax.text(0, 0, f"Total\n{sum(sizes)}", ha='center', va='center', fontsize=13, fontweight='bold', color='#c0392b')
            ax.axis('equal')
            ax.set_title('Branch-wise Overdue Distribution', fontsize=14, fontweight='bold', color='#D55E00')
            ax.legend(wedges, legend_labels, title="Branch", loc='center left', bbox_to_anchor=(1.0, 0.5), fontsize=8)

            def on_overdue_branch_click(event):
                if event.inaxes == ax:
                    for i, wedge in enumerate(wedges):
                        contains, _ = wedge.contains(event)
                        if contains:
                            self.show_branch_overdue_dialog(raw_labels[i])
                            break

            for w in self.branch_overdue_frame.winfo_children():
                w.destroy()
            canvas = FigureCanvasTkAgg(fig, self.branch_overdue_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            canvas.mpl_connect('button_press_event', on_overdue_branch_click)
            try: self._analysis_bind_wheel(canvas.get_tk_widget())
            except: pass
            self.current_charts['branch_overdue'] = (fig, raw_labels, sizes)
        except Exception as e:
            print(f"Error creating branch overdue pie: {e}")

    def show_branch_students_dialog(self, branch, days):
        """Show students who borrowed books from a specific branch in the period"""
        start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
        conn = self.db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT s.enrollment_no, s.name, s.year, COUNT(br.id) as borrows "
            "FROM students s JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
            "WHERE br.borrow_date >= ? AND s.department = ? "
            "GROUP BY s.enrollment_no, s.name, s.year ORDER BY borrows DESC",
            (start_date, branch))
        rows = cur.fetchall()
        conn.close()
        cols = ("Enrollment No", "Name", "Year", "Books Borrowed")
        self._show_table_dialog(f"Students from {branch} (Last {days}d)", cols, rows, export_name=f"branch_{branch}")

    def show_branch_overdue_dialog(self, branch):
        """Show overdue students from a specific branch"""
        today = datetime.now().strftime('%Y-%m-%d')
        conn = self.db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT s.enrollment_no, s.name, s.year, b.title, br.borrow_date, br.due_date "
            "FROM students s JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
            "JOIN books b ON br.book_id = b.book_id "
            "WHERE br.status = 'borrowed' AND br.due_date < ? AND s.department = ? "
            "ORDER BY br.due_date ASC",
            (today, branch))
        rows = cur.fetchall()
        conn.close()
        cols = ("Enrollment No", "Name", "Year", "Book Title", "Issue Date", "Due Date")
        self._show_table_dialog(f"Overdue from {branch}", cols, rows, export_name=f"overdue_{branch}")
    
    def create_daily_trend_chart(self, days, data=None):
        """Create bar chart showing daily borrowing trends"""
        try:
            # Get data
            if data is not None:
                results = data
            else:
                conn = self.db.get_connection()
                cursor = conn.cursor()
                
                start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
                
                cursor.execute("""
                    SELECT 
                        borrow_date,
                        COUNT(*) as daily_count
                    FROM borrow_records 
                    WHERE borrow_date >= ?
                    GROUP BY borrow_date
                    ORDER BY borrow_date
                """, (start_date,))
                
                results = cursor.fetchall()
                conn.close()
            
            if not results:
                no_data_label = tk.Label(
                    self.daily_trend_frame,
                    text=f"No borrowing activity\nin last {days} days",
                    font=('Segoe UI', 12),
                    bg=self.colors['primary'],
                    fg='#666666'
                )
                no_data_label.pack(expand=True)
                return
            
            dates = [row[0] for row in results]
            counts = [row[1] for row in results]
            
            # Create figure
            fig = Figure(figsize=(6, 4), dpi=100)
            ax = fig.add_subplot(111)
            
            bars = ax.bar(dates, counts, color='#45b7d1', alpha=0.7)
            ax.set_title(f'Daily Borrowing Trends (Last {days} Days)', fontsize=12, fontweight='bold')
            ax.set_xlabel('Date')
            ax.set_ylabel('Books Issued')
            
            # Rotate x-axis labels for better readability
            plt.setp(ax.get_xticklabels(), rotation=45, ha='right')
            
            # Add click interaction
            def on_bar_click(event):
                if event.inaxes == ax:
                    for i, bar in enumerate(bars):
                        contains, info = bar.contains(event)
                        if contains:
                            self.show_borrow_details_for_date(dates[i])
                            break
            
            fig.tight_layout()
            
            # Add to GUI
            canvas = FigureCanvasTkAgg(fig, self.daily_trend_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            canvas.mpl_connect('button_press_event', on_bar_click)
            
            self.current_charts['daily_trend'] = (fig, dates, counts)
            
        except Exception as e:
            print(f"Error creating daily trend chart: {e}")
    
    def create_popular_books_chart(self, days, data=None):
        """Create bar chart showing most popular books"""
        try:
            # Get data
            if data is not None:
                results = data
            else:
                conn = self.db.get_connection()
                cursor = conn.cursor()
                
                start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
                
                cursor.execute("""
                    SELECT 
                        b.title,
                        COUNT(br.id) as borrow_count
                    FROM books b
                    INNER JOIN borrow_records br ON b.book_id = br.book_id
                    WHERE br.borrow_date >= ?
                    GROUP BY b.book_id, b.title
                    ORDER BY borrow_count DESC, b.title ASC
                    LIMIT 10
                """, (start_date,))
                
                results = cursor.fetchall()
                conn.close()
            
            if not results:
                no_data_label = tk.Label(
                    self.popular_books_frame,
                    text=f"No borrowing activity\nin last {days} days",
                    font=('Segoe UI', 12),
                    bg=self.colors['primary'],
                    fg='#666666'
                )
                no_data_label.pack(expand=True)
                return
            
            titles = [row[0][:20] + ('...' if len(row[0]) > 20 else '') for row in results]
            counts = [row[1] for row in results]
            
            # Create figure
            fig = Figure(figsize=(6, 4), dpi=100)
            ax = fig.add_subplot(111)
            
            bars = ax.barh(titles, counts, color='#f9ca24', alpha=0.7)
            ax.set_title(f'Most Popular Books (Last {days} Days)', fontsize=12, fontweight='bold')
            ax.set_xlabel('Times Issued')
            
            # Add click interaction
            def on_popular_click(event):
                if event.inaxes == ax:
                    for i, bar in enumerate(bars):
                        contains, info = bar.contains(event)
                        if contains:
                            # Show all students who borrowed that book in period
                            full_title = results[i][0]
                            self.show_book_borrowers_dialog(full_title, days)
                            break
            
            fig.tight_layout()
            
            # Add to GUI
            canvas = FigureCanvasTkAgg(fig, self.popular_books_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            canvas.mpl_connect('button_press_event', on_popular_click)
            
            self.current_charts['popular_books'] = (fig, titles, counts)
            
        except Exception as e:
            print(f"Error creating popular books chart: {e}")

    def create_least_popular_books_chart(self, days, data=None):
        """Create bar chart showing least popular books (books with least borrows or zero borrows)"""
        try:
            # Get data
            if data is not None:
                results = data
            else:
                conn = self.db.get_connection()
                cursor = conn.cursor()
                
                start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
                
                # Get books with lowest borrow count in the period (prioritize zero borrows)
                # First get the minimum count that's been borrowed (to exclude books with higher counts)
                cursor.execute("""
                    SELECT 
                        b.title,
                        COALESCE(COUNT(br.id), 0) as borrow_count
                    FROM books b
                    LEFT JOIN borrow_records br ON b.book_id = br.book_id AND br.borrow_date >= ?
                    GROUP BY b.book_id, b.title
                    HAVING COALESCE(COUNT(br.id), 0) = (
                        SELECT MIN(cnt) FROM (
                            SELECT COALESCE(COUNT(br2.id), 0) as cnt
                            FROM books b2
                            LEFT JOIN borrow_records br2 ON b2.book_id = br2.book_id AND br2.borrow_date >= ?
                            GROUP BY b2.book_id
                        )
                    )
                    ORDER BY b.title ASC
                    LIMIT 10
                """, (start_date, start_date))
                
                results = cursor.fetchall()
                conn.close()
            
            if not results:
                no_data_label = tk.Label(
                    self.least_popular_books_frame,
                    text=f"No book data available",
                    font=('Segoe UI', 12),
                    bg=self.colors['primary'],
                    fg='#666666'
                )
                no_data_label.pack(expand=True)
                return
            
            titles = [row[0][:20] + ('...' if len(row[0]) > 20 else '') for row in results]
            counts = [row[1] for row in results]
            
            # Create figure
            fig = Figure(figsize=(6, 4), dpi=100)
            ax = fig.add_subplot(111)
            
            bars = ax.barh(titles, counts, color='#e74c3c', alpha=0.7)
            ax.set_title(f'Least Popular Books (Last {days} Days)', fontsize=12, fontweight='bold')
            ax.set_xlabel('Times Issued')
            
            # Add click interaction
            def on_least_popular_click(event):
                if event.inaxes == ax:
                    for i, bar in enumerate(bars):
                        contains, info = bar.contains(event)
                        if contains:
                            # Show all students who borrowed that book in period (if any)
                            full_title = results[i][0]
                            self.show_book_borrowers_dialog(full_title, days)
                            break
            
            fig.tight_layout()
            
            # Add to GUI
            canvas = FigureCanvasTkAgg(fig, self.least_popular_books_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            canvas.mpl_connect('button_press_event', on_least_popular_click)
            
            self.current_charts['least_popular_books'] = (fig, titles, counts)
            
        except Exception as e:
            print(f"Error creating least popular books chart: {e}")

    # ---------------------- Focused Insights ----------------------
    def create_student_specific_pie(self, days, enrollment_no, data=None):
        """Pie: student's borrow status in period (borrowed vs returned)."""
        try:
            if data:
                active = data.get('active', 0)
                returned = data.get('returned', 0)
                # total = data.get('total', 0)
            else:
                conn = self.db.get_connection()
                cur = conn.cursor()
                start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
                cur.execute("SELECT COUNT(*) FROM borrow_records WHERE enrollment_no=? AND borrow_date>=?", (enrollment_no, start_date))
                total = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM borrow_records WHERE enrollment_no=? AND borrow_date>=? AND status='borrowed'", (enrollment_no, start_date))
                active = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM borrow_records WHERE enrollment_no=? AND return_date>=? AND return_date IS NOT NULL", (enrollment_no, start_date))
                returned = cur.fetchone()[0]
                conn.close()
            sizes = [active, returned]
            labels = ["Currently Issued", "Returned"]
            if sum(sizes) == 0:
                lbl = tk.Label(self.student_specific_frame, text=f"No activity for {enrollment_no} in last {days} days", bg=self.colors['primary'], fg='#666', font=('Segoe UI', 11))
                lbl.pack(fill=tk.X, padx=10, pady=10)
            else:
                fig = Figure(figsize=(4.5, 3.5), dpi=100)
                ax = fig.add_subplot(111)
                ax.pie(sizes, labels=labels, colors=['#ff9f43', '#10ac84'], autopct='%1.1f%%', startangle=90)
                ax.set_title(f"Student {enrollment_no} - Status (Last {days}d)", fontsize=11, fontweight='bold')
                canvas = FigureCanvasTkAgg(fig, self.student_specific_frame)
                canvas.draw()
                canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
                self.current_charts['student_specific_status'] = (fig, labels, sizes)

            # packing handled in refresh_analysis
        except Exception as e:
            print(f"Student-specific pie error: {e}")

    def create_book_specific_pie(self, days, book_id, data=None):
        """Pie: book's copies status currently (available vs borrowed)."""
        try:
            if data:
                row = data.get('row')
            else:
                conn = self.db.get_connection()
                cur = conn.cursor()
                cur.execute("SELECT title, total_copies, available_copies FROM books WHERE book_id=?", (book_id,))
                row = cur.fetchone()
                conn.close()
            if not row:
                lbl = tk.Label(self.book_specific_frame, text=f"Book {book_id} not found", bg=self.colors['primary'], fg='#c00', font=('Segoe UI', 11, 'bold'))
                lbl.pack(fill=tk.X, padx=10, pady=10)
            else:
                title, total, avail = row
                borrowed = max(total - (avail or 0), 0)
                sizes = [avail or 0, borrowed]
                labels = ["Available", "Borrowed"]
                fig = Figure(figsize=(4.5, 3.5), dpi=100)
                ax = fig.add_subplot(111)
                ax.pie(sizes, labels=labels, colors=['#2ed573', '#ff4757'], autopct='%1.1f%%', startangle=90)
                ax.set_title(f"Book {book_id} - Copies Status", fontsize=11, fontweight='bold')
                canvas = FigureCanvasTkAgg(fig, self.book_specific_frame)
                canvas.draw()
                canvas.get_tk_widget().pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
                self.current_charts['book_specific_status'] = (fig, labels, sizes)

            # packing handled in refresh_analysis
        except Exception as e:
            print(f"Book-specific pie error: {e}")

    # ---------------------- Drill-down Dialogs ----------------------
    def show_currently_borrowed_dialog(self):
        """Show all currently borrowed items with student and book details, respecting filters if set."""
        en = self.analysis_filter.get('enrollment_no')
        bk = self.analysis_filter.get('book_id')
        conn = self.db.get_connection()
        cur = conn.cursor()
        base = (
            "SELECT br.enrollment_no, s.name, br.book_id, b.title, br.borrow_date, br.due_date "
            "FROM borrow_records br JOIN students s ON br.enrollment_no=s.enrollment_no "
            "JOIN books b ON br.book_id=b.book_id WHERE br.status='borrowed'"
        )
        params = []
        if en:
            base += " AND br.enrollment_no=?"
            params.append(en)
        if bk:
            base += " AND br.book_id=?"
            params.append(bk)
        base += " ORDER BY br.due_date"
        cur.execute(base, tuple(params))
        rows = cur.fetchall()
        conn.close()
        cols = ("Enrollment No", "Student Name", "Book ID", "Book Title", "Issue Date", "Due Date")
        self._show_table_dialog("Currently Issued", cols, rows, export_name="currently_issued")

    def show_available_books_dialog(self):
        """Show all books with available copies > 0 (optionally filtered by book_id)."""
        bk = self.analysis_filter.get('book_id')
        conn = self.db.get_connection()
        cur = conn.cursor()
        if bk:
            cur.execute("SELECT book_id, title, author, available_copies FROM books WHERE book_id=?", (bk,))
        else:
            cur.execute("SELECT book_id, title, author, available_copies FROM books WHERE available_copies>0 ORDER BY title")
        rows = cur.fetchall()
        conn.close()
        cols = ("Book ID", "Title", "Author", "Available Copies")
        self._show_table_dialog("Available Books", cols, rows, export_name="available_books")

    def show_students_by_year_dialog(self, year, days):
        start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
        conn = self.db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT s.enrollment_no, s.name, COUNT(br.id) as borrows FROM students s "
            "LEFT JOIN borrow_records br ON s.enrollment_no=br.enrollment_no AND br.borrow_date>=? "
            "WHERE s.year=? GROUP BY s.enrollment_no, s.name HAVING COUNT(br.id)>0 ORDER BY borrows DESC",
            (start_date, year)
        )
        rows = cur.fetchall()
        conn.close()
        cols = ("Enrollment No", "Name", "Issue Count")
        self._show_table_dialog(f"Students in {year} Year", cols, rows, export_name=f"students_year_{year}")

    def show_borrow_details_for_date(self, date_str):
        en = self.analysis_filter.get('enrollment_no')
        bk = self.analysis_filter.get('book_id')
        conn = self.db.get_connection()
        cur = conn.cursor()
        q = (
            "SELECT br.borrow_date, br.enrollment_no, s.name, br.book_id, b.title FROM borrow_records br "
            "JOIN students s ON br.enrollment_no=s.enrollment_no JOIN books b ON br.book_id=b.book_id "
            "WHERE br.borrow_date=?"
        )
        params = [date_str]
        if en:
            q += " AND br.enrollment_no=?"; params.append(en)
        if bk:
            q += " AND br.book_id=?"; params.append(bk)
        q += " ORDER BY s.name"
        cur.execute(q, tuple(params))
        rows = cur.fetchall()
        conn.close()
        cols = ("Issue Date", "Enrollment No", "Student Name", "Book ID", "Book Title")
        self._show_table_dialog(f"Issues on {date_str}", cols, rows, export_name=f"issues_{date_str}")

    def show_book_borrowers_dialog(self, title, days):
        start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
        conn = self.db.get_connection()
        cur = conn.cursor()
        cur.execute(
            "SELECT br.enrollment_no, s.name, COUNT(br.id) as n FROM borrow_records br "
            "JOIN students s ON br.enrollment_no=s.enrollment_no JOIN books b ON br.book_id=b.book_id "
            "WHERE b.title=? AND br.borrow_date>=? GROUP BY br.enrollment_no, s.name ORDER BY n DESC",
            (title, start_date)
        )
        rows = cur.fetchall()
        conn.close()
        cols = ("Enrollment No", "Student Name", "Times Issued")
        self._show_table_dialog(f"Issuers of '{title}' (Last {days}d)", cols, rows, export_name="book_issuers")

    # Generic dialog with export
    def _show_table_dialog(self, title, columns, rows, export_name="data_export"):
        dlg = tk.Toplevel(self.root)
        dlg.title(title)
        dlg.geometry("800x500")
        frm = tk.Frame(dlg, bg=self.colors['primary'])
        frm.pack(fill=tk.BOTH, expand=True)
        tv = ttk.Treeview(frm, columns=[f"c{i}" for i in range(len(columns))], show='headings')
        for i, c in enumerate(columns):
            tv.heading(f"c{i}", text=c)
            tv.column(f"c{i}", width=max(100, int(750/len(columns))))
        vsb = ttk.Scrollbar(frm, orient='vertical', command=tv.yview)
        tv.configure(yscroll=vsb.set)
        tv.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        for r in rows:
            # Convert sqlite3.Row objects to tuples for display
            if hasattr(r, 'keys') and not isinstance(r, (list, tuple)):
                tv.insert('', tk.END, values=tuple(r))
            else:
                tv.insert('', tk.END, values=r)
        btns = tk.Frame(dlg, bg=self.colors['primary'])
        btns.pack(fill=tk.X)
        def do_export():
            if not rows:
                messagebox.showinfo("Export", "No data to export.")
                return
            filename = filedialog.asksaveasfilename(defaultextension='.xlsx', filetypes=[["Excel", "*.xlsx"]], initialfile=f"{export_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
            if not filename:
                return
            try:
                import xlsxwriter
                wb = xlsxwriter.Workbook(filename)
                ws = wb.add_worksheet('Data')
                start = self._xlsxwriter_write_header(ws, wb, start_row=0)
                for j, c in enumerate(columns):
                    ws.write(start, j, c)
                for i, row in enumerate(rows, start=start+1):
                    for j, v in enumerate(row):
                        ws.write(i, j, v)
                wb.close()
                messagebox.showinfo("Export", f"Saved to {filename}")
                if messagebox.askyesno("Open File", "Open the exported file?"):
                    self.open_file(filename)
            except Exception as e:
                messagebox.showerror("Export", f"Failed to export: {e}")
        tk.Button(btns, text="Export to Excel", command=do_export, bg='#28a745', fg='white', relief='flat', padx=10, pady=6).pack(side=tk.RIGHT, padx=10, pady=8)

    # ---------------------- Filter handlers ----------------------
    def apply_analysis_filter(self):
        en = self.analysis_student_var.get().strip() or None
        bk = self.analysis_book_var.get().strip() or None
        # Validate existence lightly (non-blocking)
        if en and not self.get_student_name(en):
            if not messagebox.askyesno("Unknown Student", f"Enrollment {en} not found. Apply filter anyway?"):
                return
        if bk and not self.get_book_title(bk):
            if not messagebox.askyesno("Unknown Book", f"Book ID {bk} not found. Apply filter anyway?"):
                return
        self.analysis_filter = {'enrollment_no': en, 'book_id': bk}
        self.refresh_analysis()

    def clear_analysis_filter(self):
        self.analysis_student_var.set('')
        self.analysis_book_var.set('')
        if hasattr(self, 'analysis_branch_var'):
            self.analysis_branch_var.set('All')
        self.analysis_filter = {'enrollment_no': None, 'book_id': None}
        self.refresh_analysis()

    # ---------------------- Helpers ----------------------
    def get_student_name(self, enrollment_no):
        try:
            conn = self.db.get_connection()
            cur = conn.cursor()
            cur.execute("SELECT name FROM students WHERE enrollment_no=?", (enrollment_no,))
            row = cur.fetchone()
            conn.close()
            return row[0] if row else None
        except Exception:
            return None

    def get_book_title(self, book_id):
        try:
            conn = self.db.get_connection()
            cur = conn.cursor()
            cur.execute("SELECT title FROM books WHERE book_id=?", (book_id,))
            row = cur.fetchone()
            conn.close()
            return row[0] if row else None
        except Exception:
            return None
    
    def create_summary_stats(self, days, data=None):
        """Create summary statistics display"""
        try:
            # Get comprehensive stats
            if data:
                total_borrowings = data.get('total_borrowings', 0)
                total_returns = data.get('total_returns', 0)
                overdue_count = data.get('overdue_count', 0)
                active_students = data.get('active_students', 0)
                fines_raw = data.get('fines_data', [])
            else:
                conn = self.db.get_connection()
                cursor = conn.cursor()
                
                start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
                
                # Total borrowings in period
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE borrow_date >= ?", (start_date,))
                total_borrowings = cursor.fetchone()[0]
                
                # Total returns in period
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE return_date >= ? AND return_date IS NOT NULL", (start_date,))
                total_returns = cursor.fetchone()[0]
                
                # Currently overdue
                today = datetime.now().strftime('%Y-%m-%d')
                cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE status = 'borrowed' AND due_date < ?", (today,))
                overdue_count = cursor.fetchone()[0]
                
                # Active students (who borrowed in period)
                cursor.execute("SELECT COUNT(DISTINCT enrollment_no) FROM borrow_records WHERE borrow_date >= ?", (start_date,))
                active_students = cursor.fetchone()[0]
                
                # Total fines collected (approximation)
                # Calculate fines in Python to be DB-agnostic (avoid julianday vs EXTRACT differences)
                cursor.execute("""
                    SELECT return_date, due_date 
                    FROM borrow_records 
                    WHERE return_date > due_date AND return_date IS NOT NULL AND return_date >= ?
                """, (start_date,))
                fines_raw = cursor.fetchall()
                conn.close()
            
            fine_sum = 0
            fine_per_day = self.get_fine_per_day()
            for r_date, d_date in fines_raw:
                # Handle string dates vs date objects

                if isinstance(r_date, str): r_date = datetime.strptime(r_date, '%Y-%m-%d').date()
                if isinstance(d_date, str): d_date = datetime.strptime(d_date, '%Y-%m-%d').date()
                days_late = (r_date - d_date).days
                fine_sum += days_late * fine_per_day
                
            total_fines = fine_sum
            
            # Create stats display
            stats_container = tk.Frame(self.stats_summary_frame, bg=self.colors['primary'])
            stats_container.pack(fill=tk.X, padx=15, pady=15)
            
            # Show branch indicator if filtered
            br = self.analysis_branch_var.get() if hasattr(self, 'analysis_branch_var') else "All"
            branch_suffix = f" ({br})" if br and br != "All" else ""
            
            stats = [
                (f"📚 Total Borrowings{branch_suffix}", total_borrowings, "#4ecdc4"),
                (f"📥 Total Returns{branch_suffix}", total_returns, "#45b7d1"),
                (f"⚠️ Currently Overdue{branch_suffix}", overdue_count, "#ff6b6b"),
                (f"👥 Active Students{branch_suffix}", active_students, "#f9ca24"),
                (f"💰 Fines Collected (₹{self.get_fine_per_day()}/day)", f"₹{total_fines:.0f}", "#a55eea")
            ]
            
            for i, (label, value, color) in enumerate(stats):
                stat_frame = tk.Frame(stats_container, bg=color, relief='flat', bd=0)
                stat_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
                
                tk.Label(
                    stat_frame,
                    text=str(value),
                    font=('Segoe UI', 18, 'bold'),
                    bg=color,
                    fg='white'
                ).pack(pady=(15, 5))
                
                tk.Label(
                    stat_frame,
                    text=label,
                    font=('Segoe UI', 10, 'bold'),
                    bg=color,
                    fg='white'
                ).pack(pady=(0, 15))
            
        except Exception as e:
            print(f"Error creating summary stats: {e}")
    
    def export_analysis_excel(self):
        """Export analysis charts and data to Excel with embedded charts"""
        try:
            if not XLSXWRITER_AVAILABLE:
                messagebox.showerror("Export Error", "xlsxwriter package is required for Excel export with charts.\n\nPlease install: pip install xlsxwriter")
                return
            
            days = int(self.analysis_period.get())
            start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
            
            # File dialog
            filename = f"library_analysis_{days}days_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                initialfile=filename
            )
            
            if not file_path:
                return
            
            # Create workbook
            workbook = xlsxwriter.Workbook(file_path)
            
            # Add formats
            header_format = workbook.add_format({'bold': True, 'font_size': 14, 'bg_color': '#4ecdc4', 'font_color': 'white'})
            subheader_format = workbook.add_format({'bold': True, 'font_size': 12, 'bg_color': '#f0f0f0'})
            
            # Summary worksheet
            summary_ws = workbook.add_worksheet('Analysis Summary')
            # Three-line header requested by user
            row = self._xlsxwriter_write_header(summary_ws, workbook, start_row=0)
            summary_ws.write(row, 0, f'Library Analysis Report - Last {days} Days', header_format)
            summary_ws.write(row + 1, 0, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            row += 3
            
            # Get and write summary data
            conn = self.db.get_connection()
            cursor = conn.cursor()
            
            # Book status data
            summary_ws.write(row, 0, 'Book Status Distribution', subheader_format)
            row += 1
            
            cursor.execute("""
                SELECT 
                    CASE 
                        WHEN br.status = 'borrowed' THEN 'Currently Issued'
                        ELSE 'Available'
                    END as status,
                    COUNT(DISTINCT b.book_id) as count
                FROM books b
                LEFT JOIN borrow_records br ON b.book_id = br.book_id AND br.status = 'borrowed'
                GROUP BY status
            """)
            
            book_status_data = cursor.fetchall()
            for i, (status, count) in enumerate(book_status_data):
                summary_ws.write(row + i, 0, status)
                summary_ws.write(row + i, 1, count)
            
            row += len(book_status_data) + 2
            
            # Student activity data
            summary_ws.write(row, 0, f'Student Activity (Last {days} Days)', subheader_format)
            row += 1
            
            cursor.execute("""
                SELECT 
                    s.year,
                    COUNT(br.id) as borrow_count
                FROM students s
                LEFT JOIN borrow_records br ON s.enrollment_no = br.enrollment_no 
                    AND br.borrow_date >= ?
                GROUP BY s.year
                HAVING COUNT(br.id) > 0
                ORDER BY borrow_count DESC
            """, (start_date,))
            
            student_activity_data = cursor.fetchall()
            for i, (year, count) in enumerate(student_activity_data):
                summary_ws.write(row + i, 0, f"{year} Year")
                summary_ws.write(row + i, 1, count)
            
            row += len(student_activity_data) + 2

            # Branch-wise Borrowing Data
            summary_ws.write(row, 0, f'Branch-wise Borrowing (Last {days} Days)', subheader_format)
            row += 1
            summary_ws.write(row, 0, 'Branch')
            summary_ws.write(row, 1, 'Books Borrowed')
            row += 1
            
            cursor.execute(
                "SELECT s.department, COUNT(br.id) as borrow_count FROM students s "
                "JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
                "WHERE br.borrow_date >= ? AND s.department IS NOT NULL AND s.department != '' "
                "GROUP BY s.department ORDER BY borrow_count DESC", (start_date,))
            branch_data = cursor.fetchall()
            for i, (branch, count) in enumerate(branch_data):
                summary_ws.write(row + i, 0, str(branch))
                summary_ws.write(row + i, 1, count)
            
            row += len(branch_data) + 2

            # Branch-wise Overdue Data
            today = datetime.now().strftime('%Y-%m-%d')
            summary_ws.write(row, 0, 'Branch-wise Overdue', subheader_format)
            row += 1
            summary_ws.write(row, 0, 'Branch')
            summary_ws.write(row, 1, 'Overdue Count')
            row += 1

            cursor.execute(
                "SELECT s.department, COUNT(br.id) as overdue_count FROM students s "
                "JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
                "WHERE br.status = 'borrowed' AND br.due_date < ? AND s.department IS NOT NULL AND s.department != '' "
                "GROUP BY s.department ORDER BY overdue_count DESC", (today,))
            branch_overdue_data = cursor.fetchall()
            for i, (branch, count) in enumerate(branch_overdue_data):
                summary_ws.write(row + i, 0, str(branch))
                summary_ws.write(row + i, 1, count)

            conn.close()

            # Insert chart images for the current on-screen charts
            def fig_to_image_bytes(fig):
                bio = BytesIO()
                fig.savefig(bio, format='png', dpi=120, bbox_inches='tight')
                bio.seek(0)
                return bio

            img_row = row + len(branch_overdue_data) + 3
            summary_ws.write(img_row, 0, 'Charts Snapshot', subheader_format)
            img_row += 1
            col = 0
            for key in ['borrow_status', 'student_activity', 'branch_activity', 'branch_overdue', 'inventory_overdue_donut', 'daily_trend', 'popular_books', 'student_specific_status', 'book_specific_status']:
                if key in self.current_charts and isinstance(self.current_charts[key][0], Figure):
                    fig = self.current_charts[key][0]
                    try:
                        bio = fig_to_image_bytes(fig)
                        summary_ws.insert_image(img_row, col, f"{key}.png", {'image_data': bio, 'x_scale': 0.9, 'y_scale': 0.9})
                        col += 8  # move to the right for next image
                        if col > 8:
                            img_row += 20
                            col = 0
                    except Exception as e:
                        print(f"Failed to insert image for {key}: {e}")

            workbook.close()
            
            messagebox.showinfo("Export Successful", f"Analysis exported to:\n{file_path}")
            
            # Ask if user wants to open the file
            if messagebox.askyesno("Open File", "Do you want to open the exported file?"):
                self.open_file(file_path)
                
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export analysis: {str(e)}")
    
    def export_analysis_word(self):
        """Export analysis summary to Word document"""
        try:
            if Document is None:
                messagebox.showerror("Export Error", "python-docx package is required for Word export.\n\nPlease install: pip install python-docx")
                return
            
            days = int(self.analysis_period.get())
            start_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
            
            # File dialog
            filename = f"library_analysis_{days}days_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                initialfile=filename
            )
            
            if not file_path:
                return
            
            # Create document
            doc = Document()
            
            # Add logo at the top if available
            logo_path = os.path.join(os.path.dirname(__file__), 'logo.png')
            if os.path.exists(logo_path):
                try:
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(logo_path, width=Pt(60))
                except Exception as e:
                    print(f"Could not add logo: {e}")
            
            # Institutional Header with colors
            def add_colored_center(text, size, rgb):
                from docx.shared import RGBColor
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(size)
                run.font.color.rgb = RGBColor(*rgb)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            add_colored_center("Government Polytechnic Awasari (Kh)", 22, (31, 71, 136))
            add_colored_center("Main Library", 18, (46, 92, 138))
            
            # Add separator line
            doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Header
            header = doc.add_heading(f'Library Analysis Report', 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f'Analysis Period: Last {days} Days')
            doc.add_paragraph(f'Date Range: {start_date} to {datetime.now().strftime("%Y-%m-%d")}')
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('')
            
            # Get data
            conn = self.db.get_connection()
            cursor = conn.cursor()
            
            # Summary statistics
            doc.add_heading('Summary Statistics', level=1)
            
            cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE borrow_date >= ?", (start_date,))
            total_borrowings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE return_date >= ? AND return_date IS NOT NULL", (start_date,))
            total_returns = cursor.fetchone()[0]
            
            today = datetime.now().strftime('%Y-%m-%d')
            cursor.execute("SELECT COUNT(*) FROM borrow_records WHERE status = 'borrowed' AND due_date < ?", (today,))
            overdue_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(DISTINCT enrollment_no) FROM borrow_records WHERE borrow_date >= ?", (start_date,))
            active_students = cursor.fetchone()[0]
            
            stats_table = doc.add_table(rows=5, cols=2)
            stats_table.style = 'Table Grid'
            
            stats_data = [
                ('Total Borrowings', total_borrowings),
                ('Total Returns', total_returns), 
                ('Currently Overdue', overdue_count),
                ('Active Students', active_students),
                ('Analysis Period', f'{days} days')
            ]
            
            for i, (metric, value) in enumerate(stats_data):
                stats_table.cell(i, 0).text = metric
                stats_table.cell(i, 1).text = str(value)
            
            # Student activity by year
            doc.add_heading('Student Activity by Year', level=1)
            
            cursor.execute("""
                SELECT 
                    s.year,
                    COUNT(br.id) as borrow_count
                FROM students s
                LEFT JOIN borrow_records br ON s.enrollment_no = br.enrollment_no 
                    AND br.borrow_date >= ?
                GROUP BY s.year
                HAVING COUNT(br.id) > 0
                ORDER BY borrow_count DESC
            """, (start_date,))
            
            year_activity = cursor.fetchall()
            
            if year_activity:
                year_table = doc.add_table(rows=len(year_activity) + 1, cols=2)
                year_table.style = 'Table Grid'
                year_table.cell(0, 0).text = 'Academic Year'
                year_table.cell(0, 1).text = 'Books Borrowed'
                
                for i, (year, count) in enumerate(year_activity):
                    year_table.cell(i + 1, 0).text = f"{year} Year"
                    year_table.cell(i + 1, 1).text = str(count)
            else:
                doc.add_paragraph('No student activity recorded in the selected period.')
            
            # Branch-wise Borrowing Distribution
            doc.add_heading('Branch-wise Borrowing Distribution', level=1)
            
            cursor.execute(
                "SELECT s.department, COUNT(br.id) as borrow_count FROM students s "
                "JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
                "WHERE br.borrow_date >= ? AND s.department IS NOT NULL AND s.department != '' "
                "GROUP BY s.department ORDER BY borrow_count DESC", (start_date,))
            branch_activity = cursor.fetchall()
            
            if branch_activity:
                branch_table = doc.add_table(rows=len(branch_activity) + 1, cols=2)
                branch_table.style = 'Table Grid'
                branch_table.cell(0, 0).text = 'Branch'
                branch_table.cell(0, 1).text = 'Books Borrowed'
                
                for i, (branch, count) in enumerate(branch_activity):
                    branch_table.cell(i + 1, 0).text = str(branch)
                    branch_table.cell(i + 1, 1).text = str(count)
            else:
                doc.add_paragraph('No branch-wise borrowing data available.')

            # Branch-wise Overdue Distribution
            doc.add_heading('Branch-wise Overdue Distribution', level=1)
            
            cursor.execute(
                "SELECT s.department, COUNT(br.id) as overdue_count FROM students s "
                "JOIN borrow_records br ON s.enrollment_no = br.enrollment_no "
                "WHERE br.status = 'borrowed' AND br.due_date < ? AND s.department IS NOT NULL AND s.department != '' "
                "GROUP BY s.department ORDER BY overdue_count DESC", (today,))
            branch_overdue = cursor.fetchall()
            
            if branch_overdue:
                overdue_table = doc.add_table(rows=len(branch_overdue) + 1, cols=2)
                overdue_table.style = 'Table Grid'
                overdue_table.cell(0, 0).text = 'Branch'
                overdue_table.cell(0, 1).text = 'Overdue Count'
                
                for i, (branch, count) in enumerate(branch_overdue):
                    overdue_table.cell(i + 1, 0).text = str(branch)
                    overdue_table.cell(i + 1, 1).text = str(count)
            else:
                doc.add_paragraph('No overdue records across branches.')
            
            conn.close()
            
            # Save document
            doc.save(file_path)
            
            messagebox.showinfo("Export Successful", f"Analysis exported to:\n{file_path}")
            
            # Ask if user wants to open the file
            if messagebox.askyesno("Open File", "Do you want to open the exported file?"):
                self.open_file(file_path)
                
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export analysis: {str(e)}")

    def start_student_portal(self):
        """Start the Flask server in a daemon thread"""
        if not WEB_PORTAL_AVAILABLE:
            print("Web portal is not available. Install flask and waitress.")
            return
        
        if self.portal_thread and self.portal_thread.is_alive():
            return

        def run_server():
            # Use waitress for production-ready stable server
            # Listen on all interfaces (0.0.0.0) so other devices can access
            print(f"Starting Student Portal on port {self.portal_port}...")
            # increased threads to 50 to handle concurrent requests (approx 500 users @ 10% active)
            serve(flask_app, host='0.0.0.0', port=self.portal_port, threads=50)

        self.portal_thread = threading.Thread(target=run_server, daemon=True)
        self.portal_thread.start()

    def show_qr_code(self):
        """Generate and show QR code for the student portal"""
        if not WEB_PORTAL_AVAILABLE:
            messagebox.showwarning("Web Portal Unavailable", 
                                 "Web portal is not available.\n\nPlease install required packages: flask, waitress")
            return
        
        # Ensure server is running
        self.start_student_portal()
        
        # Get local IP address
        try:
            # Connect to an external server (doesn't send data) to get the preferred local IP
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            local_ip = s.getsockname()[0]
            s.close()
        except:
            local_ip = "127.0.0.1"
            
        url = f"http://{local_ip}:{self.portal_port}"
        
        # Create dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("📱 Student Mobile Portal")
        dialog.geometry("500x650")
        dialog.configure(bg='white')
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center dialog
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 250
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 325
        dialog.geometry(f"+{x}+{y}")
        
        # Header
        tk.Label(dialog, text="📱 Student Portal", font=('Segoe UI', 22, 'bold'),
                 bg='white', fg=self.colors['accent']).pack(pady=(40, 5))
        
        tk.Label(dialog, text="Scan with mobile phone to access", font=('Segoe UI', 12),
                 bg='white', fg='#666').pack(pady=(0, 30))
        
        # QR Code
        qr = qrcode.QRCode(box_size=10, border=2)
        qr.add_data(url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Convert to PhotoImage
        img_tk = ImageTk.PhotoImage(img)
        qr_label = tk.Label(dialog, image=img_tk, bg='white')
        qr_label.image = img_tk  # Keep reference
        qr_label.pack(pady=10)
        
        # URL Text
        url_frame = tk.Frame(dialog, bg='#f8f9fa', padx=20, pady=10, relief='solid', bd=1)
        url_frame.pack(pady=20, fill=tk.X, padx=50)
        
        tk.Label(url_frame, text="Local Network URL:", font=('Segoe UI', 10, 'bold'),
                 bg='#f8f9fa', fg='#666').pack()
        
        tk.Label(url_frame, text=url, font=('Segoe UI', 14, 'bold'),
                 bg='#f8f9fa', fg=self.colors['secondary']).pack(pady=(5, 0))
        
        # Instructions
        tk.Label(dialog, text="Note: Both devices must be on the same Wi-Fi network.",
                 font=('Segoe UI', 10, 'italic'), bg='white', fg='#999').pack(side=tk.BOTTOM, pady=20)


# Main application entry point
if __name__ == "__main__":
    root = tk.Tk()
    app = LibraryApp(root)
    root.mainloop()
